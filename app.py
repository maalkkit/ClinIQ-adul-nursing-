import json
import re
import ast
import os
import csv
import io
import math
import html as _html  # used for safe HTML escaping in Grade Center + AI debrief rendering
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except Exception:
    # PDF export is optional; app should still run without reportlab installed.
    REPORTLAB_AVAILABLE = False

import time
import random
import hashlib
import base64
import secrets
import string
import urllib.request
import urllib.error
from pathlib import Path
from datetime import datetime
from zoneinfo import ZoneInfo

import streamlit as st

# =============================
# Flash messages (persist across st.rerun)
# =============================
def flash_success(msg: str = "Saved."):
    try:
        st.session_state["_flash_success"] = str(msg)
    except Exception:
        pass

def _show_flash_success():
    msg = None
    try:
        msg = st.session_state.get("_flash_success")
    except Exception:
        msg = None
    if msg:
        try:
            st.success(msg)
        except Exception:
            pass
        try:
            st.toast(msg)
        except Exception:
            pass
        try:
            del st.session_state["_flash_success"]
        except Exception:
            pass


# =============================
# Safe JSON loader (must be defined before first use)
# =============================
def load_json_safe(path, default):
    """Load JSON from `path`. Return `default` if file missing or invalid."""
    try:
        if not path or not os.path.exists(path):
            return default
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        # Keep app running even if settings files are malformed
        try:
            print(f"[WARN] Failed to load JSON from {path}: {e}")
        except Exception:
            pass
        return default

# =============================
# Student Security Hardening (best-effort)
# =============================
def apply_student_hardening(student_profile: dict):
    """
    Best-effort deterrents against copying/printing in browser.
    NOTE: Cannot fully prevent screenshots/phone photos.
    Applies to the whole page for student view (practice + exam).
    """
    try:
        sid = (student_profile or {}).get("username") or (student_profile or {}).get("id") or "student"
        sname = (student_profile or {}).get("display_name") or (student_profile or {}).get("name") or "Student"
        import datetime as _dt
        ts = _dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        session_code = st.session_state.get("session_code")
        if not session_code:
            session_code = secrets.token_hex(3).upper()
            st.session_state["session_code"] = session_code
        # optional footer session code (controlled by admin NCLEX policy)
        footer_enabled = bool(load_nclex_policy().get("footer_session_code_enabled", True))
        footer_text = f"Session {session_code} • {sid} • {ts}" if footer_enabled else ""
    except Exception:
        session_code = st.session_state.get("session_code")
        if not session_code:
            session_code = secrets.token_hex(3).upper()
            st.session_state["session_code"] = session_code
        footer_enabled = bool(load_nclex_policy().get("footer_session_code_enabled", True))
        footer_text = f"Session {session_code} • Student • " + __import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M:%S") if footer_enabled else ""

    # Inject CSS/JS into the top-level document (works in Streamlit app page)
    import streamlit.components.v1 as components
    components.html(f"""
    <script>
    (function() {{
      const doc = (window.top && window.top.document) ? window.top.document : document;

      // --- CSS: disable selection globally, but allow inputs/textarea/select and streamlit widgets text entry.
      const styleId = "cliniq-student-hardening-style";
      if (!doc.getElementById(styleId)) {{
        const st = doc.createElement("style");
        st.id = styleId;
        st.textContent = `
          html, body {{
            -webkit-user-select: none !important;
            -moz-user-select: none !important;
            -ms-user-select: none !important;
            user-select: none !important;
          }}
          input, textarea, select, option, [contenteditable="true"] {{
            -webkit-user-select: text !important;
            -moz-user-select: text !important;
            -ms-user-select: text !important;
            user-select: text !important;
          }}
          @media print {{
            body {{ display: none !important; }}
          }}
        `;
        doc.head.appendChild(st);
      }}

      // --- Tiny footer session code (optional)
      const ftText = "{footer_text}";
      const footerId = "cliniq-student-footer";
      if (ftText && ftText.trim()) {{
        if (!doc.getElementById(footerId)) {{
          const ft = doc.createElement("div");
          ft.id = footerId;
          ft.style.position = "fixed";
          ft.style.bottom = "10px";
          ft.style.right = "10px";
          ft.style.zIndex = "999998";
          ft.style.pointerEvents = "none";
          ft.style.fontFamily = "system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif";
          ft.style.fontSize = "11px";
          ft.style.color = "rgba(0,0,0,0.55)";
          ft.style.background = "rgba(255,255,255,0.55)";
          ft.style.padding = "4px 6px";
          ft.style.borderRadius = "6px";
          ft.style.backdropFilter = "blur(2px)";
          ft.textContent = ftText;
          doc.body.appendChild(ft);
        }} else {{
          const ft = doc.getElementById(footerId);
          if (ft) ft.textContent = ftText;
        }}
      }} else {{
        const existing = doc.getElementById(footerId);
        if (existing) existing.remove();
      }}

      // --- Block right-click
      if (!doc.__cliniq_ctx_blocked) {{
        doc.__cliniq_ctx_blocked = true;
        doc.addEventListener("contextmenu", function(e) {{
          e.preventDefault();
          return false;
        }}, true);

        // --- Block copy/cut/drag
        ["copy","cut","dragstart"].forEach(function(evt) {{
          doc.addEventListener(evt, function(e) {{
            e.preventDefault();
            return false;
          }}, true);
        }});

        // --- Block common shortcuts
        doc.addEventListener("keydown", function(e) {{
          const key = (e.key || "").toLowerCase();
          const ctrl = e.ctrlKey || e.metaKey;
          if (!ctrl) return true;
          // block: Ctrl/Cmd + C/P/S/U/A/X
          if (["c","p","s","u","a","x"].includes(key)) {{
            e.preventDefault();
            e.stopPropagation();
            return false;
          }}
          return true;
        }}, true);

        // --- Hide on print attempt
        window.addEventListener("beforeprint", function() {{
          try {{ doc.body.style.display = "none"; }} catch(err) {{}}
        }});
        window.addEventListener("afterprint", function() {{
          try {{ doc.body.style.display = ""; }} catch(err) {{}}
        }});
      }}
    }})();
    </script>
    """, height=0)

def render_mcq_no_default(label, options, key, help_text=None, horizontal=False):
    """MCQ widget with a placeholder so no real option is preselected."""
    placeholder = "— Select an answer —"
    opts = [placeholder] + list(options or [])
    prev = st.session_state.get(key)
    # Map prev (real option) into selection
    if prev in opts:
        idx = opts.index(prev)
    else:
        idx = 0
    choice = st.radio(label, opts, index=idx, key=f"{key}__radio", help=help_text, horizontal=horizontal)
    if choice == placeholder:
        st.session_state[key] = None
        return None
    st.session_state[key] = choice
    return choice
import streamlit.components.v1 as components

# =============================
# AI (OpenAI) helper
# =============================
AI_SYSTEM_PROMPT = (
    "You are an expert NCLEX-style nursing clinical reasoning coach. "
    "Provide clear, concise, supportive guidance focused on safety and prioritization. "
    "Do NOT reveal the correct option(s) explicitly. "
    "Instead, explain what to look for, what is unsafe, and how to reason to the best answer."
)

def openai_responses_call(model: str, system_prompt: str, user_prompt: str) -> str:
    """Call OpenAI to generate a short coaching/debrief response.

    Works with:
    - OpenAI Python SDK v1 (Responses API preferred; falls back to Chat Completions).
    - Returns plain text. Raises a helpful exception if OPENAI_API_KEY is missing.
    """
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY is not set")

    # Try OpenAI Python SDK v1
    try:
        from openai import OpenAI  # type: ignore
        client = OpenAI(api_key=api_key)

        # Preferred: Responses API
        try:
            resp = client.responses.create(
                model=model,
                input=[
                    {"role": "system", "content": [{"type": "text", "text": system_prompt}]},
                    {"role": "user", "content": [{"type": "text", "text": user_prompt}]},
                ],
            )
            # Best-effort text extraction across SDK variants
            if hasattr(resp, "output_text") and resp.output_text:
                return str(resp.output_text).strip()
            # Fallback: scan output blocks
            out = getattr(resp, "output", None) or []
            chunks = []
            for item in out:
                for c in getattr(item, "content", []) or []:
                    if getattr(c, "type", None) == "output_text":
                        chunks.append(getattr(c, "text", ""))
            txt = "\n".join([t for t in chunks if t]).strip()
            if txt:
                return txt
        except Exception:
            pass

        # Fallback: Chat Completions (older but still supported in many installs)
        chat = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        )
        return (chat.choices[0].message.content or "").strip()

    except Exception as e:
        raise RuntimeError(f"OpenAI call failed: {e}")




# =============================
# Always-safe Session State init (prevents AttributeError on first submit)
# =============================
if "answers" not in st.session_state:
    st.session_state["answers"] = {"A": {}, "B": {}, "C": {}, "D": {}, "E": {}}
if "attempt_started_epoch" not in st.session_state:
    st.session_state["attempt_started_epoch"] = None
if "attempt_case_id" not in st.session_state:
    st.session_state["attempt_case_id"] = None
if "attempt_case_title" not in st.session_state:
    st.session_state["attempt_case_title"] = None
if "attempt_mode" not in st.session_state:
    st.session_state["attempt_mode"] = "practice"
# =============================
# Global UI styling (consistent colors & typography)
# =============================
GLOBAL_CSS = """
<style>
/* Titles (all pages/sections) */
.nr-title { color: #b00020 !important; font-weight: 900 !important; }
h1, h2, h3, h4, h5, h6 { color: #b00020 !important; font-weight: 900 !important; }

/* Smaller "key" labels (use .kv-key spans) */
.kv-key { color: #111 !important; font-size: 0.92rem !important; font-weight: 700 !important; }
.kv-val { color: #111 !important; font-size: 0.98rem !important; font-weight: 400 !important; }

/* Case summary headings (blue + bold) */
.case-blue { color: #b00020 !important; font-weight: 800 !important; }

/* Slightly tighter spacing for summary */
.summary-block { margin-top: 0.2rem; margin-bottom: 0.2rem; }

    /* Sidebar background */
    section[data-testid="stSidebar"] {
        background: #eaf4ff !important;
    }
    section[data-testid="stSidebar"] * {
        color: #0b2233;
    }


/* =============================
   SBAR visual styling (E domain)
   ============================= */
.sbar-box{
  background:#e8f4ff;
  border:1px solid rgba(0,0,0,.08);
  border-left:6px solid #2b78c5;
  padding:14px 16px;
  border-radius:14px;
  margin:10px 0;
}
.sbar-e-badge{
  background:#ffe066;
  color:#ffffff;
  font-weight:900;
  padding:4px 10px;
  border-radius:8px;
  display:inline-block;
  margin-right:8px;
}
.sbar-title{
  background:#ffe066;
  color:#b00020;
  font-weight:900;
  padding:4px 10px;
  border-radius:8px;
  display:inline-block;
  margin:6px 0 4px 0;
}
.sbar-key{
  background:#ffe066;
  color:#8b0000;
  font-weight:900;
  padding:6px 12px;
  border-radius:10px;
  display:inline-block;
  margin:8px 0 4px 0;
}


/* AI Debrief yellow page */
.ai-debrief-box{
  background:#fff7cc;
  border:1px solid rgba(0,0,0,.08);
  padding:16px 18px;
  border-radius:14px;
  line-height:1.6;
  margin:10px 0;
}
.ai-debrief-h2{
  color:#b00020;
  font-weight:900;
  margin:10px 0 6px 0;
  font-size:1.05rem;
}
.ai-debrief-label{
  color:#b00020;
  font-weight:900;
}
.ai-debrief-p{
  margin:6px 0;
  color:#111;
}
.ai-debrief-ul{ margin:6px 0 10px 22px; }
.ai-debrief-ul li{ margin:4px 0; color:#111; }

</style>
"""






def format_text_with_red_titles(raw: str) -> str:
    """Convert short 'Section:' lines into markdown headings so global CSS makes them red/bold.
    Keeps body text as normal black.
    """
    if not raw:
        return ""
    lines = str(raw).splitlines()
    out_lines = []
    for line in lines:
        s = line.rstrip()
        # Match short headers like "Assessment:", "Rationale:", "Key Points:"
        m = re.match(r'^\s*([A-Za-z][A-Za-z0-9 \-/&]{1,60})\s*:\s*$', s)
        if m:
            out_lines.append(f"### {m.group(1).strip()}")
            continue
        # Match "Header: content" for short headers; split into heading + paragraph
        m2 = re.match(r'^\s*([A-Za-z][A-Za-z0-9 \-/&]{1,40})\s*:\s*(.+)\s*$', s)
        if m2:
            out_lines.append(f"### {m2.group(1).strip()}")
            out_lines.append(m2.group(2).strip())
            continue
        out_lines.append(line)
    return "\n".join(out_lines)


def render_ai_debrief_html(raw: str) -> str:
    """Convert AI debrief markdown-ish text into clean HTML:
    - Wrap in a yellow box (handled by caller)
    - Major numbered headings (e.g., '1) Strengths') become red bold titles
    - Labels 'Strengths:' style become red bold label only; body stays black
    - Bullets render as <ul><li>
    """
    if not raw:
        return ""
    lines = str(raw).splitlines()
    out = []
    ul_open = False

    def _close_ul():
        nonlocal ul_open
        if ul_open:
            out.append("</ul>")
            ul_open = False

    for line in lines:
        s = (line or "").rstrip()
        if not s.strip():
            _close_ul()
            continue

        # Major headings like: '1) Strengths (3)'
        mh = re.match(r"^\s*(\d+)\)\s*(.+?)\s*$", s)
        if mh:
            _close_ul()
            title = _html.escape(mh.group(2).strip())
            out.append(f"<div class='ai-debrief-h2'>{title}</div>")
            continue

        # Bullets: '-', '•'
        mb = re.match(r"^\s*(?:[-•]\s+)(.+?)\s*$", s)
        if mb:
            if not ul_open:
                out.append("<ul class='ai-debrief-ul'>")
                ul_open = True
            item = _html.escape(mb.group(1).strip())
            out.append(f"<li>{item}</li>")
            continue

        # Label + text: 'Situation: ...'
        ml = re.match(r"^\s*([A-Za-z][A-Za-z \-/]{1,32})\s*:\s*(.+)\s*$", s)
        if ml:
            _close_ul()
            label = _html.escape(ml.group(1).strip())
            body = _html.escape(ml.group(2).strip())
            out.append(
                f"<div class='ai-debrief-p'><span class='ai-debrief-label'>{label}:</span> {body}</div>"
            )
            continue

        _close_ul()
        out.append(f"<div class='ai-debrief-p'>{_html.escape(s.strip())}</div>")

    _close_ul()
    return "".join(out)


def _strip_md_emphasis_markers(s: str) -> str:
    """Remove markdown emphasis markers (*, **) that appear as literal symbols inside HTML containers.
    Keeps all other text unchanged."""
    if not s:
        return ""
    try:
        s = str(s)
    except Exception:
        return ""
    # Remove bold/italic markers that are not rendered inside our HTML wrappers
    s = s.replace("**", "")
    s = s.replace("*", "")
    return s


def inject_basic_exam_protection(enabled: bool = True, watermark_text: str = ""):
    """Best-effort deterrents against copy/print. Not foolproof (screenshots/phone cameras still work).

    Applies to the entire student UI (practice + exam), including case triggers and A–E sections.
    Optionally overlays a faint repeated watermark with student identity + timestamp.
    """
    if not enabled:
        return
    try:
        html = r"""
        <script>
        // Set watermark text from Streamlit (string injected)
        try { window.__CLINIQ_WM_TEXT__ = "__CLINIQ_WM_VALUE__"; } catch(e) {}
        (function() {
          // Attach to the top-level document when possible (Streamlit runs components in an iframe).
          var doc = document;
          try { if (window.parent && window.parent.document) doc = window.parent.document; } catch(e) {}
          try { if (window.top && window.top.document) doc = window.top.document; } catch(e) {}

          // 1) Disable right-click context menu
          try { doc.addEventListener('contextmenu', function(e){ e.preventDefault(); }, {capture:true}); } catch(e) {}

          // 2) Block common copy/print/dev shortcuts (Ctrl/Cmd + ...)
          try {
            doc.addEventListener('keydown', function(e){
              var k = (e.key || '').toLowerCase();
              var ctrl = e.ctrlKey || e.metaKey;
              if (!ctrl) return;
              if (k === 'p' || k === 's' || k === 'c' || k === 'x' || k === 'u' || k === 'a') {
                e.preventDefault();
                e.stopPropagation();
              }
            }, {capture:true});
          } catch(e) {}

          // 3) Hide page contents during print (best-effort)
          function ensurePrintStyle(targetDoc) {
            try {
              var styleId = 'cliniq-no-print-style';
              if (targetDoc.getElementById(styleId)) return;
              var st = targetDoc.createElement('style');
              st.id = styleId;
              st.innerHTML = "@media print { body { display:none !important; } }";
              targetDoc.head.appendChild(st);
            } catch(e) {}
          }
          ensurePrintStyle(doc);

          // 4) Disable text selection globally, but allow typing/selecting inside inputs/textarea and editable elements
          function ensureSelectStyle(targetDoc) {
            try {
              var styleId = 'cliniq-no-select-style';
              if (targetDoc.getElementById(styleId)) return;
              var st = targetDoc.createElement('style');
              st.id = styleId;
              st.innerHTML = [
                "body, body * { -webkit-user-select:none !important; -moz-user-select:none !important; -ms-user-select:none !important; user-select:none !important; }",
                "input, textarea, [contenteditable='true'] { -webkit-user-select:text !important; -moz-user-select:text !important; -ms-user-select:text !important; user-select:text !important; }"
              ].join("\n");
              targetDoc.head.appendChild(st);
            } catch(e) {}
          }
          ensureSelectStyle(doc);

          // 5) Extra: block drag selection and copy events and copy events
          try {
            doc.addEventListener('copy', function(e){ e.preventDefault(); }, {capture:true});
            doc.addEventListener('cut', function(e){ e.preventDefault(); }, {capture:true});
            doc.addEventListener('dragstart', function(e){ e.preventDefault(); }, {capture:true});
          } catch(e) {}
        })();
        </script>
        """
        # Inject watermark text safely (escape backslashes/quotes/newlines)
        try:
            wm = (watermark_text or "")
            wm = wm.replace("\\", "\\\\").replace('"', '\\"').replace("\n", " ").replace("\r", " ")
            html = html.replace("__CLINIQ_WM_VALUE__", wm)
        except Exception:
            pass
        components.html(html, height=0)
    except Exception:
        pass

# =============================
# Paths
# =============================
BASE_DIR = Path(__file__).resolve().parent

CASES_PATH = BASE_DIR / "cases.json"
STUDENTS_PATH = BASE_DIR / "students.json"
ATTEMPTS_PATH = BASE_DIR / "attempts_log.jsonl"
ATTEMPTS_CSV_PATH = BASE_DIR / "attempts_export.csv"
NCLEX_ITEM_CSV_PATH = BASE_DIR / "nclex_item_analysis.csv"

# Attempts policy (your working filename)
ATTEMPT_POLICY_PATH = BASE_DIR / "attempts_policy.json"

# Persistent admin settings (mode + AI toggles)
ADMIN_SETTINGS_PATH = BASE_DIR / "admin_settings.json"
RESEARCH_POLICY_PATH = BASE_DIR / "research_policy.json"
RESEARCH_LOG_PATH = BASE_DIR / "research_log.jsonl"
RESEARCH_DATASET_PATH = BASE_DIR / "research_dataset.jsonl"
IRB_DOCS_DIR = BASE_DIR / "irb_docs"


RESEARCH_ARCHIVE_DIR = BASE_DIR / "research_archives"
# Case visibility + timer policy
CASE_POLICY_PATH = BASE_DIR / "case_policy.json"

# ✅ Exam password policy (global, auto-expire)
EXAM_ACCESS_POLICY_PATH = BASE_DIR / "exam_access_policy.json"

# ✅ NCLEX-style practice files
NCLEX_ITEMS_PATH = BASE_DIR / "nclex_items.json"
NCLEX_POLICY_PATH = BASE_DIR / "nclex_policy.json"
NCLEX_ACTIVE_SETS_PATH = BASE_DIR / "nclex_active_sets.json"

# ✅ Feature flags + optional data files
FEATURES_PATH = BASE_DIR / "features.json"
BACKUP_DIR = BASE_DIR / "backups"
AUTOSAVE_DRAFTS_PATH = BASE_DIR / "autosave_drafts.jsonl"
KPI_POLICY_PATH = BASE_DIR / "kpi_policy.json"
EXAM_OVERRIDES_PATH = BASE_DIR / "exam_overrides.json"

# UI default: keep sidebar tidy (admin tools moved to top menu + Settings)
SHOW_ADMIN_PANELS_IN_SIDEBAR = False

# Optional feature flags (features.json). Safe even if missing.
features = load_json_safe(FEATURES_PATH, {})
if not isinstance(features, dict):
    features = {}

# Default autosave ON to support resume/progress (safe; can be turned off in features.json)
features.setdefault("autosave_enabled", True)

# =============================
# Admin password (case-sensitive)
# =============================
ADMIN_PASSWORD = "Mahmoud@Nurse123"

# =============================
# Timezone helpers (Qatar)
# =============================
TZ = ZoneInfo("Asia/Qatar")


def now_local():
    return datetime.now(TZ)


def parse_iso_dt(s: str):
    try:
        if not s:
            return None
        return datetime.fromisoformat(s)
    except Exception:
        return None


# =============================
# Utilities
# =============================
def utc_now_iso() -> str:
    return datetime.utcnow().isoformat()


def sha256_hex(s: str) -> str:
    return hashlib.sha256((s or "").encode("utf-8")).hexdigest()


def safe_int(x, default=None):
    try:
        return int(x)
    except Exception:
        return default


def ensure_file(path: Path, default_obj):
    if not path.exists():
        path.write_text(json.dumps(default_obj, ensure_ascii=False, indent=2), encoding="utf-8")


def load_json_safe(path: Path, default):
    try:
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        pass
    return default


def load_features():
    # Safe default: empty dict (features off unless present)
    return load_json_safe(FEATURES_PATH, {})


def save_features(feat: dict):
    """Persist feature flags to features.json (with backup)."""
    try:
        if not isinstance(feat, dict):
            return
        backup_file(FEATURES_PATH)
        FEATURES_PATH.write_text(json.dumps(feat, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass



def backup_file(path: Path):
    try:
        if not path.exists():
            return
        BACKUP_DIR.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out = BACKUP_DIR / f"{path.stem}_{ts}{path.suffix}"
        out.write_bytes(path.read_bytes())
    except Exception:
        pass


def maybe_backup_on_start(features: dict):
    if not features.get("backup_on_start", False):
        return
    for p in [
        CASES_PATH,
        STUDENTS_PATH,
        ATTEMPTS_PATH,
        ATTEMPT_POLICY_PATH,
        ADMIN_SETTINGS_PATH,
        CASE_POLICY_PATH,
        EXAM_ACCESS_POLICY_PATH,
        NCLEX_ITEMS_PATH,
        NCLEX_POLICY_PATH,
        FEATURES_PATH,
    ]:
        backup_file(p)


# =============================
# Create required config files if missing
# =============================
ensure_file(STUDENTS_PATH, {"students": []})
ensure_file(EXAM_OVERRIDES_PATH, {"students": {}})
ensure_file(ATTEMPT_POLICY_PATH, {"default_max_attempts": "unlimited", "per_case": {}})
ensure_file(ADMIN_SETTINGS_PATH, {
    "app_mode": "Practice",
    "ai_enabled": False,
    "ai_debrief_enabled": False,
    "ai_model": "gpt-5.2",
    "research_mode": False
})
ensure_file(RESEARCH_POLICY_PATH, {
    "enabled": False,
    "require_consent": True,
    "anonymize_student_id": True,
    "collect_reflection": False,
    "collect_answer_change": True,
    "collect_section_performance": True,
    "consent_title": "Research Consent",
    "consent_text": "This activity may be used for educational research. Participation is voluntary. Your responses will be anonymized for analysis when possible. By checking the box below, you consent to your data being used for research purposes."
})
ensure_file(CASE_POLICY_PATH, {
    "default_visibility": True,
    "default_timer_minutes_practice": "unlimited",
    "default_timer_minutes_exam": 20,
    "per_case": {}
})
ensure_file(EXAM_ACCESS_POLICY_PATH, {
    "enabled": False,
    "force_exam_password_only": True,
    "exam_password_sha256": "",
    "expires_at": ""
})
ensure_file(NCLEX_POLICY_PATH, {
    "enabled": False,
    "mode_visibility": {"Practice": True, "Exam": True},
    "items_per_case": 30,
    "per_case_items": {},
    "optional_practical_section": True,
    "show_correct_answers_after_submit": True,
    "show_rationales_after_submit": True,
    "shuffle_options": True,
    "separate_grade": {
        "enabled": True,
        "max_points_per_case": 30,
        "weight_vs_reasoning_case_score": 0.5
    },
    "client_need_blueprint": {
        "Physiological Integrity": 0.35,
        "Safe and Effective Care Environment": 0.25,
        "Health Promotion and Maintenance": 0.2,
        "Psychosocial Integrity": 0.2
    },
    "enabled_types": {
        "mcq": True,
        "sata": True,
        "ordered_response": True,
        "cloze": True,
        "matrix": True,
        "evolving_case": True
    }
})
ensure_file(NCLEX_ITEMS_PATH, {"cases": {}})
ensure_file(NCLEX_ACTIVE_SETS_PATH, {"by_case": {}})

ensure_file(KPI_POLICY_PATH, {
    "total_score": True,
    "domain_profile": True,
    "nclex_score": True,
    "time_to_completion": True,
    "unsafe_flags": True,
    "attempts_per_case": True
})

# If features.json doesn't exist, we don't force-create it; app still runs safely.


# =============================
# IO
# =============================
def load_cases():
    data = load_json_safe(CASES_PATH, None)
    if data is None:
        raise ValueError("cases.json not readable")
    if not isinstance(data, list):
        raise ValueError("cases.json must contain a LIST")
    return data


def get_cases_list():
    """Return cases as a list regardless of whether cases.json loader returns list or dict."""
    try:
        data = load_cases()
    except Exception:
        return []
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        # Some legacy structures might wrap the list.
        for k in ("cases", "items", "data"):
            v = data.get(k)
            if isinstance(v, list):
                return v
        return []
    return []

def save_cases(cases_list: list):
    CASES_PATH.write_text(json.dumps(cases_list, ensure_ascii=False, indent=2), encoding="utf-8")


def load_students():
    ensure_file(STUDENTS_PATH, {"students": []})
    ensure_file(EXAM_OVERRIDES_PATH, {"students": {}})
    data = load_json_safe(STUDENTS_PATH, {"students": []})
    if not isinstance(data, dict) or "students" not in data or not isinstance(data["students"], list):
        raise ValueError("students.json must be an object with key 'students' as a list")
    return data


def save_students(data: dict):
    STUDENTS_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def load_attempt_policy():
    ensure_file(ATTEMPT_POLICY_PATH, {"default_max_attempts": "unlimited", "per_case": {}})
    data = load_json_safe(ATTEMPT_POLICY_PATH, {"default_max_attempts": "unlimited", "per_case": {}})
    if not isinstance(data, dict):
        data = {}
    data.setdefault("default_max_attempts", "unlimited")
    data.setdefault("per_case", {})
    if not isinstance(data["per_case"], dict):
        data["per_case"] = {}
    return data


def save_attempt_policy(policy: dict):
    ATTEMPT_POLICY_PATH.write_text(json.dumps(policy, ensure_ascii=False, indent=2), encoding="utf-8")


def load_admin_settings():
    ensure_file(ADMIN_SETTINGS_PATH, {
        "app_mode": "Practice",
        "ai_enabled": False,
        "ai_debrief_enabled": False,
        "ai_model": "gpt-5.2",
        "research_mode": False,
        "intro_videos": {}
    })
    data = load_json_safe(ADMIN_SETTINGS_PATH, {})
    if not isinstance(data, dict):
        data = {}
    data.setdefault("app_mode", "Practice")
    data.setdefault("ai_enabled", False)
    data.setdefault("ai_debrief_enabled", False)
    data.setdefault("ai_model", "gpt-5.2")
    data.setdefault("research_mode", False)
    data.setdefault("intro_videos", {})
    return data


def save_admin_settings(settings: dict):
    flash_success('Saved.')
    ADMIN_SETTINGS_PATH.write_text(json.dumps(settings, ensure_ascii=False, indent=2), encoding="utf-8")

def load_research_policy() -> dict:
    ensure_file(RESEARCH_POLICY_PATH, {
        "enabled": False,
        "require_consent": True,
        "anonymize_student_id": True,
        "collect_reflection": False,
        "collect_answer_change": True,
        "collect_section_performance": True,
        "consent_title": "Research Consent",
        "consent_text": "This activity may be used for educational research. Participation is voluntary. Your responses will be anonymized for analysis when possible. By checking the box below, you consent to your data being used for research purposes.",
        "irb_status": "Pending",
        "irb_reference": "",
        "irb_docs": [],
        "anonymization_salt": "",
        "retention_months": 24,
        "retention_mode": "24 months"
    })
    data = load_json_safe(RESEARCH_POLICY_PATH, {})
    if not isinstance(data, dict):
        data = {}
    data.setdefault("enabled", False)
    data.setdefault("require_consent", True)
    data.setdefault("anonymize_student_id", True)
    data.setdefault("collect_reflection", False)
    data.setdefault("collect_answer_change", True)
    data.setdefault("collect_section_performance", True)
    data.setdefault("consent_title", "Research Consent")
    data.setdefault("consent_text", "This activity may be used for educational research. Participation is voluntary. Your responses will be anonymized for analysis when possible. By checking the box below, you consent to your data being used for research purposes.")
    data.setdefault("irb_status", "Pending")
    data.setdefault("irb_reference", "")
    data.setdefault("irb_docs", [])
    data.setdefault("anonymization_salt", "")
    data.setdefault("retention_months", 24)
    data.setdefault("retention_mode", "24 months")
    # Generate a private salt if missing (used for hashed participant_id)
    if not str(data.get("anonymization_salt","")).strip():
        try:
            data["anonymization_salt"] = secrets.token_hex(16)
        except Exception:
            data["anonymization_salt"] = "NR_RESEARCH_V1"
    return data


def save_research_policy(policy: dict):
    flash_success('Saved.')
    RESEARCH_POLICY_PATH.write_text(json.dumps(policy, ensure_ascii=False, indent=2), encoding="utf-8")


def append_research_event(event: dict):
    """Append one research event line (fast)."""
    try:
        with open(RESEARCH_LOG_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(event, ensure_ascii=False) + "\n")
    except Exception:
        pass



def _hash_participant(identifier: str) -> str:
    """Hash identifier with admin-configured salt to produce participant_id."""
    try:
        rp = load_research_policy()
        salt = str(rp.get("anonymization_salt", "") or "").strip() or "NR_RESEARCH_V1"
        return hashlib.sha256((salt + "|" + (identifier or "")).encode("utf-8")).hexdigest()[:16]
    except Exception:
        return "anon"


def append_research_dataset_row(row: dict):
    """Append a de-identified research row to research_dataset.jsonl (separate from teaching attempts)."""
    try:
        # Ensure directory exists for IRB docs (used elsewhere)
        try:
            IRB_DOCS_DIR.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
        with open(RESEARCH_DATASET_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")
    except Exception:
        pass


def iter_research_dataset():
    """Yield rows from research_dataset.jsonl (best-effort)."""
    if not RESEARCH_DATASET_PATH.exists():
        return []
    out = []
    try:
        with open(RESEARCH_DATASET_PATH, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    rec = json.loads(line)
                    if isinstance(rec, dict):
                        out.append(rec)
                except Exception:
                    continue
    except Exception:
        return []
    return out


def purge_research_dataset() -> tuple[bool, str]:
    """Delete only the de-identified research dataset file (does not touch teaching attempts)."""
    try:
        if RESEARCH_DATASET_PATH.exists():
            RESEARCH_DATASET_PATH.unlink()
            return True, "De-identified research dataset purged (research_dataset.jsonl deleted)."
        return True, "No research dataset file to purge."
    except Exception as e:
        return False, f"Could not purge research dataset: {e}"

def regenerate_anonymization_salt(policy: dict | None = None) -> dict:
    """Regenerate the private anonymization salt and persist it."""
    try:
        pol = policy if isinstance(policy, dict) else load_research_policy()
        pol["anonymization_salt"] = secrets.token_hex(16)
        save_research_policy(pol)
        return pol
    except Exception:
        pol = load_research_policy()
        try:
            pol["anonymization_salt"] = secrets.token_hex(16)
            save_research_policy(pol)
        except Exception:
            pass
        return pol


def archive_research_data(note: str = "") -> tuple[bool, str]:
    """Archive research_dataset.jsonl + research_log.jsonl + research_policy snapshot into a timestamped folder."""
    try:
        RESEARCH_ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
        ts = now_local().strftime("%Y%m%d_%H%M%S")
        safe_note = re.sub(r"[^a-zA-Z0-9_-]+", "_", (note or "").strip())[:40].strip("_")
        folder = RESEARCH_ARCHIVE_DIR / (f"study_{ts}" + (f"_{safe_note}" if safe_note else ""))
        folder.mkdir(parents=True, exist_ok=True)

        moved_any = False
        # dataset
        if RESEARCH_DATASET_PATH.exists():
            try:
                (folder / RESEARCH_DATASET_PATH.name).write_bytes(RESEARCH_DATASET_PATH.read_bytes())
                moved_any = True
            except Exception:
                pass
        # log
        if RESEARCH_LOG_PATH.exists():
            try:
                (folder / RESEARCH_LOG_PATH.name).write_bytes(RESEARCH_LOG_PATH.read_bytes())
                moved_any = True
            except Exception:
                pass
        # policy snapshot
        try:
            pol = load_research_policy()
            (folder / "research_policy_snapshot.json").write_text(json.dumps(pol, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass

        if not moved_any:
            return True, f"No research dataset/log to archive. Created archive folder: {folder}"

        return True, f"Archived research files to: {folder}"
    except Exception as e:
        return False, f"Archive failed: {e}"


def start_new_study_guided(note: str = "") -> tuple[bool, str]:
    """Guided 'new study' reset: archive → purge → regenerate salt → force IRB status to Pending."""
    try:
        # 1) Archive
        ok_a, msg_a = archive_research_data(note=note)
        if not ok_a:
            return False, msg_a

        # 2) Purge dataset (de-identified only)
        ok_p, msg_p = purge_research_dataset()
        if not ok_p:
            return False, msg_p

        # 3) Reset research log too (optional but recommended for a clean study)
        try:
            if RESEARCH_LOG_PATH.exists():
                RESEARCH_LOG_PATH.unlink()
        except Exception:
            pass

        # 4) Regenerate anonymization salt + force IRB status update
        pol = load_research_policy()
        try:
            pol["anonymization_salt"] = secrets.token_hex(16)
        except Exception:
            pol["anonymization_salt"] = pol.get("anonymization_salt", "NR_RESEARCH_V1")
        pol["irb_status"] = "Pending"
        pol["irb_reference"] = ""
        pol["irb_docs"] = []
        # Optional safety: disable research collection until you confirm IRB details again
        pol["enabled"] = False
        save_research_policy(pol)

        # 5) Log event
        try:
            append_research_event({"event": "study_reset", "ts": now_local().isoformat(), "note": note or ""})
        except Exception:
            pass

        return True, "Start New Study completed: archived, purged dataset, regenerated salt, and reset IRB status to Pending (research disabled until re-enabled)."
    except Exception as e:
        return False, f"Start New Study failed: {e}"


def _anon_student(student_username: str) -> str:
    # Stable anonymization (same input -> same output) for cohort research
    try:
        salt = "NR_RESEARCH_V1"
        return hashlib.sha256((salt + "|" + (student_username or "")).encode("utf-8")).hexdigest()[:16]
    except Exception:
        return "anon"


def should_collect_research() -> bool:
    policy = load_research_policy()
    if not policy.get("enabled", False):
        return False
    if policy.get("require_consent", True):
        return bool(st.session_state.get("research_consent", False))
    return True


def track_nclex_change(qid: str, prev, ans):
    """Track initial vs final answers and change counts (per question)."""
    if "nclex_track" not in st.session_state:
        st.session_state.nclex_track = {}
    t = st.session_state.nclex_track.get(qid) or {"changes": 0, "initial": None, "final": None}
    # decide if ans is non-empty
    non_empty = ans is not None and ans != "" and ans != [] and ans != {}
    if t.get("initial") is None and non_empty:
        try:
            t["initial"] = json.loads(json.dumps(ans))
        except Exception:
            t["initial"] = ans
    if prev is not None and prev != ans:
        t["changes"] = int(t.get("changes", 0) or 0) + 1
    try:
        t["final"] = json.loads(json.dumps(ans))
    except Exception:
        t["final"] = ans
    st.session_state.nclex_track[qid] = t



def load_case_policy():
    ensure_file(CASE_POLICY_PATH, {
        "default_visibility": True,
        "default_timer_minutes_practice": "unlimited",
        "default_timer_minutes_exam": 20,
        "per_case": {}
    })
    data = load_json_safe(CASE_POLICY_PATH, {})
    if not isinstance(data, dict):
        data = {}
    data.setdefault("default_visibility", True)
    data.setdefault("default_timer_minutes_practice", "unlimited")
    data.setdefault("default_timer_minutes_exam", 20)
    data.setdefault("per_case", {})
    if not isinstance(data["per_case"], dict):
        data["per_case"] = {}

    _dirty = False

    # --- Timer compatibility layer (minutes ↔ seconds) ---
    # Canonical runtime behavior uses seconds (timer_seconds_* / default_timer_seconds_*).
    # However, multiple UI blocks may store minutes only. If both exist and disagree,
    # TRUST MINUTES (what the admin explicitly selected) and overwrite seconds.
    def _is_unlimited(v):
        if v is None:
            return True
        if isinstance(v, str):
            s = v.strip().lower()
            return s in ("", "0", "unlimited", "off", "none", "unlimit")
        try:
            return int(v) <= 0
        except Exception:
            return True

    def _min_to_sec(v):
        # v may be int/str/"unlimited"/"default"
        if v is None:
            return 0
        if isinstance(v, str):
            s = v.strip().lower()
            if s in ("default",):
                return None  # sentinel meaning "inherit"
            if s in ("unlimited", "off", "", "0"):
                return 0
            try:
                return int(s) * 60
            except Exception:
                return 0
        try:
            iv = int(v)
            return 0 if iv <= 0 else iv * 60
        except Exception:
            return 0

    def _sec_to_min(v):
        try:
            s = int(v or 0)
        except Exception:
            s = 0
        return "unlimited" if s <= 0 else int(s // 60)

    def _sync_default(mode: str):
        mk = "default_timer_minutes_practice" if mode == "Practice" else "default_timer_minutes_exam"
        sk = "default_timer_seconds_practice" if mode == "Practice" else "default_timer_seconds_exam"

        minutes_val = data.get(mk, "unlimited" if mode == "Practice" else 20)
        seconds_val = data.get(sk, None)

        sec_from_min = _min_to_sec(minutes_val)
        if sec_from_min is None:
            # defaults should never be "default" — treat as unlimited
            sec_from_min = 0

        if seconds_val is None:
            # Only minutes existed → derive seconds
            data[sk] = int(sec_from_min or 0)
            _dirty = True
        else:
            # Both exist → trust SECONDS (canonical runtime) and refresh minutes to match.
            # This prevents stale minutes keys from overwriting valid seconds.
            try:
                seconds_int = int(seconds_val or 0)
            except Exception:
                seconds_int = 0
            data[sk] = seconds_int
            data[mk] = _sec_to_min(seconds_int)
            _dirty = True

        # Ensure minutes reflect seconds for consistent display
        data[mk] = _sec_to_min(data.get(sk, 0))

    _sync_default("Practice")
    _sync_default("Exam")

    # Per-case timers
    for cid, per in list(data.get("per_case", {}).items()):
        if not isinstance(per, dict):
            continue

        def _sync_case(mode: str):
            mk = "timer_minutes_practice" if mode == "Practice" else "timer_minutes_exam"
            sk = "timer_seconds_practice" if mode == "Practice" else "timer_seconds_exam"

            minutes_val = per.get(mk, None)
            seconds_val = per.get(sk, None)

            # If minutes explicitly set:
            if minutes_val is not None:
                if isinstance(minutes_val, str) and minutes_val.strip().lower() == "default":
                    # Inherit: remove seconds override; keep minutes as "default"
                    per.pop(sk, None)
                    _dirty = True
                    return
                sec_from_min = _min_to_sec(minutes_val)
                if seconds_val is None:
                    # Only minutes existed → derive seconds
                    per[sk] = int(sec_from_min or 0)
                    _dirty = True
                else:
                    # Both exist → trust SECONDS (canonical runtime) and refresh minutes to match
                    try:
                        seconds_int = int(seconds_val or 0)
                    except Exception:
                        seconds_int = 0
                    per[sk] = seconds_int
                    per[mk] = _sec_to_min(seconds_int)
                    _dirty = True
            else:
                # Minutes absent: if seconds exist, create minutes for UI
                if seconds_val is not None:
                    per[mk] = _sec_to_min(seconds_val)

            # If we have seconds override, keep minutes in sync for UI display
            if sk in per:
                per[mk] = _sec_to_min(per.get(sk, 0))

        _sync_case("Practice")
        _sync_case("Exam")

    # If we detected a minutes↔seconds mismatch, persist the normalized values once.
    if _dirty:
        try:
            CASE_POLICY_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass

    return data


def save_case_policy(policy: dict):
    flash_success('Saved.')
    CASE_POLICY_PATH.write_text(json.dumps(policy, ensure_ascii=False, indent=2), encoding="utf-8")


def load_exam_access_policy():
    ensure_file(EXAM_ACCESS_POLICY_PATH, {
        "enabled": False,
        "force_exam_password_only": True,
        "exam_password_sha256": "",
        "expires_at": ""
    })
    data = load_json_safe(EXAM_ACCESS_POLICY_PATH, {})
    if not isinstance(data, dict):
        data = {}
    data.setdefault("enabled", False)
    data.setdefault("force_exam_password_only", True)
    data.setdefault("exam_password_sha256", "")
    data.setdefault("expires_at", "")
    return data


def save_exam_access_policy(policy: dict):
    flash_success('Saved.')
    EXAM_ACCESS_POLICY_PATH.write_text(json.dumps(policy, ensure_ascii=False, indent=2), encoding="utf-8")


def is_exam_password_active(policy: dict) -> bool:
    if not policy.get("enabled", False):
        return False

    expires_at = parse_iso_dt(policy.get("expires_at", ""))
    if expires_at is None:
        return True  # active until disabled manually

    try:
        expires_at_local = expires_at.astimezone(TZ)
    except Exception:
        expires_at_local = expires_at

    if now_local() >= expires_at_local:
        policy["enabled"] = False
        save_exam_access_policy(policy)
        return False

    return True


def save_attempt(record):
    with open(ATTEMPTS_PATH, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")


def iter_attempts():
    if not ATTEMPTS_PATH.exists():
        return
    with open(ATTEMPTS_PATH, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                yield json.loads(line)
            except Exception:
                continue



def build_attempt_record_from_state(case_id: str) -> dict:
    """Build a complete attempt record using current Streamlit session_state.

    This is used by the NCLEX Practical "Save Attempt" button and by exam auto-save.
    It must be GLOBAL (not nested) so it can be called from render_nclex_practical().
    """
    # Time/duration
    started_epoch = st.session_state.get("attempt_started_epoch") or time.time()
    try:
        started_epoch = float(started_epoch)
    except Exception:
        started_epoch = time.time()
    duration_sec = int(max(0.0, time.time() - started_epoch))

    admin_settings = load_admin_settings()
    mode = admin_settings.get("app_mode", "Practice")

    # Student identity
    student_profile = st.session_state.get("student_profile") or {}
    student_username = st.session_state.get("student_username") or student_profile.get("username") or ""

    student_display_name = st.session_state.get("student_display_name") or student_profile.get("display_name") or student_profile.get("name") or ""
    student_id = st.session_state.get("student_id") or student_profile.get("student_id") or student_profile.get("id") or ""
    cohort = student_profile.get("cohort", "") or st.session_state.get("student_cohort", "")

    # Case info (best-effort)
    case_obj = None
    try:
        # Prefer already-loaded global cases list if present
        _cases = globals().get("cases")
        if isinstance(_cases, list):
            for c in _cases:
                if isinstance(c, dict) and str(c.get("id", "")) == str(case_id):
                    case_obj = c
                    break
    except Exception:
        case_obj = None
    if case_obj is None:
        try:
            for c in load_cases():
                if isinstance(c, dict) and str(c.get("id", "")) == str(case_id):
                    case_obj = c
                    break
        except Exception:
            case_obj = None

    case_title = (case_obj or {}).get("title", "")

    # Scores / answers
    scores = st.session_state.get("scores") if isinstance(st.session_state.get("scores"), dict) else {}
    answers = st.session_state.get("answers") if isinstance(st.session_state.get("answers"), dict) else {}

    intake_score = st.session_state.get("intake_score")
    intake_breakdown = st.session_state.get("intake_breakdown") if isinstance(st.session_state.get("intake_breakdown"), dict) else {}

    # Unsafe counts (optional)
    unsafe_counts = {}
    try:
        for dom in ["A", "B", "C", "D", "E"]:
            fb = st.session_state.get("last_feedback", {}).get(dom)
            unsafe_hits = []
            if isinstance(fb, dict):
                unsafe_hits = fb.get("unsafe", []) or []
            elif isinstance(fb, (list, tuple)) and len(fb) >= 4:
                unsafe_hits = fb[3] or []
            unsafe_counts[dom] = len(unsafe_hits) if isinstance(unsafe_hits, list) else 0
    except Exception:
        unsafe_counts = {}
    unsafe_total = sum([int(v or 0) for v in unsafe_counts.values()]) if unsafe_counts else 0

    # NCLEX scored blob
    nclex_scored = st.session_state.get("nclex_scored") if isinstance(st.session_state.get("nclex_scored"), dict) else {}
    nclex_answers = st.session_state.get("nclex_answers") if isinstance(st.session_state.get("nclex_answers"), dict) else {}
    nclex_total = int(nclex_scored.get("total_max", 0) or 0)
    nclex_score = int(nclex_scored.get("total_points", 0) or 0)
    nclex_details = nclex_scored.get("details") if isinstance(nclex_scored.get("details"), list) else []

    # Research consent snapshot (optional)
    research_consent = bool(st.session_state.get("research_consent", False))

    rec = {
        # Timing
        "started_at": datetime.utcfromtimestamp(started_epoch).isoformat(),
        "submitted_at": utc_now_iso(),
        "duration_seconds": duration_sec,

        # Identity
        "student_username": student_username,
        "student_display_name": student_display_name,
        "student_id": student_id,
        "cohort": cohort,

        # Context
        "mode": mode,
        "caseId": str(case_id),
        "caseTitle": str(case_title),
        "system": safe_get_system(case_obj) if isinstance(case_obj, dict) else "",

        # Performance (reasoning)
        "intake_score": intake_score,
        "intake_breakdown": intake_breakdown,
        "scores": scores,
        "answers": answers,

        # Safety
        "unsafe_counts": unsafe_counts,
        "unsafe_total": unsafe_total,

        # NCLEX (psychometrics depends on this)
        "nclex_total": nclex_total,
        "nclex_score": nclex_score,
        "nclex_answers": nclex_answers,
        "nclex": {
            "total_points": nclex_score,
            "total_max": nclex_total,
            "details": nclex_details,
        },

        # Research (optional)
        "research_consent": research_consent,
    }
    return rec


# =============================
# Research exports & archiving (Admin-only, on-demand)
# =============================

def _safe_json_dumps(x):
    try:
        return json.dumps(x, ensure_ascii=False)
    except Exception:
        try:
            return str(x)
        except Exception:
            return ""

def _flatten_attempt_row(rec: dict) -> dict:
    """One row per attempt for 'Download Research CSV'."""
    rec = rec or {}
    scores = rec.get("scores") if isinstance(rec.get("scores"), dict) else {}
    intake_breakdown = rec.get("intake_breakdown") if isinstance(rec.get("intake_breakdown"), dict) else {}

    # time fields (handle multiple versions)
    started = rec.get("started_at") or rec.get("attempt_started_at") or rec.get("startedAt") or ""
    submitted = rec.get("submitted_at") or rec.get("attempt_submitted_at") or rec.get("submittedAt") or ""
    duration_s = rec.get("duration_seconds") or rec.get("durationSecs") or ""

    return {
        "submitted_at": submitted,
        "started_at": started,
        "duration_seconds": duration_s,

        "student_username": rec.get("student_username", ""),
        "student_id": rec.get("student_id", ""),
        "cohort": rec.get("cohort", ""),

        "caseId": rec.get("caseId", ""),
        "caseTitle": rec.get("caseTitle", ""),
        "mode": rec.get("mode", rec.get("attempt_mode", "")),

        "intake_score": rec.get("intake_score", ""),
        "intake_breakdown_json": _safe_json_dumps(intake_breakdown),

        "A_score": scores.get("A", ""),
        "B_score": scores.get("B", ""),
        "C_score": scores.get("C", ""),
        "D_score": scores.get("D", ""),
        "E_score": scores.get("E", ""),
        "total_score": rec.get("total_score", rec.get("totalScore", "")),
        "total_with_intake": rec.get("total_with_intake", rec.get("totalWithIntake", "")),

        "unsafe_flags_json": _safe_json_dumps(rec.get("unsafe_flags", rec.get("unsafeFlags", ""))),
        "unsafe_counts_json": _safe_json_dumps(rec.get("unsafe_counts", rec.get("unsafeCounts", ""))),

        "answers_json": _safe_json_dumps(rec.get("answers", {})),
        "nclex_total": rec.get("nclex_total", ""),
        "nclex_score": rec.get("nclex_score", ""),
        "nclex_answers_json": _safe_json_dumps(rec.get("nclex_answers", {})),
    }


def build_research_csv_bytes() -> bytes:
    """Download Research CSV.

    If research_dataset.jsonl exists, export the de-identified dataset (participant_id-based).
    Otherwise, fall back to the teaching attempts export.
    """
    rp = load_research_policy()
    use_dataset = bool(rp.get("enabled", False)) and RESEARCH_DATASET_PATH.exists()

    if use_dataset:
        rows = iter_research_dataset() or []
        # Stable, human-friendly column order
        headers = [
            "participant_id",
            "submitted_at",
            "caseId",
            "caseTitle",
            "mode",
            "cohort",
            "intake_score",
            "total_score",
            "total_with_intake",
            "nclex_score",
            "nclex_total",
            "duration_seconds",
            "domain_scores_json",
            "performance_by_section_json",
            "nclex_changes_json",
            "reflection",
        ]
        buf = io.StringIO()
        w = csv.DictWriter(buf, fieldnames=headers, extrasaction="ignore")
        w.writeheader()
        for r in rows:
            if not isinstance(r, dict):
                continue
            row = {
                "participant_id": r.get("participant_id",""),
                "submitted_at": r.get("submitted_at",""),
                "caseId": r.get("caseId",""),
                "caseTitle": r.get("caseTitle",""),
                "mode": r.get("mode",""),
                "cohort": r.get("cohort",""),
                "intake_score": r.get("intake_score",""),
                "total_score": r.get("total_score",""),
                "total_with_intake": r.get("total_with_intake",""),
                "nclex_score": r.get("nclex_score",""),
                "nclex_total": r.get("nclex_total",""),
                "duration_seconds": r.get("duration_seconds",""),
                "domain_scores_json": json.dumps(r.get("domain_scores", {}), ensure_ascii=False),
                "performance_by_section_json": json.dumps(r.get("performance_by_section", {}), ensure_ascii=False),
                "nclex_changes_json": json.dumps(r.get("nclex_changes", {}), ensure_ascii=False),
                "reflection": r.get("reflection",""),
            }
            w.writerow(row)
        return buf.getvalue().encode("utf-8")

    # ---- Fallback: teaching attempts export (original behavior) ----
    rows = []
    try:
        for rec in iter_attempts() or []:
            if isinstance(rec, dict):
                rows.append(_flatten_attempt_row(rec))
    except Exception:
        rows = []

    headers = [
        "submitted_at","started_at","duration_seconds",
        "student_username","student_id","cohort",
        "caseId","caseTitle","mode",
        "intake_score","intake_breakdown_json",
        "A_score","B_score","C_score","D_score","E_score",
        "A_selected_json","B_selected_json","C_selected_json","D_selected_json","E_text",
        "A_notes","B_notes","C_notes","D_notes",
        "SBAR_text",
        "unsafe_actions_count","unsafe_actions_json",
        "nclex_total","nclex_score","nclex_items_json",
        "nclex_answers_json","answers_json","unsafe_counts_json","unsafe_flags_json",
        "nclex_changes_json",
        "performance_by_section_json",
        "reflection",
        "total_score","total_with_intake",
        "duration_seconds"
    ]

    buf = io.StringIO()
    w = csv.DictWriter(buf, fieldnames=headers, extrasaction="ignore")
    w.writeheader()

    for r in rows:
        # Optional: anonymize identifiers in the export when research policy asks for it
        if bool(rp.get("enabled", False)) and bool(rp.get("anonymize_student_id", True)):
            r["student_username"] = ""
            r["student_id"] = _hash_participant(str(r.get("student_id","") or r.get("student_username","") or ""))
        w.writerow(r)

    return buf.getvalue().encode("utf-8")


def build_attempt_summary_csv_bytes() -> bytes:
    """Download Attempt Summary CSV (latest attempt per student+case)."""
    latest = {}
    try:
        for rec in iter_attempts() or []:
            if not isinstance(rec, dict):
                continue
            key = (str(rec.get("student_username","")), str(rec.get("caseId","")))
            sub = str(rec.get("submitted_at") or rec.get("submittedAt") or "")
            prev = latest.get(key)
            if not prev:
                latest[key] = rec
            else:
                prev_sub = str(prev.get("submitted_at") or prev.get("submittedAt") or "")
                if sub and (not prev_sub or sub >= prev_sub):
                    latest[key] = rec
                elif not sub:
                    latest[key] = rec
    except Exception:
        latest = {}

    headers = [
        "student_username","student_id","cohort",
        "caseId","caseTitle","mode",
        "latest_submitted_at",
        "intake_score",
        "A_score","B_score","C_score","D_score","E_score",
        "total_score","total_with_intake",
        "nclex_score","nclex_total",
        "duration_seconds"
    ]

    buf = io.StringIO()
    w = csv.DictWriter(buf, fieldnames=headers, extrasaction="ignore")
    w.writeheader()

    for (stu, cid), rec in sorted(latest.items(), key=lambda x: (x[0][0], x[0][1])):
        scores = rec.get("scores") if isinstance(rec.get("scores"), dict) else {}
        row = {
            "student_username": rec.get("student_username",""),
            "student_id": rec.get("student_id",""),
            "cohort": rec.get("cohort",""),
            "caseId": rec.get("caseId",""),
            "caseTitle": rec.get("caseTitle",""),
            "mode": rec.get("mode", rec.get("attempt_mode","")),
            "latest_submitted_at": rec.get("submitted_at", rec.get("submittedAt","")),
            "intake_score": rec.get("intake_score",""),
            "A_score": scores.get("A",""),
            "B_score": scores.get("B",""),
            "C_score": scores.get("C",""),
            "D_score": scores.get("D",""),
            "E_score": scores.get("E",""),
            "total_score": rec.get("total_score", rec.get("totalScore","")),
            "total_with_intake": rec.get("total_with_intake", rec.get("totalWithIntake","")),
            "nclex_score": rec.get("nclex_score",""),
            "nclex_total": rec.get("nclex_total",""),
            "duration_seconds": rec.get("duration_seconds", rec.get("durationSecs","")),
        }
        w.writerow(row)

    return buf.getvalue().encode("utf-8")


# =============================
# Psychometrics / Item Analysis (NCLEX) — computed export (Excel)
# =============================
def _safe_float(x, default=None):
    try:
        if x is None or x == "":
            return default
        return float(x)
    except Exception:
        return default

def _point_biserial(item_scores: list[int], total_scores: list[float]):
    """Point-biserial correlation between item (0/1) and total score."""
    try:
        n = len(item_scores)
        if n < 5:
            return None
        # totals variance
        mt = sum(total_scores) / n
        vt = sum((t - mt) ** 2 for t in total_scores)
        if vt <= 0:
            return None
        st = math.sqrt(vt / (n - 1))
        if st <= 0:
            return None

        p = sum(item_scores) / n
        q = 1 - p
        if p <= 0 or p >= 1:
            return None

        # mean total for correct vs incorrect
        tot1 = [total_scores[i] for i in range(n) if item_scores[i] == 1]
        tot0 = [total_scores[i] for i in range(n) if item_scores[i] == 0]
        if len(tot1) < 2 or len(tot0) < 2:
            return None
        m1 = sum(tot1) / len(tot1)
        m0 = sum(tot0) / len(tot0)
        r_pb = (m1 - m0) / st * math.sqrt(p * q)
        return r_pb
    except Exception:
        return None

def _top_bottom_discrimination(item_scores: list[int], total_scores: list[float], frac: float = 0.27):
    """Top–bottom discrimination (difference in p between top and bottom groups)."""
    try:
        n = len(item_scores)
        if n < 10:
            return None
        idx = list(range(n))
        idx.sort(key=lambda i: total_scores[i])
        g = max(1, int(round(n * frac)))
        bottom = idx[:g]
        top = idx[-g:]
        p_top = sum(item_scores[i] for i in top) / len(top)
        p_bottom = sum(item_scores[i] for i in bottom) / len(bottom)
        return p_top - p_bottom
    except Exception:
        return None

def _kr20_from_matrix(matrix: list[list[int]]):
    """KR-20 from a 0/1 matrix (rows=examinees, cols=items)."""
    try:
        if not matrix:
            return None
        k = len(matrix[0])
        if k < 2:
            return None
        # total scores
        totals = [sum(row) for row in matrix]
        n = len(totals)
        if n < 3:
            return None
        mean_t = sum(totals) / n
        var_t = sum((t - mean_t) ** 2 for t in totals) / (n - 1) if n > 1 else 0.0
        if var_t <= 0:
            return None

        # sum p*q across items
        pq_sum = 0.0
        for j in range(k):
            col = [row[j] for row in matrix]
            p = sum(col) / n
            q = 1 - p
            pq_sum += p * q

        kr20 = (k / (k - 1)) * (1 - (pq_sum / var_t))
        return kr20
    except Exception:
        return None

def _index_sheet(ws):
    # Basic readability
    ws.freeze_panes = "A2"
    try:
        ws.auto_filter.ref = ws.dimensions
    except Exception:
        pass

def _write_table(ws, headers, rows):
    from openpyxl.styles import Font, Alignment
    bold = Font(bold=True)
    ws.append(list(headers))
    for cell in ws[1]:
        cell.font = bold
        cell.alignment = Alignment(wrap_text=True, vertical="top")
    for r in rows:
        ws.append([r.get(h, "") for h in headers])
    _index_sheet(ws)

def build_nclex_psychometrics_excel_bytes(min_attempts_per_item: int = 10, min_items_intersection: int = 10) -> bytes:
    """Build a multi-sheet Excel psychometrics report from attempts_log.jsonl.

    This exports COMPUTED results (difficulty, discrimination, KR-20) — not raw logs.
    """
    from io import BytesIO
    from openpyxl import Workbook

    attempts = _load_attempts_records(ATTEMPTS_PATH)

    # Build lookup for item metadata (difficulty/client_need/topic/type/correct)
    try:
        bank = load_nclex_items()
        by_case = (bank.get("cases") or {}) if isinstance(bank, dict) else {}
    except Exception:
        by_case = {}

    item_meta = {}
    for cid, pack in (by_case or {}).items():
        if not isinstance(pack, dict):
            continue
        for it in (pack.get("items") or []):
            if not isinstance(it, dict):
                continue
            qid = str(it.get("id","") or "")
            if not qid:
                continue
            item_meta[qid] = {
                "caseId": str(cid),
                "type": it.get("type",""),
                "difficulty_tag": it.get("difficulty",""),
                "client_need": it.get("client_need",""),
                "topic": it.get("topic",""),
                "correct_answer": it.get("correct",""),
                "max_points": 1,
            }

    # Parse attempts into per-attempt dictionaries
    per_attempt_rows = []
    # per item: list of (score, total_excluding_item, total, student, caseId)
    per_item_obs = {}
    # per case: list of attempt vectors {qid:0/1} for KR20 computation
    per_case_attempts = {}

    for rec in attempts:
        if not isinstance(rec, dict):
            continue
        case_id = str(rec.get("caseId","") or "")
        student = str(rec.get("student_username","") or "")
        ts = rec.get("submitted_at") or rec.get("submittedAt") or rec.get("timestamp") or ""
        nblob = rec.get("nclex") or {}
        ndetails = nblob.get("details") or []
        if not isinstance(ndetails, list) or not ndetails:
            continue

        # scores by qid
        scores = {}
        for d in ndetails:
            if not isinstance(d, dict):
                continue
            qid = str(d.get("qid","") or "")
            if not qid:
                continue
            mx = _safe_float(d.get("max"), default=1.0)
            pts = _safe_float(d.get("points"), default=0.0)
            if mx is None or mx <= 0:
                mx = 1.0
            # treat as dichotomous 0/1 for psychometrics
            sc01 = 1 if (pts >= mx) else 0
            scores[qid] = sc01

        if not scores:
            continue

        total = float(sum(scores.values()))
        k = int(len(scores))
        per_attempt_rows.append({
            "submitted_at": ts,
            "student_username": student,
            "caseId": case_id,
            "items_answered": k,
            "total_score_0_1": total,
            "pct": round((total / k) * 100, 1) if k else "",
        })

        # store per item observations
        for qid, sc in scores.items():
            per_item_obs.setdefault(qid, []).append((int(sc), float(total - sc), float(total), student, case_id))

        # for KR-20 per case
        per_case_attempts.setdefault(case_id or "—", []).append(scores)

    # Compute item stats
    item_rows = []
    for qid, obs in per_item_obs.items():
        n = len(obs)
        if n < int(min_attempts_per_item or 0):
            continue

        item_scores = [o[0] for o in obs]
        total_minus = [o[1] for o in obs]  # discrimination should exclude the item
        total_full = [o[2] for o in obs]
        p = sum(item_scores) / n if n else None
        r_pb = _point_biserial(item_scores, total_minus)
        d_top = _top_bottom_discrimination(item_scores, total_minus)

        meta = item_meta.get(qid, {})
        # flags (common QA thresholds)
        flag_low_p = (p is not None and p < 0.20)
        flag_high_p = (p is not None and p > 0.90)
        flag_low_disc = (r_pb is not None and r_pb < 0.20)

        item_rows.append({
            "qid": qid,
            "caseId": meta.get("caseId") or (obs[0][4] if obs else ""),
            "type": meta.get("type",""),
            "difficulty_tag": meta.get("difficulty_tag",""),
            "client_need": meta.get("client_need",""),
            "topic": meta.get("topic",""),
            "n_attempts": n,
            "p_value": round(p, 4) if p is not None else "",
            "discrimination_rpb": round(r_pb, 4) if r_pb is not None else "",
            "disc_top27_minus_bottom27": round(d_top, 4) if d_top is not None else "",
            "flag_low_p(<0.20)": "YES" if flag_low_p else "",
            "flag_high_p(>0.90)": "YES" if flag_high_p else "",
            "flag_low_disc(<0.20)": "YES" if flag_low_disc else "",
        })

    # KR-20 per case (intersection-based)
    kr_rows = []
    for case_id, score_dicts in per_case_attempts.items():
        if not score_dicts:
            continue
        # intersection items across attempts
        sets = [set(d.keys()) for d in score_dicts if isinstance(d, dict)]
        if not sets:
            continue
        common = set.intersection(*sets) if sets else set()
        if len(common) < int(min_items_intersection or 0):
            kr_rows.append({
                "caseId": case_id,
                "attempts": len(score_dicts),
                "common_items": len(common),
                "kr20": "",
                "note": f"Not enough common items for KR-20 (need ≥ {int(min_items_intersection)})."
            })
            continue
        # build matrix
        common_sorted = sorted(common)
        matrix = []
        for d in score_dicts:
            row = []
            for qid in common_sorted:
                row.append(1 if int(d.get(qid, 0) or 0) == 1 else 0)
            matrix.append(row)
        kr = _kr20_from_matrix(matrix)
        kr_rows.append({
            "caseId": case_id,
            "attempts": len(score_dicts),
            "common_items": len(common_sorted),
            "kr20": round(kr, 4) if kr is not None else "",
            "note": "" if kr is not None else "KR-20 not computable (zero variance or insufficient attempts)."
        })

    # Build workbook
    wb = Workbook()
    ws0 = wb.active
    ws0.title = "README"
    ws0.append(["NCLEX Psychometrics Report"])
    ws0.append(["Generated at", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    ws0.append(["Source", "attempts_log.jsonl (NCLEX details stored per attempt)"])
    ws0.append(["Filters", f"min_attempts_per_item={int(min_attempts_per_item)}, min_items_intersection={int(min_items_intersection)}"])
    ws0.append([])
    ws0.append(["Sheets"])
    ws0.append(["ItemStats", "Difficulty (p-value) + discrimination + flags"])
    ws0.append(["KR20_by_case", "KR-20 computed per case using common items across attempts"])
    ws0.append(["Attempts_NCLEX", "One row per attempt with NCLEX total"])
    ws0.append([])
    ws0.append(["Notes"])
    ws0.append(["- Difficulty (p-value): proportion correct (0–1). Ideal range often ~0.3–0.8."])
    ws0.append(["- Discrimination (r_pb): point-biserial between item score and total score excluding that item."])
    ws0.append(["- KR-20: internal consistency for dichotomous items, computed only when enough common items exist."])

    # ItemStats
    ws1 = wb.create_sheet("ItemStats")
    headers1 = [
        "qid","caseId","type","difficulty_tag","client_need","topic",
        "n_attempts","p_value","discrimination_rpb","disc_top27_minus_bottom27",
        "flag_low_p(<0.20)","flag_high_p(>0.90)","flag_low_disc(<0.20)"
    ]
    _write_table(ws1, headers1, item_rows)

    # KR20
    ws2 = wb.create_sheet("KR20_by_case")
    headers2 = ["caseId","attempts","common_items","kr20","note"]
    _write_table(ws2, headers2, kr_rows)

    # Attempts
    ws3 = wb.create_sheet("Attempts_NCLEX")
    headers3 = ["submitted_at","student_username","caseId","items_answered","total_score_0_1","pct"]
    _write_table(ws3, headers3, per_attempt_rows)

    # Finalize
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# =============================
# Student credential exports (Excel / Word)
# =============================
def build_credentials_xlsx_bytes(rows: list, title: str = "Student Credentials") -> bytes:
    """Build an Excel .xlsx file in memory with student credentials."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment
    except Exception:
        # openpyxl should be available, but fail safely
        return b""

    wb = Workbook()
    ws = wb.active
    ws.title = "Credentials"

    ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:E1")

    headers = ["display_name", "username", "password", "student_id", "cohort"]
    ws.append(headers)
    for cell in ws[2]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")

    for r in (rows or []):
        ws.append([
            str(r.get("display_name","") or ""),
            str(r.get("username","") or ""),
            str(r.get("password","") or ""),
            str(r.get("student_id","") or ""),
            str(r.get("cohort","") or ""),
        ])

    # Auto width (simple)
    for col in ["A","B","C","D","E"]:
        max_len = 10
        for cell in ws[col]:
            try:
                max_len = max(max_len, len(str(cell.value or "")))
            except Exception:
                pass
        ws.column_dimensions[col].width = min(40, max_len + 2)

    bio = io.BytesIO()
    wb.save(bio)
    return bio.getvalue()


def build_credentials_docx_bytes(rows: list, title: str = "Student Credentials") -> bytes:
    """Build a Word .docx file in memory with a table of credentials."""
    try:
        from docx import Document
    except Exception:
        return b""

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    headers = ["Display name", "Username", "Password", "Student ID", "Cohort"]
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for r in (rows or []):
        row_cells = table.add_row().cells
        row_cells[0].text = str(r.get("display_name","") or "")
        row_cells[1].text = str(r.get("username","") or "")
        row_cells[2].text = str(r.get("password","") or "")
        row_cells[3].text = str(r.get("student_id","") or "")
        row_cells[4].text = str(r.get("cohort","") or "")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def archive_attempt_logs(clear_after: bool = True) -> tuple[bool, str]:
    """Archive attempts_log.jsonl into BACKUP_DIR with date stamp; optionally clear."""
    try:
        if not ATTEMPTS_PATH.exists():
            return False, "No attempts_log.jsonl to archive."
        BACKUP_DIR.mkdir(parents=True, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        out = BACKUP_DIR / f"attempts_log_{ts}.jsonl"
        out.write_bytes(ATTEMPTS_PATH.read_bytes())
        if clear_after:
            ATTEMPTS_PATH.write_text("", encoding="utf-8")
        return True, f"Archived to: {out.name}" + (" (log cleared)" if clear_after else "")
    except Exception as e:
        return False, f"Archive failed: {e}"

def attempts_count_for(student_username: str, case_id: str) -> int:
    if not student_username or not case_id:
        return 0
    c = 0
    for rec in iter_attempts():
        if str(rec.get("student_username", "")) == str(student_username) and str(rec.get("caseId", "")) == str(case_id):
            c += 1
    return c


def load_nclex_items():
    """
    Load NCLEX-style practice bank and normalize to the structure this app expects:

        {"cases": {case_id: {"items": [ ... ]}}}

    Supports generator outputs like:
        {"items": {case_id: [ ... ]}}
        {"practical": {case_id: [ ... ]}}
        {"cases": {case_id: [ ... ]}}  (list values)
        {"cases": {case_id: {"items": [ ... ]}}}

    This prevents errors like:
        AttributeError: 'list' object has no attribute 'get'
    """
    data = load_json_safe(NCLEX_ITEMS_PATH, {})
    if not isinstance(data, dict):
        data = {}

    # Pick the best available container
    cases = data.get("cases")
    if not isinstance(cases, dict) or not cases:
        practical = data.get("practical")
        items = data.get("items")
        if isinstance(practical, dict) and practical:
            cases = practical
        elif isinstance(items, dict) and items:
            cases = items
        else:
            cases = {}

    # Normalize each case pack to {"items": [...]}
    norm_cases = {}
    for cid, pack in (cases or {}).items():
        if isinstance(pack, dict):
            its = pack.get("items", [])
            if isinstance(its, list):
                norm_cases[str(cid)] = {"items": its}
            elif isinstance(pack.get("items"), dict) and isinstance(list(pack["items"].values())[0], list):
                # rare nested shape; fall back to empty list
                norm_cases[str(cid)] = {"items": []}
            else:
                # If dict but no items list, keep but ensure items key exists
                norm_cases[str(cid)] = dict(pack)
                if not isinstance(norm_cases[str(cid)].get("items"), list):
                    norm_cases[str(cid)]["items"] = []
        elif isinstance(pack, list):
            norm_cases[str(cid)] = {"items": pack}
        else:
            norm_cases[str(cid)] = {"items": []}

    # Keep aliases in sync so any lookup style works
    data["cases"] = norm_cases
    data["items"] = norm_cases
    data["practical"] = norm_cases
    return data


def save_nclex_items(data: dict):
    NCLEX_ITEMS_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def load_nclex_policy():
    data = load_json_safe(NCLEX_POLICY_PATH, {"enabled": False})
    if not isinstance(data, dict):
        data = {"enabled": False}

    # Backward-compatible defaults
    data.setdefault("enabled", False)
    data.setdefault("mode_visibility", {"Practice": True, "Exam": True})
    data.setdefault("items_per_case", 30)
    data.setdefault("per_case_items", {})
    if not isinstance(data.get("per_case_items"), dict):
        data["per_case_items"] = {}
    data.setdefault("optional_practical_section", True)
    data.setdefault("show_correct_answers_after_submit", True)
    data.setdefault("show_rationales_after_submit", True)
    data.setdefault("shuffle_options", True)
    data.setdefault("show_review_after_finalize", True)
    data.setdefault("ai_explanations_after_finalize", True)

    # Security/variation enhancements
    data.setdefault("watermark_enabled", False)  # (deprecated; watermark removed)
    data.setdefault("footer_session_code_enabled", True)
    data.setdefault("one_question_at_a_time", False)
    data.setdefault("randomize_per_student_session", True)

    # Admin-controlled rotation: only used when admin generates a new active set
    data.setdefault("rotation_enabled", False)

    # Enabled types guard
    if not isinstance(data.get("enabled_types"), dict):
        data["enabled_types"] = {
            "mcq": True,
            "sata": True,
            "ordered_response": True,
            "cloze": True,
            "matrix": True,
            "evolving_case": True
        }
    return data



def nclex_items_per_case(policy: dict, case_id: str, default_k: int = 30) -> int:
    """Return effective NCLEX item count for a case (global items_per_case + optional per_case override)."""
    try:
        pol = policy if isinstance(policy, dict) else {}
        k = int(pol.get("items_per_case", default_k) or default_k)
        pcm = pol.get("per_case_items") or {}
        if isinstance(pcm, dict):
            ov = pcm.get(str(case_id)) or pcm.get(str(case_id).strip())
            if ov is not None:
                ov_s = str(ov).strip()
                if ov_s and ov_s.lower() not in ("none", "default", "unlimited", "0"):
                    ov_i = int(float(ov_s))
                    if ov_i > 0:
                        k = ov_i
        return max(1, int(k))
    except Exception:
        try:
            return max(1, int(default_k))
        except Exception:
            return 30



def save_nclex_policy(policy: dict):
    """Persist NCLEX policy settings to nclex_policy.json (safe, minimal)."""
    flash_success('Saved.')
    if not isinstance(policy, dict):
        return
    NCLEX_POLICY_PATH.write_text(
        json.dumps(policy, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )



def load_nclex_active_sets():
    """Stores admin-generated 'active' NCLEX question sets per case (rotation).
    Shape:
      {
        "by_case": {
          "<case_id>": {
            "qids": [...],
            "generated_at": "...",
            "generated_by": "admin",
            "seed": "...",
            "history": [
              {"qids":[...], "generated_at":"...", "generated_by":"admin", "seed":"..."},
              ...
            ]
          }
        }
      }
    """
    ensure_file(NCLEX_ACTIVE_SETS_PATH, {"by_case": {}})
    data = load_json_safe(NCLEX_ACTIVE_SETS_PATH, {"by_case": {}})
    if not isinstance(data, dict):
        data = {"by_case": {}}
    if not isinstance(data.get("by_case"), dict):
        data["by_case"] = {}

    # ensure per-case shapes
    for cid, rec in list((data.get("by_case") or {}).items()):
        if not isinstance(rec, dict):
            rec = {}
            data["by_case"][cid] = rec
        if "history" not in rec or not isinstance(rec.get("history"), list):
            rec["history"] = []

    return data

def save_nclex_active_sets(data: dict):
    try:
        if not isinstance(data, dict):
            data = {"by_case": {}}
        if not isinstance(data.get("by_case"), dict):
            data["by_case"] = {}
        NCLEX_ACTIVE_SETS_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


def build_instructor_key_pdf_bytes(
    case_obj: dict,
    intake_gold: dict,
    gold: dict,
    sbar_expected: str,
    nclex_items: list,
    sections: list | None = None,
    nclex_type_allow: list | None = None,
) -> bytes:
    """Generate Instructor Key PDF (admin-only) and return bytes.

    Enhancements:
    - `sections`: choose which parts to include: ["ae", "intake", "nclex"]
    - `nclex_type_allow`: filter NCLEX items by type (e.g., ["mcq","sata"]).
      If None/empty/"all" present => include all provided items.
    """
    buf = io.BytesIO()
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(buf, pagesize=letter, title="ClinIQ Nurse - Instructor Key")
    story = []

    title = str(case_obj.get("title", "") or "Instructor Key")
    cid = str(case_obj.get("id", "") or "")
    system = str(case_obj.get("system", case_obj.get("category", "")) or "")
    setting = str(case_obj.get("setting", "") or "")
    patient = str(case_obj.get("patient", "") or case_obj.get("patientDesc", "") or "")

    # Normalize section selection
    sec_set = set([s.strip().lower() for s in (sections or ["ae", "intake", "nclex"]) if str(s).strip()])
    if not sec_set:
        sec_set = {"ae", "intake", "nclex"}

    # Normalize NCLEX type filter
    allow_types = None
    if isinstance(nclex_type_allow, list) and nclex_type_allow:
        allow_types = set([str(t).strip().lower() for t in nclex_type_allow if str(t).strip()])
        if "all" in allow_types:
            allow_types = None

    # Apply type filtering (only if NCLEX section included)
    use_nclex = list(nclex_items or [])
    if "nclex" in sec_set and allow_types:
        use_nclex = [it for it in use_nclex if str((it or {}).get("type","")).strip().lower() in allow_types]

    story.append(Paragraph(f"<b>{title}</b>", styles["Title"]))
    meta = f"Case ID: {cid}<br/>System: {system}<br/>Setting: {setting}"
    if patient:
        meta += f"<br/>Patient: {patient}"
    story.append(Paragraph(meta, styles["Normal"]))
    story.append(Spacer(1, 12))

    def add_list_section(label: str, items):
        story.append(Paragraph(f"<b>{label}</b>", styles["Heading2"]))
        if not items:
            story.append(Paragraph("—", styles["Normal"]))
            story.append(Spacer(1, 8))
            return
        for it in items:
            story.append(Paragraph(f"• {str(it)}", styles["Normal"]))
        story.append(Spacer(1, 10))

    # A–E keys
    if "ae" in sec_set:
        story.append(Paragraph("<b>Clinical Reasoning Key (A–E)</b>", styles["Heading1"]))
        add_list_section("Assessment (Gold)", (gold or {}).get("keyAssessments") or (gold or {}).get("assessments") or [])
        add_list_section("Prioritize (Gold)", (gold or {}).get("priorities") or [])
        add_list_section("Interventions (Gold)", (gold or {}).get("interventions") or [])
        add_list_section("Reassess (Gold)", (gold or {}).get("reassessment") or (gold or {}).get("reassess") or [])
        story.append(Paragraph("<b>SBAR (Expected)</b>", styles["Heading2"]))
        story.append(Paragraph((sbar_expected or "—").replace("\n", "<br/>"), styles["Normal"]))
        story.append(Spacer(1, 12))

    # Intake key
    if "intake" in sec_set:
        if story:
            story.append(PageBreak())
        story.append(Paragraph("<b>Student Intake Key</b>", styles["Heading1"]))
        if isinstance(intake_gold, dict) and intake_gold:
            for k, v in intake_gold.items():
                story.append(Paragraph(f"<b>{k}:</b> {str(v)}", styles["Normal"]))
        else:
            story.append(Paragraph("—", styles["Normal"]))
        story.append(Spacer(1, 12))

    # NCLEX key
    if "nclex" in sec_set:
        if story:
            story.append(PageBreak())
        story.append(Paragraph("<b>NCLEX Key</b>", styles["Heading1"]))
        if allow_types:
            story.append(Paragraph(f"Types included: {', '.join(sorted(allow_types))}", styles["Normal"]))
            story.append(Spacer(1, 8))
        if not use_nclex:
            story.append(Paragraph("No NCLEX items available for the selected filter.", styles["Normal"]))
        else:
            for i, it in enumerate(use_nclex, start=1):
                if not isinstance(it, dict):
                    continue
                qid = str(it.get("id","") or "")
                qtype = str(it.get("type","") or "")
                diff = str(it.get("difficulty","") or "")
                stem = (it.get("stem") or it.get("prompt") or "")
                correct = it.get("correct")
                rationale = it.get("rationale") or ""
                story.append(Paragraph(f"<b>Q{i}. {qid}</b> ({qtype}, {diff})", styles["Heading3"]))
                story.append(Paragraph(str(stem), styles["Normal"]))
                story.append(Paragraph(f"<b>Correct:</b> {str(correct)}", styles["Normal"]))
                if rationale:
                    story.append(Paragraph(f"<b>Rationale:</b> {str(rationale)}", styles["Normal"]))
                story.append(Spacer(1, 10))

    doc.build(story)
    return buf.getvalue()

def load_kpi_policy():
    ensure_file(KPI_POLICY_PATH, {
        # Core (original)
        "total_score": True,
        "domain_profile": True,
        "nclex_score": True,
        "time_to_completion": True,
        "unsafe_flags": True,
        "attempts_per_case": True,

        # Expanded KPIs (restored catalog; safe defaults)
        "intake_score": True,
        "completion_rate": True,
        "domain_missed_profile": True,
        "unsafe_by_domain": True,
        "time_per_section": False,
        "nclex_by_type_accuracy": True,
        "nclex_difficulty_accuracy": True,
        "top_wrong_items": True,
        "attempts_over_time": False,
        "student_rank_summary": False,
        "case_rank_summary": True
    })
    try:
        with open(KPI_POLICY_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        data = {}

    if not isinstance(data, dict):
        data = {}

    data.setdefault("total_score", True)
    data.setdefault("domain_profile", True)
    data.setdefault("nclex_score", True)
    data.setdefault("time_to_completion", True)
    data.setdefault("unsafe_flags", True)
    data.setdefault("attempts_per_case", True)
    return data

def save_kpi_policy(policy: dict):
    flash_success('Saved.')
    KPI_POLICY_PATH.write_text(json.dumps(policy, ensure_ascii=False, indent=2), encoding="utf-8")


# =============================
# Students auth
# =============================
def generate_password(length=12):
    alphabet = string.ascii_letters + string.digits + "@#_-!"
    return "".join(secrets.choice(alphabet) for _ in range(length))


def upsert_students(existing: dict, new_students: list):
    students = existing.get("students", [])
    by_username = {str(s.get("username", "")).strip(): s for s in students}
    for ns in new_students:
        u = str(ns.get("username", "")).strip()
        if not u:
            continue
        by_username[u] = ns
    existing["students"] = list(by_username.values())
    return existing


def verify_student(username: str, password: str):
    """
    Login logic:
    - If Exam password policy is ACTIVE and has hash:
        - If entered password matches exam password => allow login for ANY valid username
        - If force_exam_password_only is True => reject personal passwords during exam
        - Else fallback to personal password
    - Otherwise personal password only
    """
    username = (username or "").strip()
    password = (password or "")
    if not username or not password:
        return None

    data = load_students()
    student_obj = None
    for s in data.get("students", []):
        if str(s.get("username", "")).strip() == username:
            student_obj = s
            break
    if not student_obj:
        return None

    policy = load_exam_access_policy()
    exam_active = is_exam_password_active(policy)
    exam_hash = (policy.get("exam_password_sha256") or "").strip().lower()
    force_only = bool(policy.get("force_exam_password_only", True))

    if exam_active and exam_hash:
        if sha256_hex(password).lower() == exam_hash:
            return {
                "username": username,
                "display_name": student_obj.get("display_name", username),
                "student_id": student_obj.get("student_id", ""),
                "cohort": student_obj.get("cohort", ""),
            }
        if force_only:
            return None

    stored = (student_obj.get("password_sha256") or "").strip().lower()
    if stored and sha256_hex(password).lower() == stored:
        return {
            "username": username,
            "display_name": student_obj.get("display_name", username),
            "student_id": student_obj.get("student_id", ""),
            "cohort": student_obj.get("cohort", ""),
        }

    return None


# =============================
# Safe getters
# =============================
def safe_get_system(case: dict) -> str:
    return (case.get("system") or case.get("category") or "Uncategorized").strip()


def safe_get_setting(case: dict) -> str:
    return (case.get("setting") or "Unspecified").strip()

def fmt_patient(p: dict) -> str:
    if not isinstance(p, dict):
        return ""
    age = p.get("age", "")
    sex = p.get("sex", "")
    parts = []
    if age != "":
        parts.append(f"{age}-year-old")
    if sex:
        sx = str(sex).strip().upper()
        if sx.startswith("F"):
            parts.append("female")
        elif sx.startswith("M"):
            parts.append("male")
        else:
            parts.append(str(sex).strip())
    return " ".join(parts).strip()

def _join_list(x) -> str:
    if isinstance(x, list):
        return "; ".join(str(i) for i in x if str(i).strip())
    if x is None:
        return ""
    s = str(x).strip()
    return s

def fmt_history(h: dict) -> str:
    if not isinstance(h, dict):
        return ""
    pmh = _join_list(h.get("pmh", []))
    meds = _join_list(h.get("meds", []))
    allergies = _join_list(h.get("allergies", []))
    lines = []
    if pmh:
        lines.append(f"**PMH:** {pmh}")
    if meds:
        lines.append(f"**Meds:** {meds}")
    if allergies:
        lines.append(f"**Allergies:** {allergies}")
    return "\n\n".join(lines)

def fmt_kv_block(d: dict) -> str:
    if not isinstance(d, dict) or not d:
        return ""
    lines = []
    for k, v in d.items():
        if isinstance(v, (dict, list)):
            vv = json.dumps(v, ensure_ascii=False)
        else:
            vv = str(v)
        lines.append(f"- **{k}**: {vv}")
    return "\n".join(lines)

def fmt_findings(lst) -> str:
    if not isinstance(lst, list) or not lst:
        return ""
    return "\n".join(f"- {str(x)}" for x in lst if str(x).strip())


def get_ui_list(case: dict, path: list, fallback: list):
    cur = case
    for p in path:
        if not isinstance(cur, dict) or p not in cur:
            return fallback
        cur = cur[p]
    if isinstance(cur, list) and len(cur) > 0:
        return cur
    return fallback


def get_gs_list(case: dict, key_options: list):
    gs = case.get("goldStandard", {}) or {}
    for k in key_options:
        v = gs.get(k)
        if isinstance(v, list) and any(str(i).strip() for i in v):
            return v
    return []


def get_gs_sbar(case: dict):
    gs = case.get("goldStandard", {}) or {}
    sbar = gs.get("sbar", {}) or {}
    return {
        "S": (sbar.get("S", "") or "").strip(),
        "B": (sbar.get("B", "") or "").strip(),
        "A": (sbar.get("A", "") or "").strip(),
        "R": (sbar.get("R", "") or "").strip(),
    }


# =============================
# Case visibility + timer helpers
# =============================
def case_is_visible(case_policy: dict, case_id: str, mode: str = None, now_dt: datetime = None) -> bool:
    """Return True if the case should be shown.

    Backward compatible:
    - Supports old keys: per_case[case_id]['visible']
    - Supports selected/all policy: visibility_mode + visible_case_ids + default_visibility
    - Supports per-mode visibility: per_case[case_id]['mode_visibility'] (Practice/Exam)
    - NEW: Supports per-case availability windows: per_case[case_id]['availability'][mode]['start'|'end'] (ISO)
    """
    case_id = str(case_id or "").strip()
    policy = case_policy or {}
    per = (policy.get("per_case") or {}).get(case_id, {}) or {}
    if not isinstance(per, dict):
        per = {}

    # Base visibility (selected/all)
    vis_mode = str(policy.get("visibility_mode", "all") or "all").lower()
    if vis_mode == "selected":
        visible_ids = policy.get("visible_case_ids", []) or []
        base_visible = case_id in set([str(x).strip() for x in visible_ids])
    else:
        base_visible = bool(policy.get("default_visibility", True))

    # Legacy override: explicit visible flag per case
    # IMPORTANT:
    # - In 'selected' visibility mode, the global selected list is the source of truth.
    #   A per-case 'visible=True' should NOT expand visibility beyond the selected list.
    #   However, 'visible=False' is allowed to hide a case even if selected.
    if "visible" in per:
        try:
            vflag = bool(per.get("visible"))
            if vis_mode == "selected":
                if not vflag:
                    base_visible = False
            else:
                base_visible = vflag
        except Exception:
            pass

    if not base_visible:
        return False

    # Per-mode visibility (if mode provided)
    if mode:
        mv = per.get("mode_visibility", {}) or {}
        if isinstance(mv, dict):
            if not bool(mv.get(str(mode), True)):
                return False

        # Availability window (if configured)
        avail = per.get("availability", {}) or {}
        if isinstance(avail, dict):
            win = avail.get(str(mode), {}) or {}
            if isinstance(win, dict):
                start_s = str(win.get("start", "") or "").strip()
                end_s = str(win.get("end", "") or "").strip()
                if now_dt is None:
                    now_dt = now_local()
                start_dt = parse_iso_dt(start_s) if start_s else None
                end_dt = parse_iso_dt(end_s) if end_s else None
                # Normalize to TZ if naive
                try:
                    if start_dt and start_dt.tzinfo is None:
                        start_dt = start_dt.replace(tzinfo=TZ)
                    if end_dt and end_dt.tzinfo is None:
                        end_dt = end_dt.replace(tzinfo=TZ)
                except Exception:
                    pass
                if start_dt and now_dt < start_dt:
                    return False
                if end_dt and now_dt > end_dt:
                    return False

    return True
def timer_minutes_for(case_policy: dict, case_id: str, mode: str):
    """Return timer minutes for a case/mode.

    Canonical storage is seconds (timer_seconds_* / default_timer_seconds_*).
    Minutes keys are kept for backward compatibility and UI display.
    """
    per = case_policy.get("per_case", {}).get(case_id, {})
    if not isinstance(per, dict):
        per = {}

    # Prefer seconds
    if mode == "Practice":
        sec = per.get("timer_seconds_practice", case_policy.get("default_timer_seconds_practice", 0))
        min_fallback = per.get("timer_minutes_practice", case_policy.get("default_timer_minutes_practice", "unlimited"))
    else:
        sec = per.get("timer_seconds_exam", case_policy.get("default_timer_seconds_exam", 0))
        min_fallback = per.get("timer_minutes_exam", case_policy.get("default_timer_minutes_exam", 20))

    try:
        sec = int(sec or 0)
    except Exception:
        sec = 0

    if sec > 0:
        return int(round(sec / 60))

    # Seconds say unlimited (0) → honor minutes fallback if it's a positive int
    if isinstance(min_fallback, str) and min_fallback.strip().lower() == "unlimited":
        return "unlimited"
    n = safe_int(min_fallback, None)
    if n is None or n <= 0:
        return "unlimited"
    return n


def format_seconds(secs: int) -> str:
    secs = max(0, int(secs))
    m = secs // 60
    s = secs % 60
    return f"{m:02d}:{s:02d}"



def _render_nclex_correct_pretty(correct):
    """Render 'correct answer' in a clean, no-symbols format (bullets / key-value lines)."""
    try:
        if correct is None:
            st.write("—")
            return
        # Dict (e.g., matrix) -> bullet lines "Key — Value"
        if isinstance(correct, dict):
            if not correct:
                st.write("—")
                return
            lines = []
            for k, v in correct.items():
                ks = str(k).strip()
                vs = str(v).strip()
                if not ks and not vs:
                    continue
                if ks and vs:
                    lines.append(f"{ks} — {vs}")
                elif ks:
                    lines.append(ks)
                else:
                    lines.append(vs)
            if not lines:
                st.write("—")
                return
            st.markdown("\n".join([f"- {ln}" for ln in lines]))
            return
        # List/tuple/set -> bullets
        if isinstance(correct, (list, tuple, set)):
            lines = [str(x).strip() for x in list(correct) if str(x).strip()]
            if not lines:
                st.write("—")
                return
            st.markdown("\n".join([f"- {ln}" for ln in lines]))
            return
        # Everything else -> plain
        s = str(correct).strip()
        st.write(s if s else "—")
    except Exception:
        try:
            st.write(str(correct))
        except Exception:
            st.write("—")



# =============================
# Step 8 Rubric Scoring (matching)
# =============================
STOPWORDS = {
    "the", "a", "an", "and", "or", "to", "of", "in", "on", "for", "with", "without", "at", "by",
    "is", "are", "was", "were", "be", "been", "being", "as", "from", "that", "this", "it", "its",
    "patient", "pt"
}


# =============================
# Attempt state reset (safe init)
# =============================
def reset_attempt_state():
    """Reset *attempt-level* state when the selected case changes.

    IMPORTANT:
    - This should NOT wipe section lock flags or widget inputs; it only resets the scoring containers
      and (crucially) re-initializes the attempt timer anchors so a previous case's expired deadline
      cannot lock a new case.
    """
    # Domain scores
    st.session_state["scores"] = {"A": 0, "B": 0, "C": 0, "D": 0, "E": 0}
    # Raw answers payload (saved to attempts_log)
    st.session_state["answers"] = {"A": {}, "B": {}, "C": {}, "D": {}, "E": {}}

    # Timer anchors (reset on case change)
    st.session_state["attempt_started_epoch"] = time.time()
    st.session_state["attempt_deadline_epoch"] = None


def reset_full_attempt_for_case(case_id: str, student_username: str = ""):
    """Full reset for *starting a new attempt* on the same case.

    This clears section submissions, lock/focus state, NCLEX states, and autosave drafts for this student+case.
    """
    case_id = str(case_id or "").strip()
    if not case_id:
        return

    # Core attempt containers
    reset_attempt_state()

    # Intake state
    st.session_state["intake_submitted"] = False
    st.session_state["intake_score"] = 0
    st.session_state["intake_breakdown"] = {}
    st.session_state["intake_answers"] = {}
    st.session_state["show_intake_grade_table"] = False

    # Ensure the case restarts from the very beginning (scenario + video gate)
    st.session_state["nclex_in_progress"] = False
    st.session_state["nclex_finalized"] = False
    st.session_state["show_nclex_review"] = False

    # Reset "I watched the video" confirmation for this case (if present)
    k_ok = f"intro_video_confirmed__{case_id}"
    if k_ok in st.session_state:
        del st.session_state[k_ok]

    # Clear Intake widget keys for this case (so the form is truly fresh)
    for _wk in [f"intake_age_{case_id}", f"intake_setting_{case_id}", f"intake_cc_{case_id}",
                f"intake_sx_{case_id}", f"intake_findings_{case_id}", f"intake_hist_{case_id}",
                f"intake_save_{case_id}", f"intake_submit_{case_id}"]:
        if _wk in st.session_state:
            del st.session_state[_wk]

    # A–E reasoning workflow
    st.session_state["last_feedback"] = {}
    st.session_state["ae_focus"] = "A"
    st.session_state["ae_next_ready"] = False
    st.session_state["ae_next_domain"] = None

    # Clear widget inputs (avoid stale frozen widgets)
    for k in ["A_selected","A_notes","B_selected","B_rationale","C_selected","C_rationale","D_selected","D_rationale","E_S","E_B","E_A","E_R"]:
        if k in st.session_state:
            del st.session_state[k]

    # NCLEX / Practical
    st.session_state["nclex_answers"] = {}
    st.session_state["nclex_scored"] = {}
    st.session_state["nclex_finalized"] = False
    st.session_state["practical_submitted"] = False
    st.session_state["show_nclex_review"] = False
    st.session_state["attempt_saved"] = False
    st.session_state["show_save_attempt"] = False

    # Debrief gate flags (if present)
    for k in ["debrief_ready", "debrief_shown", "ai_debrief_generated", "show_ai_debrief"]:
        if k in st.session_state:
            st.session_state[k] = False

    # Clear per-case lock flags if they exist
    for dom in ["A","B","C","D","E"]:
        lk = f"lock_{case_id}_{dom}"
        if lk in st.session_state:
            st.session_state[lk] = False

    # Reset timer anchors
    st.session_state["attempt_started_epoch"] = time.time()
    st.session_state["attempt_deadline_epoch"] = None

    # Remove autosave drafts for this student+case (prevents immediate re-freeze)
    if student_username:
        delete_autosaves_for(student_username, case_id)


def ensure_section_locks(case_id: str):
    """Ensure per-case section lock flags exist. Does NOT reset inputs."""
    for dom in ["A", "B", "C", "D", "E"]:
        k = f"lock_{case_id}_{dom}"
        if k not in st.session_state:
            st.session_state[k] = False


def ensure_widget_defaults():
    """Ensure widget/session keys exist without overwriting user progress."""
    defaults = {
        # A–E widgets
        "A_selected": [],
        "A_notes": "",
        "B_selected": [],
        "B_rationale": "",
        "C_selected": [],
        "C_rationale": "",
        "D_selected": [],
        "D_notes": "",
        "D_time": "15 minutes",
        "E_S": "",
        "E_B": "",
        "E_A": "",
        "E_R": "",
        "E_selected": [],
        "E_selected_elements": [],  # backward-compat

        # Feedback + AI
        "last_feedback": {"A": None, "B": None, "C": None, "D": None, "E": None},
        "ai_coach": {"A": None, "B": None, "C": None, "D": None, "E": None},
        "ai_debrief": None,

        # Intake
        "intake": {"age": "", "setting": "", "chief_complaint": "", "signs_symptoms": "", "history": ""},
        "intake_submitted": False,
        "intake_score": 0,
        "intake_breakdown": {},

        # NCLEX
        "practical_submitted": False,
        "nclex_answers": {},
        "nclex_score": 0,
        "nclex_total": 0,
        "nclex_scored": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def init_case_state_if_needed(case_id: str, case_title: str = ""):
    # normalize to string to avoid accidental resets (e.g., int vs str ids)
    case_id = str(case_id)
    """
    Initialize/reset attempt state ONLY when the selected case changes.
    This prevents B–E from 'disappearing' after submitting A (no accidental resets).
    """
    if str(st.session_state.get("attempt_case_id") or "") != str(case_id):
        # New case selected → reset attempt state
        st.session_state["attempt_case_id"] = str(case_id)
        st.session_state["attempt_case_title"] = case_title
        st.session_state["attempt_started_epoch"] = time.time()

        # Locks (per-case)
        for dom in ["A", "B", "C", "D", "E"]:
            st.session_state[f"lock_{case_id}_{dom}"] = False

        # Scores + answers containers
        st.session_state["scores"] = {"A": 0, "B": 0, "C": 0, "D": 0, "E": 0}
        st.session_state["answers"] = {"A": {}, "B": {}, "C": {}, "D": {}, "E": {}}

        # Reset UI inputs + feedback for this attempt
        for k in ["A_selected","A_notes","B_selected","B_rationale","C_selected","C_rationale",
                  "D_selected","D_notes","D_time","E_S","E_B","E_A","E_R","E_selected","E_selected_elements"]:
            if k == "D_time":
                st.session_state[k] = st.session_state.get(k, "15 minutes") or "15 minutes"
            else:
                st.session_state[k] = [] if k.endswith("_selected") or k in ("E_selected_elements",) else ""

        st.session_state["last_feedback"] = {"A": None, "B": None, "C": None, "D": None, "E": None}
        st.session_state["ai_coach"] = {"A": None, "B": None, "C": None, "D": None, "E": None}
        st.session_state["ai_debrief"] = None

        # Intake reset
        st.session_state["intake"] = {"age": "", "setting": "", "chief_complaint": "", "signs_symptoms": "", "history": ""}
        st.session_state["intake_submitted"] = False
        st.session_state["intake_score"] = 0
        st.session_state["intake_breakdown"] = {}

        # NCLEX reset
        st.session_state["practical_submitted"] = False
        st.session_state["nclex_answers"] = {}
        st.session_state["nclex_score"] = 0
        st.session_state["nclex_total"] = 0
        st.session_state["nclex_scored"] = None

def norm(s: str) -> str:
    return (s or "").lower().strip()


def tokenize(s: str):
    s = norm(s)
    tokens = re.findall(r"[a-z0-9]+", s)
    return [t for t in tokens if t and t not in STOPWORDS]

# =============================
# Intake (5 marks) + scenario trigger
# =============================
def build_scenario_trigger(case: dict) -> str:
    """Create a more realistic clinical scenario (trigger) from the case.

    Design goals:
    - Reads like a short, real-world handover/triage vignette (not a metadata list).
    - Avoids markdown styling (no **bold** or bullet labels).
    - Includes enough clinical detail to support reasoning, but requires the student to extract key intake elements.
    - Uses only information present in the case object (plus neutral connective phrasing).
    """
    if not isinstance(case, dict):
        return ""

    patient = case.get("patient", {}) if isinstance(case.get("patient", {}), dict) else {}

    age = str(patient.get("age", "") or "").strip()
    sex_raw = str(patient.get("sex", "") or patient.get("gender", "") or "").strip()

    def _sex_full(s: str) -> str:
        s2 = (s or "").strip().lower()
        if s2 in {"f", "female", "woman", "girl"}:
            return "female"
        if s2 in {"m", "male", "man", "boy"}:
            return "male"
        return (s or "").strip()

    sex = _sex_full(sex_raw)

    # We'll mention the location/setting naturally in the story (student must extract it).
    setting = (safe_get_setting(case) or "").strip()
    setting_lower = setting.lower()

    # Chief complaint / presenting problem
    cc = (case.get("chiefComplaint", "") or case.get("chief_complaint", "") or "").strip()

    # Findings (symptoms/exam findings) – take first few
    findings = case.get("findings", [])
    f_list: list[str] = []
    if isinstance(findings, list):
        for f in findings:
            if isinstance(f, str) and f.strip():
                f_list.append(f.strip())
            elif isinstance(f, dict):
                txt = f.get("text") or f.get("finding") or f.get("value")
                if isinstance(txt, str) and txt.strip():
                    f_list.append(txt.strip())
            if len(f_list) >= 8:
                break

    # History (PMH/meds)
    hist = case.get("history", {})
    pmh, meds = [], []
    if isinstance(hist, dict):
        pmh = hist.get("pmh", []) if isinstance(hist.get("pmh", []), list) else []
        meds = hist.get("meds", []) if isinstance(hist.get("meds", []), list) else []
    pmh_line = ", ".join([str(x).strip() for x in pmh if str(x).strip()]).strip()
    meds_line = ", ".join([str(x).strip() for x in meds if str(x).strip()]).strip()

    # Allergies (only state "no known" if explicitly present)
    allergies = case.get("allergies", None)
    if allergies is None and isinstance(patient, dict):
        allergies = patient.get("allergies", None)
    allergy_line = ""
    if isinstance(allergies, str) and allergies.strip():
        allergy_line = allergies.strip()
    elif isinstance(allergies, list):
        allergy_line = ", ".join([str(a).strip() for a in allergies if str(a).strip()]).strip()

    # Vitals (prefer common keys)
    vitals = case.get("vitals", {}) if isinstance(case.get("vitals", {}), dict) else {}

    def _v(key, aliases=()):
        for k in (key, *aliases):
            if k in vitals and str(vitals.get(k)).strip():
                return str(vitals.get(k)).strip()
        return ""

    hr = _v("HR", ("heart_rate", "Pulse"))
    bp = _v("BP", ("blood_pressure",))
    rr = _v("RR", ("resp_rate", "RespiratoryRate"))
    spo2 = _v("SpO2", ("spo2", "O2Sat", "O2"))
    temp = _v("Temp", ("temperature",))
    pain = _v("Pain", ("pain", "PainScore"))

    def _clean(s: str) -> str:
        s = (s or "").strip()
        # Remove accidental markdown markers if they exist in stored text
        s = s.replace("**", "").replace("__", "")
        s = re.sub(r"\s+", " ", s)
        return s.strip()

    def _sentence(s: str) -> str:
        s = _clean(s)
        if not s:
            return ""
        if s[-1] not in ".!?":
            s += "."
        return s

    # --- Narrative construction ---
    lines: list[str] = []

    # Opening: place the nurse in a setting without giving "Setting:" label
    if "er" in setting_lower or "ed" in setting_lower or "emergency" in setting_lower or "triage" in setting_lower:
        lines.append("You are the registered nurse working in a busy emergency department triage area.")
    elif "icu" in setting_lower or "intensive" in setting_lower:
        lines.append("You are the registered nurse caring for a patient in an intensive care environment.")
    elif setting:
        lines.append(f"You are the registered nurse on a {setting} unit.")
    else:
        lines.append("You are the registered nurse receiving a new patient for assessment.")

    # Arrival + demographics (age/sex may remain, but woven in)
    demo_bits = []
    if age:
        demo_bits.append(f"{age}-year-old")
    if sex:
        demo_bits.append(sex)
    demo = " ".join([b for b in demo_bits if b]).strip()
    if demo:
        lines.append(f"A {demo} arrives for evaluation.")
    else:
        lines.append("A patient arrives for evaluation.")

    # Context: use a slight delay/trajectory to force extraction
    if cc:
        lines.append(_sentence(f"The patient describes {cc.lower().rstrip('.')}, stating it began earlier and has been getting worse"))
    elif f_list:
        lines.append(_sentence(f"The patient describes {f_list[0].rstrip('.')}, stating it began earlier and has been getting worse"))
    else:
        lines.append("The patient reports a new problem that is worsening and requires prompt assessment.")

    # Symptoms / assessment cues (reported vs observed) – phrased naturally
    observed = [_clean(x) for x in f_list if _clean(x)]
    # If pain score exists, add it as a detail without spoon-feeding
    if pain and not any("/10" in x or "pain" in x.lower() for x in observed):
        observed = [f"rates the discomfort as {pain}"] + observed

    if observed:
        # Keep some detail but not in a labelled list
        first = observed[:3]
        rest = observed[3:6]
        if first:
            lines.append(_sentence("While answering your questions, the patient mentions " + "; ".join([x.rstrip(".") for x in first])))
        if rest:
            lines.append(_sentence("You also notice " + "; ".join([x.rstrip(".") for x in rest])))

    # History: still present, but embedded
    if pmh_line:
        lines.append(_sentence(f"When you ask about background health issues, the patient reports a history of {pmh_line.rstrip('.')}"))
    if meds_line:
        lines.append(_sentence(f"Regarding home medications, the patient reports {meds_line.rstrip('.')}"))

    if allergy_line:
        al = _clean(allergy_line)
        if al.lower() in {"nkda", "nka", "no known drug allergies", "no known allergies"}:
            lines.append("The patient reports no known drug allergies.")
        else:
            lines.append(_sentence(f"The patient reports allergies to {al.rstrip('.')}"))

    # Vitals: keep numeric values if present, in one sentence
    vit_parts = []
    if hr: vit_parts.append(f"heart rate {hr}")
    if bp: vit_parts.append(f"blood pressure {bp}")
    if rr: vit_parts.append(f"respiratory rate {rr}")
    if spo2: vit_parts.append(f"oxygen saturation {spo2}")
    if temp: vit_parts.append(f"temperature {temp}")
    if vit_parts:
        lines.append(_sentence("Initial vital signs show " + ", ".join(vit_parts)))

    # Closing cue: urgency without giving away the template fields
    lines.append("During your first assessment, the patient indicates the symptoms are continuing to worsen.")

    scenario = " ".join([_clean(s) for s in lines if _clean(s)]).strip()
    scenario = re.sub(r"\s+", " ", scenario)
    return scenario

def _token_set(s: str):
    return set(tokenize(s or ""))

def _overlap_ratio(a: str, b: str) -> float:
    A = _token_set(a)
    B = _token_set(b)
    if not A or not B:
        return 0.0
    return len(A & B) / max(1, len(B))

def score_intake(case: dict, intake: dict) -> tuple[int, dict]:
    """Score student intake out of 5 using overlap against the case ground truth."""
    intake = intake or {}
    patient = case.get("patient", {}) if isinstance(case.get("patient", {}), dict) else {}
    expected_age = str(patient.get("age", "")).strip()
    expected_setting = str(safe_get_setting(case) or "")
    expected_cc = str(case.get("chiefComplaint", "") or "")
    expected_hist = " ".join([
        " ".join((case.get("history", {}) or {}).get("pmh", []) if isinstance((case.get("history", {}) or {}).get("pmh", []), list) else []),
        " ".join((case.get("history", {}) or {}).get("meds", []) if isinstance((case.get("history", {}) or {}).get("meds", []), list) else []),
        " ".join((case.get("history", {}) or {}).get("allergies", []) if isinstance((case.get("history", {}) or {}).get("allergies", []), list) else []),
        str((case.get("history", {}) or {}).get("social", "") or ""),
        str((case.get("history", {}) or {}).get("hpi", "") or ""),
    ]).strip()

    vitals = case.get("vitals", {}) if isinstance(case.get("vitals", {}), dict) else {}
    vit_text = " ".join([f"{k} {vitals.get(k,'')}" for k in vitals.keys() if str(vitals.get(k,'')).strip()])
    findings = case.get("findings", []) if isinstance(case.get("findings", []), list) else []
    findings_text = " ".join([str(x) for x in findings if str(x).strip()])

    expected_sx = (expected_cc + " " + vit_text + " " + findings_text).strip()

    # Student inputs
    s_age = str(intake.get("age", "")).strip()
    s_setting = str(intake.get("setting", "")).strip()
    s_cc = str(intake.get("chief_complaint", "")).strip()
    s_sx = str(intake.get("signs_symptoms", "")).strip()
    s_findings = str(intake.get("findings", "")).strip()
    s_hist = str(intake.get("history", "")).strip()

    breakdown = {}
    score = 0

    # 1) Age (exact or close numeric)
    age_ok = False
    try:
        ea = int(re.findall(r"\d+", expected_age)[0]) if expected_age else None
        sa = int(re.findall(r"\d+", s_age)[0]) if s_age else None
        if ea is not None and sa is not None and abs(ea - sa) <= 5:
            age_ok = True
    except Exception:
        age_ok = False
    if age_ok or (expected_age and s_age and expected_age == s_age):
        score += 1
        breakdown["age"] = 1
    else:
        breakdown["age"] = 0

    # 2) Setting (token overlap)
    setting_ratio = _overlap_ratio(s_setting, expected_setting)
    if expected_setting and setting_ratio >= 0.45:
        score += 1
        breakdown["setting"] = 1
    else:
        breakdown["setting"] = 0

    # 3) Chief complaint
    cc_ratio = _overlap_ratio(s_cc, expected_cc)
    if expected_cc and cc_ratio >= 0.45:
        score += 1
        breakdown["chief_complaint"] = 1
    else:
        breakdown["chief_complaint"] = 0

    # 4) Major signs/symptoms + findings (combined)
    sx_student = (s_sx + " " + s_findings).strip()
    sx_ratio = _overlap_ratio(sx_student, expected_sx)
    if expected_sx and sx_ratio >= 0.30:
        score += 1
        breakdown["signs_symptoms_findings"] = 1
    else:
        breakdown["signs_symptoms_findings"] = 0

    # 5) History
    hist_ratio = _overlap_ratio(s_hist, expected_hist)
    if expected_hist and hist_ratio >= 0.25:
        score += 1
        breakdown["history"] = 1
    else:
        breakdown["history"] = 0

    return int(score), breakdown


# =============================
# Gold-4 options + scoring helpers (4 correct + distractors)
# =============================
SYSTEM_DISTRACTORS = {
    "Cardiovascular": {
        "A": [
            "Ask about long-term diet goals only and delay immediate assessment",
            "Complete discharge planning before stabilizing the patient",
            "Assess pain only and skip focused cardiopulmonary assessment",
            "Perform a full psychosocial interview before obtaining vital signs",
            "Delay vital signs because the patient is awake and talking",
            "Focus on sodium education before reassessing instability",
        ],
        "B": [
            "Prioritize completing documentation before reassessment",
            "Prioritize patient education first despite red flags",
            "Wait for the next scheduled vital signs without reassessment",
            "Focus on comfort measures only despite instability",
            "Arrange routine outpatient follow-up as the first action",
            "Delay escalation to avoid overreacting",
        ],
        "C": [
            "Administer a medication without verifying orders/allergies",
            "Change prescribed doses independently",
            "Allow unassisted ambulation while unstable",
            "Delay escalation and recheck in 1 hour despite red flags",
            "Stop oxygen abruptly without targets/orders",
            "Give extra dose because values are abnormal without protocol",
        ],
        "D": [
            "Reassess only at end of shift regardless of symptoms",
            "Stop monitoring to reduce alarms",
            "Do not trend vital signs after interventions",
            "Reassess only if the patient asks for help",
            "Skip neuro checks despite headache/visual symptoms",
            "Document later without trending",
        ],
        "E": [
            "Leave out objective vital signs to keep SBAR brief",
            "State opinions without objective assessment data",
            "Recommend medication changes without orders",
            "Delay calling provider until next round despite red flags",
            "Do not mention response to interventions",
            "Request discharge planning as the main recommendation",
        ],
    },
    "Respiratory": {
        "A": [
            "Check diet history first and delay respiratory assessment",
            "Do a full skin assessment before assessing airway",
            "Focus on family history only and skip SpO2 assessment",
            "Assess bowel sounds before lung sounds",
            "Delay vital signs to avoid disturbing the patient",
            "Ask about exercise routine before checking work of breathing",
        ],
        "B": [
            "Provide education first before treating hypoxia",
            "Complete documentation before applying oxygen",
            "Encourage ambulation while dyspneic",
            "Wait for respiratory therapist without reassessment",
            "Delay escalation despite increased work of breathing",
            "Address diet as first priority over breathing",
        ],
        "C": [
            "Discontinue oxygen abruptly without targets/orders",
            "Administer sedatives without verifying respiratory status",
            "Delay protocol actions despite wheeze",
            "Change prescribed doses independently",
            "Allow oral intake in severe respiratory distress",
            "Stop monitoring to reduce alarms",
        ],
        "D": [
            "Reassess only once daily regardless of condition",
            "Stop pulse oximetry to reduce alarms",
            "Do not reassess after oxygen/nebulizer treatments",
            "Skip respiratory rate trending",
            "Ignore accessory muscle use during reassessment",
            "Document later without trending",
        ],
        "E": [
            "Do not report SpO2 or oxygen requirements",
            "Provide a vague recommendation with no clear ask",
            "Leave out response to interventions",
            "Delay provider call despite deterioration",
            "Focus SBAR on social history only",
            "Recommend discharge despite unstable status",
        ],
    },
    "Neurological": {
        "A": [
            "Assess diet history only and skip neuro checks",
            "Delay vital signs to complete full psychosocial interview",
            "Skip LOC assessment because patient is talking",
            "Check skin integrity before neuro status",
            "Focus on sleep hygiene before urgent assessment",
            "Assess bowel habits first",
        ],
        "B": [
            "Provide reassurance only despite neuro red flags",
            "Delay reassessment for 1 hour",
            "Complete paperwork before escalation",
            "Focus on discharge planning first",
            "Wait for next scheduled vitals",
            "Address diet first over neuro status",
        ],
        "C": [
            "Give medication without checking contraindications/orders",
            "Allow ambulation without fall precautions",
            "Delay escalation despite acute neuro change",
            "Change prescribed doses independently",
            "Ignore seizure precautions when indicated",
            "Stop monitoring to reduce alarms",
        ],
        "D": [
            "Skip neuro checks during reassessment",
            "Stop monitoring to reduce alarms",
            "Reassess only at end of shift",
            "Do not trend LOC/vitals after interventions",
            "Ignore headache/vision changes",
            "Document later without trending",
        ],
        "E": [
            "Do not mention neuro status/LOC in SBAR",
            "Provide no objective data (vitals, GCS) to provider",
            "Recommend medication changes without orders",
            "Delay provider call despite deterioration",
            "Omit time course/onset details",
            "Ask for non-urgent consult as main recommendation",
        ],
    },
}

GENERIC_DOMAIN_DISTRACTORS = {
    "A": [
        "Complete a full head-to-toe exam before checking ABCs",
        "Focus on documentation first before reassessment",
        "Discuss discharge planning before stabilizing the patient",
        "Assess only pain and ignore abnormal vital signs",
        "Delay vital signs to avoid disturbing the patient",
        "Ask about lifestyle goals only and delay urgent assessment",
    ],
    "B": [
        "Complete documentation before reassessing the patient",
        "Provide education first despite instability",
        "Delay escalation and wait for next scheduled vitals",
        "Focus on comfort only despite red flags",
        "Arrange routine follow-up before stabilization",
        "Address diet first over acute symptoms",
    ],
    "C": [
        "Administer medication without verifying orders/allergies",
        "Change prescribed doses independently",
        "Delay escalation despite red flags",
        "Allow unassisted ambulation while unstable",
        "Stop monitoring to reduce alarms",
        "Provide reassurance only and no interventions",
    ],
    "D": [
        "Reassess only at the end of shift regardless of condition",
        "Do not trend vital signs after interventions",
        "Stop monitoring to reduce alarms",
        "Reassess only if the patient asks",
        "Document later without trending",
        "Skip reassessment of response to interventions",
    ],
    "E": [
        "Leave out objective data to keep SBAR short",
        "State opinions without objective assessment",
        "Recommend medication changes without orders",
        "Delay provider call despite deterioration",
        "Omit time course/response to interventions",
        "Ask for non-urgent tasks as the main recommendation",
    ],
}

def infer_system_key(case: dict) -> str:
    s = (safe_get_system(case) or "").lower()
    if "card" in s or "cv" in s or "hypert" in s or "heart" in s:
        return "Cardiovascular"
    if "resp" in s or "pulm" in s or "asth" in s or "copd" in s:
        return "Respiratory"
    if "neuro" in s or "stroke" in s or "seiz" in s:
        return "Neurological"
    return "Generic"

def build_domain_options(domain_key: str, case: dict, gold_list: list, total: int = 10, distractors: int = 6):
    """Build multi-select options for a domain (A–D).

    IMPORTANT (fairness):
    - We do NOT force '4 correct' anymore.
    - The required selection count is the number of gold-standard targets available for that domain in the case.
    - If a case has 2/3/4 gold targets, the student must select exactly 2/3/4 respectively.
    """
    gold = [str(x).strip() for x in (gold_list or []) if str(x).strip()]
    # Cap very large lists just to keep UI usable (typical cases use 2–4)
    cap = min(len(gold), 6)
    goldN = gold[:cap] if cap > 0 else []
    req_n = len(goldN) if goldN else 4  # safe fallback if a case is missing gold targets

    sys_key = infer_system_key(case)

    pool = []
    if sys_key != "Generic":
        pool += SYSTEM_DISTRACTORS.get(sys_key, {}).get(domain_key, [])
    pool += GENERIC_DOMAIN_DISTRACTORS.get(domain_key, [])

    pool = [p for p in pool if str(p).strip() and str(p).strip() not in set(goldN)]

    picked = []
    seen = set()
    for p in pool:
        p = str(p).strip()
        if p and p not in seen:
            seen.add(p)
            picked.append(p)
        if len(picked) >= distractors:
            break

    # Fill if pool is short
    while len(picked) < distractors:
        filler = f"Non-priority action ({domain_key}) — routine follow-up"
        if filler not in seen and filler not in set(goldN):
            seen.add(filler)
            picked.append(filler)
        else:
            break

    # Ensure enough total options for reasonable choice
    total = max(int(total or 10), req_n + int(distractors or 6))
    options = (goldN + picked)[:total]

    # Deterministic shuffle by case_id+domain
    seed = sha256_hex(f"{case.get('id','')}-{domain_key}-options")[:8]
    try:
        r = random.Random(int(seed, 16))
        r.shuffle(options)
    except Exception:
        pass

    return options, goldN


def diff_selected_vs_gold(selected: list, gold4: list):
    selected = [str(x).strip() for x in (selected or []) if str(x).strip()]
    gold4 = [str(x).strip() for x in (gold4 or []) if str(x).strip()]
    correct = [x for x in selected if x in gold4]
    wrong = [x for x in selected if x not in gold4]
    missed = [x for x in gold4 if x not in selected]
    return correct, wrong, missed

def score_select4(selected: list, gold4: list) -> int:
    """Backwards-compatible wrapper for score_selectN (kept name to avoid breaking other code)."""
    return score_selectN(selected, gold4, max_points=len([x for x in (gold4 or []) if str(x).strip()]) or 4)

def score_selectN(selected: list, gold_list: list, max_points: int) -> int:
    """Domain scoring with partial credit and variable max.

    +1 for each correct gold-standard option selected.
    No negative marking for wrong selections.

    The score is floored at 0 and capped at max_points (typically 4, but may vary by session/question settings).
    """
    selected = [str(x).strip() for x in (selected or []) if str(x).strip()]
    gold_list = [str(x).strip() for x in (gold_list or []) if str(x).strip()]
    correct = [x for x in selected if x in gold_list]

    score = len(correct)

    if score < 0:
        score = 0
    if max_points is None or int(max_points) <= 0:
        max_points = len(gold_list) or 4
    if score > int(max_points):
        score = int(max_points)
    return int(score)


def resolve_ae_rationale(case: dict | None, domain_key: str, gold_items: list | None = None, sbar_expected: str | None = None) -> str:
    """Return an activity-specific, case-aware rationale line for A–E.

    Priority order:
      1) Case-authored rationale (if present)
      2) Auto-generated rationale using activity type + a few gold items as context
    """
    domain_key = str(domain_key or "").strip().upper()
    # 1) Case-authored rationale locations (non-breaking optional fields)
    custom = None
    if isinstance(case, dict):
        # Common optional nesting patterns:
        # case["clinical_reasoning"]["rationales"]["A"]
        cr = case.get("clinical_reasoning") or case.get("clinical_reasoning_key") or {}
        if isinstance(cr, dict):
            rr = cr.get("rationales") or cr.get("rationale") or {}
            if isinstance(rr, dict):
                custom = rr.get(domain_key) or rr.get(domain_key.lower())
        # case["ae_rationales"]["A"]
        if not custom:
            rr2 = case.get("ae_rationales") or {}
            if isinstance(rr2, dict):
                custom = rr2.get(domain_key) or rr2.get(domain_key.lower())
        # case["gold_answers"]["ae_rationales"]["A"]
        if not custom:
            ga = case.get("gold_answers") or case.get("gold") or {}
            if isinstance(ga, dict):
                rr3 = ga.get("ae_rationales") or ga.get("rationales") or {}
                if isinstance(rr3, dict):
                    custom = rr3.get(domain_key) or rr3.get(domain_key.lower())

    if isinstance(custom, str) and custom.strip():
        return custom.strip()

    # 2) Auto-generate (activity-specific)
    gold_items = [str(x).strip() for x in (gold_items or []) if str(x).strip()]
    gold_hint = ""
    if gold_items:
        # Use only a couple hints, without exposing full "gold answers"
        gold_hint = " Key themes include: " + "; ".join(gold_items[:2]) + "."

    base = {
        "A": "Assessment is about rapid safety-first data gathering: identify instability, confirm the problem pattern, and collect the most decision-making cues before acting." + gold_hint,
        "B": "Prioritize focuses on what could harm the patient first: address ABC threats, time-critical risks, and escalation triggers before secondary problems." + gold_hint,
        "C": "Interventions should be protocol- and evidence-based: implement immediate safety actions, ordered therapies, and monitoring steps that change outcomes." + gold_hint,
        "D": "Reassess is about closing the loop: evaluate response to interventions, trend vitals/symptoms, watch for adverse effects, and decide whether to escalate or de-escalate." + gold_hint,
        "E": "SBAR is structured escalation: communicate the critical situation, relevant context, your current assessment, and a clear recommendation/request." + (" Expected SBAR should align with: " + str(sbar_expected).strip() if (sbar_expected and str(sbar_expected).strip()) else "") + gold_hint,
    }.get(domain_key, "Use a safety-first, protocol-based approach with timely escalation when needed." + gold_hint)

    # Keep it short and readable
    return base.strip()


def render_select_feedback(domain_name: str, domain_key: str, score: int, correct_selected, wrong_selected, missed_correct, max_points: int,
                          case: dict | None = None, gold_items: list | None = None, sbar_expected: str | None = None):
    lines = [f"### ✅ {domain_name} Result", f"**Score:** **{score}/{int(max_points)}**", ""]
    if wrong_selected:
        lines.append("❌ **Wrong selections:**")
        for w in (wrong_selected or [])[:10]:
            lines.append(f"- {w}")
        lines.append("")
    if correct_selected:
        lines.append("✅ **Correct selections:**")
        for c in (correct_selected or [])[:10]:
            lines.append(f"- {c}")
        lines.append("")
    if missed_correct:
        lines.append("⚠️ **Correct options you missed:**")
        for m in (missed_correct or [])[:10]:
            lines.append(f"- {m}")
        lines.append("")
    rationale = resolve_ae_rationale(case, domain_key, gold_items=gold_items, sbar_expected=sbar_expected)
    lines.append(f"**Rationale:** {rationale}")
    return "\n".join(lines)


def render_ae_section_header_html(letter: str, title: str) -> str:
    letter = (letter or "").strip()
    title = (title or "").strip()
    safe_letter = _html.escape(letter)
    safe_title = _html.escape(title)
    return (
        "<div class='ae-section-header'>"
        f"<span class='ae-letter-badge'>{safe_letter}</span>"
        f"<span class='ae-section-title'>{safe_title}</span>"
        "</div>"
    )


def render_select_feedback_html(domain_name: str, domain_key: str, score: int, correct_selected, wrong_selected, missed_correct, max_points: int,
                               case: dict | None = None, gold_items: list | None = None, sbar_expected: str | None = None) -> str:
    # Build rationale text (plain)
    rationale = resolve_ae_rationale(case, domain_key, gold_items=gold_items, sbar_expected=sbar_expected)
    rationale = _html.escape(str(rationale or "").strip())

    def _li(items):
        items = items or []
        out = []
        for it in list(items)[:20]:
            s = str(it).strip()
            if s:
                out.append(f"<li>{_html.escape(s)}</li>")
        return "".join(out) if out else "<li>—</li>"

    score_str = f"{int(score)}/{int(max_points)}"

    html_parts = []
    html_parts.append("<div class='ae-review-box'>")
    html_parts.append(f"<div><span class='ae-score-badge'>Score: { _html.escape(score_str) }</span></div>")
    html_parts.append(f"<div class='ae-review-title'>{_html.escape(domain_name)} Review</div>")
    html_parts.append("<div class='ae-review-text'>")

    # Wrong / Correct / Missed sections
    html_parts.append("<div class='ae-review-title'>Wrong selections</div>")
    html_parts.append(f"<ul>{_li(wrong_selected)}</ul>")

    html_parts.append("<div class='ae-review-title'>Correct selections</div>")
    html_parts.append(f"<ul>{_li(correct_selected)}</ul>")

    html_parts.append("<div class='ae-review-title'>Correct options you missed</div>")
    html_parts.append(f"<ul>{_li(missed_correct)}</ul>")

    html_parts.append("<div class='ae-review-title'>Rationale</div>")
    html_parts.append(f"<div class='ae-review-text'>{rationale or '—'}</div>")

    html_parts.append("</div></div>")
    return "".join(html_parts)




def item_match(student_text: str, gold_item: str) -> bool:
    stxt = norm(student_text)
    g = norm(str(gold_item))
    if not stxt or not g:
        return False
    if g in stxt:
        return True
    s_tokens = set(tokenize(stxt))
    g_tokens = set(tokenize(g))
    if not g_tokens:
        return False
    return len(s_tokens.intersection(g_tokens)) >= 1


def rubric_match_report(student_text: str, gold_list):
    gold_list = gold_list or []
    matched, missed = [], []
    for g in gold_list:
        if item_match(student_text, str(g)):
            matched.append(str(g))
        else:
            missed.append(str(g))
    return matched, missed


def score_from_matched(matched_count: int, has_any_text: bool) -> int:
    if matched_count >= 3:
        return 4
    if matched_count == 2:
        return 3
    if matched_count == 1:
        return 2
    if has_any_text:
        return 1
    return 0


UNSAFE_PATTERNS = [
    (r"\blower\s+bp\s+quick(ly)?\b", "Rapid BP lowering without orders can cause ischemia/stroke."),
    (r"\bdouble\s+(the\s+)?dose\b", "Dose changes without orders are unsafe."),
    (r"\bgive\s+extra\s+dose\b", "Extra dosing without verification/orders is unsafe."),
    (r"\bpush\s+potassium\b", "IV potassium must NEVER be IV push."),
    (r"\bbolus\s+insulin\b", "Insulin dosing must follow protocol/orders."),
    (r"\bdiscontinue\s+oxygen\b", "Stopping oxygen abruptly can worsen hypoxia—follow orders/targets."),
]


def detect_unsafe(text: str):
    t = norm(text)
    hits = []
    for pat, msg in UNSAFE_PATTERNS:
        if re.search(pat, t):
            hits.append(msg)
    return hits


def apply_unsafe_penalty(score: int, unsafe_hits):
    return min(score, 1) if unsafe_hits else score


def score_explainer_markdown() -> str:
    return (
        "**How scoring works (0–4, partial credit):\n\n"
        "- Choose **up to 4** options (goal: pick the 4 gold-standard actions).\n"
        "- You earn **+1 point** for each correct (gold) option you select.\n"
        "- The score is **floored at 0** and **capped at 4** (based on the session and question settings).\n\n"
        "Examples:\n"
        "- 3 correct → **3/4**\n"
        "- 4 correct → **4/4**\n"
        "- 1 correct → **1/4**\n\n"
        "🚫 **Safety rule:** If an unsafe statement is detected in your free-text fields, the score is capped at **1/4** for that domain.\n"
        "✅ Tip: Focus on selecting the best actions—avoid over-selecting just to fill 4."
    )


def build_case_context(case: dict) -> str:
    """
    Build a compact, AI-safe case context string from a case dict.
    IMPORTANT: We intentionally DO NOT include any explicit 'correct' answers or rationales.
    """
    if not isinstance(case, dict):
        return "Case context unavailable (invalid case object)."

    title = case.get("title") or case.get("case_title") or case.get("name") or ""
    system = case.get("system") or case.get("category") or case.get("domain") or ""
    age = case.get("age") or case.get("patient_age") or ""
    sex = case.get("sex") or case.get("gender") or ""
    setting = case.get("setting") or ""

    # Narrative / stem
    narrative = (
        case.get("scenario")
        or case.get("stem")
        or case.get("presentation")
        or case.get("chief_complaint")
        or case.get("summary")
        or ""
    )

    def fmt_kv_block(label: str, obj) -> str:
        """Format dict/list/str blocks into bullet text."""
        if not obj:
            return ""
        if isinstance(obj, dict):
            lines = [f"- {k}: {v}" for k, v in obj.items() if v not in (None, "", [], {})]
            return f"{label}:\n" + "\n".join(lines) + "\n" if lines else ""
        if isinstance(obj, list):
            lines = [f"- {x}" for x in obj if x not in (None, "", [], {})]
            return f"{label}:\n" + "\n".join(lines) + "\n" if lines else ""
        return f"{label}: {obj}\n"

    vitals = case.get("vitals") or case.get("vital_signs")
    labs = case.get("labs") or case.get("lab_results")
    meds = case.get("meds") or case.get("medications")
    hx = case.get("history") or case.get("pmh") or case.get("medical_history")
    assessment = case.get("assessment") or case.get("nursing_assessment")
    orders = case.get("orders") or case.get("provider_orders")

    # Items / questions (stems + options only, no correct/rationale)
    items = case.get("items") or []
    items_text = ""
    if isinstance(items, list) and items:
        q_lines = []
        for i, it in enumerate(items, start=1):
            if not isinstance(it, dict):
                continue
            stem = it.get("stem") or it.get("question") or ""
            qtype = it.get("type") or ""
            options = it.get("options") if isinstance(it.get("options"), list) else []
            if stem:
                q_lines.append(f"Q{i} ({qtype}): {stem}".strip())
            if options:
                for opt in options:
                    q_lines.append(f"  - {opt}")
        if q_lines:
            items_text = "Questions (no answers provided):\n" + "\n".join(q_lines) + "\n"

    header_bits = []
    if title:
        header_bits.append(f"Title: {title}")
    if system:
        header_bits.append(f"System/Category: {system}")

    demo_bits = []
    if age:
        demo_bits.append(str(age))
    if sex:
        demo_bits.append(str(sex))
    if setting:
        demo_bits.append(str(setting))
    if demo_bits:
        header_bits.append("Patient/Setting: " + " | ".join(demo_bits))

    header = "\n".join(header_bits).strip()
    if header:
        header += "\n"

    parts = [
        header,
        f"Scenario:\n{narrative}\n" if narrative else "",
        fmt_kv_block("History", hx),
        fmt_kv_block("Assessment", assessment),
        fmt_kv_block("Vitals", vitals),
        fmt_kv_block("Labs", labs),
        fmt_kv_block("Medications", meds),
        fmt_kv_block("Orders", orders),
        items_text,
    ]
    context = "".join(parts).strip()
    return context if context else "Case context unavailable."


def build_domain_coach_prompt(domain_key: str, case: dict, student_text: str, matched, missed, unsafe_hits):
    domain_names = {"A": "Assessment", "B": "Prioritize", "C": "Interventions", "D": "Reassess", "E": "SBAR"}
    name = domain_names.get(domain_key, domain_key)
    context = build_case_context(case)

    missed_themes = missed[:3] if missed else []
    matched_examples = matched[:2] if matched else []

    rubric_targets = {
        "A": "Include: ABC/primary survey focus + system-focused checks + relevant history/med timing + trend monitoring.",
        "B": "Include: top risks first (unstable/ABC threats) + rationale tied to symptoms/vitals + escalation if needed.",
        "C": "Include: immediate safety actions + monitoring + protocol steps and what to prepare for orders.",
        "D": "Include: when to reassess + what to monitor + thresholds for escalation + document trends.",
        "E": "Include: numbers + symptoms + why dangerous + what you need done now (provider review, monitoring, protocol).",
    }

    style_instructions = {
        "A": "Ask 2–3 coaching questions + 1 hint. Do NOT give full answers.",
        "B": "Coach ordering/rationale and escalation triggers. Do NOT give full answers.",
        "C": "Coach immediate nursing actions and monitoring. Do NOT give full answers.",
        "D": "Coach timing + what to trend + escalation thresholds. Do NOT give full answers.",
        "E": (
            "Start with 4 coaching lines labeled Situation/Background/Assessment/Recommendation.\n"
            "These must be coaching instructions, NOT a completed SBAR.\n"
        ),
    }

    safety_line = ""
    if unsafe_hits:
        safety_line = "SAFETY ALERT: Student included unsafe actions. Emphasize protocol/orders and escalation.\n"

    return (
        f"{context}\n"
        f"Domain: {name}\n"
        f"Student input:\n{student_text}\n\n"
        f"Rubric goal (high-level): {rubric_targets.get(domain_key,'')}\n"
        f"Matched examples (do not repeat as answers): {matched_examples}\n"
        f"Missing themes (do not reveal full answers): {missed_themes}\n"
        f"{safety_line}\n"
        f"Coaching style rules:\n{style_instructions.get(domain_key,'')}\n"
        "Now produce coaching using the required global output format.\n"
    )


def build_debrief_prompt(case: dict, scores: dict, missed_by_domain: dict, unsafe_by_domain: dict):
    context = build_case_context(case)
    return (
        f"{context}\n"
        f"Student domain scores (A-E): {scores}\n"
        f"Missed items by domain (summarize themes): {missed_by_domain}\n"
        f"Safety flags by domain: {unsafe_by_domain}\n\n"
        "Write a short end-of-case debrief with:\n"
        "1) 3 strengths\n"
        "2) 2 improvement priorities\n"
        "3) 1 patient-safety takeaway\n"
        "4) 1 next-step practice suggestion\n"
        "Do NOT reveal full answers. Keep <= 170 words."
    )


# =============================
# NCLEX-style practice (Step 8–10 improvements)
# =============================
# =============================
# NCLEX display formatting helpers (UI)
# =============================
def nclex_qnum_from_id(qid: str) -> str:
    """Extract Q number from IDs like adult_htn_emerg_01-Q10"""
    qid = str(qid or "")
    if "-Q" in qid:
        return qid.split("-Q")[-1].lstrip("0") or "0"
    m = re.search(r"Q(\d+)", qid)
    return (m.group(1).lstrip("0") or "0") if m else qid

def fmt_nclex_header(item: dict, display_num: int | None = None) -> str:
    qnum = int(display_num) if display_num is not None else nclex_qnum_from_id(item.get("id"))
    qtype = str(item.get("type", "")).strip()
    client_need = str(item.get("client_need", "")).strip()
    return (
        "<span style='color:#1f4e79; font-weight:800; font-size:18px;'>"
        f"Q{qnum} ({qtype} / {client_need})"
        "</span>"
    )

def fmt_nclex_stem(stem: str) -> str:
    stem = (stem or "").strip()
    return (
        "<span style='color:#1f4e79; font-weight:700; font-size:16px;'>"
        f"• {stem}"
        "</span>"
    )


def normalize_text_basic(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"[^\w\s]", "", s)  # remove punctuation
    return s


def shuffle_if_needed(options: list, seed_str: str, enabled: bool) -> list:
    if not enabled:
        return options
    # deterministic shuffle per item id
    rnd = hashlib.sha256(seed_str.encode("utf-8")).hexdigest()
    # simple deterministic shuffle: sort by hash of (rnd + option)
    return sorted(list(options), key=lambda x: hashlib.sha256((rnd + "||" + str(x)).encode("utf-8")).hexdigest())


def nclex_score_item(item: dict, answer, policy: dict, features: dict):
    qtype = item.get("type")
    correct = item.get("correct")
    max_points = 1
    points = 0
    detail = {"qid": item.get("id"), "type": qtype, "points": 0, "max": 1, "correct": False}

    partial_credit = bool(features.get("nclex_partial_credit", False))

    if qtype == "mcq":
        points = 1 if answer == correct else 0

    elif qtype == "sata":
        # answer expected list
        correct_set = set(correct or [])
        ans_set = set(answer or [])
        if not partial_credit:
            points = 1 if ans_set == correct_set else 0
        else:
            # simple partial: +1 for each correct chosen, -1 for each extra wrong chosen, floor at 0, normalize to 1
            # if you want a stricter rule, disable partial_credit.
            hit = len(ans_set.intersection(correct_set))
            extra = len(ans_set - correct_set)
            raw = max(0, hit - extra)
            points = 1 if (raw > 0 and hit == len(correct_set) and extra == 0) else (1 if raw >= len(correct_set) else 0)
            # We keep it conservative: either correct=1 or 0. Turn on partial later if you want fractional.

    elif qtype == "ordered_response":
        # correct expected list
        corr = list(correct or [])
        ans = list(answer or [])
        if not corr:
            points = 0
        else:
            if not partial_credit:
                points = 1 if ans == corr else 0
            else:
                # partial per position (0..1)
                max_points = len(corr)
                points = sum(1 for i in range(min(len(ans), len(corr))) if ans[i] == corr[i])
                detail["max"] = max_points
                detail["points"] = points
                detail["correct"] = (points == max_points)
                return detail

    elif qtype == "cloze":
        # correct_text OR correct + acceptable
        correct_text = item.get("correct_text") or correct or ""
        acceptable = item.get("acceptable") or []
        ans = normalize_text_basic(answer or "")
        ok = ans == normalize_text_basic(correct_text)
        if not ok:
            for a in acceptable:
                if ans == normalize_text_basic(a):
                    ok = True
                    break
        points = 1 if ok else 0

    elif qtype == "matrix":
        # correct expected dict: row -> col
        corr = item.get("correct") or {}
        ans = answer or {}
        rows = item.get("rows") or list(corr.keys())
        if not partial_credit:
            ok = True
            for r in rows:
                if ans.get(r) != corr.get(r):
                    ok = False
                    break
            points = 1 if ok else 0
        else:
            max_points = len(rows) if rows else 1
            points = 0
            for r in rows:
                if ans.get(r) == corr.get(r):
                    points += 1
            detail["max"] = max_points
            detail["points"] = points
            detail["correct"] = (points == max_points)
            return detail

    elif qtype == "evolving_case":
        stages = item.get("stages") or []
        # answer expected dict stage_index -> stage_answer
        ans = answer or {}
        if not partial_credit:
            ok_all = True
            for si, stage in enumerate(stages):
                q = stage.get("question", {}) or {}
                stype = q.get("type")
                scorrect = q.get("correct")
                given = ans.get(str(si))
                if stype == "mcq":
                    if given != scorrect:
                        ok_all = False
                elif stype == "sata":
                    if set(given or []) != set(scorrect or []):
                        ok_all = False
                else:
                    ok_all = False
            points = 1 if ok_all and stages else 0
        else:
            max_points = 0
            points = 0
            for si, stage in enumerate(stages):
                q = stage.get("question", {}) or {}
                stype = q.get("type")
                scorrect = q.get("correct")
                given = ans.get(str(si))
                if stype == "mcq":
                    max_points += 1
                    points += 1 if given == scorrect else 0
                elif stype == "sata":
                    max_points += 1
                    points += 1 if set(given or []) == set(scorrect or []) else 0
            detail["max"] = max_points if max_points else 1
            detail["points"] = points
            detail["correct"] = (points == detail["max"])
            return detail

    else:
        points = 0

    detail["max"] = max_points
    detail["points"] = points
    detail["correct"] = (points == max_points)
    return detail



def render_nclex_practical(case_id: str, policy: dict, nclex: dict, features: dict, mode: str, timer_lock: bool, student_username: str = "", case_title: str = ""):
    """
    Renders NCLEX-style practice AFTER A–E completion only.
    Adds:
    - Per-question "Save" (writes to autosave file; helps against crashes)
    - Finalize + Review page that shows correct/wrong + rationale (and optional AI explanation)
    """
    if not policy.get("enabled", False):
        return

    vis = policy.get("mode_visibility", {"Practice": True, "Exam": True})
    if not bool(vis.get(mode, True)):
        return

    # Policy defaults (safe if older nclex_policy.json is missing keys)
    show_correct_after = False  # per-question feedback disabled; review only after final submit
    show_rationales_after = False  # per-question feedback disabled; review only after final submit
    shuffle_opts = bool(policy.get("shuffle_options", True))
    show_review_after_finalize = bool(policy.get("show_review_after_finalize", True))
    ai_explain_after_finalize = bool(policy.get("ai_explanations_after_finalize", True))

    st.divider()
    st.subheader("🧾 Practical (NCLEX-Style)")

    case_pack = (nclex.get("cases") or {}).get(case_id)
    if not case_pack:
        st.error(
            f"No NCLEX-style practice items found for this case ID in nclex_items.json.\n\n"
            f"Looking for case id: {case_id}"
        )
        st.info("Admin: open ✅ NCLEX Validator to auto-fix missing packs (IDs).")
        return

    items = list(case_pack.get("items", []) or [])
    enabled_types = (policy.get("enabled_types") or {})
    items = [it for it in items if enabled_types.get(it.get("type"), True)]

    if not items:
        st.info("This case has an NCLEX pack, but no items are inside (or all types are disabled).")
        return

    # Apply items_per_case + optional Admin-controlled active set rotation
    items_per_case = nclex_items_per_case(policy, case_id)
    try:
        if bool(policy.get("rotation_enabled", False)):
            active_sets = load_nclex_active_sets()
            active = (active_sets.get("by_case") or {}).get(str(case_id))
            qids = (active or {}).get("qids") if isinstance(active, dict) else None
            if isinstance(qids, list) and qids:
                by_id = {str(it.get("id", "")): it for it in items}
                ordered = [by_id[q] for q in qids if q in by_id]
                remaining = [it for it in items if str(it.get("id", "")) not in set(qids)]
                items = (ordered + remaining)[:items_per_case]
            else:
                # No active set yet: fall back to first N
                items = items[:items_per_case]
        else:
            items = items[:items_per_case]
    except Exception:
        items = items[:items_per_case]

    # Randomize order per student/session (optional) AFTER rotation is applied
    try:
        if bool(policy.get("randomize_per_student_session", False)):
            seed_base = f"{student_username}|{case_id}|{st.session_state.get('attempt_started_epoch') or ''}|{mode}"
            items = sorted(list(items), key=lambda it: hashlib.sha256((seed_base + '||' + str(it.get('id',''))).encode('utf-8')).hexdigest())
    except Exception:
        pass

    # Persist the presented items order for stable scoring/review (prevents missing/out-of-order questions on reruns)
    if st.session_state.get("nclex_presented_case_id") != str(case_id):
        st.session_state["nclex_presented_case_id"] = str(case_id)
        st.session_state["nclex_presented_items"] = list(items)
        # Reset one-at-a-time index for a new case/session
        if "nclex_one_idx" in st.session_state:
            st.session_state.nclex_one_idx = 0
    else:
        # Reuse the same presented list once established
        pi = st.session_state.get("nclex_presented_items")
        if isinstance(pi, list) and pi:
            items = list(pi)
        else:
            st.session_state["nclex_presented_items"] = list(items)

    st.caption(f"{len(items)} items loaded for this case.")

    # Session defaults
    if "practical_submitted" not in st.session_state:
        st.session_state.practical_submitted = False
    if "nclex_answers" not in st.session_state:
        st.session_state.nclex_answers = {}
    if "nclex_scored" not in st.session_state:
        st.session_state.nclex_scored = None
    if "nclex_finalized" not in st.session_state:
        st.session_state.nclex_finalized = False
    if "nclex_ai_explanations" not in st.session_state:
        st.session_state.nclex_ai_explanations = {}  # qid -> text

    # Reset NCLEX state when switching to a different case
    # (Prevents "frozen" inputs when a prior case was finalized/locked.)
    if str(st.session_state.get("nclex_case_id", "")).strip() != str(case_id).strip():
        st.session_state["nclex_case_id"] = str(case_id).strip()
        st.session_state.practical_submitted = False
        st.session_state.nclex_answers = {}
        st.session_state.nclex_scored = None
        st.session_state.nclex_finalized = False
        st.session_state.nclex_ai_explanations = {}
        # Reset one-at-a-time index if it exists
        if "nclex_one_idx" in st.session_state:
            st.session_state.nclex_one_idx = 0


    # When finalized, lock inputs completely
    disabled_inputs = bool(timer_lock or bool(st.session_state.get('nclex_finalized', False)) or bool(st.session_state.get('practical_submitted', False)))

    one_at_a_time = bool(policy.get("one_question_at_a_time", False))
    # In "one question at a time" mode, we keep a per-session index
    if one_at_a_time:
        if "nclex_one_idx" not in st.session_state:
            st.session_state.nclex_one_idx = 0
        # Clamp index
        st.session_state.nclex_one_idx = max(0, min(int(st.session_state.nclex_one_idx), max(0, len(items) - 1)))

        # Show scenario trigger above each question (helps context without revealing the answer keys)
        try:
            case_obj = next((c for c in load_cases() if str(c.get("id","")).strip() == str(case_id).strip()), None)
            if case_obj:
                trig = build_scenario_trigger(case_obj)
                if trig:
                    st.markdown(trig)
                    st.divider()
        except Exception:
            pass

        # Navigation
        cols_nav = st.columns([1,2,1])
        with cols_nav[0]:
            if st.button("⬅️ Prev", disabled=(disabled_inputs or st.session_state.nclex_one_idx <= 0), key="nclex_prev"):
                st.session_state.nclex_one_idx = max(0, st.session_state.nclex_one_idx - 1)
                st.rerun()
        with cols_nav[1]:
            st.caption(f"Question {st.session_state.nclex_one_idx + 1} of {len(items)}")
        with cols_nav[2]:
            if st.button("Next ➡️", disabled=(disabled_inputs or st.session_state.nclex_one_idx >= len(items) - 1), key="nclex_next"):
                st.session_state.nclex_one_idx = min(len(items) - 1, st.session_state.nclex_one_idx + 1)
                st.rerun()

        items_to_render = [items[st.session_state.nclex_one_idx]] if items else []
        start_number = st.session_state.nclex_one_idx + 1
    else:
        items_to_render = list(items)
        start_number = 1

    # Render each item
    for display_i, item in enumerate(items_to_render, start=start_number):
        qid = item.get("id", "")
        qtype = item.get("type", "")

        st.markdown(fmt_nclex_header(item, display_i), unsafe_allow_html=True)
        stem = item.get("stem", "")
        if stem:
            st.markdown(fmt_nclex_stem(stem), unsafe_allow_html=True)
        # === NCLEX IMAGE DISABLED ===
        # Images for NCLEX items are intentionally not rendered to avoid unrelated/generated visuals.
        # (Rotation, randomization, one-question-at-a-time, scoring, and all other NCLEX behavior remains unchanged.)
        # === END NCLEX IMAGE DISABLED ===


        # MCQ
        if qtype == "mcq":
            opts = item.get("options", []) or []
            opts = shuffle_if_needed(opts, seed_str=qid, enabled=shuffle_opts)
            key = f"nclex_{qid}"
            prev = st.session_state.nclex_answers.get(qid)
            idx = opts.index(prev) if (prev in opts) else 0
            ans = st_radio_no_preselect("Select one", opts, prev, key=key, disabled=disabled_inputs)
            try:
                track_nclex_change(qid, prev, ans)
            except Exception:
                pass
            try:
                track_nclex_change(qid, prev, ans)
            except Exception:
                pass
            st.session_state.nclex_answers[qid] = ans
            try:
                nclex_autosave_now()
            except Exception:
                pass

        # SATA
        elif qtype == "sata":
            opts = item.get("options", []) or []
            opts = shuffle_if_needed(opts, seed_str=qid, enabled=shuffle_opts)
            key = f"nclex_{qid}"
            prev = st.session_state.nclex_answers.get(qid, [])
            ans = st.multiselect(
                "Select all that apply",
                opts,
                default=[p for p in prev if p in opts],
                key=key,
                disabled=disabled_inputs
            )
            try:
                track_nclex_change(qid, prev, ans)
            except Exception:
                pass
            try:
                track_nclex_change(qid, prev, ans)
            except Exception:
                pass
            st.session_state.nclex_answers[qid] = ans
            try:
                nclex_autosave_now()
            except Exception:
                pass

        # ORDERED
        elif qtype == "ordered_response":
            opts = item.get("options", []) or []
            prev = st.session_state.nclex_answers.get(qid, [])
            if not prev:
                prev = opts[:]

            st.caption("Order them from first → last. (Simple version: choose position for each.)")
            ordered = []
            remaining = [o for o in opts]
            for i in range(len(opts)):
                default_i = 0
                if i < len(prev) and prev[i] in remaining:
                    default_i = remaining.index(prev[i])
                choice = st.selectbox(
                    f"Position {i+1}",
                    remaining,
                    index=default_i if remaining else 0,
                    key=f"nclex_{qid}_pos_{i}",
                    disabled=disabled_inputs
                )
                ordered.append(choice)
                if choice in remaining:
                    remaining.remove(choice)
            try:
                track_nclex_change(qid, prev, ordered)
            except Exception:
                pass
            st.session_state.nclex_answers[qid] = ordered
            try:
                nclex_autosave_now()
            except Exception:
                pass

        # CLOZE
        elif qtype == "cloze":
            prev = st.session_state.nclex_answers.get(qid, "")
            # Prefer dropdown-style cloze if options are provided (shows the selection arrow)
            cloze_opts = item.get("options") or item.get("choices") or item.get("dropdown_options") or item.get("answer_choices") or []
            if isinstance(cloze_opts, (list, tuple)) and len(cloze_opts) > 0:
                options = ["— Select —"] + [str(x) for x in cloze_opts]
                prev_str = str(prev) if prev is not None else ""
                idx = 0
                if prev_str in options:
                    idx = options.index(prev_str)
                ans = st.selectbox(
                    "Select the best answer",
                    options,
                    index=idx,
                    key=f"nclex_{qid}_cloze_sel",
                    disabled=disabled_inputs
                )
                ans = "" if ans == "— Select —" else str(ans)
            else:
                # Fallback to free-text cloze when no options exist
                ans = st.text_input("Your answer", value=str(prev), key=f"nclex_{qid}_text", disabled=disabled_inputs)

            try:
                track_nclex_change(qid, prev, ans)
            except Exception:
                pass
            st.session_state.nclex_answers[qid] = ans
            try:
                nclex_autosave_now()
            except Exception:
                pass

        # MATRIX
        elif qtype == "matrix":
            rows = item.get("rows", []) or []
            cols = item.get("cols", []) or []
            prev = st.session_state.nclex_answers.get(qid, {}) or {}
            ans_map = {}
            for r in rows:
                default = prev.get(r, cols[0] if cols else "")
                if cols and default not in cols:
                    default = cols[0]
                choice = st_radio_no_preselect(str(r), cols, default, key=f"nclex_{qid}_row_{hash(r)}", disabled=disabled_inputs, horizontal=True)
                ans_map[r] = choice
            try:
                track_nclex_change(qid, prev, ans_map)
            except Exception:
                pass
            st.session_state.nclex_answers[qid] = ans_map
            try:
                nclex_autosave_now()
            except Exception:
                pass

        # EVOLVING CASE
        elif qtype == "evolving_case":
            stages = item.get("stages", []) or []
            prev = st.session_state.nclex_answers.get(qid, {}) or {}
            ans_stage = {}
            for si, stage in enumerate(stages):
                st.markdown(f"**{stage.get('stage','Stage')}** — {stage.get('update','')}")
                q = stage.get("question", {}) or {}
                stype = q.get("type")
                qstem = q.get("stem", "")
                if qstem:
                    st.write(qstem)

                if stype == "mcq":
                    opts = q.get("options", []) or []
                    opts = shuffle_if_needed(opts, seed_str=f"{qid}:{si}", enabled=shuffle_opts)
                    prev_ans = prev.get(str(si))
                    idx = opts.index(prev_ans) if prev_ans in opts else 0
                    a = st_radio_no_preselect("Select one", opts, prev_ans, key=f"nclex_{qid}_stage_{si}_mcq", disabled=disabled_inputs)
                    ans_stage[str(si)] = a

                elif stype == "sata":
                    opts = q.get("options", []) or []
                    opts = shuffle_if_needed(opts, seed_str=f"{qid}:{si}", enabled=shuffle_opts)
                    prev_ans = prev.get(str(si), [])
                    a = st.multiselect(
                        "Select all that apply",
                        opts,
                        default=[p for p in prev_ans if p in opts],
                        key=f"nclex_{qid}_stage_{si}_sata",
                        disabled=disabled_inputs
                    )
                    ans_stage[str(si)] = a
                else:
                    st.warning("Unsupported evolving stage type in this item.")
            st.session_state.nclex_answers[qid] = ans_stage
            try:
                nclex_autosave_now()
            except Exception:
                pass

        else:
            st.warning(f"This item type will be enabled next: {qtype}")

        # Per-question save (fallback against sudden shutdown)
        colS1, colS2 = st.columns([1, 3])
        with colS1:
            if st.button("💾 Save", key=f"nclex_save_{qid}", disabled=disabled_inputs):
                autosave_draft(
                    features,
                    student_username,
                    case_id,
                    {
                        "kind": "nclex_question",
                        "caseTitle": case_title,
                        "qid": qid,
                        "answer": st.session_state.nclex_answers.get(qid),
                        "nclex_answers": st.session_state.nclex_answers,
                        "practical_submitted": bool(st.session_state.practical_submitted),
                        "nclex_finalized": bool(st.session_state.nclex_finalized),
                        "mode": mode,
                    },
                )
                st.success("Saved.")

        # Lightweight feedback after submit (existing behavior)
        if st.session_state.practical_submitted and show_correct_after:
            correct = item.get("correct", None)
            if qtype == "cloze":
                correct = item.get("correct_text") or correct
            st.markdown("**✅ Correct answer:**")
            _render_nclex_correct_pretty(correct)
            if show_rationales_after:
                rat = item.get("rationale", "")
                if rat:
                    st.markdown("**Explanation:**")
                    st.write(rat)
        st.divider()

    # Submit + score + finalize controls
    col1, col2, col3 = st.columns([1, 1, 2])

    with col1:
        if st.button("✅ Submit Practical Section", disabled=(timer_lock or bool(st.session_state.get('practical_submitted', False)) or bool(st.session_state.get('nclex_finalized', False)))):
            st.session_state.practical_submitted = True
            st.session_state["attempt_saved"] = False
            st.session_state["show_save_attempt"] = True

            score_details = []
            total_points = 0
            total_max = 0
            presented_items = list((st.session_state.get("nclex_presented_items") or items or []) or [])
            for item in presented_items:
                qid = item.get("id")
                ans = st.session_state.nclex_answers.get(qid)
                d = nclex_score_item(item, ans, policy, features)
                score_details.append(d)
                total_points += int(d.get("points", 0))
                total_max += int(d.get("max", 1))

            st.session_state.nclex_scored = {
                "total_points": total_points,
                "total_max": total_max,
                "details": score_details,
            }
            st.rerun()

    with col2:
        if st.button("↩ Reset Practical Answers", disabled=(disabled_inputs or bool(st.session_state.get('nclex_finalized', False)))):
            st.session_state.practical_submitted = False
            st.session_state.nclex_answers = {}
            st.session_state.nclex_scored = None
            st.session_state.nclex_ai_explanations = {}
            st.session_state.nclex_finalized = False
            st.session_state["show_save_attempt"] = False
            st.session_state["attempt_saved"] = False
            st.rerun()

    with col3:
        scored = st.session_state.get("nclex_scored")
        if scored:
            st.metric("Practical Score", f"{scored['total_points']}/{scored['total_max']}")


    # Save attempt (appears after NCLEX is scored via "Submit Practical Section")
    if st.session_state.get("nclex_scored") and st.session_state.get("show_save_attempt", False):
        if not st.session_state.get("attempt_saved", False):
            if st.button("💾 Save Attempt", disabled=(locked_out or timer_lock), key="save_attempt_after_nclex"):
                save_attempt(build_attempt_record_from_state(case_id))
                st.session_state["attempt_saved"] = True
                st.success("Saved to attempts_log.jsonl")
                st.rerun()

    # Finalize + Review (what you asked for)
    if st.session_state.get("nclex_scored") and show_review_after_finalize:
        st.divider()
        cA, cB = st.columns([1, 3])
        with cA:
            if st.button("🏁 End Practical (Finalize)", disabled=(timer_lock or bool(st.session_state.get("nclex_finalized", False)))):
                st.session_state.nclex_finalized = True
                st.session_state.practical_submitted = True
                st.rerun()
        with cB:
            if st.session_state.nclex_finalized:
                st.success("Practical finalized. Answers are locked. Review is available below.")

        if st.session_state.nclex_finalized:
            # Review is shown only when the student clicks the Review button (after finalize)
            if "nclex_show_review" not in st.session_state:
                st.session_state.nclex_show_review = False

            if not bool(st.session_state.nclex_show_review):
                if st.button("📋 Show Review", key=f"nclex_show_review_btn_{case_id}"):
                    st.session_state.nclex_show_review = True
                    st.rerun()
                st.info("Review is hidden until you click **Show Review**. Answers remain locked after finalize.")
                return

            # Build quick lookup for scoring detail
            detail_map = {}
            for d in (st.session_state.get("nclex_scored") or {}).get("details", []) or []:
                detail_map[str(d.get("qid"))] = d

            st.subheader("🧾 Practical Review (Correct / Wrong + Rationale)")
            st.caption("This review is shown after the student ends the test (finalize).")

            AI_SYSTEM_RATIONALE_PROMPT = (
                "You are an expert NCLEX item rationales writer. "
                "Explain why the correct answer is correct and why the student's answer is wrong (if wrong). "
                "Be concise (<= 120 words). Use nursing safety and prioritization language."
            )

            presented_items = list((st.session_state.get("nclex_presented_items") or items or []) or [])
            for qi, item in enumerate(presented_items, start=1):
                qid = str(item.get("id", ""))
                d = detail_map.get(qid, {})
                qtype = item.get("type", "")
                stem = item.get("stem", "")
                correct = item.get("correct", None)
                if qtype == "cloze":
                    correct = item.get("correct_text") or correct

                student_ans = st.session_state.nclex_answers.get(qid)

                is_correct = bool(d.get("correct", False))
                icon = "✅" if is_correct else "❌"
                with st.expander(f"{icon} Q{qi} — {qtype}"):
                    if stem:
                        st.write(stem)
                    st.markdown("**Your answer:**")
                    st.write(student_ans)
                    st.markdown("**Correct answer:**")
                    st.write(correct)
                    rat = item.get("rationale", "")
                    if rat:
                        st.markdown("**Official rationale:**")
                        st.write(rat)

                    if ai_explain_after_finalize:
                        # Only run AI explanation if AI key is available (avoid breaking exam flow)
                        if qid in st.session_state.nclex_ai_explanations:
                            st.markdown("**AI explanation:**")
                            st.write(st.session_state.nclex_ai_explanations[qid])
                        else:
                            if st.button("🧠 Generate AI explanation", key=f"nclex_ai_{qid}"):
                                try:
                                    opts = item.get("options", []) or []
                                    user_prompt = (
                                        f"Question: {stem}\n"
                                        f"Options: {opts}\n"
                                        f"Correct: {correct}\n"
                                        f"Student answered: {student_ans}\n"
                                        "Write an explanation."
                                    )
                                    with st.spinner("Generating AI explanation..."):
                                        txt = openai_responses_call(
                                            (admin_settings.get("ai_model") or "gpt-5.2").strip() or "gpt-5.2",
                                            AI_SYSTEM_RATIONALE_PROMPT,
                                            user_prompt,
                                        )
                                    st.session_state.nclex_ai_explanations[qid] = txt
                                    st.rerun()
                                except Exception as e:
                                    st.warning(f"AI explanation unavailable: {e}")
def autosave_draft(features: dict, student_username: str, case_id: str, payload: dict):
    if not features.get("autosave_enabled", False):
        return
    if not student_username or not case_id:
        return
    try:

        # Normalize and enrich payload with current session state so resume/progress works reliably
        if not isinstance(payload, dict):
            payload = {}
        payload = dict(payload)
        # Always snapshot the key workflow state
        payload["answers"] = st.session_state.get("answers", {}) or {}
        payload["scores"] = st.session_state.get("scores", {}) or {}
        payload["last_feedback"] = st.session_state.get("last_feedback", {}) or {}
        payload["ae_focus"] = st.session_state.get("ae_focus", "A") or "A"
        payload["nclex_answers"] = st.session_state.get("nclex_answers", {}) or {}
        payload["practical_submitted"] = bool(st.session_state.get("practical_submitted", False))
        payload["nclex_scored"] = st.session_state.get("nclex_scored")
        payload["nclex_finalized"] = bool(st.session_state.get("nclex_finalized", False))
        payload["intake"] = st.session_state.get("intake", {}) or {}
        payload["intake_score"] = int(st.session_state.get("intake_score", 0) or 0)
        payload["intake_breakdown"] = st.session_state.get("intake_breakdown", {}) or {}
        payload["intake_submitted"] = bool(st.session_state.get("intake_submitted", False))

        rec = {
            "timestamp": utc_now_iso(),
            "student_username": student_username,
            "caseId": case_id,
            "draft": payload
        }
        with open(AUTOSAVE_DRAFTS_PATH, "a", encoding="utf-8") as f:
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")
    except Exception:
        pass


def load_last_autosave(student_username: str, case_id: str):
    """Return the latest autosave draft dict for this student+case (or None)."""
    if not AUTOSAVE_DRAFTS_PATH.exists():
        return None
    student_username = (student_username or "").strip()
    case_id = (case_id or "").strip()
    if not student_username or not case_id:
        return None

    last = None
    try:
        # Read from end for speed (file can grow). Best-effort: scan all.
        with open(AUTOSAVE_DRAFTS_PATH, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    rec = json.loads(line)
                except Exception:
                    continue
                if str(rec.get("student_username", "")).strip() != student_username:
                    continue
                if str(rec.get("caseId", "")).strip() != case_id:
                    continue
                d = rec.get("draft")
                if isinstance(d, dict):
                    last = d
    except Exception:
        return None
    return last


def delete_autosaves_for(student_username: str, case_id: str) -> int:
    """Delete autosave drafts for a specific student+case. Returns number removed."""
    if not AUTOSAVE_DRAFTS_PATH.exists():
        return 0
    student_username = (student_username or "").strip()
    case_id = (case_id or "").strip()
    if not student_username or not case_id:
        return 0
    kept = []
    removed = 0
    try:
        with open(AUTOSAVE_DRAFTS_PATH, "r", encoding="utf-8") as f:
            for line in f:
                s = (line or "").strip()
                if not s:
                    continue
                try:
                    rec = json.loads(s)
                except Exception:
                    # Keep malformed lines (don't destroy data)
                    kept.append(line)
                    continue
                if str(rec.get("student_username", "")).strip() == student_username and str(rec.get("caseId", "")).strip() == case_id:
                    removed += 1
                    continue
                kept.append(line)
        with open(AUTOSAVE_DRAFTS_PATH, "w", encoding="utf-8") as f:
            for line in kept:
                f.write(line if line.endswith("\n") else (line + "\n"))
    except Exception:
        return 0
    return removed


def delete_all_autosaves_for_student(student_username: str) -> int:
    """Delete ALL autosave drafts for a student. Returns number removed."""
    if not AUTOSAVE_DRAFTS_PATH.exists():
        return 0
    student_username = (student_username or "").strip()
    if not student_username:
        return 0
    kept = []
    removed = 0
    try:
        with open(AUTOSAVE_DRAFTS_PATH, "r", encoding="utf-8") as f:
            for line in f:
                s = (line or "").strip()
                if not s:
                    continue
                try:
                    rec = json.loads(s)
                except Exception:
                    kept.append(line)
                    continue
                if str(rec.get("student_username", "")).strip() == student_username:
                    removed += 1
                    continue
                kept.append(line)
        with open(AUTOSAVE_DRAFTS_PATH, "w", encoding="utf-8") as f:
            for line in kept:
                f.write(line if line.endswith("\n") else (line + "\n"))
    except Exception:
        return 0
    return removed


def load_exam_overrides() -> dict:
    d = load_json_safe(EXAM_OVERRIDES_PATH, {"students": {}})
    if not isinstance(d, dict):
        d = {"students": {}}
    if not isinstance(d.get("students"), dict):
        d["students"] = {}
    return d


def save_exam_overrides(d: dict):
    try:
        EXAM_OVERRIDES_PATH.write_text(json.dumps(d, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        pass


def bump_student_reset_token(student_username: str) -> int:
    student_username = (student_username or "").strip()
    if not student_username:
        return 0
    d = load_exam_overrides()
    s = d.get("students", {}).get(student_username, {})
    try:
        token = int(s.get("reset_token") or 0) + 1
    except Exception:
        token = 1
    s["reset_token"] = token
    d["students"][student_username] = s
    save_exam_overrides(d)
    return token


def set_student_force_unlock(student_username: str, value: bool = True):
    student_username = (student_username or "").strip()
    if not student_username:
        return
    d = load_exam_overrides()
    s = d.get("students", {}).get(student_username, {})
    s["force_unlock"] = bool(value)
    d["students"][student_username] = s
    save_exam_overrides(d)


def consume_student_force_unlock(student_username: str):
    student_username = (student_username or "").strip()
    if not student_username:
        return
    d = load_exam_overrides()
    s = d.get("students", {}).get(student_username, {})
    if s.get("force_unlock"):
        s["force_unlock"] = False
        d["students"][student_username] = s
        save_exam_overrides(d)


def apply_exam_overrides_for_student_session(student_username: str):
    """Apply server-side overrides (force unlock / reset token) for the CURRENT student session."""
    student_username = (student_username or "").strip()
    if not student_username:
        return

    d = load_exam_overrides()
    s = d.get("students", {}).get(student_username, {}) if isinstance(d, dict) else {}
    if not isinstance(s, dict):
        s = {}

    # 1) Force unlock case switching (consumed once)
    if s.get("force_unlock"):
        for _k in [
            "locked_exam_case_id",
            "exam_in_progress",
            "exam_started",
            "exam_case_id",
            "current_case_id",
            "active_case_id",
            "_exam_started_at",
            "_exam_timer_start",
            "_exam_timer_deadline",
        ]:
            st.session_state.pop(_k, None)
        for _k in ["case_pick_main", "rot_case_pick_main", "case_system_main"]:
            st.session_state.pop(_k, None)
        try:
            reset_attempt_state()
        except Exception:
            pass
        consume_student_force_unlock(student_username)

    # 2) Reset token (idempotent per session)
    token = s.get("reset_token", 0)
    try:
        token_int = int(token or 0)
    except Exception:
        token_int = 0

    last_seen = st.session_state.get("_last_seen_reset_token", None)
    try:
        last_seen_int = int(last_seen or 0)
    except Exception:
        last_seen_int = 0

    if token_int and token_int != last_seen_int:
        # Clear locks/state, clear ALL autosaves, reset attempt state
        for _k in [
            "locked_exam_case_id",
            "exam_in_progress",
            "exam_started",
            "exam_case_id",
            "current_case_id",
            "active_case_id",
            "_exam_started_at",
            "_exam_timer_start",
            "_exam_timer_deadline",
        ]:
            st.session_state.pop(_k, None)
        for _k in ["case_pick_main", "rot_case_pick_main", "case_system_main"]:
            st.session_state.pop(_k, None)

        delete_all_autosaves_for_student(student_username)

        try:
            reset_attempt_state()
        except Exception:
            pass

        st.session_state["_last_seen_reset_token"] = token_int


def index_latest_autosaves_for_student(student_username: str):
    """Return dict(caseId -> latest draft dict) for a student. Best-effort; safe if file missing."""
    out = {}
    if not AUTOSAVE_DRAFTS_PATH.exists():
        return out
    student_username = (student_username or "").strip()
    if not student_username:
        return out
    try:
        with open(AUTOSAVE_DRAFTS_PATH, "r", encoding="utf-8") as f:
            for line in f:
                line = (line or "").strip()
                if not line:
                    continue
                try:
                    rec = json.loads(line)
                except Exception:
                    continue
                if str(rec.get("student_username", "")).strip() != student_username:
                    continue
                cid = str(rec.get("caseId", "")).strip()
                d = rec.get("draft")
                if cid and isinstance(d, dict):
                    out[cid] = d
    except Exception:
        return out
    return out


def compute_resume_hint_from_draft(draft: dict):
    """Return (status, resume_hint) where resume_hint in {'intake','A','B','C','D','E','nclex','review'}"""
    if not isinstance(draft, dict) or not draft:
        return ("Not started", "intake")
    # Completed = finalized NCLEX (or submitted practical with scoring)
    if bool(draft.get("nclex_finalized", False)):
        return ("Completed", "review")
    if bool(draft.get("practical_submitted", False)):
        return ("Submitted", "review")
    intake_sub = bool(draft.get("intake_submitted", False))
    if not intake_sub:
        return ("In progress", "intake")
    lf = draft.get("last_feedback", {}) if isinstance(draft.get("last_feedback", {}), dict) else {}
    # Determine next needed A–E
    for dom in ["A","B","C","D","E"]:
        if lf.get(dom) is None:
            return ("In progress", dom)
    # A–E done; continue to NCLEX
    return ("In progress", "nclex")

def apply_restored_draft(draft: dict):
    """Apply restored autosave draft into Streamlit session state safely."""
    if not isinstance(draft, dict):
        return

    # Restore A–E (if present)
    answers = draft.get("answers")
    scores = draft.get("scores")
    if isinstance(scores, dict):
        st.session_state["scores"] = scores
    if isinstance(answers, dict):
        st.session_state["answers"] = answers

        # Re-hydrate widget keys from saved answers (best-effort)
        A = answers.get("A", {}) if isinstance(answers.get("A"), dict) else {}
        st.session_state["A_selected"] = A.get("selected", []) or []
        st.session_state["A_notes"] = A.get("notes", "") or ""

        B = answers.get("B", {}) if isinstance(answers.get("B"), dict) else {}
        st.session_state["B_selected"] = B.get("selected", []) or []
        st.session_state["B_rationale"] = B.get("rationale", "") or ""

        C = answers.get("C", {}) if isinstance(answers.get("C"), dict) else {}
        st.session_state["C_selected"] = C.get("selected", []) or []
        st.session_state["C_rationale"] = C.get("rationale", "") or ""

        D = answers.get("D", {}) if isinstance(answers.get("D"), dict) else {}
        st.session_state["D_selected"] = D.get("selected", []) or []
        st.session_state["D_notes"] = D.get("notes", "") or ""
        st.session_state["D_time"] = D.get("timing", st.session_state.get("D_time", "15 minutes")) or st.session_state.get("D_time", "15 minutes")

        E = answers.get("E", {}) if isinstance(answers.get("E"), dict) else {}
        st.session_state["E_S"] = E.get("S", "") or ""
        st.session_state["E_B"] = E.get("B", "") or ""
        st.session_state["E_A"] = E.get("A", "") or ""
        st.session_state["E_R"] = E.get("R", "") or ""
        st.session_state["E_selected"] = E.get("selected_elements", []) or []

    # Restore NCLEX
    nclex_ans = draft.get("nclex_answers")
    if isinstance(nclex_ans, dict):
        # Clear existing per-widget keys so the UI rehydrates from nclex_answers cleanly
        for k in list(st.session_state.keys()):
            if str(k).startswith("nclex_"):
                try:
                    del st.session_state[k]
                except Exception:
                    pass
        st.session_state["nclex_answers"] = nclex_ans
        st.session_state["practical_submitted"] = bool(draft.get("practical_submitted", False))
        st.session_state["nclex_scored"] = draft.get("nclex_scored")



    
    # Restore A–E progress markers
    lf = draft.get("last_feedback")
    if isinstance(lf, dict):
        st.session_state["last_feedback"] = lf
    st.session_state["ae_focus"] = str(draft.get("ae_focus", st.session_state.get("ae_focus", "A")) or "A")

    # Restore NCLEX finalized/lock state
    if "nclex_finalized" in draft:
        st.session_state["nclex_finalized"] = bool(draft.get("nclex_finalized", False))

# Restore Intake (5 marks)
    intake = draft.get("intake")
    if isinstance(intake, dict):
        st.session_state["intake"] = {
            "age": intake.get("age", ""),
            "setting": intake.get("setting", ""),
            "chief_complaint": intake.get("chief_complaint", ""),
            "signs_symptoms": intake.get("signs_symptoms", ""),
            "findings": intake.get("findings", ""),
            "history": intake.get("history", ""),
        }
    st.session_state["intake_score"] = int(draft.get("intake_score", 0) or 0)
    st.session_state["intake_breakdown"] = draft.get("intake_breakdown", {}) if isinstance(draft.get("intake_breakdown", {}), dict) else {}
    st.session_state["intake_submitted"] = bool(draft.get("intake_submitted", False))

def nclex_autosave_now(features: dict, student_username: str, case_id: str, mode: str):
    """Immediate autosave for NCLEX answers (called on each change)."""
    try:
        autosave_draft(features, student_username, case_id, {
            "answers": st.session_state.get("answers", {}),
            "scores": st.session_state.get("scores", {}),
            "nclex_answers": st.session_state.get("nclex_answers", {}),
            "practical_submitted": bool(st.session_state.get("practical_submitted", False)),
            "nclex_finalized": bool(st.session_state.get("nclex_finalized", False)),
            "caseId": case_id,
            "mode": mode,
            "intake": st.session_state.get("intake", {}),
            "intake_score": int(st.session_state.get("intake_score", 0) or 0),
        })
    except Exception:
        pass


# =============================
# UI
# =============================
st.set_page_config(page_title="ClinIQ Nurse Adult-NURS-Reason", layout="wide")
_show_flash_success()



st.markdown(
    '''
    <style>
    /* NCLEX UI CLEANUP */
    div[role="radiogroup"] svg,
    div[role="radiogroup"] circle {
        display: none !important;
    }

    div[data-testid="stCheckbox"] svg,
    div[data-testid="stCheckbox"] circle {
        display: none !important;
    }

    input[type="radio"], input[type="checkbox"] {
        accent-color: transparent !important;
    }
    </style>
    ''',
    unsafe_allow_html=True
)


def st_radio_no_preselect(label: str, options: list, prev_value, key: str, disabled: bool = False, horizontal: bool = False):
    """Radio with EMPTY circles until user clicks. Uses index=None when possible, with a safe fallback."""
    opts = list(options or [])
    if not opts:
        return None
    idx = opts.index(prev_value) if (prev_value in opts) else None
    try:
        # Streamlit supports index=None (no default selection) in recent versions
        return st.radio(label, opts, index=idx, key=key, disabled=disabled, horizontal=horizontal)
    except TypeError:
        # Fallback for older Streamlit: add a placeholder and treat it as None
        placeholder = "—"
        opts2 = [placeholder] + opts
        idx2 = (opts2.index(prev_value) if (prev_value in opts2) else 0)
        val = st.radio(label, opts2, index=idx2, key=key, disabled=disabled, horizontal=horizontal)
        return None if val == placeholder else val



features = load_features()
maybe_backup_on_start(features)
APP_TITLE = "ClinIQ Nurse Adult-NURS-Reason"
APP_SUBTITLE = "AI Clinical Reasoning • Critical Thinking • Improvement"
LOGO_FILENAME = "logo_cliniq.png"  # put the logo file next to app.py

def render_centered_app_header(title: str, subtitle: str = ""):
    """Centered app header with optional logo (logo_cliniq.png next to app.py)."""
    logo_path = (BASE_DIR / LOGO_FILENAME) if "BASE_DIR" in globals() else (Path(__file__).resolve().parent / LOGO_FILENAME)
    b64 = ""
    try:
        if logo_path.exists():
            b64 = base64.b64encode(logo_path.read_bytes()).decode("utf-8")
    except Exception:
        b64 = ""

    img_html = ""
    if b64:
        img_html = f'<img src="data:image/png;base64,{b64}" style="width:280px; height:auto; margin:0 auto 0.35rem auto; display:block;" />'

    sub_html = f'<div style="font-size:1.05rem; font-weight:700; color:#1f4e79; margin-top:0.1rem;">{subtitle}</div>' if subtitle else ""

    st.markdown(
        f"""<div style="text-align:center; margin-top:0.1rem; margin-bottom:0.4rem;">{img_html}<div style="font-size:2.15rem; font-weight:900; color:#b00020; line-height:1.1;">{title}</div>{sub_html}</div>""",
        unsafe_allow_html=True
    )

render_centered_app_header(APP_TITLE, APP_SUBTITLE)
st.markdown(GLOBAL_CSS, unsafe_allow_html=True)

# =============================
# Session initialization (safe)
# =============================
if ("scores" not in st.session_state) or ("attempt_started_epoch" not in st.session_state) or ("answers" not in st.session_state):
    reset_attempt_state()


# =============================
# AUTH
# =============================
if "is_admin" not in st.session_state:
    st.session_state["is_admin"] = False
if "student_profile" not in st.session_state:
    st.session_state["student_profile"] = None

# =============================
# AUTH (Unified Sidebar Login)
# =============================
if "is_admin" not in st.session_state:
    st.session_state["is_admin"] = False
if "student_profile" not in st.session_state:
    st.session_state["student_profile"] = None

with st.sidebar.expander("🔐 Login", expanded=True):
    role = st.selectbox("Login as", ["Student", "Admin"], index=0, key="login_role_sel")

    if role == "Student":
        if st.session_state.get("student_profile"):
            prof = st.session_state["student_profile"]
            st.success(f"Logged in: {prof.get('display_name','Student')} ({prof.get('username','')})")
            if st.button("Logout Student", key="student_logout_btn"):
                # Clear student identity + any exam/case locks for a clean logout
                st.session_state["student_profile"] = None
                for _k in [
                    "locked_exam_case_id",
                    "exam_in_progress",
                    "exam_started",
                    "exam_case_id",
                    "current_case_id",
                    "active_case_id",
                    "_exam_started_at",
                    "_exam_timer_start",
                    "_exam_timer_deadline",
                    "_resume_pending",
                    "_resume_case_id",
                    "_resume_hint",
                    "_resume_loaded",
                    "_last_seen_reset_token",
                ]:
                    st.session_state.pop(_k, None)
                st.rerun()
        else:
            su = st.text_input("Username", key="stu_user")
            sp = st.text_input("Password", type="password", key="stu_pass")
            if st.button("Login Student", key="stu_login_btn"):
                prof = verify_student(su, sp)
                if prof:
                    st.session_state["student_profile"] = prof
                    st.sidebar.success(f"Welcome, {prof['display_name']}")
                    st.rerun()
                else:
                    st.sidebar.error("Wrong username/password (or students.json has no matching user / exam password mismatch).")

    else:
        # Admin login (password gate for Admin Mode)
        admin_pw = st.text_input("Admin password", type="password", key="admin_pw")
        colA, colB = st.columns(2)
        with colA:
            if st.button("Enable Admin Mode", key="admin_enable_btn"):
                if admin_pw == ADMIN_PASSWORD:
                    st.session_state["is_admin"] = True
                    st.sidebar.success("🧑‍🏫 Admin mode enabled")
                    st.rerun()
                else:
                    st.sidebar.error("Wrong admin password")
        with colB:
            if st.button("Disable Admin Mode", key="admin_disable_btn"):
                st.session_state["is_admin"] = False
                st.sidebar.info("Admin mode disabled")
                st.rerun()

is_admin = bool(st.session_state.get("is_admin", False))
is_student = bool(st.session_state.get("student_profile"))
# Effective admin: never treat a logged-in student session as admin, even if Admin Mode was previously enabled.
effective_admin = bool(is_admin and (not is_student))

# UI: keep sidebar minimal (login + status); admin controls are shown in the top Settings page
SHOW_ADMIN_PANELS_IN_SIDEBAR = False

# Apply student hardening (practice + exam) when a student is logged in.
if is_student:
    apply_student_hardening(st.session_state.get("student_profile") or {})

student_profile = st.session_state.get("student_profile")

is_student_logged_in = bool(student_profile)

# Required: student OR admin
if not is_student_logged_in and not is_admin:
    st.warning("Please login as **Student** (sidebar).")
    st.stop()


# =============================
# Research consent (students only, opt-in, no penalty to opt out)
# =============================
if is_student_logged_in and (not is_admin):
    rp = load_research_policy()
    if bool(rp.get("enabled", False)) and bool(rp.get("require_consent", True)):
        with st.expander(f"📝 {rp.get('consent_title','Research Consent')}", expanded=True):
            st.write(str(rp.get("consent_text","") or ""))
            choice = st.radio(
                "Do you consent to your de-identified data being used for research exports?",
                ["Yes, I consent (opt-in)", "No, I do not consent (opt-out)"],
                index=0 if bool(st.session_state.get("research_consent", False)) else 1,
                key="research_consent_choice",
            )
            st.session_state["research_consent"] = (choice.startswith("Yes"))
            st.caption("No penalty if you opt out: you can still use ClinIQ normally; your data will be excluded from research_dataset/exports.")
    else:
        # Teaching mode default: clear any old consent flag so research is not collected accidentally
        st.session_state["research_consent"] = False

# Badges
if effective_admin:
    st.sidebar.success("🧑‍🏫 Admin mode ON")
else:
    st.sidebar.info("👩‍⚕️ Admin mode OFF")

if is_student_logged_in:
    st.sidebar.success("🎓 Student logged in")
else:
    st.sidebar.warning("🎓 No student logged in (Admin can still test)")

# =============================
# Research consent (students) is shown in the main page under the title in tidy UI.

# =============================
# Load persistent settings
# =============================
admin_settings = load_admin_settings()
research_policy = load_research_policy()
case_policy = load_case_policy()

IS_PRACTICE = (admin_settings.get("app_mode") == "Practice")
IS_EXAM = not IS_PRACTICE



# =============================
# Faculty-only Exam Control (Option 2): End Exam Session
# Visible ONLY when Mode = Exam and user is Admin/Faculty.
# Clears only exam/session lock + attempt session keys for the current Streamlit session,
# then returns to case selection (prevents "stuck in exam" issues).
# =============================
if is_admin and IS_EXAM and SHOW_ADMIN_PANELS_IN_SIDEBAR:
    st.sidebar.markdown(
        "<div style='background:#e8f4ff;border:1px solid rgba(0,0,0,.08);padding:10px 12px;border-radius:14px;'>"
        "<span style='color:#b00020;font-weight:900;'>🔒 Exam Control</span>"
        "</div>",
        unsafe_allow_html=True,
    )
    # End Exam Session scope (Option B)
    _end_scope = st.sidebar.selectbox(
        "End Exam Session for",
        ["This device", "Selected student"],
        index=0,
        key="faculty_end_exam_scope_sel",
    )

    _sel_student = None
    if _end_scope == "Selected student":
        try:
            _sd = load_students()
            _usernames = sorted(list({str(s.get("username","")).strip() for s in (_sd.get("students", []) or []) if str(s.get("username","")).strip()}))
        except Exception:
            _usernames = []
        _sel_student = st.sidebar.selectbox(
            "Student username",
            _usernames if _usernames else [""],
            index=0,
            key="faculty_end_exam_student_sel",
        )

    if st.sidebar.button("🔒 End Exam Session", key="faculty_end_exam_session_btn"):
        if _end_scope == "This device":
            # Clear exam/case-switch lock for this session only (no logs are modified).
            for _k in [
                "locked_exam_case_id",
                "exam_in_progress",
                "exam_started",
                "exam_case_id",
                "current_case_id",
                "active_case_id",
                "_exam_started_at",
                "_exam_timer_start",
                "_exam_timer_deadline",
            ]:
                st.session_state.pop(_k, None)

            # Reset case selectors to allow switching cases immediately.
            for _k in ["case_pick_main", "rot_case_pick_main", "case_system_main"]:
                st.session_state.pop(_k, None)

            # Reset attempt-level state using the app's existing helper.
            try:
                reset_attempt_state()
            except Exception:
                pass

            st.sidebar.success("Exam session ended for this device. You can now select a different case.")
            st.rerun()

        else:
            _sel_student = (_sel_student or "").strip()
            if not _sel_student:
                st.sidebar.error("Please select a student username.")
            else:
                # New behavior: End/Reset Selected Student (server-side)
                try:
                    set_student_force_unlock(_sel_student, True)
                    bump_student_reset_token(_sel_student)  # also clears autosaves on next run
                    st.sidebar.success(f"End/Reset requested for student: {_sel_student}. Ask them to refresh the app.")
                except Exception:
                    st.sidebar.error("Could not request reset for the selected student.")

    # Student Session Manager (delete unfinished autosaves + force unlock/reset)
    with st.sidebar.expander("👤 Student Session Manager", expanded=False):
        try:
            _sd2 = load_students()
            _usernames2 = sorted(list({str(s.get("username","")).strip() for s in (_sd2.get("students", []) or []) if str(s.get("username","")).strip()}))
        except Exception:
            _usernames2 = []
        _mgr_student = st.selectbox(
            "Select student",
            _usernames2 if _usernames2 else [""],
            index=0,
            key="admin_student_mgr_sel",
        )
        _mgr_student = (_mgr_student or "").strip()

        if _mgr_student:
            drafts = index_latest_autosaves_for_student(_mgr_student) or {}
            case_ids = sorted(list(drafts.keys()))
            st.caption(f"In-progress autosaves detected: {len(case_ids)}")
            if case_ids:
                _pick = st.multiselect(
                    "Autosaved cases (select to delete)",
                    case_ids,
                    default=[],
                    key="admin_student_mgr_case_pick",
                )
            else:
                _pick = []

            c1, c2, c3 = st.columns(3)
            with c1:
                if st.button("🧹 Delete selected autosaves", key="admin_del_sel_autosaves_btn"):
                    if not _pick:
                        st.warning("Select at least one case ID to delete.")
                    else:
                        removed_total = 0
                        for cid in _pick:
                            removed_total += delete_autosaves_for(_mgr_student, cid)
                        st.success(f"Deleted {removed_total} autosave record(s).")
                        st.rerun()
            with c2:
                if st.button("🗑️ Delete ALL autosaves", key="admin_del_all_autosaves_btn"):
                    removed_total = delete_all_autosaves_for_student(_mgr_student)
                    st.success(f"Deleted {removed_total} autosave record(s).")
                    st.rerun()
            with c3:
                if st.button("🔓 Force unlock", key="admin_force_unlock_btn"):
                    set_student_force_unlock(_mgr_student, True)
                    st.success("Force unlock requested. Student should refresh the app.")
                    st.rerun()

            st.markdown("---")
            if st.button("🔁 Full reset (unlock + clear autosaves)", key="admin_full_reset_btn"):
                set_student_force_unlock(_mgr_student, True)
                bump_student_reset_token(_mgr_student)
                st.success("Full reset requested. Student should refresh the app.")
                st.rerun()
        else:
            st.info("Select a student to manage their unfinished progress or unlock/reset their exam.")

# =============================
# Load cases
# =============================
if not CASES_PATH.exists():
    st.error("❌ cases.json not found in the same folder as app.py.")
    st.info("Fix: Put cases.json in this folder: " + str(BASE_DIR))
    st.stop()

try:
    cases = load_cases()
except Exception as e:
    st.error("❌ Could not load cases.json (JSON may be invalid).")
    st.write(e)
    st.stop()

st.sidebar.success(f"Loaded cases: {len(cases)}")

# =============================
# Admin-only: Student Generator (moved to Settings in tidy UI)
# (kept here for compatibility; hidden unless SHOW_ADMIN_PANELS_IN_SIDEBAR is True)
# =============================
if effective_admin and SHOW_ADMIN_PANELS_IN_SIDEBAR:
    with st.sidebar.expander("👥 Student Accounts (Generate list + passwords)", expanded=False):
        st.caption("This creates students in students.json and gives you a CSV to distribute.")
        prefix = st.text_input("Username prefix", value="nurse")
        start_num = st.number_input("Start number", min_value=1, value=1, step=1)
        count = st.number_input("How many students to generate", min_value=1, value=30, step=1)
        cohort = st.text_input("Cohort (optional)", value="Adult Nursing Y2")
        pw_len = st.number_input("Password length", min_value=8, max_value=24, value=12, step=1)

        if st.button("➕ Generate & Save Students"):
            data = load_students()
            created_rows = []
            new_students = []
            for i in range(int(count)):
                n = int(start_num) + i
                username = f"{prefix}{n:03d}"
                password_plain = generate_password(int(pw_len))
                student_id = f"{n:03d}"
                display_name = f"Student {n:03d}"
                new_students.append({
                    "username": username,
                    "display_name": display_name,
                    "student_id": student_id,
                    "cohort": cohort,
                    "password_sha256": sha256_hex(password_plain),
                })
                created_rows.append({
                    "username": username,
                    "password": password_plain,
                    "display_name": display_name,
                    "student_id": student_id,
                    "cohort": cohort,
                })

            data = upsert_students(data, new_students)
            save_students(data)

            out = io.StringIO()
            w = csv.DictWriter(out, fieldnames=["username", "password", "display_name", "student_id", "cohort"])
            w.writeheader()
            for r in created_rows:
                w.writerow(r)

            st.download_button(
                "⬇️ Download NEW credentials CSV",
                data=out.getvalue().encode("utf-8"),
                file_name=f"student_credentials_{utc_now_iso().replace(':','-')}.csv",
                mime="text/csv"
            )
            st.success(f"Saved {len(created_rows)} accounts into students.json")
            # Store last generated credentials for downloads
            st.session_state["last_generated_credentials_rows"] = created_rows


# =============================
# Admin-only: Password Management (sidebar)
# =============================
if effective_admin and SHOW_ADMIN_PANELS_IN_SIDEBAR:
    with st.sidebar.expander("🔑 Password Management (Reset / Exam / Expiry)", expanded=False):
        st.caption("Reset individual student passwords + enable exam password override with auto-expiry (Qatar time).")

        policy_exam = load_exam_access_policy()
        active_exam = is_exam_password_active(policy_exam)

        st.markdown("### 🧑‍🎓 Individual student reset")
        data_students = load_students()
        usernames = sorted([str(s.get("username", "")).strip() for s in data_students.get("students", []) if str(s.get("username", "")).strip()])
        if usernames:
            target_user = st.selectbox("Select student username", usernames, key="pw_reset_user")
            new_pw = st.text_input("New password for this student", type="password", key="pw_reset_newpw")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Reset this student password", key="pw_reset_btn"):
                    if not new_pw.strip():
                        st.error("Enter a new password.")
                    else:
                        for s in data_students.get("students", []):
                            if str(s.get("username", "")).strip() == target_user:
                                s["password_sha256"] = sha256_hex(new_pw.strip())
                                break
                        save_students(data_students)
                        st.success(f"Password updated for {target_user}")
            with col2:
                if st.button("Generate random password", key="pw_reset_gen_btn"):
                    rp = generate_password(12)
                    st.session_state["pw_reset_newpw"] = rp
                    st.info(f"Generated: {rp}")
        else:
            st.warning("No students found in students.json")

        st.divider()
        st.markdown("### 📝 Exam password (everyone)")
        st.write("Status:", "✅ ACTIVE" if active_exam else "⛔ INACTIVE")

        policy_exam["force_exam_password_only"] = st.toggle(
            "Force exam password only (recommended for exams)",
            value=bool(policy_exam.get("force_exam_password_only", True)),
            key="force_exam_pw_only"
        )

        exam_pw_plain = st.text_input("Set exam password (e.g., NURS@2026)", type="password", key="exam_pw_plain")
        expires_enabled = st.toggle("Auto-expire exam password", value=True, key="exam_expire_toggle")

        default_exp = now_local().replace(hour=12, minute=0, second=0, microsecond=0)
        exp_date = st.date_input("Expiry date (Qatar time)", value=default_exp.date(), key="exam_expire_date")
        exp_time = st.time_input("Expiry time (Qatar time)", value=default_exp.time(), key="exam_expire_time")

        if st.button("✅ Enable exam password now", key="exam_enable_btn"):
            if not exam_pw_plain.strip():
                st.error("Enter an exam password.")
            else:
                policy_exam["enabled"] = True
                policy_exam["exam_password_sha256"] = sha256_hex(exam_pw_plain.strip())

                if expires_enabled:
                    exp_dt = datetime.combine(exp_date, exp_time).replace(tzinfo=TZ)
                    policy_exam["expires_at"] = exp_dt.isoformat()
                else:
                    policy_exam["expires_at"] = ""

                save_exam_access_policy(policy_exam)
                st.success("Exam password enabled.")
                st.rerun()

        colA, colB = st.columns(2)
        with colA:
            if st.button("⛔ Disable exam password", key="exam_disable_btn"):
                policy_exam["enabled"] = False
                save_exam_access_policy(policy_exam)
                st.success("Exam password disabled.")
                st.rerun()

        with colB:
            if st.button("📄 Show exam policy details", key="exam_show_policy_btn"):
                st.json(load_exam_access_policy())

# =============================
# Admin-only: Mode + AI Settings
# =============================
if effective_admin and SHOW_ADMIN_PANELS_IN_SIDEBAR:
    with st.sidebar.expander("🧪 Mode (Practice / Exam)", expanded=True):
        new_mode = st.selectbox(
            "Select mode (saved for everyone)",
            ["Practice", "Exam"],
            index=0 if admin_settings.get("app_mode") == "Practice" else 1,
        )
        if new_mode != admin_settings.get("app_mode"):
            admin_settings["app_mode"] = new_mode
            save_admin_settings(admin_settings)
            st.success(f"Mode saved: {new_mode}")
            st.rerun()

    with st.sidebar.expander("🧠 AI Coach (Step 9)", expanded=False):
        ai_enabled = st.toggle(
            "Enable AI coaching after each submission (Practice only)",
            value=bool(admin_settings.get("ai_enabled", False)),
            disabled=(new_mode == "Exam")
        )
        ai_debrief_enabled = st.toggle(
            "Enable end-of-case AI debrief (Practice only)",
            value=bool(admin_settings.get("ai_debrief_enabled", False)),
            disabled=(new_mode == "Exam")
        )
        ai_model = st.text_input("Model", value=str(admin_settings.get("ai_model", "gpt-5.2")))

        st.caption("AI requires OPENAI_API_KEY on the machine/server running Streamlit.")
        st.caption("If AI still doesn't run, open Debug panel and check OPENAI_API_KEY set: True.")

        if st.button("💾 Save AI Settings"):
            admin_settings["ai_enabled"] = bool(ai_enabled)
            admin_settings["ai_debrief_enabled"] = bool(ai_debrief_enabled)
            admin_settings["ai_model"] = (ai_model or "gpt-5.2").strip()
            save_admin_settings(admin_settings)
            st.success("Saved AI settings for everyone.")
            st.rerun()
else:
    st.sidebar.markdown(f"**Mode:** `{admin_settings.get('app_mode')}`")



# =============================
# ✅ Admin-only: Research & Teaching KPIs
# =============================
if effective_admin and SHOW_ADMIN_PANELS_IN_SIDEBAR:
    with st.sidebar.expander("📊 Research & Teaching KPIs", expanded=False):
        kpi = load_kpi_policy()

        kpi["total_score"] = st.checkbox("Track total score (A–E)", value=bool(kpi.get("total_score", True)))
        kpi["domain_profile"] = st.checkbox("Track domain profile (A–E)", value=bool(kpi.get("domain_profile", True)))
        kpi["nclex_score"] = st.checkbox("Track NCLEX-style practice score", value=bool(kpi.get("nclex_score", True)))
        kpi["time_to_completion"] = st.checkbox("Track time to completion", value=bool(kpi.get("time_to_completion", True)))
        kpi["unsafe_flags"] = st.checkbox("Track unsafe actions count", value=bool(kpi.get("unsafe_flags", True)))
        kpi["attempts_per_case"] = st.checkbox("Track attempts per case", value=bool(kpi.get("attempts_per_case", True)))

        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Enable ALL KPIs", key="kpi_enable_all"):
                kpi = {
                    "total_score": True,
                    "domain_profile": True,
                    "nclex_score": True,
                    "time_to_completion": True,
                    "unsafe_flags": True,
                    "attempts_per_case": True
                }
                save_kpi_policy(kpi)
                st.success("All KPIs enabled.")
                st.rerun()


# =============================
# Admin-only: Research Mode + Exports
# =============================
if effective_admin and SHOW_ADMIN_PANELS_IN_SIDEBAR:
    with st.sidebar.expander("🔬 Research Mode", expanded=False):
        rp = load_research_policy()
        rp["enabled"] = st.checkbox("Enable research collection", value=bool(rp.get("enabled", False)))
        rp["require_consent"] = st.checkbox("Require student consent", value=bool(rp.get("require_consent", True)))
        rp["anonymize_student_id"] = st.checkbox("Anonymize student IDs in exports", value=bool(rp.get("anonymize_student_id", True)))
        rp["collect_reflection"] = st.checkbox("Collect reflection (post-test)", value=bool(rp.get("collect_reflection", False)))
        rp["collect_answer_change"] = st.checkbox("Collect answer-change behavior (NCLEX)", value=bool(rp.get("collect_answer_change", True)))
        rp["collect_section_performance"] = st.checkbox("Collect section-level performance summary", value=bool(rp.get("collect_section_performance", True)))

        rp["consent_title"] = st.text_input("Consent title", value=str(rp.get("consent_title","Research Consent")))
        rp["consent_text"] = st.text_area("Consent text", value=str(rp.get("consent_text","")), height=120)

        if st.button("💾 Save research settings", key="save_research_policy_btn"):
            save_research_policy(rp)
            st.success("Saved research_policy.json")
            st.rerun()

    
        # Quick access to psychometrics (computed on-demand inside Research Reports)
        st.info("Psychometrics (KR-20 / item difficulty / discrimination) are available under Admin Pages → 📊 Research Reports.")
        if st.button("🧪 Open Research Reports (Psychometrics)", key="open_research_reports_btn"):
            st.session_state["admin_pages_v6"] = "📊 Research Reports"
            st.rerun()
        st.markdown("---")
        st.download_button(
            "⬇️ Download Research CSV",
            data=build_research_csv_bytes(),
            file_name=f"research_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv",
            key="dl_research_csv"
        )
        st.download_button(
            "⬇️ Download Attempt Summary CSV",
            data=build_attempt_summary_csv_bytes(),
            file_name=f"attempt_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv",
            key="dl_attempt_summary_csv"
        )

        clear_after = st.checkbox("Clear logs after archive", value=False, key="archive_clear_after")
        if st.button("🗄️ Archive research logs", key="archive_logs_btn"):
            ok, msg = archive_research_logs(clear_after=bool(clear_after))
            if ok:
                st.success(msg)
            else:
                st.warning(msg)
            with col2:
                if st.button("💾 Save KPI settings", key="kpi_save"):
                    save_kpi_policy(kpi)
                    st.success("KPI settings saved.")
        

# =============================
# Attempts policy
# =============================
def max_attempts_for_case(case_id: str):
    """Return allowed attempts for a case (student-facing).

    Priority (newest -> oldest):
    1) attempts_policy.json per_case override (legacy)
    2) case_policy.json default_attempts_allowed (current admin setting)
    3) attempts_policy.json default_max_attempts (legacy)
    """
    case_id = str(case_id or "").strip()

    # Always re-load (so changes in Settings reflect immediately after Save)
    pol_legacy = load_attempt_policy()
    per = pol_legacy.get("per_case", {}).get(case_id)
    if per in (None, "", "default"):
        per = None
    if per not in (None, ""):
        return per

    # Current settings file (used by the Settings page)
    try:
        cp = load_case_policy()
        if isinstance(cp, dict) and cp.get("default_attempts_allowed") is not None:
            return int(cp.get("default_attempts_allowed", 1))
    except Exception:
        pass

    return pol_legacy.get("default_max_attempts", "unlimited")


if effective_admin and SHOW_ADMIN_PANELS_IN_SIDEBAR:
    with st.sidebar.expander("🎯 Attempts Policy (per case)", expanded=False):
        default_choice = st.selectbox(
            "Default attempts for all cases",
            ["unlimited", "1", "2", "3", "5"],
            index=["unlimited", "1", "2", "3", "5"].index(str(policy_attempts.get("default_max_attempts", "unlimited")))
        )
        if st.button("💾 Save default attempts"):
            policy_attempts["default_max_attempts"] = default_choice
            save_attempt_policy(policy_attempts)
            st.success("Saved attempt policy.")
    

# =============================
# Case visibility + timer policy (Admin controls)
# =============================
if effective_admin and SHOW_ADMIN_PANELS_IN_SIDEBAR:
    with st.sidebar.expander("👁️ Case Visibility + ⏱ Timers", expanded=False):
        st.caption("Controls what students can see + time limits per case.")

        default_vis = st.toggle(
            "Default: cases visible to students",
            value=bool(case_policy.get("default_visibility", True))
        )

        d_prac = case_policy.get("default_timer_minutes_practice", "unlimited")
        d_exam = case_policy.get("default_timer_minutes_exam", 20)

        _timer_opts = ["unlimited", "10", "15", "20", "30", "45", "60", "120", "180"]

        default_timer_practice = st.selectbox(
            "Default timer (Practice)",
            _timer_opts,
            index=_timer_opts.index(str(d_prac)) if str(d_prac) in _timer_opts else 0
        )

        default_timer_exam = st.selectbox(
            "Default timer (Exam)",
            _timer_opts,
            index=_timer_opts.index(str(d_exam)) if str(d_exam) in _timer_opts else 3
        )

        case_ids = [str(c.get("id", "")) for c in cases]
        pick_case_id = st.selectbox("Edit a case", case_ids)

        per = case_policy.get("per_case", {}).get(pick_case_id, {})
        if not isinstance(per, dict):
            per = {}

        per_visible = st.toggle("This case visible to students", value=bool(per.get("visible", default_vis)))

        per_prac = per.get("timer_minutes_practice", default_timer_practice)
        per_exam = per.get("timer_minutes_exam", default_timer_exam)

        _per_timer_opts = ["default"] + _timer_opts
        per_timer_practice = st.selectbox(
            "Timer for this case (Practice)",
            _per_timer_opts,
            index=_per_timer_opts.index(str(per_prac)) if str(per_prac) in _per_timer_opts else 0
        )
        per_timer_exam = st.selectbox(
            "Timer for this case (Exam)",
            _per_timer_opts,
            index=_per_timer_opts.index(str(per_exam)) if str(per_exam) in _per_timer_opts else 0
        )

        if st.button("💾 Save visibility + timers"):
            case_policy["default_visibility"] = bool(default_vis)
            case_policy["default_timer_minutes_practice"] = default_timer_practice
            case_policy["default_timer_minutes_exam"] = default_timer_exam

            # Keep seconds representation in sync (prevents UI showing stale/doubled values)
            def _to_sec(v):
                if v in (None, "", 0, "0", "Unlimited", "unlimited"):
                    return 0
                try:
                    return int(v) * 60
                except Exception:
                    return 0

            case_policy["default_timer_seconds_practice"] = _to_sec(default_timer_practice)
            case_policy["default_timer_seconds_exam"] = _to_sec(default_timer_exam)

            case_policy.setdefault("per_case", {})
            per_out = {
                "visible": bool(per_visible),
                "timer_minutes_practice": per_timer_practice,
                "timer_minutes_exam": per_timer_exam,
            }

            # Per-case seconds timers: if set to "default", remove seconds override
            if str(per_timer_practice) == "default":
                per_out.pop("timer_seconds_practice", None)
            else:
                per_out["timer_seconds_practice"] = _to_sec(per_timer_practice)

            if str(per_timer_exam) == "default":
                per_out.pop("timer_seconds_exam", None)
            else:
                per_out["timer_seconds_exam"] = _to_sec(per_timer_exam)

            case_policy["per_case"][pick_case_id] = per_out
            save_case_policy(case_policy)
            st.success("Saved case_policy.json")
            st.rerun()


# =============================
# ✅ Admin-only Research / Reports / Data Tools (v6)
# (Read-only by default; does NOT modify student workflow)
# =============================
def _safe_iso_to_dt(s: str):
    from datetime import datetime
    try:
        return datetime.fromisoformat(str(s).replace("Z",""))
    except Exception:
        return None

def _load_attempts_records(path: Path):
    records = []
    if not path.exists():
        return records
    try:
        with path.open("r", encoding="utf-8") as f:
            for line in f:
                line = (line or "").strip()
                if not line:
                    continue
                try:
                    rec = json.loads(line)
                    if isinstance(rec, dict):
                        records.append(rec)
                except Exception:
                    # skip malformed line
                    continue
    except Exception:
        return records
    return records

def _attempt_case_id(a: dict) -> str:
    return (a.get("caseId") or a.get("case_id") or a.get("case") or "").strip()

def _attempt_student(a: dict) -> str:
    return (a.get("student_username") or a.get("student_id") or a.get("student_display_name") or "").strip()

def _attempt_system(a: dict) -> str:
    # prefer stored value (what student saw at the time)
    s = (a.get("system") or "").strip()
    if s:
        return s
    cid = _attempt_case_id(a)
    if cid and isinstance(CASES_BY_ID.get(cid), dict):
        return safe_get_system(CASES_BY_ID[cid]) or ""
    return ""

def _attempt_total(a: dict) -> float:
    t = a.get("total")
    if isinstance(t, (int, float)):
        return float(t)
    # sometimes stored as sum of A-E
    sc = a.get("scores") if isinstance(a.get("scores"), dict) else {}
    if sc:
        try:
            return float(sum(float(sc.get(k,0) or 0) for k in ["A","B","C","D","E"]))
        except Exception:
            return 0.0
    return 0.0

def _attempt_intake_score(a: dict) -> float:
    v = a.get("intake_score")
    if isinstance(v, (int, float)):
        return float(v)
    return 0.0

def _attempt_total_with_intake(a: dict) -> float:
    # prefer stored
    v = a.get("total_with_intake")
    if isinstance(v, (int, float)):
        return float(v)
    return _attempt_total(a) + _attempt_intake_score(a)

def _attempt_domain_scores(a: dict) -> dict:
    sc = a.get("scores")
    return sc if isinstance(sc, dict) else {}

def _attempt_unsafe_counts(a: dict) -> dict:
    uc = a.get("unsafe_counts")
    return uc if isinstance(uc, dict) else {}

def _attempt_nclex_summary(a: dict):
    n = a.get("nclex")
    if not isinstance(n, dict):
        return (0.0, 0.0)
    pts = n.get("total_points")
    mx = n.get("total_max")
    if isinstance(pts,(int,float)) and isinstance(mx,(int,float)):
        return (float(pts), float(mx))
    return (0.0, 0.0)

def _attempt_nclex_details(a: dict):
    n = a.get("nclex")
    if not isinstance(n, dict):
        return []
    det = n.get("details")
    return det if isinstance(det, list) else []

def _build_cases_index(cases_list):
    by_id = {}
    for c in (cases_list or []):
        if isinstance(c, dict):
            cid = str(c.get("id") or c.get("caseId") or c.get("case_id") or "").strip()
            if cid:
                by_id[cid] = c
    return by_id

CASES_BY_ID = _build_cases_index(cases)

def _build_nclex_index(nclex_data):
    # expected structure: {"cases": {"case_id": {"items":[...]}}} OR {"case_id": {"items":[...]}}
    idx = {}
    if isinstance(nclex_data, dict):
        if isinstance(nclex_data.get("cases"), dict):
            for cid, blob in nclex_data["cases"].items():
                items = (blob or {}).get("items") if isinstance(blob, dict) else None
                if isinstance(items, list):
                    for it in items:
                        if isinstance(it, dict) and it.get("id"):
                            idx[str(it["id"])] = it
        else:
            # try flat
            for cid, blob in nclex_data.items():
                if isinstance(blob, dict) and isinstance(blob.get("items"), list):
                    for it in blob["items"]:
                        if isinstance(it, dict) and it.get("id"):
                            idx[str(it["id"])] = it
    return idx

try:
    _NCLEX_DATA = load_nclex_items()
except Exception:
    _NCLEX_DATA = {}
NCLEX_BY_QID = _build_nclex_index(_NCLEX_DATA)

def _filter_attempts(attempts, *, systems_sel=None, cases_sel=None, student_q="", dt_start=None, dt_end=None, score_min=None, score_max=None):
    out = []
    for a in attempts:
        if not isinstance(a, dict):
            continue
        cid = _attempt_case_id(a)
        sys = _attempt_system(a)
        stu = _attempt_student(a)
        ts = _safe_iso_to_dt(a.get("timestamp",""))
        total25 = _attempt_total_with_intake(a)

        if systems_sel and "All" not in systems_sel:
            if (sys or "—") not in systems_sel:
                continue
        if cases_sel and "All" not in cases_sel:
            if (cid or "—") not in cases_sel:
                continue
        if student_q:
            q = student_q.strip().lower()
            if q not in (stu or "").lower() and q not in (a.get("student_display_name","") or "").lower():
                continue
        if dt_start and ts and ts < dt_start:
            continue
        if dt_end and ts and ts > dt_end:
            continue
        if score_min is not None and total25 < float(score_min):
            continue
        if score_max is not None and total25 > float(score_max):
            continue
        out.append(a)
    return out

def _to_csv_bytes(rows, fieldnames):
    buf = io.StringIO()
    w = csv.DictWriter(buf, fieldnames=fieldnames)
    w.writeheader()
    for r in rows:
        w.writerow({k: r.get(k, "") for k in fieldnames})
    return buf.getvalue().encode("utf-8")

def admin_page_research_reports():
    st.header("📊 Research Reports (Admin)")
    st.caption("Filters + summary + exports. Read-only unless you click export.")

    attempts = _load_attempts_records(ATTEMPTS_PATH)

    cases_list_all = get_cases_list()
    all_systems = sorted({safe_get_system(c) for c in cases_list_all}) or ["—"]
    all_cases = sorted({str(c.get("id","")).strip() for c in cases_list_all if str(c.get("id","")).strip()}) or ["—"]
    # Include unknown values observed in attempts (legacy records)
    all_systems = sorted(set(all_systems) | {(_attempt_system(a) or "—") for a in attempts})
    all_cases = sorted(set(all_cases) | {(_attempt_case_id(a) or "—") for a in attempts})

    c1, c2, c3 = st.columns([2,2,2])
    with c1:
        systems_sel = st.multiselect("Systems", ["All"] + all_systems, default=["All"])
    with c2:
        cases_sel = st.multiselect("Cases", ["All"] + all_cases, default=["All"])
    with c3:
        student_q = st.text_input("Student (username/name contains)", value="")

    c4, c5, c6 = st.columns([2,2,2])
    with c4:
        start_date = st.date_input("Start date", value=None)
    with c5:
        end_date = st.date_input("End date", value=None)
    with c6:
        pass_threshold = st.number_input("Pass threshold (/25)", min_value=0.0, max_value=25.0, value=17.5, step=0.5)

    from datetime import datetime, time
    dt_start = datetime.combine(start_date, time.min) if start_date else None
    dt_end = datetime.combine(end_date, time.max) if end_date else None

    c7, c8 = st.columns(2)
    with c7:
        score_min = st.number_input("Score min (/25)", min_value=0.0, max_value=25.0, value=0.0, step=0.5)
    with c8:
        score_max = st.number_input("Score max (/25)", min_value=0.0, max_value=25.0, value=25.0, step=0.5)

    filtered = _filter_attempts(attempts, systems_sel=systems_sel, cases_sel=cases_sel, student_q=student_q, dt_start=dt_start, dt_end=dt_end, score_min=score_min, score_max=score_max)

    # Summary
    import statistics
    totals20 = [_attempt_total(a) for a in filtered]
    totals25 = [_attempt_total_with_intake(a) for a in filtered]
    n_unique_students = len({(_attempt_student(a) or "—") for a in filtered})

    st.subheader("Summary")
    cA, cB, cC, cD = st.columns(4)
    cA.metric("Attempts", len(filtered))
    cB.metric("Unique students", n_unique_students)
    cC.metric("Avg total (/25)", f"{(sum(totals25)/len(totals25)):.2f}" if totals25 else "—")
    cD.metric("Median (/25)", f"{statistics.median(totals25):.2f}" if totals25 else "—")

    if totals25:
        passed = sum(1 for x in totals25 if x >= float(pass_threshold))
        st.metric("Pass rate", f"{(passed/len(totals25))*100:.1f}%")

    # Attempts per student
    st.subheader("Attempts per student")
    per_student = {}
    for a in filtered:
        stu = _attempt_student(a) or "—"
        per_student[stu] = per_student.get(stu, 0) + 1
    top_rows = [{"student": k, "attempts": v} for k,v in sorted(per_student.items(), key=lambda kv: (-kv[1], kv[0]))]
    st.dataframe(top_rows, width="stretch", hide_index=True)

    # Domain performance
    st.subheader("Domain performance (A–E)")
    doms = ["A","B","C","D","E"]
    dom_rows = []
    for d in doms:
        vals = []
        for a in filtered:
            sc = _attempt_domain_scores(a)
            v = sc.get(d)
            if isinstance(v, (int,float)):
                vals.append(float(v))
        if vals:
            dom_rows.append({"domain": d, "avg(/4)": round(sum(vals)/len(vals), 2), "miss_rate": round((4-(sum(vals)/len(vals)))/4*100, 1)})
        else:
            dom_rows.append({"domain": d, "avg(/4)": "—", "miss_rate": "—"})
    st.dataframe(dom_rows, width="stretch", hide_index=True)

    # NCLEX summary
    st.subheader("NCLEX summary")
    n_pts = []
    n_max = []
    for a in filtered:
        pts, mx = _attempt_nclex_summary(a)
        if mx > 0:
            n_pts.append(pts)
            n_max.append(mx)
    if n_max:
        total_pts = sum(n_pts)
        total_mx = sum(n_max)
        st.write(f"NCLEX coverage: {len(n_max)} attempts include NCLEX scoring.")
        st.metric("Avg NCLEX %", f"{(total_pts/total_mx)*100:.1f}%")
    else:
        st.info("No NCLEX scoring found in the filtered attempts.")

    # Export: filtered attempts (flat)

    # Psychometrics (NCLEX) — computed on demand to avoid slowing the app
    st.subheader("🧪 Psychometrics / QA (NCLEX)")
    st.caption("Compute KR-20, item difficulty, and discrimination from aggregated attempts. Files saved next to app.py when computed.")
    with st.expander("🧪 Psychometrics (NCLEX) — KR-20 / difficulty / discrimination", expanded=True):
        st.caption("Runs only when you click the button (no runtime slowdown). Uses per-item correctness stored in attempts_log.jsonl.")
        min_attempts = st.number_input("Minimum attempts per item (filter)", min_value=2, max_value=500, value=10, step=1, key="psy_min_attempts")
        min_items_intersection = st.number_input("Minimum common items required for KR-20 (intersection)", min_value=2, max_value=200, value=10, step=1, key="psy_min_items_intersection")

        def _corr(x, y):
            try:
                n = len(x)
                if n < 3:
                    return None
                mx = sum(x)/n
                my = sum(y)/n
                vx = sum((xi-mx)**2 for xi in x)
                vy = sum((yi-my)**2 for yi in y)
                if vx <= 0 or vy <= 0:
                    return None
                cov = sum((x[i]-mx)*(y[i]-my) for i in range(n))
                return cov / math.sqrt(vx*vy)
            except Exception:
                return None

        def _kr20(scores_by_attempt):
            # scores_by_attempt: list of lists/tuples of 0/1 with same length k
            try:
                if not scores_by_attempt:
                    return None
                k = len(scores_by_attempt[0])
                if k < 2:
                    return None
                totals = [sum(row) for row in scores_by_attempt]
                if len(totals) < 3:
                    return None
                mean_total = sum(totals)/len(totals)
                var_total = sum((t-mean_total)**2 for t in totals) / (len(totals)-1)
                if var_total <= 0:
                    return None
                ps = []
                for j in range(k):
                    col = [row[j] for row in scores_by_attempt]
                    p = sum(col)/len(col)
                    ps.append(p)
                sum_pq = sum(p*(1-p) for p in ps)
                return (k/(k-1)) * (1 - (sum_pq/var_total))
            except Exception:
                return None

        def _extract_item_correct_map(attempt):
            # Returns {qid: 0/1} or {}
            det = _attempt_nclex_details(attempt)
            out = {}
            for d in (det or []):
                qid = d.get("qid")
                if not qid:
                    continue
                c = d.get("correct")
                if c is True:
                    out[qid] = 1
                elif c is False:
                    out[qid] = 0
            return out

        def _compute_psychometrics_for_case(attempts_case):
            # returns: item_rows(list[dict]), kr20_value, k_items_intersection, n_attempts_used
            maps = []
            for a in attempts_case:
                m = _extract_item_correct_map(a)
                if m:
                    maps.append(m)
            if len(maps) < 3:
                return [], None, 0, len(maps)

            # Union and counts
            all_qids = {}
            for m in maps:
                for qid, v in m.items():
                    all_qids[qid] = all_qids.get(qid, 0) + 1

            # Per-item stats (difficulty + discrimination)
            item_rows = []
            for qid, cnt in sorted(all_qids.items(), key=lambda kv: (-kv[1], kv[0])):
                if cnt < int(min_attempts):
                    continue
                xs = []
                totals_excl = []
                for m in maps:
                    if qid not in m:
                        continue
                    x = m[qid]
                    # total excluding this item (over the attempt's available items)
                    total_ex = sum(v for k,v in m.items() if k != qid)
                    xs.append(x)
                    totals_excl.append(total_ex)
                p = (sum(xs)/len(xs)) if xs else None
                disc = _corr(xs, totals_excl) if xs else None
                it = NCLEX_BY_QID.get(qid, {}) if isinstance(globals().get("NCLEX_BY_QID", {}), dict) else {}
                item_rows.append({
                    "qid": qid,
                    "n": len(xs),
                    "difficulty_index_p": round(p, 4) if isinstance(p, (int,float)) else None,
                    "discrimination_r": round(disc, 4) if isinstance(disc, (int,float)) else None,
                    "type": it.get("type",""),
                    "designed_difficulty": it.get("difficulty",""),
                    "client_need": it.get("client_need",""),
                })

            # KR-20 on intersection items (stable k across attempts)
            qids_intersection = None
            for m in maps:
                s = set(m.keys())
                qids_intersection = s if qids_intersection is None else (qids_intersection & s)
            qids_intersection = sorted(list(qids_intersection or []))
            if len(qids_intersection) < int(min_items_intersection):
                kr = None
                kint = len(qids_intersection)
                n_used = len(maps)
            else:
                matrix = []
                for m in maps:
                    row = [m.get(qid, 0) for qid in qids_intersection]
                    matrix.append(row)
                kr = _kr20(matrix)
                kint = len(qids_intersection)
                n_used = len(matrix)

            return item_rows, kr, kint, n_used

        if st.button("Compute psychometrics for filtered attempts", key="psy_compute_btn"):
            by_case = {}
            for a in filtered:
                cid = _attempt_case_id(a) or "—"
                by_case.setdefault(cid, []).append(a)

            out_rows = []
            summary_rows = []
            for cid, group in sorted(by_case.items(), key=lambda kv: kv[0]):
                item_rows, kr, kint, n_used = _compute_psychometrics_for_case(group)
                if item_rows:
                    for r in item_rows:
                        r2 = dict(r)
                        r2["case_id"] = cid
                        out_rows.append(r2)
                summary_rows.append({
                    "case_id": cid,
                    "attempts_with_item_details": n_used,
                    "items_reported(>=min_attempts)": sum(1 for r in item_rows if r.get("n",0) >= int(min_attempts)),
                    "kr20_intersection": round(kr, 4) if isinstance(kr, (int,float)) else None,
                    "k_items_intersection": kint,
                })

            # Write psychometrics outputs to disk (optional but useful for QA/audit trails)
            try:
                ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                items_path = Path("psychometrics_items.json")
                cases_path = Path("psychometrics_cases.json")
                summary_path = Path("psychometrics_summary.json")

                items_payload = {
                    "generated_at": ts,
                    "min_attempts_per_item": int(min_attempts),
                    "min_items_intersection": int(min_items_intersection),
                    "rows": out_rows,
                }
                cases_payload = {
                    "generated_at": ts,
                    "min_attempts_per_item": int(min_attempts),
                    "min_items_intersection": int(min_items_intersection),
                    "rows": summary_rows,
                }
                summary_payload = {
                    "generated_at": ts,
                    "note": "KR-20 is computed on the intersection of common items across attempts within each case.",
                    "attempts_filtered": len(filtered),
                    "cases_count": len(summary_rows),
                    "items_rows": len(out_rows),
                }

                items_path.write_text(json.dumps(items_payload, ensure_ascii=False, indent=2), encoding="utf-8")
                cases_path.write_text(json.dumps(cases_payload, ensure_ascii=False, indent=2), encoding="utf-8")
                summary_path.write_text(json.dumps(summary_payload, ensure_ascii=False, indent=2), encoding="utf-8")

                st.success("Saved psychometrics outputs: psychometrics_items.json, psychometrics_cases.json, psychometrics_summary.json")
            except Exception as e:
                st.warning(f"Psychometrics computed, but saving JSON outputs failed: {e}")

            if summary_rows:
                st.markdown("**KR-20 summary by case (intersection items only):**")
                st.dataframe(summary_rows, width="stretch", hide_index=True)

            if out_rows:
                st.markdown("**Item-level psychometrics (filtered attempts):**")
                # Sort by most informative first: low p or high discrimination etc.
                out_rows_sorted = sorted(out_rows, key=lambda r: (r.get("case_id",""), -(r.get("n") or 0), -(r.get("discrimination_r") or -999)))
                st.dataframe(out_rows_sorted, width="stretch", hide_index=True)

                # Download CSV
                try:
                    csv_bytes = _to_csv_bytes(out_rows_sorted, list(out_rows_sorted[0].keys()))
                    st.download_button("⬇️ Download psychometrics CSV", data=csv_bytes, file_name="psychometrics_nclex.csv", mime="text/csv")

                except Exception as e:
                    st.error(f"Could not prepare CSV: {e}")
            else:
                st.info("No NCLEX item details found in the filtered attempts (or not enough attempts per item).")
    st.subheader("NCLEX Psychometrics Excel")
    cxl1, cxl2 = st.columns(2)
    with cxl1:
        min_attempts_xlsx = int(st.number_input("Min attempts per item (Excel)", min_value=1, max_value=500, value=3, step=1, key="min_attempts_xlsx_global"))
    with cxl2:
        min_items_xlsx = int(st.number_input("Min common items for KR-20 (Excel)", min_value=2, max_value=500, value=5, step=1, key="min_items_xlsx_global"))

    try:
        xlsx_bytes = build_nclex_psychometrics_excel_bytes(
            min_attempts_per_item=min_attempts_xlsx,
            min_items_intersection=min_items_xlsx,
        )
        st.download_button(
            "⬇️ Download NCLEX Psychometrics Excel (KR-20 + ItemStats)",
            data=xlsx_bytes,
            file_name="NCLEX_Psychometrics_Report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="download_nclex_psychometrics_xlsx_global",
        )
    except Exception as e:
        st.error(f"Could not prepare NCLEX psychometrics Excel: {e}")

    st.subheader("Exports")
    export_rows = []
    for a in filtered:
        cid = _attempt_case_id(a)
        sys = _attempt_system(a)
        stu = _attempt_student(a)
        ts = a.get("timestamp","")
        total20 = _attempt_total(a)
        intake = _attempt_intake_score(a)
        total25 = _attempt_total_with_intake(a)
        pts, mx = _attempt_nclex_summary(a)
        export_rows.append({
            "timestamp": ts,
            "student": stu,
            "case_id": cid,
            "system": sys,
            "mode": a.get("mode",""),
            "total_20": total20,
            "intake_5": intake,
            "total_25": total25,
            "nclex_points": pts,
            "nclex_max": mx,
            "unsafe_total": a.get("unsafe_total", 0),
            "duration_seconds": a.get("duration_seconds", ""),
        })
    if export_rows:
        csv_bytes = _to_csv_bytes(export_rows, list(export_rows[0].keys()))
        st.download_button("⬇️ Download filtered attempts CSV", data=csv_bytes, file_name="filtered_attempts.csv", mime="text/csv")
    else:
        st.warning("Nothing to export (no attempts match filters).")

def admin_page_attempt_search():
    st.header("🔎 Attempt Search (Admin)")
    attempts = _load_attempts_records(ATTEMPTS_PATH)

    cases_list_all = get_cases_list()
    all_cases = sorted({str(c.get("id","")).strip() for c in cases_list_all if str(c.get("id","")).strip()}) or ["—"]
    all_systems = sorted({safe_get_system(c) for c in cases_list_all}) or ["—"]
    # Include unknown values observed in attempts (legacy records)
    all_cases = sorted(set(all_cases) | {(_attempt_case_id(a) or "—") for a in attempts})
    all_systems = sorted(set(all_systems) | {(_attempt_system(a) or "—") for a in attempts})

    c1, c2, c3 = st.columns([2,2,2])
    with c1:
        case_sel = st.selectbox("Case", ["All"] + all_cases)
    with c2:
        sys_sel = st.selectbox("System", ["All"] + all_systems)
    with c3:
        student_q = st.text_input("Student contains", value="")

    c4, c5 = st.columns(2)
    with c4:
        score_min = st.number_input("Score min (/25)", min_value=0.0, max_value=25.0, value=0.0, step=0.5)
    with c5:
        score_max = st.number_input("Score max (/25)", min_value=0.0, max_value=25.0, value=25.0, step=0.5)

    filtered = _filter_attempts(
        attempts,
        systems_sel=[sys_sel] if sys_sel != "All" else ["All"],
        cases_sel=[case_sel] if case_sel != "All" else ["All"],
        student_q=student_q,
        score_min=score_min,
        score_max=score_max,
    )

    # Build table
    rows = []
    for i, a in enumerate(filtered):
        rows.append({
            "_idx": i,
            "timestamp": a.get("timestamp",""),
            "student": _attempt_student(a) or "—",
            "case_id": _attempt_case_id(a) or "—",
            "system": _attempt_system(a) or "—",
            "total_25": _attempt_total_with_intake(a),
            "mode": a.get("mode",""),
        })
    st.dataframe(rows, width="stretch", hide_index=True)

    pick = st.number_input("Open attempt row # (_idx)", min_value=0, max_value=max(0, len(rows)-1), value=0, step=1) if rows else None
    if rows and pick is not None:
        a = filtered[int(pick)]
        st.subheader("Attempt details")

        details = {
            "Timestamp": a.get("timestamp",""),
            "Student username": a.get("student_username",""),
            "Student display name": a.get("student_display_name",""),
            "Student ID": a.get("student_id",""),
            "Cohort": a.get("student_cohort",""),
            "Case ID": a.get("caseId",""),
            "Case title": a.get("caseTitle",""),
            "System": a.get("system",""),
            "Setting": a.get("setting",""),
            "Mode": a.get("mode",""),
            "Total (/20)": _attempt_total(a),
            "Intake (/5)": _attempt_intake_score(a),
            "Total (/25)": _attempt_total_with_intake(a),
            "Unsafe total": a.get("unsafe_total", 0),
            "Duration (seconds)": a.get("duration_seconds",""),
        }
        st.markdown("\n".join([f"- **{k}:** {v}" for k, v in details.items()]))

        intake = a.get("intake")
        if isinstance(intake, dict):
            st.markdown("**Student intake:**")
            for k, v in intake.items():
                st.markdown(f"- **{k}:** {v}")

        scores = a.get("scores")
        if isinstance(scores, dict):
            st.markdown("**Domain scores (A–E):**")
            for dom in ["A","B","C","D","E"]:
                if dom in scores:
                    st.markdown(f"- **{dom}:** {scores.get(dom)}")
            for k, v in scores.items():
                if k not in {"A","B","C","D","E"}:
                    st.markdown(f"- **{k}:** {v}")

        det = _attempt_nclex_details(a)
        if det:
            st.markdown("**NCLEX (per-item correctness):**")
            mini = []
            for d in det:
                qid = d.get("qid")
                it = NCLEX_BY_QID.get(qid, {})
                mini.append({
                    "qid": qid,
                    "type": d.get("type",""),
                    "difficulty": it.get("difficulty",""),
                    "correct": d.get("correct"),
                    "points": d.get("points"),
                    "max": d.get("max"),
                    "stem": (it.get("stem","")[:90] + "…") if it.get("stem") and len(it.get("stem",""))>90 else it.get("stem",""),
                })
            st.dataframe(mini, width="stretch", hide_index=True)

def admin_page_item_analytics():
    st.header("📈 Item Analytics (Admin)")
    st.caption("Ranks: most missed domains, unsafe patterns, most-wrong NCLEX items (only where item details exist).")

    attempts = _load_attempts_records(ATTEMPTS_PATH)
    all_systems = sorted({(_attempt_system(a) or "—") for a in attempts})
    all_cases = sorted({(_attempt_case_id(a) or "—") for a in attempts})

    c1, c2 = st.columns(2)
    with c1:
        systems_sel = st.multiselect("Systems", ["All"] + all_systems, default=["All"])
    with c2:
        cases_sel = st.multiselect("Cases", ["All"] + all_cases, default=["All"])

    filtered = _filter_attempts(attempts, systems_sel=systems_sel, cases_sel=cases_sel)

    # A) Most missed domains
    st.subheader("Most missed rubric domains (A–E)")
    doms = ["A","B","C","D","E"]
    dom_stats = []
    for d in doms:
        vals = []
        for a in filtered:
            sc = _attempt_domain_scores(a)
            v = sc.get(d)
            if isinstance(v,(int,float)):
                vals.append(float(v))
        if vals:
            avg = sum(vals)/len(vals)
            dom_stats.append({"domain": d, "avg(/4)": round(avg,2), "missed(/4)": round(4-avg,2), "n": len(vals)})
    dom_stats.sort(key=lambda r: (-r["missed(/4)"], r["domain"]))
    st.dataframe(dom_stats if dom_stats else [{"note":"No domain score data found in filtered attempts."}], width="stretch", hide_index=True)

    # B) Unsafe patterns
    st.subheader("Unsafe patterns (by domain)")
    unsafe_totals = {}
    for a in filtered:
        uc = _attempt_unsafe_counts(a)
        for k, v in (uc or {}).items():
            try:
                unsafe_totals[k] = unsafe_totals.get(k, 0) + int(v or 0)
            except Exception:
                continue
    unsafe_rows = [{"domain": k, "unsafe_count": v} for k,v in sorted(unsafe_totals.items(), key=lambda kv: (-kv[1], kv[0]))]
    st.dataframe(unsafe_rows if unsafe_rows else [{"note":"No unsafe_counts found in filtered attempts."}], width="stretch", hide_index=True)

    # C) NCLEX most wrong
    st.subheader("Most frequently wrong NCLEX items")
    wrong = {}
    seen = {}
    for a in filtered:
        det = _attempt_nclex_details(a)
        for d in det:
            qid = d.get("qid")
            if not qid:
                continue
            seen[qid] = seen.get(qid, 0) + 1
            if not bool(d.get("correct", False)):
                wrong[qid] = wrong.get(qid, 0) + 1

    rows = []
    for qid, n_seen in seen.items():
        n_wrong = wrong.get(qid, 0)
        it = NCLEX_BY_QID.get(qid, {})
        qnum_str = nclex_qnum_from_id(qid)
        try:
            qnum = int(qnum_str)
        except Exception:
            qnum = 9999
        rows.append({
            "q#": f"Q{qnum}" if qnum != 9999 else qnum_str,
            "qid": qid,
            "wrong_%": round((n_wrong/n_seen)*100, 1) if n_seen else 0.0,
            "n_seen": n_seen,
            "difficulty": it.get("difficulty",""),
            "type": it.get("type",""),
            "stem": (it.get("stem","")[:110] + "…") if it.get("stem") and len(it.get("stem",""))>110 else it.get("stem",""),
        })
    # Show in question order (Q1..Qn)
    rows.sort(key=lambda r: (int(str(r.get("q#","Q9999")).lstrip("Q").strip() or 9999), -r["wrong_%"], -r["n_seen"]))
    st.dataframe(rows[:50] if rows else [{"note":"No NCLEX per-item details found (older attempts may not include nclex.details)."}], width="stretch", hide_index=True)

    if rows:
        st.download_button("⬇️ Download NCLEX item analytics CSV", data=_to_csv_bytes(rows, list(rows[0].keys())), file_name="nclex_item_analytics.csv", mime="text/csv")

def admin_page_data_tools():
    st.header("🧹 Data Tools (Admin)")
    st.caption("Backup + safe delete (with preview, typed confirmation, audit log).")

    attempts = _load_attempts_records(ATTEMPTS_PATH)
    st.write(f"Attempts file: {ATTEMPTS_PATH.name} • Records: {len(attempts)}")

    # Backup download
    st.subheader("Backup")
    raw = "\n".join(json.dumps(a, ensure_ascii=False) for a in attempts).encode("utf-8")
    st.download_button("⬇️ Download backup (attempts_log.jsonl)", data=raw, file_name=f"attempts_backup_{utc_now_iso().replace(':','-')}.jsonl", mime="application/jsonl")

    # Delete filters
    st.subheader("Safe delete / purge")
    # Build dropdown choices.
    # NOTE: Data Tools is attempt-centric (for filtering/deleting records), but we also show ALL cases/systems
    # from cases.json so admins can pick any case/system even if there are currently zero attempts.
    all_cases_attempts = sorted({(_attempt_case_id(a) or "—") for a in attempts})
    all_systems_attempts = sorted({(_attempt_system(a) or "—") for a in attempts})

    cases_master = get_cases_list()
    master_case_ids = sorted({str(c.get("id", "")).strip() for c in cases_master if str(c.get("id", "")).strip()})
    master_systems = sorted({safe_get_system(c) for c in cases_master if safe_get_system(c)})

    all_cases = sorted(set(master_case_ids) | set(all_cases_attempts))
    all_systems = sorted(set(master_systems) | set(all_systems_attempts))

    c1, c2, c3 = st.columns([2,2,2])
    with c1:
        systems_sel = st.multiselect("Systems to delete", ["All"] + all_systems, default=["All"])
    with c2:
        cases_sel = st.multiselect("Cases to delete", ["All"] + all_cases, default=["All"])
    with c3:
        student_q = st.text_input("Student contains (optional)", value="")

    c4, c5 = st.columns(2)
    with c4:
        start_date = st.date_input("Start date (optional)", value=None, key="del_start_date")
    with c5:
        end_date = st.date_input("End date (optional)", value=None, key="del_end_date")

    from datetime import datetime, time
    dt_start = datetime.combine(start_date, time.min) if start_date else None
    dt_end = datetime.combine(end_date, time.max) if end_date else None

    to_delete = _filter_attempts(attempts, systems_sel=systems_sel, cases_sel=cases_sel, student_q=student_q, dt_start=dt_start, dt_end=dt_end)
    # Build keep list safely (no dict 'in' comparisons)
    to_delete_set = set()
    for i, a in enumerate(attempts):
        # mark deletions by index
        if a in to_delete:
            to_delete_set.add(i)
    keep = [a for i, a in enumerate(attempts) if i not in to_delete_set]


    st.warning(f"Preview: {len(to_delete)} records will be deleted. {len(keep)} will remain.")

    confirm = st.text_input('Type DELETE to confirm', value="", key="delete_confirm_text")
    if st.button("🗑️ Permanently delete selected records", disabled=(confirm.strip() != "DELETE")):
        # Write audit log first
        audit_path = BASE_DIR / "audit_log.jsonl"
        audit_event = {
            "timestamp": utc_now_iso(),
            "actor": st.session_state.get("student_profile", {}).get("username","") if st.session_state.get("student_profile") else "admin",
            "action": "delete_attempts",
            "filter": {
                "systems": systems_sel,
                "cases": cases_sel,
                "student_contains": student_q,
                "start_date": str(start_date) if start_date else "",
                "end_date": str(end_date) if end_date else "",
            },
            "deleted": len(to_delete),
            "remaining": len(keep),
        }
        try:
            with audit_path.open("a", encoding="utf-8") as f:
                f.write(json.dumps(audit_event, ensure_ascii=False) + "\n")
        except Exception as e:
            st.error(f"Could not write audit log: {e}")
            st.stop()

        # Backup file on disk (copy; do NOT move)
        import shutil
        backup_path = BASE_DIR / f"attempts_backup_before_delete_{utc_now_iso().replace(':','-')}.jsonl"
        try:
            shutil.copy2(ATTEMPTS_PATH, backup_path)
        except Exception as e:
            st.error(f"Could not create on-disk backup: {e}")
            st.stop()

        # Write new attempts file
        try:
            with ATTEMPTS_PATH.open("w", encoding="utf-8") as f:
                for a in keep:
                    f.write(json.dumps(a, ensure_ascii=False) + "\n")
            st.success(f"Deleted {len(to_delete)} records. Backup saved as {backup_path.name}.")
            st.rerun()
        except Exception as e:
            st.error(f"Failed to write new attempts_log.jsonl: {e}")
            st.stop()



def run_validator_ui():
    """Main-page NCLEX validator UI (mirrors sidebar validator)."""
    nclex = load_nclex_items()
    nclex_case_keys = sorted(list((nclex.get("cases") or {}).keys()))
    case_ids_all = sorted([str(c.get("id", "")).strip() for c in cases if str(c.get("id", "")).strip()])

    st.write("Cases in cases.json:", len(case_ids_all))
    st.write("Cases in nclex_items.json:", len(nclex_case_keys))

    missing = [cid for cid in case_ids_all if cid not in (nclex.get("cases") or {})]
    extra = [cid for cid in nclex_case_keys if cid not in set(case_ids_all)]

    if missing:
        st.error(f"Missing NCLEX packs for {len(missing)} case IDs")
        st.caption("Example missing IDs:")
        st.code("\n".join(missing[:10]))
    else:
        st.success("No missing NCLEX case packs ✅")

    if extra:
        st.warning(f"Extra NCLEX packs not found in cases.json: {len(extra)}")
        st.caption("Example extra IDs:")
        st.code("\n".join(extra[:10]))

    if st.button("🛠 Auto-create missing NCLEX packs (empty)", key="nclex_autofix_btn_main"):
        nclex.setdefault("cases", {})
        for cid in missing:
            nclex["cases"][cid] = {"items": []}
        save_nclex_items(nclex)
        st.success("Created empty packs for missing cases. Now paste/generate items for those cases.")
        st.rerun()


def render_nclex_rotation_admin_ui(pol: dict):
    """Simplified main-page rotation generator + history (keeps core behavior intact)."""
    nclex = load_nclex_items()
    cases_list = get_cases_list()
    case_ids_all = sorted([str(c.get("id", "")).strip() for c in cases_list if str(c.get("id", "")).strip()])

    cid_pick = st.selectbox("Select case to manage", options=[""] + case_ids_all, index=0, key="rot_case_pick_main")
    if not cid_pick:
        st.info("Pick a case to view bank size, current active set, generate a new active set, or download history.")
        return

    pack = (nclex.get("cases") or {}).get(cid_pick) or {}
    bank_items = list(pack.get("items", []) or [])
    enabled_types = pol.get("enabled_types") or {}
    if isinstance(enabled_types, dict) and enabled_types:
        bank_items = [it for it in bank_items if enabled_types.get(it.get("type"), True)]

    k = nclex_items_per_case(pol, cid_pick)
    st.write(f"Bank size for {cid_pick}: **{len(bank_items)}** items (after type filters). Target active set size: **{k}**")

    active_sets = load_nclex_active_sets()
    by_case = active_sets.get("by_case") or {}
    active = by_case.get(cid_pick) or {}
    active_qids = active.get("qids", []) if isinstance(active, dict) else []
    if active_qids:
        st.success(f"Active set: {len(active_qids)} items (generated_at: {active.get('generated_at','')})")
    else:
        st.warning("No active set generated yet for this case (students will use default selection).")

    # Preview active
    if active_qids:
        with st.expander("👀 Preview active set", expanded=False):
            by_id = {str(it.get("id","")): it for it in bank_items if str(it.get("id","")).strip()}
            items = [by_id[q] for q in active_qids if q in by_id][:k]
            for i, it in enumerate(items, start=1):
                st.markdown(f"**{i}. {it.get('id','')}** ({it.get('type','')}, {it.get('difficulty','')})")
                st.write(it.get('stem') or it.get('prompt') or '')
                st.divider()

    # Generate new active set
    if st.button("🎲 Generate NEW active set now (random)", key="gen_active_set_btn_main"):
        qids_all = [str(it.get("id","")).strip() for it in bank_items if str(it.get("id","")).strip()]
        if not qids_all:
            st.error("No NCLEX items available for this case (after filters).")
            return
        rnd_seed = sha256_hex(f"active|{cid_pick}|{time.time()}|{secrets.token_hex(6)}")
        rnd = random.Random(int(rnd_seed[:8], 16))
        rnd.shuffle(qids_all)
        new_qids = qids_all[:k]

        now_ts = utc_now_iso()
        by_case.setdefault(cid_pick, {})
        by_case[cid_pick]["qids"] = list(new_qids)
        by_case[cid_pick]["generated_at"] = now_ts
        by_case[cid_pick]["seed"] = rnd_seed

        hist = by_case[cid_pick].setdefault("history", [])
        if isinstance(hist, list):
            hist.append({"generated_at": now_ts, "seed": rnd_seed, "count": len(new_qids), "qids": list(new_qids)})

        active_sets["by_case"] = by_case
        save_nclex_active_sets(active_sets)
        st.success("Generated and saved a new active set.")
        st.rerun()

    # Download history
    hist = (by_case.get(cid_pick) or {}).get("history") or []
    if isinstance(hist, list) and hist:
        with st.expander("📜 Rotation history", expanded=False):
            st.write(f"History entries: {len(hist)}")
            # show last 5
            for h in hist[-5:][::-1]:
                st.write(h.get("generated_at",""), "|", h.get("count",""), "items")
            # download CSV
            import csv as _csv
            import io as _io
            out = _io.StringIO()
            w = _csv.DictWriter(out, fieldnames=["generated_at","count","seed","qids_json"])
            w.writeheader()
            for h in hist:
                w.writerow({
                    "generated_at": h.get("generated_at",""),
                    "count": h.get("count",""),
                    "seed": h.get("seed",""),
                    "qids_json": json.dumps(h.get("qids", []), ensure_ascii=False),
                })
            st.download_button(
                "⬇️ Download rotation history (CSV)",
                data=out.getvalue().encode("utf-8"),
                file_name=f"rotation_history_{cid_pick}.csv",
                mime="text/csv",
                key="dl_rot_hist_main"
            )
    else:
        st.caption("No rotation history yet for this case.")

def admin_page_settings():
    """Main-page settings hub for admins (keeps sidebar minimal)."""
    st.subheader("⚙️ Settings")
    st.caption("All admin configuration panels are here to keep the sidebar tidy. Changes save to the same JSON policies used by the original app.")

    # --- Case switching / Exam lock ---
    with st.expander("🧭 Case Switching (Exam/Practice)", expanded=False):
        st.caption("Controls whether students can switch cases freely while in Exam mode.")
        allow_switch_exam = st.toggle(
            "Allow case switching in Exam mode",
            value=not bool(features.get("lock_case_switch_exam", False)),
            key="allow_case_switch_exam_main",
            help="If ON: students can change cases in Exam mode. If OFF: case switching is locked once a case is started."
        )
        # Internally we store the lock flag (True = locked)
        features["lock_case_switch_exam"] = (not bool(allow_switch_exam))
        if st.button("💾 Save case switching settings", key="save_case_switch_settings_main"):
            save_features(features)
            flash_success("Saved case switching settings.")
            st.rerun()


    # --- Student Accounts ---
    with st.expander("👥 Student Accounts (Generate list + passwords)", expanded=False):
        st.caption("Create student accounts in students.json and download credentials (CSV / Excel / Word).")
        data_students = load_students()

        col1, col2, col3 = st.columns(3)
        with col1:
            prefix = st.text_input("Username prefix", value="student", key="gen_students_prefix_main").strip() or "student"
        with col2:
            start_num = st.number_input("Starting number", min_value=1, max_value=999999, value=1, step=1, key="gen_students_start_main")
        with col3:
            pad3 = st.toggle("Pad numbers (001, 002…)", value=True, key="gen_students_pad3_main")

        cohort = st.text_input("Cohort label (optional)", value="Adult Nursing Y2", key="gen_students_cohort_main")

        st.markdown("**Option A — Generate by count**")
        n_new = st.number_input("How many accounts to create?", min_value=1, max_value=500, value=30, step=1, key="gen_students_n_main")
        pw_len = st.slider("Password length", min_value=6, max_value=18, value=10, step=1, key="gen_students_pwlen_main")

        st.markdown("**Option B — Paste student names (one per line)**")
        names_text = st.text_area("Student names (optional)", value="", height=120, key="gen_students_names_main",
                                  help="If you paste names here, the app will create one account per name and put the name in the export sheets.")

        def _make_username(i: int) -> str:
            n = int(start_num) + i
            return f"{prefix}{n:03d}" if pad3 else f"{prefix}{n}"

        # =============================
        # NEW: Import students from Excel (.xlsx) and generate passwords
        # =============================
        st.divider()
        st.markdown("### 📥 Import student list from Excel (.xlsx)")
        st.caption("Upload an Excel file, choose the columns, and the app will create accounts + generate passwords, then let you download the credentials.")

        xlsx_up = st.file_uploader("Upload Excel file (.xlsx)", type=["xlsx"], key="students_xlsx_upload_main")
        if xlsx_up is not None:
            try:
                # Read with openpyxl (keeps dependencies light; already used elsewhere)
                from openpyxl import load_workbook
                wb = load_workbook(filename=io.BytesIO(xlsx_up.getvalue()), data_only=True)
                sheet_names = wb.sheetnames
                sheet_name = st.selectbox("Sheet", sheet_names, index=0, key="students_xlsx_sheet_main")
                ws = wb[sheet_name]

                # Extract rows (first non-empty row is treated as header)
                rows = list(ws.iter_rows(values_only=True))
                # Find header row
                header_idx = None
                for i, r in enumerate(rows[:20]):  # look at first 20 rows
                    if r and any((str(c).strip() not in ("", "None", "nan")) for c in r):
                        header_idx = i
                        break
                if header_idx is None:
                    st.warning("Excel sheet appears empty.")
                else:
                    header = [str(c).strip() if c is not None else "" for c in rows[header_idx]]
                    # If header is mostly empty, auto-label columns
                    if sum(1 for h in header if h) < max(1, len(header)//3):
                        header = [f"Column {i+1}" for i in range(len(rows[header_idx]))]

                    data_rows = rows[header_idx+1:]

                    # Build a small preview table (first 10 rows)
                    preview = []
                    for r in data_rows[:10]:
                        if r is None:
                            continue
                        preview.append({header[i]: ("" if i >= len(r) or r[i] is None else str(r[i])) for i in range(len(header))})
                    if preview:
                        st.write("Preview (first 10 rows):")
                        st.dataframe(preview, use_container_width=True)
                    else:
                        st.info("No data rows found under the header.")

                    cols = [h for h in header if h]
                    name_col = st.selectbox("Student name column", cols, index=0, key="students_xlsx_namecol_main")
                    username_col = st.selectbox("Username column (optional)", ["(none)"] + cols, index=0, key="students_xlsx_usercol_main")
                    id_col = st.selectbox("Student ID column (optional)", ["(none)"] + cols, index=0, key="students_xlsx_idcol_main")
                    cohort_col = st.selectbox("Cohort column (optional)", ["(none)"] + cols, index=0, key="students_xlsx_cohortcol_main")

                    pw_len_x = st.slider("Password length (Excel import)", min_value=6, max_value=18, value=int(pw_len), step=1, key="students_xlsx_pwlen_main")

                    st.warning("This will create NEW accounts only. If a username already exists, that row will be skipped.")
                    if st.button("➕ Generate accounts from Excel", key="students_xlsx_genbtn_main"):
                        existing = {str(s.get("username","")).strip() for s in data_students.get("students", []) if str(s.get("username","")).strip()}
                        created_rows_x = []

                        # helper: safe get by column name
                        col_index = {h: i for i, h in enumerate(header)}
                        def _cell(row, colname):
                            i = col_index.get(colname, None)
                            if i is None or i >= len(row):
                                return ""
                            v = row[i]
                            if v is None:
                                return ""
                            return str(v).strip()

                        # Build list of entries
                        total_added = 0
                        for i, r in enumerate(data_rows):
                            if not r or all(c is None or str(c).strip()=="" for c in r):
                                continue
                            display_name_x = _cell(r, name_col)
                            if not display_name_x or display_name_x.lower() in ("nan", "none"):
                                continue

                            uname_x = _cell(r, username_col) if username_col != "(none)" else ""
                            sid_x = _cell(r, id_col) if id_col != "(none)" else ""
                            cohort_x = _cell(r, cohort_col) if cohort_col != "(none)" else ""

                            # Auto username if not provided
                            if not uname_x:
                                uname_x = _make_username(total_added)

                            # Skip if exists
                            if uname_x in existing:
                                continue

                            pw_x = generate_password(int(pw_len_x))
                            student_id_x = sid_x or (f"{int(start_num)+total_added:03d}" if pad3 else str(int(start_num)+total_added))

                            data_students.setdefault("students", []).append({
                                "username": uname_x,
                                "display_name": display_name_x,
                                "student_id": student_id_x,
                                "cohort": cohort_x or cohort,
                                "password_sha256": sha256_hex(pw_x),
                                "created_at": utc_now_iso(),
                            })
                            existing.add(uname_x)
                            created_rows_x.append({
                                "display_name": display_name_x,
                                "username": uname_x,
                                "password": pw_x,
                                "student_id": student_id_x,
                                "cohort": cohort_x or cohort,
                            })
                            total_added += 1

                        save_students(data_students)

                        if not created_rows_x:
                            st.warning("No new accounts created (all usernames already existed or file had no valid rows).")
                        else:
                            # Reuse same export builders
                            out_x = io.StringIO()
                            writer_x = csv.DictWriter(out_x, fieldnames=["display_name","username","password","student_id","cohort"])
                            writer_x.writeheader()
                            for r in created_rows_x:
                                writer_x.writerow(r)

                            xlsx_bytes_x = build_credentials_xlsx_bytes(created_rows_x)
                            docx_bytes_x = build_credentials_docx_bytes(created_rows_x)

                            st.success(f"Imported and created {len(created_rows_x)} accounts from Excel.")
                            # Store last generated credentials for downloads
                            st.session_state["last_generated_credentials_rows"] = created_rows_x


                            d1, d2, d3 = st.columns(3)
                            with d1:
                                st.download_button(
                                    "⬇️ Download CSV (Excel import)",
                                    data=out_x.getvalue().encode("utf-8"),
                                    file_name=f"student_credentials_import_{utc_now_iso().replace(':','-')}.csv",
                                    mime="text/csv",
                                    key="dl_students_csv_import_main"
                                )
                            with d2:
                                st.download_button(
                                    "⬇️ Download Excel (.xlsx) (Excel import)",
                                    data=xlsx_bytes_x,
                                    file_name=f"student_credentials_import_{utc_now_iso().replace(':','-')}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    key="dl_students_xlsx_import_main"
                                )
                            with d3:
                                st.download_button(
                                    "⬇️ Download Word (.docx) (Excel import)",
                                    data=docx_bytes_x,
                                    file_name=f"student_credentials_import_{utc_now_iso().replace(':','-')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="dl_students_docx_import_main"
                                )
                        st.rerun()

            except Exception as e:
                st.error(f"Failed to read Excel: {e}")


        if st.button("➕ Generate accounts now", key="gen_students_btn_main"):
            existing = {str(s.get("username","")).strip() for s in data_students.get("students", []) if str(s.get("username","")).strip()}
            created_rows = []

            names = [ln.strip() for ln in (names_text or "").splitlines() if ln.strip()]
            total = len(names) if names else int(n_new)

            for i in range(total):
                uname = _make_username(i)
                if uname in existing:
                    continue
                pw = generate_password(int(pw_len))
                display_name = names[i] if names else f"Student {int(start_num)+i:03d}" if pad3 else f"Student {int(start_num)+i}"
                student_id = f"{int(start_num)+i:03d}" if pad3 else str(int(start_num)+i)

                data_students.setdefault("students", []).append({
                    "username": uname,
                    "display_name": display_name,
                    "student_id": student_id,
                    "cohort": cohort,
                    "password_sha256": sha256_hex(pw),
                    "created_at": utc_now_iso(),
                })
                existing.add(uname)
                created_rows.append({
                    "display_name": display_name,
                    "username": uname,
                    "password": pw,
                    "student_id": student_id,
                    "cohort": cohort,
                })

            save_students(data_students)

            if not created_rows:
                st.warning("No new accounts created (usernames already exist).")
            else:
                # CSV
                out = io.StringIO()
                writer = csv.DictWriter(out, fieldnames=["display_name","username","password","student_id","cohort"])
                writer.writeheader()
                for r in created_rows:
                    writer.writerow(r)

                # Excel / Word
                xlsx_bytes = build_credentials_xlsx_bytes(created_rows, title=f"Student Credentials — {cohort}".strip(" —"))
                docx_bytes = build_credentials_docx_bytes(created_rows, title=f"Student Credentials — {cohort}".strip(" —"))

                cA, cB, cC = st.columns(3)
                with cA:
                    st.download_button(
                        "⬇️ Download CSV",
                        data=out.getvalue().encode("utf-8"),
                        file_name=f"student_credentials_{utc_now_iso().replace(':','-')}.csv",
                        mime="text/csv",
                        key="dl_students_csv_main"
                    )
                with cB:
                    st.download_button(
                        "⬇️ Download Excel (.xlsx)",
                        data=xlsx_bytes,
                        file_name=f"student_credentials_{utc_now_iso().replace(':','-')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="dl_students_xlsx_main"
                    )
                with cC:
                    st.download_button(
                        "⬇️ Download Word (.docx)",
                        data=docx_bytes,
                        file_name=f"student_credentials_{utc_now_iso().replace(':','-')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_students_docx_main"
                    )

                st.success(f"Saved {len(created_rows)} accounts into students.json")
                # Store last generated credentials for downloads
                st.session_state["last_generated_credentials_rows"] = created_rows


                # =============================
                # NEW: Delete selected / all student accounts
                # =============================
                st.divider()
                st.markdown("### 🗑️ Delete student accounts")
                students_list = data_students.get("students", []) if isinstance(data_students, dict) else []
                if not students_list:
                    st.info("No student accounts found in students.json.")
                else:
                    def _label_student(s: dict) -> str:
                        dn = str(s.get("display_name","")).strip()
                        un = str(s.get("username","")).strip()
                        sid = str(s.get("student_id","")).strip()
                        if dn and un:
                            return f"{dn} ({un})"
                        if un:
                            return un
                        if dn:
                            return dn
                        return sid or "unknown"

                    labels = [_label_student(s) for s in students_list]
                    selected_labels = st.multiselect("Select accounts to delete", options=labels, key="students_delete_select_main")

                    cdel1, cdel2 = st.columns(2)
                    with cdel1:
                        if st.button("Delete selected accounts", disabled=(len(selected_labels)==0), key="students_delete_selected_btn_main"):
                            to_delete = set(selected_labels)
                            kept = []
                            deleted_n = 0
                            for s in students_list:
                                if _label_student(s) in to_delete:
                                    deleted_n += 1
                                    continue
                                kept.append(s)
                            data_students["students"] = kept
                            save_students(data_students)
                            st.success(f"Deleted {deleted_n} account(s).")
                            st.rerun()

                    with cdel2:
                        confirm_all = st.text_input('Type "DELETE ALL" to remove ALL student accounts', key="students_delete_all_confirm_main")
                        if st.button("Delete ALL accounts", disabled=(confirm_all.strip() != "DELETE ALL"), key="students_delete_all_btn_main"):
                            data_students["students"] = []
                            save_students(data_students)
                            st.success("All student accounts deleted.")
                            st.rerun()

    # --- Password Management ---

        # =============================
        # ALWAYS AVAILABLE: Download last generated credentials + delete accounts
        # =============================
        st.divider()
        st.markdown("### ⬇️ Download newly generated credentials")
        st.caption("Downloads are available only for the most recently generated batch in this session (because existing accounts store only hashed passwords).")

        last_rows = st.session_state.get("last_generated_credentials_rows")
        if isinstance(last_rows, list) and last_rows:
            # Build CSV text
            try:
                out_csv = io.StringIO()
                fieldnames = ["display_name", "username", "password", "student_id", "cohort"]
                writer = csv.DictWriter(out_csv, fieldnames=fieldnames)
                writer.writeheader()
                for r in last_rows:
                    if isinstance(r, dict):
                        writer.writerow({k: r.get(k, "") for k in fieldnames})
                csv_bytes = out_csv.getvalue().encode("utf-8")
            except Exception:
                csv_bytes = b""

            xlsx_bytes = build_credentials_xlsx_bytes(last_rows, title="Student Credentials")
            docx_bytes = build_credentials_docx_bytes(last_rows, title="Student Credentials")

            d1, d2, d3 = st.columns(3)
            with d1:
                st.download_button(
                    "⬇️ Download CSV",
                    data=csv_bytes,
                    file_name=f"student_credentials_{utc_now_iso().replace(':','-')}.csv",
                    mime="text/csv",
                    key="dl_students_csv_last_main",
                )
            with d2:
                st.download_button(
                    "⬇️ Download Excel (.xlsx)",
                    data=xlsx_bytes,
                    file_name=f"student_credentials_{utc_now_iso().replace(':','-')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_students_xlsx_last_main",
                )
            with d3:
                st.download_button(
                    "⬇️ Download Word (.docx)",
                    data=docx_bytes,
                    file_name=f"student_credentials_{utc_now_iso().replace(':','-')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="dl_students_docx_last_main",
                )
        else:
            st.info("Generate accounts (count / paste names / import Excel) to enable downloads here.")

        st.divider()
        st.markdown("### 🗑️ Delete student accounts")
        current_students = load_students()
        students_list2 = current_students.get("students", []) if isinstance(current_students, dict) else []
        if not students_list2:
            st.info("No student accounts found in students.json.")
        else:
            def _label_student2(s: dict) -> str:
                dn = str(s.get("display_name","") or "").strip()
                un = str(s.get("username","") or "").strip()
                sid = str(s.get("student_id","") or "").strip()
                if dn and un:
                    return f"{dn} ({un})"
                if un:
                    return un
                if dn:
                    return dn
                return sid or "unknown"

            labels2 = [_label_student2(s) for s in students_list2]
            selected_labels2 = st.multiselect("Select accounts to delete", options=labels2, key="students_delete_select_main_v5")

            cdel1, cdel2 = st.columns(2)
            with cdel1:
                if st.button("🗑️ Delete selected", disabled=(len(selected_labels2)==0), key="students_delete_selected_btn_main_v5"):
                    to_delete = set(selected_labels2)
                    kept = []
                    deleted_n = 0
                    for s in students_list2:
                        if _label_student2(s) in to_delete:
                            deleted_n += 1
                            continue
                        kept.append(s)
                    save_students({"students": kept})
                    st.success(f"Deleted {deleted_n} account(s).")
                    st.rerun()

            with cdel2:
                confirm_all = st.text_input('Type "DELETE ALL" to delete ALL student accounts', key="students_delete_all_confirm_main_v5")
                if st.button("🗑️ Delete ALL", disabled=(confirm_all.strip() != "DELETE ALL"), key="students_delete_all_btn_main_v5"):
                    save_students({"students": []})
                    st.success("All student accounts deleted.")
                    st.rerun()


    with st.expander("🔑 Password Management (Reset / Exam / Expiry)", expanded=False):
        st.caption("Reset individual student passwords + enable exam password override with auto-expiry (Qatar time).")

        policy_exam = load_exam_access_policy()
        active_exam = is_exam_password_active(policy_exam)

        st.markdown("### 🧑‍🎓 Individual student reset")
        data_students = load_students()
        usernames = sorted([str(s.get("username", "")).strip() for s in data_students.get("students", []) if str(s.get("username", "")).strip()])
        if usernames:
            target_user = st.selectbox("Select student username", usernames, key="pw_reset_user_main")
            new_pw = st.text_input("New password for this student", type="password", key="pw_reset_newpw_main")
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Reset this student password", key="pw_reset_btn_main"):
                    if not new_pw.strip():
                        st.error("Enter a new password.")
                    else:
                        for s in data_students.get("students", []):
                            if str(s.get("username", "")).strip() == target_user:
                                s["password_sha256"] = sha256_hex(new_pw.strip())
                                break
                        save_students(data_students)
                        st.success(f"Password updated for {target_user}")
            with col2:
                if st.button("Generate random password", key="pw_reset_gen_btn_main"):
                    rp = generate_password(12)
                    st.session_state["pw_reset_newpw_main"] = rp
                    st.info(f"Generated: {rp}")
        else:
            st.warning("No students found in students.json")

        st.divider()
        st.markdown("### 📝 Exam password (everyone)")
        st.write("Status:", "✅ ACTIVE" if active_exam else "⛔ INACTIVE")

        policy_exam["force_exam_password_only"] = st.toggle(
            "Force exam password only (recommended for exams)",
            value=bool(policy_exam.get("force_exam_password_only", True)),
            key="force_exam_pw_only_main"
        )

        exam_pw_plain = st.text_input("Set exam password (e.g., NURS@2026)", type="password", key="exam_pw_plain_main")
        expires_enabled = st.toggle("Auto-expire exam password", value=True, key="exam_expire_toggle_main")

        default_exp = now_local().replace(hour=12, minute=0, second=0, microsecond=0)
        exp_date = st.date_input("Expiry date (Qatar time)", value=default_exp.date(), key="exam_expire_date_main")
        exp_time = st.time_input("Expiry time (Qatar time)", value=default_exp.time(), key="exam_expire_time_main")

        if st.button("✅ Enable exam password now", key="exam_enable_btn_main"):
            if not exam_pw_plain.strip():
                st.error("Enter an exam password.")
            else:
                policy_exam["enabled"] = True
                policy_exam["exam_password_sha256"] = sha256_hex(exam_pw_plain.strip())

                if expires_enabled:
                    exp_dt = datetime.combine(exp_date, exp_time).replace(tzinfo=TZ)
                    policy_exam["expires_at"] = exp_dt.isoformat()
                else:
                    policy_exam["expires_at"] = ""

                save_exam_access_policy(policy_exam)
                st.success("Exam password enabled.")
                st.rerun()

        colA, colB = st.columns(2)
        with colA:
            if st.button("⛔ Disable exam password", key="exam_disable_btn_main"):
                policy_exam["enabled"] = False
                save_exam_access_policy(policy_exam)
                st.success("Exam password disabled.")
                st.rerun()

        with colB:
            if st.button("📄 Show exam policy details", key="exam_show_policy_btn_main"):
                st.json(load_exam_access_policy())

    # --- Mode + AI ---
    with st.expander("🧪 Mode (Practice / Exam)", expanded=True):
        admin_settings = load_admin_settings()
        new_mode = st.selectbox(
            "Select mode (saved for everyone)",
            ["Practice", "Exam"],
            index=0 if admin_settings.get("app_mode") == "Practice" else 1,
            key="mode_select_main"
        )
        if new_mode != admin_settings.get("app_mode"):
            admin_settings["app_mode"] = new_mode
            save_admin_settings(admin_settings)
            st.success(f"Mode saved: {new_mode}")
            st.rerun()

    with st.expander("🧠 AI Coach (Practice only)", expanded=False):
        admin_settings = load_admin_settings()
        mode_now = admin_settings.get("app_mode", "Practice")
        ai_enabled = st.toggle(
            "Enable AI coaching after each submission (Practice only)",
            value=bool(admin_settings.get("ai_enabled", False)),
            disabled=(mode_now == "Exam"),
            key="ai_enabled_main"
        )
        ai_debrief_enabled = st.toggle(
            "Enable end-of-case AI debrief (Practice only)",
            value=bool(admin_settings.get("ai_debrief_enabled", False)),
            disabled=(mode_now == "Exam"),
            key="ai_debrief_enabled_main"
        )
        ai_model = st.text_input("Model", value=str(admin_settings.get("ai_model", "gpt-5.2")), key="ai_model_main")
        if st.button("💾 Save AI settings", key="save_ai_settings_main"):
            admin_settings["ai_enabled"] = bool(ai_enabled)
            admin_settings["ai_debrief_enabled"] = bool(ai_debrief_enabled)
            admin_settings["ai_model"] = str(ai_model).strip() or "gpt-5.2"
            save_admin_settings(admin_settings)
            st.success("Saved.")
            st.rerun()

    # --- KPIs ---
    with st.expander("📊 Research & Teaching KPIs", expanded=False):
        kpi_policy = load_kpi_policy()

        # Catalog of KPIs (includes original + expanded)
        KPI_CATALOG = [
            ("total_score", "Track total score (A–E)"),
            ("intake_score", "Track intake score"),
            ("domain_profile", "Track domain profile (A–E)"),
            ("domain_missed_profile", "Track missed-correct profile (A–E)"),
            ("unsafe_flags", "Track unsafe actions count"),
            ("unsafe_by_domain", "Track unsafe actions by domain"),
            ("time_to_completion", "Track time to completion"),
            ("time_per_section", "Track time per section (A–E / NCLEX)"),
            ("attempts_per_case", "Track attempts per case"),
            ("attempts_over_time", "Track attempts over time"),
            ("nclex_score", "Track NCLEX-style practice score"),
            ("nclex_by_type_accuracy", "Track NCLEX accuracy by item type"),
            ("nclex_difficulty_accuracy", "Track NCLEX accuracy by difficulty"),
            ("top_wrong_items", "Track most frequently wrong NCLEX items"),
            ("case_rank_summary", "Enable case ranking summary"),
            ("student_rank_summary", "Enable student ranking summary"),
            ("completion_rate", "Track completion rate"),
        ]

        st.caption("Enable/disable any KPI. These power research reports and faculty exports without affecting student workflow.")

        colA, colB = st.columns(2)
        with colA:
            if st.button("✅ Enable ALL", key="kpi_enable_all_btn_main"):
                for k, _ in KPI_CATALOG:
                    kpi_policy[k] = True
                save_kpi_policy(kpi_policy)
                st.success("All KPIs enabled.")
                st.rerun()
        with colB:
            if st.button("🚫 Disable ALL", key="kpi_disable_all_btn_main"):
                for k, _ in KPI_CATALOG:
                    kpi_policy[k] = False
                save_kpi_policy(kpi_policy)
                st.success("All KPIs disabled.")
                st.rerun()

        st.markdown("---")

        # Render checkboxes (2 columns)
        c1, c2 = st.columns(2)
        for i, (key, label) in enumerate(KPI_CATALOG):
            target_col = c1 if i % 2 == 0 else c2
            with target_col:
                kpi_policy[key] = st.checkbox(label, value=bool(kpi_policy.get(key, True)), key=f"kpi_{key}_main")

        if st.button("💾 Save KPI settings", key="kpi_save_policy_btn_main"):
            save_kpi_policy(kpi_policy)
            st.success("Saved KPI settings.")
            st.rerun()



    # --- Research Mode ---
    with st.expander("🔬 Research Mode", expanded=False):
        # Reports moved here (requested): Psychometrics + Excel exports
        with st.expander("📊 Research Reports (Psychometrics / Excel)", expanded=False):
            admin_page_research_reports()

        research_policy = load_research_policy()
        enabled = st.toggle("Enable research collection", value=bool(research_policy.get("enabled", False)), key="research_enabled_main")
        require_consent = st.checkbox("Require student consent", value=bool(research_policy.get("require_consent", True)), key="research_require_consent_main")
        anonymize_student_id = st.checkbox("Anonymize student IDs in exports", value=bool(research_policy.get("anonymize_student_id", True)), key="research_anonymize_main")
        collect_reflection = st.checkbox("Collect reflection (post-test)", value=bool(research_policy.get("collect_reflection", False)), key="research_reflection_main")
        collect_answer_change = st.checkbox("Collect answer-change behavior (NCLEX)", value=bool(research_policy.get("collect_answer_change", True)), key="research_ans_change_main")
        collect_section_perf = st.checkbox("Collect section-level performance summary", value=bool(research_policy.get("collect_section_performance", True)), key="research_section_perf_main")

        consent_title = st.text_input("Consent title", value=str(research_policy.get("consent_title","Research Consent")), key="consent_title_main")
        consent_text = st.text_area("Consent text", value=str(research_policy.get("consent_text","")), height=180, key="consent_text_main")


        st.markdown("#### 🧾 Ethics / IRB")
        irb_status = st.selectbox(
            "Ethics/IRB status",
            ["Pending", "Approved", "Exempt", "Not required (determination on file)"],
            index=["Pending", "Approved", "Exempt", "Not required (determination on file)"].index(
                str(research_policy.get("irb_status", "Pending"))
                if str(research_policy.get("irb_status", "Pending")) in ["Pending", "Approved", "Exempt", "Not required (determination on file)"]
                else "Pending"
            ),
            key="irb_status_main",
        )
        irb_ref = st.text_input(
            "IRB reference number (optional)",
            value=str(research_policy.get("irb_reference", "")),
            key="irb_ref_main",
        )

        st.caption("Optional: upload IRB/ethics determination letter (PDF). Stored locally in irb_docs/ (server machine).")
        irb_pdf = st.file_uploader("Upload IRB PDF", type=["pdf"], key="irb_pdf_upload_main")

        st.markdown("#### 🔐 De-identification")
        anonym_salt = st.text_input(
            "Anonymization salt (private)",
            value=str(research_policy.get("anonymization_salt", "")),
            type="password",
            help="Used to hash participant_id. Keep this private; changing it will change hashes.",
            key="anon_salt_main",
        )

        st.markdown("#### 🗓 Data retention")
        retention_mode = st.selectbox(
            "Retention period",
            ["12 months", "24 months", "36 months", "Indefinite"],
            index=["12 months", "24 months", "36 months", "Indefinite"].index(
                str(research_policy.get("retention_mode", "24 months"))
                if str(research_policy.get("retention_mode", "24 months")) in ["12 months", "24 months", "36 months", "Indefinite"]
                else "24 months"
            ),
            key="retention_mode_main",
        )
        retention_months_map = {"12 months": 12, "24 months": 24, "36 months": 36, "Indefinite": 0}
        retention_months = int(retention_months_map.get(retention_mode, 24))

        st.markdown("#### 🧹 Dataset tools")
        st.caption("Purge deletes ONLY the de-identified research dataset (research_dataset.jsonl). Teaching attempts remain unchanged.")
        col_p1, col_p2 = st.columns(2)
        with col_p1:
            purge_clicked = st.button("🧹 Purge research dataset", key="purge_research_dataset_main")
        with col_p2:
            show_path = st.checkbox("Show dataset paths", value=False, key="show_dataset_paths_main")
        # dataset paths shown below (after archive/new-study tools)
        if purge_clicked:
            ok, msg = purge_research_dataset()
            if ok:
                st.success(msg)
            else:
                st.error(msg)

        # --- New: Archive + Guided Start New Study (safe, with confirmations)
        st.divider()
        st.markdown("**Archive / New study tools**")
        st.caption("Archive keeps a copy of the current research dataset/log in a timestamped folder before you reset anything.")

        col_a1, col_a2 = st.columns(2)
        with col_a1:
            archive_note = st.text_input("Archive note (optional)", value="", key="archive_note_main")
        with col_a2:
            archive_clicked = st.button("📦 Archive research data", key="archive_research_btn_main")

        if archive_clicked:
            ok, msg = archive_research_data(note=str(archive_note or ""))
            if ok:
                st.success(msg)
            else:
                st.error(msg)

        with st.expander("🧭 Start New Study (guided)", expanded=False):
            st.warning(
                "This will: (1) archive the current research dataset/log, (2) purge the de-identified dataset, "
                "(3) reset the research log, (4) regenerate the anonymization salt, and (5) set IRB status to Pending. "
                "Research collection will be turned OFF until you re-enable it."
            )
            note2 = st.text_input("Optional note for the archive folder", value="", key="new_study_note_main")
            ack = st.checkbox("I understand this action cannot be undone.", value=False, key="new_study_ack_main")
            confirm = st.checkbox("Yes — start a NEW study now.", value=False, key="new_study_confirm_main")
            if st.button("🚀 Start New Study Now", key="new_study_run_btn_main", disabled=not (ack and confirm)):
                ok, msg = start_new_study_guided(note=str(note2 or ""))
                if ok:
                    st.success(msg)
                    st.rerun()
                else:
                    st.error(msg)

        if show_path:
            st.code(
                f"Dataset: {RESEARCH_DATASET_PATH}\nResearch log: {RESEARCH_LOG_PATH}\nArchives: {RESEARCH_ARCHIVE_DIR}\nIRB docs: {IRB_DOCS_DIR}",
                language="text"
            )
        if st.button("💾 Save research settings", key="save_research_policy_main"):
            research_policy["enabled"] = bool(enabled)
            research_policy["require_consent"] = bool(require_consent)
            research_policy["anonymize_student_id"] = bool(anonymize_student_id)
            research_policy["collect_reflection"] = bool(collect_reflection)
            research_policy["collect_answer_change"] = bool(collect_answer_change)
            research_policy["collect_section_performance"] = bool(collect_section_perf)
            research_policy["consent_title"] = consent_title
            research_policy["consent_text"] = consent_text

            # New: IRB / ethics + anonymization + retention
            research_policy["irb_status"] = str(irb_status)
            research_policy["irb_reference"] = str(irb_ref or "").strip()
            research_policy["anonymization_salt"] = str(anonym_salt or "").strip()
            research_policy["retention_mode"] = str(retention_mode)
            research_policy["retention_months"] = int(retention_months)

            # Optional: save uploaded IRB PDF locally
            if irb_pdf is not None:
                try:
                    IRB_DOCS_DIR.mkdir(parents=True, exist_ok=True)
                    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", getattr(irb_pdf, "name", "irb.pdf"))
                    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    out_path = IRB_DOCS_DIR / f"{stamp}_{safe_name}"
                    with open(out_path, "wb") as f:
                        f.write(irb_pdf.getbuffer())
                    docs = research_policy.get("irb_docs", [])
                    if not isinstance(docs, list):
                        docs = []
                    docs.append(str(out_path.name))
                    research_policy["irb_docs"] = docs[-10:]  # keep last 10 filenames
                except Exception:
                    pass

            save_research_policy(research_policy)
            st.success("Saved research_policy.json")
            st.rerun()

        st.markdown("---")
        st.info("")
        colr1, colr2 = st.columns(2)
        with colr1:
            st.download_button(
                "⬇️ Download Research CSV",
                data=build_research_csv_bytes(),
                file_name=f"research_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                key="dl_research_csv_main",
            )
        with colr2:
            st.download_button(
                "⬇️ Download Attempt Summary CSV",
                data=build_attempt_summary_csv_bytes(),
                file_name=f"attempt_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv",
                key="dl_attempt_summary_csv_main",
            )

        clear_after = st.checkbox("Clear logs after archive", value=False, key="archive_clear_after_main")
        if st.button("🗄️ Archive research logs", key="archive_logs_btn_main"):
            ok, msg = archive_research_logs(clear_after=bool(clear_after))
            st.success(msg) if ok else st.warning(msg)

    # --- Attempts Policy ---
    with st.expander("🎯 Attempts Policy (per case)", expanded=False):
        case_policy = load_case_policy()
        st.caption("Configure attempts allowed per case + what gets shown after submission (Practice policy).")
        default_attempts = int(case_policy.get("default_attempts_allowed", 1))
        new_default_attempts = st.number_input("Default attempts allowed per case", min_value=1, max_value=50, value=default_attempts, step=1, key="default_attempts_main")

        show_correct = st.checkbox(
            "Show correct answers after practical submit (Practice policy)",
            value=bool(case_policy.get("show_correct_after_submit", True)),
            key="show_correct_main"
        )
        show_rationales = st.checkbox(
            "Show rationales after practical submit (Practice policy)",
            value=bool(case_policy.get("show_rationales_after_submit", True)),
            key="show_rationales_main"
        )
        if st.button("💾 Save attempts policy", key="save_attempts_policy_main"):
            case_policy["default_attempts_allowed"] = int(new_default_attempts)
            case_policy["show_correct_after_submit"] = bool(show_correct)
            case_policy["show_rationales_after_submit"] = bool(show_rationales)
            save_case_policy(case_policy)
            st.success("Saved attempts policy.")
            st.rerun()

    # --- Case Visibility + Timers ---
    with st.expander("👁️ Case Visibility + ⏱ Timers", expanded=False):
        st.caption("Control which cases are visible to students + per-mode timers (in minutes).")
        case_policy = load_case_policy()
        cases = get_cases_list()
        cid_list = [str(c.get("id","")).strip() for c in cases if str(c.get("id","")).strip()]
        cid_list = sorted(list(dict.fromkeys(cid_list)))

        # Visibility mode
        vis_mode = case_policy.get("visibility_mode", "all")  # "all" or "selected"
        vis_mode = st.radio(
            "Case visibility policy",
            ["Show all cases", "Show only selected cases"],
            index=0 if vis_mode == "all" else 1,
            key="case_visibility_mode_main",
            horizontal=True
        )

        minutes_opts = ["Unlimited", 10, 15, 20, 30, 45, 60, 120, 180]  # minutes

        def _sec_to_min(s):
            try:
                s = int(s or 0)
            except Exception:
                s = 0
            return "Unlimited" if s <= 0 else int(round(s / 60))

        def _min_to_sec(mn):
            if mn in ("Unlimited", None, "", 0):
                return 0
            try:
                return int(mn) * 60
            except Exception:
                return 0

        # Defaults (stored as seconds for backward compatibility)
        timer_practice_default_min = _sec_to_min(case_policy.get("default_timer_seconds_practice", 0))
        timer_exam_default_min = _sec_to_min(case_policy.get("default_timer_seconds_exam", 0))

        colp, cole = st.columns(2)
        with colp:
            tpd_min = st.selectbox(
                "Default timer (Practice)",
                minutes_opts,
                index=minutes_opts.index(timer_practice_default_min) if timer_practice_default_min in minutes_opts else 0 if timer_practice_default_min in minutes_opts else 0,
                format_func=lambda x: "Unlimited / Off" if x == "Unlimited" else f"{x} min",
                key="timer_practice_default_min_main",
            )
        with cole:
            ted_min = st.selectbox(
                "Default timer (Exam)",
                minutes_opts,
                index=minutes_opts.index(timer_exam_default_min) if timer_exam_default_min in minutes_opts else 0 if timer_exam_default_min in minutes_opts else 0,
                format_func=lambda x: "Unlimited / Off" if x == "Unlimited" else f"{x} min",
                key="timer_exam_default_min_main",
            )

        # Selected-only visibility: choose visible cases
        selected_visible = case_policy.get("visible_case_ids", [])
        if not isinstance(selected_visible, list):
            selected_visible = []

        if vis_mode == "Show only selected cases":
            selected_visible = st.multiselect(
                "Select cases to show to students",
                options=cid_list,
                default=[c for c in selected_visible if c in cid_list],
                key="visible_case_ids_main",
            )
            st.info("Students will only see the cases you select here.")
        else:
            st.success("Students will see all cases (subject to per-case overrides).")

        st.divider()
        st.subheader("Per-case overrides")

        pick = st.selectbox("Select case to edit", ["(none)"] + cid_list, key="case_pick_visibility_main")
        if pick != "(none)":
            per_case = case_policy.setdefault("per_case", {})
            cp = per_case.setdefault(pick, {})
            if not isinstance(cp, dict):
                cp = {}
                per_case[pick] = cp

            # visibility overrides by mode (backward compatible keys)
            mv = cp.get("mode_visibility", {})
            if not isinstance(mv, dict):
                mv = {}
            col1, col2 = st.columns(2)
            with col1:
                vis_practice = st.checkbox("Visible in Practice", value=bool(mv.get("Practice", True)), key="vis_practice_main")
                vis_exam = st.checkbox("Visible in Exam", value=bool(mv.get("Exam", True)), key="vis_exam_main")
            with col2:
                tp_min = _sec_to_min(cp.get("timer_seconds_practice", 0))
                te_min = _sec_to_min(cp.get("timer_seconds_exam", 0))
                tp = st.selectbox(
                    "Timer for this case (Practice)",
                    minutes_opts,
                    index=minutes_opts.index(tp_min) if tp_min in minutes_opts else 0,
                    format_func=lambda x: "Unlimited / Off" if x == 0 else f"{x} min",
                    key="case_timer_practice_min_main",
                )
                te = st.selectbox(
                    "Timer for this case (Exam)",
                    minutes_opts,
                    index=minutes_opts.index(te_min) if te_min in minutes_opts else 0,
                    format_func=lambda x: "Unlimited / Off" if x == 0 else f"{x} min",
                    key="case_timer_exam_min_main",
                )


            # Availability windows (start/end) per mode (Qatar time)
            st.markdown("### ⏳ Practice / Exam availability window (optional)")
            st.caption("If set, students can only open this case within the start–end window for that mode. Leave off for always-available.")

            avail = cp.get("availability", {}) if isinstance(cp.get("availability", {}), dict) else {}
            def _get_win(mode_name: str):
                w = avail.get(mode_name, {}) if isinstance(avail.get(mode_name, {}), dict) else {}
                return str(w.get("start","") or "").strip(), str(w.get("end","") or "").strip()

            def _dt_parts(dt_obj):
                if not dt_obj:
                    base = now_local().replace(minute=0, second=0, microsecond=0)
                    return base.date(), base.time()
                try:
                    loc = dt_obj.astimezone(TZ) if dt_obj.tzinfo else dt_obj.replace(tzinfo=TZ)
                except Exception:
                    loc = dt_obj
                return loc.date(), loc.time()

            p_start_s, p_end_s = _get_win("Practice")
            e_start_s, e_end_s = _get_win("Exam")
            p_start_dt = parse_iso_dt(p_start_s) if p_start_s else None
            p_end_dt = parse_iso_dt(p_end_s) if p_end_s else None
            e_start_dt = parse_iso_dt(e_start_s) if e_start_s else None
            e_end_dt = parse_iso_dt(e_end_s) if e_end_s else None

            enable_p = st.toggle("Enable Practice window", value=bool(p_start_s or p_end_s), key="case_win_practice_enable_main")
            enable_e = st.toggle("Enable Exam window", value=bool(e_start_s or e_end_s), key="case_win_exam_enable_main")

            colw1, colw2 = st.columns(2)
            with colw1:
                st.markdown("**Practice window**")
                ps_d, ps_t = _dt_parts(p_start_dt)
                pe_d, pe_t = _dt_parts(p_end_dt)
                p_start_date = st.date_input("Practice start date", value=ps_d, disabled=not enable_p, key="p_start_date_main")
                p_start_time = st.time_input("Practice start time", value=ps_t, disabled=not enable_p, key="p_start_time_main")
                p_end_date = st.date_input("Practice end date", value=pe_d, disabled=not enable_p, key="p_end_date_main")
                p_end_time = st.time_input("Practice end time", value=pe_t, disabled=not enable_p, key="p_end_time_main")
            with colw2:
                st.markdown("**Exam window**")
                es_d, es_t = _dt_parts(e_start_dt)
                ee_d, ee_t = _dt_parts(e_end_dt)
                e_start_date = st.date_input("Exam start date", value=es_d, disabled=not enable_e, key="e_start_date_main")
                e_start_time = st.time_input("Exam start time", value=es_t, disabled=not enable_e, key="e_start_time_main")
                e_end_date = st.date_input("Exam end date", value=ee_d, disabled=not enable_e, key="e_end_date_main")
                e_end_time = st.time_input("Exam end time", value=ee_t, disabled=not enable_e, key="e_end_time_main")

            def _mk_iso(d, t):
                try:
                    return datetime.combine(d, t).replace(tzinfo=TZ).isoformat()
                except Exception:
                    return ""

            if st.button("💾 Save this case policy", key="save_case_policy_case_main"):
                cp["mode_visibility"] = {"Practice": bool(vis_practice), "Exam": bool(vis_exam)}
                cp["timer_seconds_practice"] = _min_to_sec(tp)
                cp["timer_seconds_exam"] = _min_to_sec(te)


                # Keep minutes keys in sync too (prevents stale minutes overriding seconds)
                cp["timer_minutes_practice"] = "unlimited" if (tp == 0) else int(tp)
                cp["timer_minutes_exam"] = "unlimited" if (te == 0) else int(te)
                # Save availability windows
                availability = cp.get("availability", {}) if isinstance(cp.get("availability", {}), dict) else {}
                if bool(enable_p):
                    availability["Practice"] = {
                        "start": _mk_iso(p_start_date, p_start_time),
                        "end": _mk_iso(p_end_date, p_end_time),
                    }
                else:
                    availability.pop("Practice", None)
                if bool(enable_e):
                    availability["Exam"] = {
                        "start": _mk_iso(e_start_date, e_start_time),
                        "end": _mk_iso(e_end_date, e_end_time),
                    }
                else:
                    availability.pop("Exam", None)
                cp["availability"] = availability


                # Save defaults + visibility policy
                case_policy["default_timer_seconds_practice"] = _min_to_sec(tpd_min)
                case_policy["default_timer_seconds_exam"] = _min_to_sec(ted_min)


                # Keep minutes keys in sync (prevents UI mismatch)
                case_policy["default_timer_minutes_practice"] = "unlimited" if (tpd_min == 0) else int(tpd_min)
                case_policy["default_timer_minutes_exam"] = "unlimited" if (ted_min == 0) else int(ted_min)
                if vis_mode == "Show only selected cases":
                    case_policy["visibility_mode"] = "selected"
                    case_policy["visible_case_ids"] = list(selected_visible)
                    case_policy["default_visibility"] = False
                else:
                    case_policy["visibility_mode"] = "all"
                    case_policy["visible_case_ids"] = []
                    case_policy["default_visibility"] = True

                save_case_policy(case_policy)
                st.success("Saved case visibility/timers.")
                st.rerun()
        else:
            if st.button("💾 Save defaults", key="save_case_policy_defaults_main"):
                case_policy["default_timer_seconds_practice"] = _min_to_sec(tpd_min)
                case_policy["default_timer_seconds_exam"] = _min_to_sec(ted_min)


                # Keep minutes keys in sync (prevents UI mismatch)
                case_policy["default_timer_minutes_practice"] = "unlimited" if (tpd_min == 0) else int(tpd_min)
                case_policy["default_timer_minutes_exam"] = "unlimited" if (ted_min == 0) else int(ted_min)
                if vis_mode == "Show only selected cases":
                    case_policy["visibility_mode"] = "selected"
                    case_policy["visible_case_ids"] = list(selected_visible)
                    case_policy["default_visibility"] = False
                else:
                    case_policy["visibility_mode"] = "all"
                    case_policy["visible_case_ids"] = []
                    case_policy["default_visibility"] = True

                save_case_policy(case_policy)
                st.success("Saved defaults.")
                st.rerun()

    # --- NCLEX Validator ---

    with st.expander("✅ NCLEX Validator (IDs + Auto-fix)", expanded=False):
        st.caption("Validates NCLEX items IDs + structure and can auto-fix common issues.")
        run_validator_ui()

    # --- NCLEX Rotation / Security ---
    with st.expander("🎛 NCLEX practice + rotation + security", expanded=False):
        st.caption("Restore the original NCLEX controls: practice policy, item types, rotation, question-by-question view, footer code, and copy/print deterrents.")
        pol = load_nclex_policy()

        # --- Core enablement + types ---
        st.subheader("NCLEX Practice Policy")
        pol["enabled"] = st.toggle("Enable NCLEX-style practice", value=bool(pol.get("enabled", False)), key="nclex_enabled_main")

        et = pol.get("enabled_types") or {}
        if not isinstance(et, dict):
            et = {}
        st.caption("Enable/disable NCLEX item types")
        cols = st.columns(3)
        for i, tname in enumerate(["mcq", "sata", "ordered_response", "cloze", "matrix", "evolving_case"]):
            with cols[i % 3]:
                et[tname] = st.toggle(f"Enable {tname}", value=bool(et.get(tname, True)), key=f"nclex_type_{tname}_main")
        pol["enabled_types"] = et

        colA, colB, colC = st.columns(3)
        with colA:
            pol["show_correct_answers_after_submit"] = st.toggle(
                "Show correct answers after practical submit (Practice policy)",
                value=bool(pol.get("show_correct_answers_after_submit", True)),
                key="nclex_show_correct_main"
            )
        with colB:
            pol["show_rationales_after_submit"] = st.toggle(
                "Show rationales after practical submit (Practice policy)",
                value=bool(pol.get("show_rationales_after_submit", True)),
                key="nclex_show_rationales_main"
            )
        with colC:
            pol["shuffle_options"] = st.toggle(
                "Shuffle options (where safe)",
                value=bool(pol.get("shuffle_options", True)),
                key="nclex_shuffle_options_main"
            )


        st.divider()
        st.subheader("Per-case question count (timed compatibility)")
        st.caption("Override the NCLEX questions shown for a specific case (Practice + Exam). Leave blank to use the global items_per_case.")
        if not isinstance(pol.get("per_case_items"), dict):
            pol["per_case_items"] = {}
        _case_opts = [""] + [str(c.get("id","")).strip() for c in (cases or []) if isinstance(c, dict) and str(c.get("id","")).strip()]
        _cid_sel = st.selectbox("Select case (override)", options=_case_opts, index=0, key="nclex_case_override_pick")
        if _cid_sel:
            _cur = pol["per_case_items"].get(_cid_sel, "")
            cA, cB, cC = st.columns([1,1,2])
            with cA:
                _override_val = st.number_input("Questions", min_value=1, max_value=200, value=int(_cur) if str(_cur).strip().isdigit() else int(pol.get("items_per_case", 30) or 30), step=1, key="nclex_case_override_num")
            with cB:
                if st.button("➕ Save override", key="nclex_case_override_save"):
                    pol["per_case_items"][_cid_sel] = int(_override_val)
                    save_nclex_policy(pol)
                    st.success(f"Saved override: {_cid_sel} → {int(_override_val)} questions.")
                    st.rerun()
            with cC:
                if st.button("🗑 Remove override (use global)", key="nclex_case_override_remove"):
                    try:
                        pol["per_case_items"].pop(_cid_sel, None)
                    except Exception:
                        pass
                    save_nclex_policy(pol)
                    st.success(f"Removed override for {_cid_sel}.")
                    st.rerun()

        st.divider()
        st.subheader("Rotation + Student View")
        rot_enabled = st.checkbox("Enable rotation using an admin-generated active set (per case)", value=bool(pol.get("rotation_enabled", False)), key="nclex_rotation_enabled_main")
        rand_student = st.checkbox("Randomize NCLEX item order per student/session", value=bool(pol.get("randomize_per_student_session", True)), key="nclex_rand_student_main")
        one_at_time = st.checkbox("Student view: 1 NCLEX question at a time", value=bool(pol.get("one_question_at_a_time", False)), key="nclex_one_at_time_main")
        footer_enabled = st.checkbox("Show tiny footer session code (student view)", value=bool(pol.get("footer_session_code_enabled", True)), key="nclex_footer_enabled_main")

        st.divider()
        st.subheader("Exam integrity (deterrent)")
        # New toggle (does not change scoring/exports) — just controls whether copy/print blockers are injected.
        exam_protect = st.checkbox("Prevent copy/print/right-click (best effort)", value=bool(pol.get("exam_protection_enabled", True)), key="exam_protect_main")
        pol["exam_protection_enabled"] = bool(exam_protect)

        if st.button("💾 Save NCLEX settings", key="save_nclex_settings_main"):
            pol["rotation_enabled"] = bool(rot_enabled)
            pol["randomize_per_student_session"] = bool(rand_student)
            pol["one_question_at_a_time"] = bool(one_at_time)
            pol["footer_session_code_enabled"] = bool(footer_enabled)
            # keep watermark disabled (legacy)
            pol["watermark_enabled"] = False
            save_nclex_policy(pol)
            st.success("Saved nclex_policy.json")
            st.rerun()

        st.caption("Detailed active-set generator/history (original).")
        st.divider()
        render_nclex_rotation_admin_ui(pol)


    # --- Introductory Case Videos ---
    with st.expander("🎬 Introductory Case Videos", expanded=False):
        st.caption("Upload a 4–6 min intro video per case (MP4). Students will see it before the clinical scenario.")
        admin_settings = load_admin_settings()
        intro_videos = admin_settings.get("intro_videos", {}) if isinstance(admin_settings.get("intro_videos", {}), dict) else {}

        cases = get_cases_list()
        cid_list = [str(c.get("id","")).strip() for c in cases if str(c.get("id","")).strip()]
        cid_list = sorted(list(dict.fromkeys(cid_list)))

        if not cid_list:
            st.info("No cases found to attach videos to.")
        else:
            sel_case = st.selectbox("Select case", cid_list, key="intro_video_case_sel")
            cur = intro_videos.get(str(sel_case), {}) if isinstance(intro_videos.get(str(sel_case), {}), dict) else {}
            cur_path = str(cur.get("path","") or "").strip()
            cur_require = bool(cur.get("require_watch", True))

            st.markdown("**Current video**")
            if cur_path:
                st.write(f"Saved path: `{cur_path}`")
                try:
                    p = Path(cur_path)
                    st.write(f"Exists on server: {'✅' if p.exists() else '❌'}")
                except Exception:
                    pass
            else:
                st.write("None")

            require_watch = st.checkbox("Require student confirmation ('I watched the video') before proceeding",
                                        value=cur_require, key="intro_video_require_watch")

            up = st.file_uploader("Upload MP4 for this case", type=["mp4"], key=f"intro_video_uploader__{sel_case}")

            colA, colB = st.columns(2)
            with colA:
                if st.button("Save / Replace video for this case", key="btn_save_intro_video"):
                    if up is None:
                        st.warning("Please upload an MP4 file first.")
                    else:
                        try:
                            vids_dir = Path("case_videos")
                            vids_dir.mkdir(parents=True, exist_ok=True)
                            dest = vids_dir / f"{str(sel_case)}.mp4"
                            dest.write_bytes(up.getvalue())

                            intro_videos[str(sel_case)] = {"path": str(dest), "require_watch": bool(require_watch)}
                            admin_settings["intro_videos"] = intro_videos
                            save_admin_settings(admin_settings)

                            st.success(f"Saved intro video for {sel_case}.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Failed to save video: {e}")

            with colB:
                if st.button("Remove video for this case", key="btn_remove_intro_video"):
                    # Remove mapping; do not delete file by default (safer)
                    if str(sel_case) in intro_videos:
                        intro_videos.pop(str(sel_case), None)
                        admin_settings["intro_videos"] = intro_videos
                        save_admin_settings(admin_settings)
                        st.success("Removed mapping (video will no longer appear for students).")
                        st.rerun()
                    else:
                        st.info("No video mapping exists for this case.")

# =============================
# Admin Navigation (Top bar)
# =============================

# =============================
# Faculty: Grade Center (Exam attempts only)
# =============================

def _format_dt_local(iso_str: str) -> str:
    try:
        if not iso_str:
            return ""
        dt = datetime.fromisoformat(str(iso_str).replace("Z",""))
        # If naive, treat as UTC then convert to Qatar TZ if available
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=datetime.timezone.utc)
        try:
            return dt.astimezone(TZ).strftime("%Y-%m-%d %H:%M")
        except Exception:
            return dt.astimezone(datetime.timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    except Exception:
        return str(iso_str or "")

def _attempt_total_score(rec: dict) -> int:
    # Reasoning A-E total + Intake + NCLEX (if present). Best-effort.
    try:
        scores = rec.get("scores") if isinstance(rec.get("scores"), dict) else {}
        ae_total = sum([int(scores.get(k, 0) or 0) for k in ["A","B","C","D","E"]])
        intake = int(rec.get("intake_score", 0) or 0)
        nclex = int(rec.get("nclex_score", 0) or 0)
        return int(ae_total + intake + nclex)
    except Exception:
        return 0

def _intake_max(rec: dict) -> int:
    try:
        ib = rec.get("intake_breakdown") if isinstance(rec.get("intake_breakdown"), dict) else {}
        n = len([k for k in ib.keys() if str(k).strip()])
        if n > 0:
            return int(n)
    except Exception:
        pass
    # fallback default used in most modules
    return int(rec.get("intake_total", 5) or 5)

def _ae_max(rec: dict) -> int:
    # Best-effort: prefer stored total, otherwise default 16 (A–E rubric)
    try:
        v = rec.get("ae_total_max", None)
        if v is not None:
            return int(v)
    except Exception:
        pass
    return 16

def _fmt_score(num: int, den: int) -> str:
    try:
        return f"{int(num)}/{int(den)}"
    except Exception:
        return f"{num}/{den}"

def _csv_text_score(s: str) -> str:
    # Prevent Excel from auto-parsing 2/22 as a date by prefixing apostrophe.
    # Excel displays it without the apostrophe.
    return "'" + str(s)


def _build_attempt_pdf_bytes(rec: dict) -> bytes:
    # PDF report for ONE attempt (faculty download)
    try:
        from io import BytesIO
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch

        buf = BytesIO()
        c = canvas.Canvas(buf, pagesize=letter)
        width, height = letter

        def draw_line(text, x, y, size=10, bold=False):
            try:
                c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
            except Exception:
                c.setFont("Helvetica", size)
            c.drawString(x, y, str(text)[:2000])

        y = height - 0.8*inch
        draw_line("ClinIQ / NursApp — Exam Attempt Report", 0.8*inch, y, size=14, bold=True); y -= 0.3*inch

        draw_line(f"Student: {rec.get('student_display_name','')}  ({rec.get('student_username','')})", 0.8*inch, y); y -= 0.2*inch
        draw_line(f"Case: {rec.get('caseTitle','')}  |  Case ID: {rec.get('caseId','')}", 0.8*inch, y); y -= 0.2*inch
        draw_line(f"Submitted: {_format_dt_local(rec.get('submitted_at',''))}   Duration: {int(rec.get('duration_seconds',0) or 0)} sec", 0.8*inch, y); y -= 0.25*inch

        scores = rec.get("scores") if isinstance(rec.get("scores"), dict) else {}
        ae_total = sum([int(scores.get(k, 0) or 0) for k in ["A","B","C","D","E"]])
        intake = int(rec.get("intake_score", 0) or 0)
        nclex_score = int(rec.get("nclex_score", 0) or 0)
        nclex_total = int(rec.get("nclex_total", 0) or 0)
        total = _attempt_total_score(rec)

        draw_line("Scores Summary", 0.8*inch, y, bold=True); y -= 0.18*inch
        draw_line(f"A–E total: {_fmt_score(ae_total, _ae_max(rec))}", 0.95*inch, y); y -= 0.16*inch
        draw_line(f"Intake: {_fmt_score(intake, _intake_max(rec))}", 0.95*inch, y); y -= 0.16*inch
        draw_line(f"NCLEX: {nclex_score}/{nclex_total}", 0.95*inch, y); y -= 0.16*inch
        overall_max = _intake_max(rec) + _ae_max(rec) + nclex_total
        draw_line(f"Overall: {_fmt_score(total, overall_max)}", 0.95*inch, y); y -= 0.25*inch

        # Intake breakdown
        ib = rec.get("intake_breakdown") if isinstance(rec.get("intake_breakdown"), dict) else {}
        if ib:
            draw_line("Intake Breakdown", 0.8*inch, y, bold=True); y -= 0.18*inch
            for k, v in list(ib.items())[:20]:
                draw_line(f"- {k}: {v}", 0.95*inch, y); y -= 0.14*inch
                if y < 1.0*inch:
                    c.showPage(); y = height - 0.8*inch

        # NCLEX item details (compact)
        details = (rec.get("nclex") or {}).get("details") if isinstance((rec.get("nclex") or {}).get("details"), list) else []
        if details:
            draw_line("NCLEX Details", 0.8*inch, y, bold=True); y -= 0.18*inch
            for it in details[:60]:
                qid = it.get("id","")
                ok = it.get("is_correct")
                pts = it.get("points", "")
                draw_line(f"- {qid} | correct={ok} | points={pts}", 0.95*inch, y); y -= 0.14*inch
                if y < 1.0*inch:
                    c.showPage(); y = height - 0.8*inch

        c.showPage()
        c.save()
        return buf.getvalue()
    except Exception:
        return b""



# =============================
# Admin Pages: Student Manager + Exam Control (Main screen)
# =============================
def admin_page_exam_control():
    st.subheader("🔒 Exam Control")
    admin_settings = load_admin_settings()
    IS_PRACTICE = (admin_settings.get("app_mode") == "Practice")
    IS_EXAM = not IS_PRACTICE

    if not IS_EXAM:
        st.info("App mode is currently **Practice**. Switch to **Exam** in ⚙️ Settings to use exam session controls.")
        return

    st.caption("End the exam session for **this device** (current browser session) or request an **End/Reset** for a selected student (server-side).")

    end_scope = st.selectbox(
        "End Exam Session for",
        ["This device", "Selected student"],
        index=0,
        key="main_end_exam_scope_sel",
    )

    sel_student = None
    if end_scope == "Selected student":
        try:
            sd = load_students()
            usernames = sorted(list({str(s.get("username","")).strip() for s in (sd.get("students", []) or []) if str(s.get("username","")).strip()}))
        except Exception:
            usernames = []
        sel_student = st.selectbox(
            "Student username",
            usernames if usernames else [""],
            index=0,
            key="main_end_exam_student_sel",
        )
        sel_student = (sel_student or "").strip()

    if st.button("🔒 End Exam Session", key="main_end_exam_session_btn"):
        if end_scope == "This device":
            # Clear exam/case-switch lock for this session only (no logs are modified).
            for _k in [
                "locked_exam_case_id",
                "exam_in_progress",
                "exam_started",
                "exam_case_id",
                "current_case_id",
                "active_case_id",
                "_exam_started_at",
                "_exam_timer_start",
                "_exam_timer_deadline",
            ]:
                st.session_state.pop(_k, None)

            # Reset case selectors to allow switching cases immediately.
            for _k in ["case_pick_main", "rot_case_pick_main", "case_system_main"]:
                st.session_state.pop(_k, None)

            # Reset attempt-level state using the app's existing helper.
            try:
                reset_attempt_state()
            except Exception:
                pass

            st.success("Exam session ended for this device.")
            st.rerun()
        else:
            if not sel_student:
                st.error("Please select a student username.")
                return
            try:
                set_student_force_unlock(sel_student, True)
                bump_student_reset_token(sel_student)  # also clears autosaves on next run
                st.success(f"End/Reset requested for student: {sel_student}. Ask them to refresh the app.")
            except Exception:
                st.error("Could not request reset for the selected student.")


def admin_page_student_session_manager():
    st.subheader("👤 Student Session Manager")
    st.caption("View and clear unfinished autosaves (progress drafts), and remotely force-unlock/reset a student exam session.")

    try:
        sd2 = load_students()
        usernames2 = sorted(list({str(s.get("username","")).strip() for s in (sd2.get("students", []) or []) if str(s.get("username","")).strip()}))
    except Exception:
        usernames2 = []

    mgr_student = st.selectbox(
        "Select student",
        usernames2 if usernames2 else [""],
        index=0,
        key="main_student_mgr_sel",
    )
    mgr_student = (mgr_student or "").strip()

    if not mgr_student:
        st.info("Select a student to manage their in-progress drafts and exam lock.")
        return

    drafts = index_latest_autosaves_for_student(mgr_student) or {}
    case_ids = sorted(list(drafts.keys()))
    st.write(f"**In-progress autosaves detected:** {len(case_ids)}")

    if case_ids:
        pick = st.multiselect(
            "Autosaved cases (select to delete)",
            case_ids,
            default=[],
            key="main_student_mgr_pick",
        )
        col1, col2, col3 = st.columns(3)

        with col1:
            if st.button("🗑 Delete selected autosaves", key="main_mgr_del_selected"):
                for cid in (pick or []):
                    try:
                        delete_autosaves_for(mgr_student, cid)
                    except Exception:
                        pass
                st.success("Selected autosaves deleted.")
                st.session_state.pop("_progress_index_loaded", None)
                st.rerun()

        with col2:
            if st.button("🧹 Delete ALL autosaves", key="main_mgr_del_all"):
                for cid in case_ids:
                    try:
                        delete_autosaves_for(mgr_student, cid)
                    except Exception:
                        pass
                st.success("All autosaves deleted for this student.")
                st.session_state.pop("_progress_index_loaded", None)
                st.rerun()

        with col3:
            if st.button("🔓 Force unlock", key="main_mgr_force_unlock"):
                try:
                    set_student_force_unlock(mgr_student, True)
                    st.success("Force unlock requested. Ask the student to refresh.")
                except Exception:
                    st.error("Could not set force unlock.")

        st.divider()
        st.subheader("Full reset (recommended if student is stuck)")

        if st.button("♻️ Full reset: unlock + clear autosaves", key="main_mgr_full_reset"):
            try:
                set_student_force_unlock(mgr_student, True)
                bump_student_reset_token(mgr_student)
                # Also clear drafts now (so Progress panel updates immediately)
                for cid in case_ids:
                    try:
                        delete_autosaves_for(mgr_student, cid)
                    except Exception:
                        pass
                st.success("Full reset requested + autosaves cleared. Ask the student to refresh.")
                st.rerun()
            except Exception:
                st.error("Could not perform full reset.")
    else:
        st.info("No autosaves found for this student.")
        colA, colB = st.columns(2)
        with colA:
            if st.button("🔓 Force unlock", key="main_mgr_force_unlock_noautosave"):
                try:
                    set_student_force_unlock(mgr_student, True)
                    st.success("Force unlock requested. Ask the student to refresh.")
                except Exception:
                    st.error("Could not set force unlock.")
        with colB:
            if st.button("♻️ Full reset (unlock + bump token)", key="main_mgr_full_reset_noautosave"):
                try:
                    set_student_force_unlock(mgr_student, True)
                    bump_student_reset_token(mgr_student)
                    st.success("Full reset requested. Ask the student to refresh.")
                except Exception:
                    st.error("Could not perform full reset.")

def render_grade_center_page():
    # Title in the same light-blue header style used elsewhere
    st.markdown(
        "<div style='background:#e8f4ff;border:1px solid rgba(0,0,0,.08);padding:12px 14px;border-radius:14px;'>"
        "<span style='color:#b00020;font-weight:900;'>🎓 Grade Center</span>"
        "</div>",
        unsafe_allow_html=True,
    )

    # Faculty-only note (no stars), bold black bullet points
    st.markdown(
        "<ul style='margin-top:10px;margin-bottom:8px;'>"
        "<li style='font-weight:800;color:#111;'>Faculty-only</li>"
        "<li style='font-weight:800;color:#111;'>Shows EXAM submissions only (Practice is excluded)</li>"
        "</ul>",
        unsafe_allow_html=True,
    )

    # Load attempts (JSONL) — EXAMS ONLY
    rows = []
    for rec in (iter_attempts() or []):
        if not isinstance(rec, dict):
            continue
        if str(rec.get("mode", "")) != "Exam":
            continue
        rows.append(rec)

    if not rows:
        st.info("No EXAM attempts found yet. Run an exam as a student, finalize/submit, then return here.")
        return

    # Filters
    all_students = sorted({str(r.get("student_username", "")) for r in rows if str(r.get("student_username", "")).strip()})
    all_cases = sorted({str(r.get("caseId", "")) for r in rows if str(r.get("caseId", "")).strip()})

    c1, c2, c3 = st.columns([1, 1, 1])
    with c1:
        f_student = st.selectbox("Filter: Student", options=["(All)"] + all_students, index=0, key="gc_f_student")
    with c2:
        f_case = st.selectbox("Filter: Case ID", options=["(All)"] + all_cases, index=0, key="gc_f_case")
    with c3:
        f_text = st.text_input("Search (name / case title)", value="", key="gc_f_text")

    def _keep(r):
        if f_student != "(All)" and str(r.get("student_username", "")) != f_student:
            return False
        if f_case != "(All)" and str(r.get("caseId", "")) != f_case:
            return False
        if f_text.strip():
            t = (str(r.get("student_display_name", "")) + " " + str(r.get("caseTitle", ""))).lower()
            if f_text.strip().lower() not in t:
                return False
        return True

    filt = [r for r in rows if _keep(r)]
    filt.sort(key=lambda r: str(r.get("submitted_at", "")), reverse=True)

    # Build table with score/total formatting
    table = []
    for r in filt:
        scores = r.get("scores") if isinstance(r.get("scores"), dict) else {}
        ae_score = sum([int(scores.get(k, 0) or 0) for k in ["A", "B", "C", "D", "E"]])
        ae_den = _ae_max(r)
        intake_score = int(r.get("intake_score", 0) or 0)
        intake_den = _intake_max(r)
        nclex_score = int(r.get("nclex_score", 0) or 0)
        nclex_den = int(r.get("nclex_total", 0) or 0)
        overall_score = _attempt_total_score(r)
        overall_den = intake_den + ae_den + nclex_den

        table.append({
            "Submitted (Qatar)": _format_dt_local(r.get("submitted_at", "")),
            "Student": str(r.get("student_username", "")),
            "Case ID": str(r.get("caseId", "")),
            "Case Title": str(r.get("caseTitle", "")),
            "Intake": _fmt_score(intake_score, intake_den),
            "A–E": _fmt_score(ae_score, ae_den),
            "NCLEX": _fmt_score(nclex_score, nclex_den),
            "Overall": _fmt_score(overall_score, overall_den),
        })

    # Section title
    st.markdown(
        "<div style='margin-top:6px;margin-bottom:6px;'><span style='color:#b00020;font-weight:900;'>All Exam Submissions</span></div>",
        unsafe_allow_html=True,
    )

    # Render table (light green background, red/bold column headers)
    def _render_table_html(rows_list):
        if not rows_list:
            return "<div style='color:#111;'>No records.</div>"
        cols = list(rows_list[0].keys())
        th = "".join([f"<th style='padding:8px 10px;border:1px solid rgba(0,0,0,.08);color:#b00020;font-weight:900;background:#dcf3dc;text-align:left;'>{_html.escape(c)}</th>" for c in cols])
        trs = []
        for rr in rows_list:
            tds = "".join([f"<td style='padding:8px 10px;border:1px solid rgba(0,0,0,.08);background:#e9f8e9;color:#111;'>{_html.escape(str(rr.get(c, '')))}</td>" for c in cols])
            trs.append(f"<tr>{tds}</tr>")
        return (
            "<div style='border-radius:14px;overflow:hidden;border:1px solid rgba(0,0,0,.10);'>"
            "<table style='width:100%;border-collapse:collapse;font-size:0.95rem;'>"
            f"<thead><tr>{th}</tr></thead>"
            f"<tbody>{''.join(trs)}</tbody>"
            "</table></div>"
        )

    st.markdown(_render_table_html(table), unsafe_allow_html=True)

    # Export to Excel (XLSX) with text formatting to prevent date parsing like 2/22 -> 22-Feb
    try:
        from io import BytesIO
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = "Grade Center"

        headers = list(table[0].keys())
        ws.append(headers)

        header_font = Font(bold=True, color="B00020")
        header_fill = PatternFill(start_color="DCF3DC", end_color="DCF3DC", fill_type="solid")
        for col_i, h in enumerate(headers, start=1):
            cell = ws.cell(row=1, column=col_i)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="left", vertical="center")
            ws.column_dimensions[get_column_letter(col_i)].width = 18 if col_i < 5 else 14

        # data rows (force text format for score columns)
        for rr in table:
            ws.append([rr.get(h, "") for h in headers])

        # set table background light green and set score columns as TEXT
        data_fill = PatternFill(start_color="E9F8E9", end_color="E9F8E9", fill_type="solid")
        score_cols = set(["Intake", "A–E", "NCLEX", "Overall"])
        for row_i in range(2, ws.max_row + 1):
            for col_i, h in enumerate(headers, start=1):
                cell = ws.cell(row=row_i, column=col_i)
                cell.fill = data_fill
                cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
                if h in score_cols:
                    cell.number_format = "@"  # text
                    cell.value = str(cell.value) if cell.value is not None else ""

        bio = BytesIO()
        wb.save(bio)
        xlsx_bytes = bio.getvalue()

        st.download_button(
            "⬇️ Download Excel (filtered)",
            data=xlsx_bytes,
            file_name="grade_center_exam_attempts.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    except Exception:
        pass

    st.markdown(
        "<div style='background:#e8f4ff;border:1px solid rgba(0,0,0,.08);padding:10px 12px;border-radius:14px;margin-top:10px;'>"
        "<span style='color:#b00020;font-weight:900;'>Open Attempt Details</span>"
        "</div>",
        unsafe_allow_html=True,
    )

    pick_labels = [
        f"{_format_dt_local(r.get('submitted_at',''))} | {r.get('student_username','')} | {r.get('caseTitle','')}"
        for r in filt
    ]
    picked = st.selectbox("", options=[""] + pick_labels, index=0, key="gc_pick", label_visibility="collapsed")

    if not picked:
        return

    idx = pick_labels.index(picked)
    rec = filt[idx]

    # Attempt Details — clean, readable (titles bold black, details normal black)
    scores = rec.get("scores") if isinstance(rec.get("scores"), dict) else {}
    ae_score = sum([int(scores.get(k, 0) or 0) for k in ["A", "B", "C", "D", "E"]])
    ae_den = _ae_max(rec)
    intake_score = int(rec.get("intake_score", 0) or 0)
    intake_den = _intake_max(rec)
    nclex_score = int(rec.get("nclex_score", 0) or 0)
    nclex_den = int(rec.get("nclex_total", 0) or 0)
    overall_score = _attempt_total_score(rec)
    overall_den = intake_den + ae_den + nclex_den

    st.markdown(
        "<div style='margin-top:10px;margin-bottom:6px;'><span style='color:#b00020;font-weight:900;'>Attempt Details</span></div>",
        unsafe_allow_html=True,
    )

    details_html = (
        "<ul style='margin-top:6px;'>"
        f"<li><b>Student:</b> {_html.escape(str(rec.get('student_display_name','')))} ({_html.escape(str(rec.get('student_username','')) )})</li>"
        f"<li><b>Case:</b> {_html.escape(str(rec.get('caseTitle','')))} <span style='color:#666;'>({_html.escape(str(rec.get('caseId','')) )})</span></li>"
        f"<li><b>Submitted:</b> {_html.escape(_format_dt_local(rec.get('submitted_at','')))}</li>"
        f"<li><b>Duration:</b> {_html.escape(str(int(rec.get('duration_seconds',0) or 0)))} sec</li>"
        "</ul>"
        "<div style='margin-top:6px;'>"
        f"<div><b>Intake:</b> {_html.escape(_fmt_score(intake_score, intake_den))}</div>"
        f"<div><b>A–E total:</b> {_html.escape(_fmt_score(ae_score, ae_den))}</div>"
        f"<div><b>NCLEX:</b> {_html.escape(_fmt_score(nclex_score, nclex_den))}</div>"
        f"<div><b>Overall:</b> {_html.escape(_fmt_score(overall_score, overall_den))}</div>"
        "</div>"
    )
    st.markdown(details_html, unsafe_allow_html=True)

    # Download PDF attempt report (uses existing generator)
    try:
        pdf_bytes = _build_attempt_pdf_bytes(rec)
        if pdf_bytes:
            st.download_button(
                "⬇️ Download PDF Attempt Report",
                data=pdf_bytes,
                file_name=f"exam_attempt_{rec.get('student_username','student')}_{rec.get('caseId','case')}.pdf",
                mime="application/pdf",
            )
    except Exception:
        pass

    # =============================
    # Saved Answers & Performance Report (Faculty-only, EXAM attempt)
    # =============================
    st.markdown(
        "<div style='background:#e8f4ff;border:1px solid rgba(0,0,0,.08);padding:10px 12px;border-radius:14px;margin-top:12px;'>"
        "<span style='color:#b00020;font-weight:900;'>Show saved answers and performance report</span>"
        "</div>",
        unsafe_allow_html=True,
    )

    with st.expander("Open report (saved answers + scoring summary)", expanded=False):
        # Load case (best-effort) to show scenario + gold targets
        case_obj = None
        try:
            for c in load_cases():
                if isinstance(c, dict) and str(c.get("id","")) == str(rec.get("caseId","")):
                    case_obj = c
                    break
        except Exception:
            case_obj = None

        # --- Clinical Scenario ---
        scenario_txt = ""
        try:
            if isinstance(case_obj, dict):
                scenario_txt = str(case_obj.get("scenario") or case_obj.get("clinical_scenario") or case_obj.get("prompt") or "")
        except Exception:
            scenario_txt = ""
        if scenario_txt.strip():
            st.markdown(
                "<div style='border:1px solid rgba(0,0,0,.10);border-radius:14px;overflow:hidden;margin-top:6px;'>"
                "<div style='background:#e8f4ff;padding:10px 12px;'><span style='color:#b00020;font-weight:900;'>Clinical Scenario</span></div>"
                "<div style='background:#fff7d6;padding:12px 12px;color:#111;'><b>" + _html.escape(scenario_txt) + "</b></div>"
                "</div>",
                unsafe_allow_html=True,
            )

        # Helper: render a green table (supports HTML cells via dict {'_html': '<ul>..'})
        def _green_table(title: str, headers: list, rows2: list):
            def _cell(v):
                try:
                    if isinstance(v, dict) and v.get("_html") is not None:
                        return str(v.get("_html"))
                except Exception:
                    pass
                return _html.escape(str(v))
            th = "".join([f"<th style='padding:8px 10px;border:1px solid rgba(0,0,0,.08);background:#dcf3dc;color:#b00020;font-weight:900;text-align:left;'>{_html.escape(str(h))}</th>" for h in headers])
            trs = []
            for rr in (rows2 or []):
                tds = "".join([f"<td style='padding:8px 10px;border:1px solid rgba(0,0,0,.08);background:#e9f8e9;color:#111;vertical-align:top;'>{_cell(x)}</td>" for x in rr])
                trs.append(f"<tr>{tds}</tr>")
            html = (
                "<div style='margin-top:10px;border-radius:14px;overflow:hidden;border:1px solid rgba(0,0,0,.10);'>"
                f"<div style='background:#e8f4ff;padding:10px 12px;'><span style='color:#b00020;font-weight:900;'>{_html.escape(title)}</span></div>"
                "<table style='width:100%;border-collapse:collapse;font-size:0.95rem;'>"
                f"<thead><tr>{th}</tr></thead><tbody>{''.join(trs)}</tbody></table></div>"
            )
            st.markdown(html, unsafe_allow_html=True)

        # Helpers: clean saved structures into bullet lists (removes symbols like [{ }], quotes, dict keys)
        def _to_list(v):
            """Normalize stored values (lists/dicts/strings) into a flat list of human-readable items.

            This is used for Grade Center evidence views and exports, so it must aggressively
            strip programmatic artifacts (dict wrappers, JSON text, brackets, etc.) without
            touching any grading logic.
            """
            if v is None:
                return []
            if isinstance(v, (list, tuple, set)):
                return [x for x in v]

            if isinstance(v, dict):
                # Common wrappers produced by graders / serializers
                for k in ("selected", "selection", "chosen", "student_selected", "student_selection", "answers_selected"):
                    if k in v:
                        return _to_list(v.get(k))
                for k in ("items", "item", "options", "choices", "answers", "value", "values"):
                    if k in v:
                        return _to_list(v.get(k))

                # Generic single-field wrappers
                for k in ("text", "label", "name", "title"):
                    if k in v:
                        return _to_list(v.get(k))

                # Fall back: collect any string/list-like values, ignore empty rationale keys
                out = []
                for kk, vv in v.items():
                    if str(kk).lower() in ("rationale", "why") and (vv is None or str(vv).strip() == ""):
                        continue
                    if isinstance(vv, (list, tuple, set, dict, str, int, float, bool)):
                        out.extend(_to_list(vv))
                return out

            if isinstance(v, str):
                s = v.strip()
                if not s:
                    return []
                # If it's a serialized list/dict, try to parse it.
                if s.startswith("[") or s.startswith("{"):
                    try:
                        obj = ast.literal_eval(s)
                        return _to_list(obj)
                    except Exception:
                        try:
                            obj = json.loads(s)
                            return _to_list(obj)
                        except Exception:
                            pass
                return [s]

            return [str(v)]

        def _clean_items(v):
            """Return a de-duplicated list of clean bullet items.

            Rules:
            - Remove bracket/quote/dict artifacts.
            - Split on commas into separate bullet items (as requested).
            """
            out = []
            for it in _to_list(v):
                if it is None:
                    continue
                # Recurse if nested containers remain
                if isinstance(it, (list, tuple, set, dict)):
                    out.extend(_clean_items(it))
                    continue

                s = str(it).strip()
                if not s:
                    continue

                # Strip obvious artifact wrappers
                s = s.replace("**", "").replace("__", "")
                s = re.sub(r"^[\[\]\{\}\(\)\"\']+|[\[\]\{\}\(\)\"\']+$", "", s).strip()

                # If it still looks like a dict wrapper, try to extract quoted segments
                # e.g. {"selected":["a","b"],"rationale":""}
                if ("selected" in s and "[" in s and "]" in s) or (s.startswith("{") and s.endswith("}")):
                    try:
                        obj = json.loads(s)
                        out.extend(_clean_items(obj))
                        continue
                    except Exception:
                        pass

                # Split by commas into separate bullets (user requirement)
                parts = [p.strip() for p in s.split(",") if p.strip()]
                out.extend(parts if parts else [s])

            # de-duplicate while preserving order
            seen = set(); cleaned = []
            for s in out:
                if s not in seen:
                    seen.add(s)
                    cleaned.append(s)
            return cleaned

        def _bullets_html(v, bold=True):
            items = _clean_items(v)
            if not items:
                return {"_html": "—"}
            lis = []
            for s in items:
                esc = _html.escape(s)
                if bold:
                    lis.append(f"<li><b>{esc}</b></li>")
                else:
                    lis.append(f"<li>{esc}</li>")
            return {"_html": "<ul style='margin:0.1rem 0 0.1rem 1.2rem;'>" + "".join(lis) + "</ul>"}

        # --- Intake selections (saved) ---
        intake_bd = rec.get("intake_breakdown") if isinstance(rec.get("intake_breakdown"), dict) else {}
        intake_rows = []
        if isinstance(intake_bd, dict):
            # flexible structure: key -> value (list/str/bool)
            for k, v in intake_bd.items():
                if isinstance(v, (list, tuple)):
                    vv = "; ".join([str(x) for x in v if str(x).strip()])
                else:
                    vv = str(v)
                intake_rows.append([str(k), vv])
        if intake_rows:
            _green_table("Intake (student selections)", ["Component", "Student entry"], intake_rows)

                # --- A–E review (saved answers vs gold; no rationale) ---
        ans = rec.get("answers") if isinstance(rec.get("answers"), dict) else {}
        ae_tables = []  # keep structured for export (title, headers, rows)
        if isinstance(case_obj, dict):
            dom_meta = [("A", "Assessment"), ("B", "Prioritize"), ("C", "Intervention"), ("D", "Reassess")]
            for dk, dname in dom_meta:
                student_sel = ans.get(dk, [])
                # normalize list
                if not isinstance(student_sel, (list, tuple, set)):
                    student_sel = [student_sel] if str(student_sel).strip() else []
                gold = get_gs_list(case_obj, ["keyAssessments","assessment"]) if dk=="A" else \
                       get_gs_list(case_obj, ["priorities","prioritize"]) if dk=="B" else \
                       get_gs_list(case_obj, ["interventions"]) if dk=="C" else \
                       get_gs_list(case_obj, ["reassessment","reassess"])
                _, goldN = build_domain_options(dk, case_obj, gold, total=10, distractors=6)

                correct, wrong, missed = diff_selected_vs_gold(_clean_items(student_sel), _clean_items(goldN))
                rows2 = [
                    ["Correct selections", _bullets_html(correct, bold=False)],
                    ["Wrong selections", _bullets_html(wrong, bold=False)],
                    ["Correct options you missed", _bullets_html(missed, bold=False)],
                    ["Student selections (all)", _bullets_html(student_sel, bold=False)],
                    ["Gold standard targets", _bullets_html(goldN, bold=False)],
                ]
                title = f"{dk} — {dname} (review)"
                _green_table(title, ["Section", "Details"], rows2)
                ae_tables.append((title, ["Section", "Details"], rows2))

            # Domain E (SBAR)
            sbar_txt = ans.get("E") or ans.get("SBAR") or ""
            sbar_expected_obj = get_gs_sbar(case_obj)
            rows2 = [
                ["Student SBAR submission", {"_html": "<div style='white-space:pre-wrap;'>" + _html.escape(str(sbar_txt).strip() or "—") + "</div>"}],
                ["Expected / Typical SBAR (reference)", {"_html": "<div style='white-space:pre-wrap;'>" + _html.escape(str(sbar_expected_obj).strip() or "—") + "</div>"}],
            ]
            title = "E — SBAR (review)"
            _green_table(title, ["Section", "Details"], rows2)
            ae_tables.append((title, ["Section", "Details"], rows2))

        # --- NCLEX review (best-effort via nclex_items.json lookup) ---
        nclex_answers = rec.get("nclex_answers") if isinstance(rec.get("nclex_answers"), dict) else {}
        nclex_details = (rec.get("nclex") or {}).get("details") if isinstance((rec.get("nclex") or {}).get("details"), list) else []

        det_map = {}
        for d in nclex_details:
            if isinstance(d, dict) and d.get("qid"):
                det_map[str(d.get("qid"))] = d


        # Load item bank (case-specific) and JOIN by question_id (qid)
        # This is the professional LMS approach: attempt log stores qid; we resolve stem + golden options at runtime.
        bank_map = {}
        try:
            nclex_bank = load_nclex_items()  # normalized to {"cases": {case_id: {"items":[...]}}}
            case_pack = (nclex_bank.get("cases") or {}).get(str(rec.get("caseId",""))) or {}
            items_list = case_pack.get("items") if isinstance(case_pack, dict) else []
            if not isinstance(items_list, list):
                items_list = []
            bank_map = {
                str(it.get("id")): it
                for it in items_list
                if isinstance(it, dict) and str(it.get("id","")).strip()
            }
        except Exception:
            bank_map = {}


        def _resolve_opt_text(item: dict, val):
            # Try to map a letter/id/index to option text; otherwise return as-is string.
            opts = item.get("options") or item.get("choices") or item.get("answers") or []
            opt_map = {}
            if isinstance(opts, list):
                # list of dicts (id/text)
                if opts and isinstance(opts[0], dict):
                    for o in opts:
                        oid = str(o.get("id") or o.get("key") or o.get("label") or "").strip()
                        otx = str(o.get("text") or o.get("value") or o.get("option") or "").strip()
                        if oid:
                            opt_map[oid] = otx or oid
                else:
                    # list of strings -> map A,B,C...
                    letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
                    for i, s in enumerate(opts):
                        if i < len(letters):
                            opt_map[letters[i]] = str(s)
            def one(x):
                if x is None:
                    return ""
                if isinstance(x, (list, tuple, set)):
                    return ""  # handled elsewhere
                if isinstance(x, int) and opt_map:
                    # 0-based or 1-based index
                    if 0 <= x < len(opt_map):
                        pass
                sx = str(x).strip()
                if sx in opt_map and opt_map[sx]:
                    return str(opt_map[sx]).strip()
                # sometimes stored as {"id":"A","text":"..."}
                if isinstance(x, dict):
                    if x.get("text"):
                        return str(x.get("text")).strip()
                    if x.get("id") and str(x.get("id")).strip() in opt_map:
                        return str(opt_map[str(x.get("id")).strip()]).strip()
                return sx
            return one(val)

        def _resolve_multi(item: dict, val):
            items = _clean_items(val)
            out = []
            for v in items:
                txtv = _resolve_opt_text(item, v)
                if txtv:
                    out.append(txtv)
            return out

        qids = list(nclex_answers.keys())
        qids.sort()
        mcq_rows = []
        n = 1

        def _norm_answer(v):
            # normalize answer values for correctness comparison
            if isinstance(v, (list, tuple, set)):
                return sorted([str(x).strip() for x in v if str(x).strip() != ""])
            if isinstance(v, dict):
                for k in ("selected", "selection", "answer", "answers", "value", "values"):
                    if k in v:
                        return _norm_answer(v.get(k))
                return sorted([str(x).strip() for x in _clean_items(v)])
            if isinstance(v, bool):
                return ["True"] if v else ["False"]
            if v is None:
                return []
            s = str(v).strip()
            if s.lower() in ("true", "false"):
                return ["True"] if s.lower() == "true" else ["False"]
            return [s]

        for qid in qids:
            item = bank_map.get(str(qid)) or {}
            det = det_map.get(str(qid)) or {}

            # Prefer true stem text, avoid internal IDs
            stem = str(
                item.get("stem")
                or item.get("question")
                or item.get("prompt")
                or det.get("stem")
                or det.get("question")
                or det.get("prompt")
                or det.get("stem_text")
                or det.get("question_text")
                or det.get("text")
                or ""
            ).strip()
            if stem.lower().startswith("question id") or "question id" in stem.lower():
                # If a bad stem slipped in, try alternate keys
                stem = str(det.get("question_text") or det.get("stem_text") or det.get("text") or "").strip()

            # selected + correct (gold) values
            selected_raw = nclex_answers.get(qid)
            correct_raw = item.get("correct") if item else None
            if correct_raw in (None, "", []):
                correct_raw = det.get("correct") or det.get("answer") or det.get("correct_answer") or det.get("gold") or det.get("golden")

            sel_txts = _resolve_multi(item, selected_raw)
            gold_txts = _resolve_multi(item, correct_raw)

            # correctness flag (True/False) for quick review
            is_correct = (_norm_answer(selected_raw) == _norm_answer(correct_raw))
            corr_flag = "True" if is_correct else "False"

            mcq_rows.append([
                f"{n}",
                stem if stem else "—",
                _bullets_html(sel_txts, bold=False),
                corr_flag,
                _bullets_html(gold_txts, bold=False),
            ])
            n += 1

        if mcq_rows:
            _green_table("NCLEX review (questions)", ["#", "Question stem", "Student selection", "Correct?", "Golden option(s)"], mcq_rows)

        # Export: Evidence PDF / Word (matches on-screen "Saved answers and performance report")
        def _plain_from_cell(v):
            # Convert our cell representation to readable plain text with bullets.
            if isinstance(v, dict) and v.get("_html") is not None:
                s = str(v.get("_html"))
                # turn <li> into bullet lines
                s = re.sub(r"(?is)</li>\s*", "\n", s)
                s = re.sub(r"(?is)<li[^>]*>", "• ", s)
                s = re.sub(r"(?is)<br\s*/?>", "\n", s)
                s = re.sub(r"(?is)<[^>]+>", "", s)
                s = s.replace("&nbsp;", " ")
                try:
                    s = _html.unescape(s)
                except Exception:
                    pass
                return s.strip() or "—"
            return str(v).strip() if str(v).strip() else "—"

        def _make_evidence_tables():
            blocks = []
            if intake_rows:
                blocks.append(("Intake (student selections)", ["Component", "Student entry"], [[r[0], r[1]] for r in intake_rows]))
            for (title, hdrs, rows2) in (ae_tables or []):
                blocks.append((title, hdrs, [[_plain_from_cell(a), _plain_from_cell(b)] for (a,b) in rows2]))
            if mcq_rows:
                blocks.append(("NCLEX review (questions)", ["#", "Question stem", "Student selection", "Correct?", "Golden option(s)"],
                               [[_plain_from_cell(r[0]), _plain_from_cell(r[1]), _plain_from_cell(r[2]), _plain_from_cell(r[3]), _plain_from_cell(r[4])] for r in mcq_rows]))
            return blocks

        report_blocks = _make_evidence_tables()

        c1, c2 = st.columns([1, 1])
        with c1:
            try:
                # PDF with green tables + red headers (ReportLab)

                from io import BytesIO
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib import colors
                from reportlab.lib.pagesizes import letter, landscape

                def _rl_escape_multiline(t: str) -> str:
                    # ReportLab Paragraph supports a small HTML subset. Keep it safe + preserve line breaks.
                    t = "" if t is None else str(t)
                    t = _html.escape(t)
                    t = t.replace("\n", "<br/>")
                    return t

                bio = BytesIO()
                doc = SimpleDocTemplate(
                    bio,
                    pagesize=landscape(letter),
                    rightMargin=28,
                    leftMargin=28,
                    topMargin=28,
                    bottomMargin=28,
                )
                styles = getSampleStyleSheet()

                # Compact styles so wide tables fit on the page (while keeping the on-screen look).
                body_style = ParagraphStyle(
                    "BodySmall",
                    parent=styles["BodyText"],
                    fontName="Helvetica",
                    fontSize=8.2,
                    leading=10,
                    spaceAfter=0,
                )
                hdr_style = ParagraphStyle(
                    "HdrRed",
                    parent=styles["BodyText"],
                    fontName="Helvetica-Bold",
                    fontSize=8.6,
                    leading=10,
                    textColor=colors.HexColor("#B00020"),
                )

                story = []

                story.append(Paragraph("<b>Exam Attempt Evidence Report</b>", styles["Title"]))
                story.append(Spacer(1, 10))

                stu_user = str(rec.get("student_username", "") or "")
                stu_id = str(rec.get("student_id", "") or "")
                case_title = str(rec.get("caseTitle", "") or "")
                case_id = str(rec.get("caseId", "") or "")
                submitted = _format_dt_local(rec.get("submitted_at", "") or "")

                # Header block (as requested: name, ID, case, date)
                if stu_user.strip():
                    story.append(Paragraph(f"<b>Student:</b> {_html.escape(stu_user)}", styles["Normal"]))
                if stu_id.strip():
                    story.append(Paragraph(f"<b>ID:</b> {_html.escape(stu_id)}", styles["Normal"]))
                story.append(Paragraph(f"<b>Case:</b> {_html.escape(case_title)} ({_html.escape(case_id)})", styles["Normal"]))
                story.append(Paragraph(f"<b>Submitted:</b> {_html.escape(submitted)}", styles["Normal"]))
                story.append(Spacer(1, 10))

                if scenario_txt.strip():
                    story.append(Paragraph("<b>Clinical Scenario</b>", styles["Heading2"]))
                    story.append(Paragraph(_rl_escape_multiline(scenario_txt), body_style))
                    story.append(Spacer(1, 8))

                # Table style (match the green table theme)
                hdr_bg = colors.HexColor("#DCF3DC")
                row_bg = colors.HexColor("#E9F8E9")
                red = colors.HexColor("#B00020")

                def _col_widths(ncols: int, avail_w: float):
                    # Proportional widths tuned to common tables so nothing is clipped.
                    if ncols == 2:
                        props = [0.28, 0.72]
                    elif ncols == 5:
                        props = [0.05, 0.31, 0.21, 0.08, 0.35]
                    else:
                        props = [1.0 / max(1, ncols)] * ncols
                    # Normalize in case of rounding errors
                    tot = sum(props) or 1.0
                    props = [p / tot for p in props]
                    return [avail_w * p for p in props]

                avail_w = doc.width

                for (title, hdrs, rowsx) in report_blocks:
                    story.append(Paragraph(f"<b>{_html.escape(title)}</b>", styles["Heading2"]))

                    # Convert to Paragraphs for wrapping (prevents cutoff/truncation)
                    hdr_row = [Paragraph(_rl_escape_multiline(h), hdr_style) for h in (hdrs or [])]
                    body_rows = []
                    for r in (rowsx or []):
                        prow = []
                        for cell in (r or []):
                            prow.append(Paragraph(_rl_escape_multiline(cell), body_style))
                        body_rows.append(prow)

                    data = [hdr_row] + body_rows
                    ncols = len(hdr_row) if hdr_row else (len(data[0]) if data else 1)
                    cw = _col_widths(ncols, avail_w)

                    tbl = Table(
                        data,
                        colWidths=cw,
                        hAlign="LEFT",
                        repeatRows=1,
                        splitByRow=1,
                    )
                    ts = TableStyle([
                        ("BACKGROUND", (0, 0), (-1, 0), hdr_bg),
                        ("TEXTCOLOR", (0, 0), (-1, 0), red),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("BACKGROUND", (0, 1), (-1, -1), row_bg),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.Color(0, 0, 0, alpha=0.15)),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ])
                    tbl.setStyle(ts)
                    story.append(tbl)
                    story.append(Spacer(1, 10))

                doc.build(story)
                pdf_bytes2 = bio.getvalue()
                if pdf_bytes2:
                    st.download_button(
                        "⬇️ Download Evidence PDF (attempt)",
                        data=pdf_bytes2,
                        file_name=f"evidence_{rec.get('student_username','student')}_{rec.get('caseId','case')}.pdf",
                        mime="application/pdf",
                    )
            except Exception:
                pass

        with c2:
            try:
                from io import BytesIO
                from docx import Document
                from docx.shared import Pt
                from docx.oxml.ns import qn

                doc = Document()
                doc.add_heading("Exam Attempt Evidence Report", level=1)
                doc.add_paragraph(f"Student: {rec.get('student_username','')}")
                doc.add_paragraph(f"Case: {rec.get('caseTitle','')} ({rec.get('caseId','')})")
                doc.add_paragraph(f"Submitted: {_format_dt_local(rec.get('submitted_at',''))}")
                doc.add_paragraph("")

                if scenario_txt.strip():
                    doc.add_heading("Clinical Scenario", level=2)
                    p = doc.add_paragraph(scenario_txt)
                    if p.runs:
                        p.runs[0].bold = True

                for (title, hdrs, rowsx) in report_blocks:
                    doc.add_heading(title, level=2)
                    t = doc.add_table(rows=1, cols=len(hdrs))
                    hdr = t.rows[0].cells
                    for i, h in enumerate(hdrs):
                        hdr[i].text = str(h)
                    for rr in rowsx:
                        row = t.add_row().cells
                        for i in range(len(hdrs)):
                            row[i].text = str(rr[i])
                    doc.add_paragraph("")

                bio = BytesIO()
                doc.save(bio)
                docx_bytes = bio.getvalue()
                if docx_bytes:
                    st.download_button(
                        "⬇️ Download Evidence Word (attempt)",
                        data=docx_bytes,
                        file_name=f"evidence_{rec.get('student_username','student')}_{rec.get('caseId','case')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
            except Exception:
                pass

if effective_admin:
    # Backward-compatible mapping (older saved values -> new labels)
    _old = st.session_state.get("admin_pages_v6")
    _map = {
        "Run App": "🏠 Run App",
                "🔎 Attempt Search": "🔎 Attempt Search",
        "📈 Item Analytics": "📈 Item Analytics",
        "🧹 Data Tools": "🧹 Data Tools",
        "Grade Center": "🎓 Grade Center",
    }
    if _old in _map:
        st.session_state["admin_pages_v6"] = _map[_old]

    admin_page = st.radio(
        "Admin",
        ["🏠 Run App", "🎓 Grade Center", "🔎 Attempt Search", "👤 Student Manager", "🔒 Exam Control", "📈 Item Analytics", "🧹 Data Tools", "⚙️ Settings"],
        horizontal=True,
        key="admin_pages_v6",
        label_visibility="collapsed",
    )
    st.markdown("---")
    if admin_page != "🏠 Run App":
        if admin_page.startswith("🎓"):
            render_grade_center_page()
        elif admin_page.startswith("🔎"):
            admin_page_attempt_search()
        elif admin_page.startswith("👤"):
            admin_page_student_session_manager()
        elif admin_page.startswith("🔒"):
            admin_page_exam_control()
        elif admin_page.startswith("📈"):
            admin_page_item_analytics()
        elif admin_page.startswith("🧹"):
            admin_page_data_tools()
        else:
            admin_page_settings()  # defined below
        st.stop()



# =============================
# Step 6 navigation filters (Main panel)
# =============================
# NOTE: student_username is used by the Progress/Resume panel below and must be
# defined before rendering the Case Selection UI.
student_username = (
    (st.session_state.get("student_username") or "")
    or (student_profile.get("username", "") if isinstance(student_profile, dict) and student_profile else "")
    or ""
).strip()

with st.expander("🗂 Case Selection", expanded=True):

    # =============================
    # 📌 Student Progress + Continue (autosave-based)
    # =============================
    if student_username and not st.session_state.get("_progress_index_loaded", False):
        st.session_state["_autosave_index"] = index_latest_autosaves_for_student(student_username)
        st.session_state["_progress_index_loaded"] = True

    autosave_index = st.session_state.get("_autosave_index", {}) if isinstance(st.session_state.get("_autosave_index", {}), dict) else {}
    # Build progress lists for current visible cases (after visibility filter is applied below).
    c1, c2, c3 = st.columns([2,1,1])
    with c1:
        search = st.text_input("Search case title", "", key="case_search_main")
    systems = sorted({safe_get_system(c) for c in cases})
    settings = sorted({safe_get_setting(c) for c in cases})
    with c2:
        system_sel = st.selectbox("System", ["All"] + systems, key="case_system_main")
    with c3:
        setting_sel = st.selectbox("Setting", ["All"] + settings, key="case_setting_main")

    # Visibility filter for students
    visible_filtered_cases = []
    for c in cases:
        cid = str(c.get("id", "")).strip()
        if is_admin:
            visible_filtered_cases.append(c)
        else:
            if case_is_visible(case_policy, cid, mode=admin_settings.get('app_mode','Practice')):
                visible_filtered_cases.append(c)

    filtered = visible_filtered_cases


    def _display_label(c):
        cid = str(c.get("id", "")).strip()
        hidden_badge = ""
        if is_admin and not case_is_visible(case_policy, cid, mode=admin_settings.get("app_mode", "Practice")):
            hidden_badge = " (HIDDEN)"
        title = str(c.get("title", "(No title)"))
        return f"{title} — [{cid}]{hidden_badge}"

    # Progress UI (students + admins) based on autosave drafts
    # Note: uses only autosave data; does NOT reveal correct answers.
    in_progress_opts = []
    completed_count = 0
    inprog_count = 0
    for c in visible_filtered_cases:
        cid0 = str(c.get("id","")).strip()
        d0 = autosave_index.get(cid0)
        status0, hint0 = compute_resume_hint_from_draft(d0 or {})
        if status0 == "Completed":
            completed_count += 1
        elif status0 in ["In progress", "Submitted"]:
            inprog_count += 1
            in_progress_opts.append((cid0, hint0, _display_label(c)))

    with st.expander("📌 My Progress", expanded=False):
        st.caption(f"Completed: **{completed_count}**  •  In progress: **{inprog_count}**")
        if in_progress_opts:
            labels = [lab for (_cid,_hint,lab) in in_progress_opts]
            pick_lab = st.selectbox("Continue an unfinished case", labels, key="progress_continue_pick")
            if st.button("▶ Continue", key="progress_continue_btn"):
                # Request resume: switch the selected case and restore last autosave on next run
                sel_i = labels.index(pick_lab)
                sel_cid, sel_hint, sel_label = in_progress_opts[sel_i]
                st.session_state["_resume_case_id"] = sel_cid
                st.session_state["_resume_hint"] = sel_hint
                st.session_state["case_pick_main"] = sel_label
                # refresh autosave index so we use the latest snapshot
                st.session_state["_autosave_index"] = index_latest_autosaves_for_student(student_username)
                st.session_state["_progress_index_loaded"] = True
                st.session_state["_resume_pending"] = True
                st.rerun()
        else:
            st.info("No unfinished cases yet. Start any case below to create autosave progress.")
    if system_sel != "All":
        filtered = [c for c in filtered if safe_get_system(c) == system_sel]
    if setting_sel != "All":
        filtered = [c for c in filtered if safe_get_setting(c) == setting_sel]
    if search.strip():
        s = search.lower().strip()
        filtered = [c for c in filtered if s in (c.get("title", "") or "").lower()]

    if not filtered:
        st.warning("No cases match your filters (or all are hidden).")
        if not is_admin:
            st.info("Ask the administrator to set cases as visible in case_policy.json (Settings → Case Visibility).")
        st.stop()

    def _display_label(c):
        cid = str(c.get("id", "")).strip()
        hidden_badge = ""
        if is_admin and not case_is_visible(case_policy, cid, mode=admin_settings.get('app_mode','Practice')):
            hidden_badge = " (HIDDEN)"
        return f'{c.get("title", "(No title)")} — [{cid}]{hidden_badge}'

    selected_label = st.selectbox("Case", [_display_label(c) for c in filtered], key="case_pick_main")
    case = next(c for c in filtered if _display_label(c) == selected_label)
case_id = str(case.get("id", "")).strip() or "no_id"

# =============================
# Resume from autosave (if requested via Progress panel)
# =============================
if st.session_state.get("_resume_pending", False) and str(st.session_state.get("_resume_case_id","")).strip() == str(case_id):
    draft = load_last_autosave(student_username, case_id)
    if draft:
        apply_restored_draft(draft)
        # If we can infer the next focus, set it (keeps UI aligned)
        hint = st.session_state.get("_resume_hint")
        if hint in ["A","B","C","D","E"]:
            st.session_state["ae_focus"] = hint
        st.session_state["_resume_pending"] = False
        st.session_state["_resume_case_id"] = ""
        st.session_state["_resume_hint"] = ""
        st.success("Resumed your saved progress for this case.")
        st.rerun()
    else:
        st.session_state["_resume_pending"] = False


# Enforce availability/visibility after selection (defense in depth)
if not is_admin:
    mode_now = admin_settings.get("app_mode", "Practice")
    if not case_is_visible(case_policy, case_id, mode=mode_now, now_dt=now_local()):
        st.error("This case is not currently available to students (hidden or outside the allowed time window).")
        st.stop()


# =============================
# Step 6 (Exam integrity): lock case switching in EXAM (students)
# =============================
# Apply admin overrides (force unlock / reset) for this student session (server-side)
student_username = student_profile.get("username", "") if student_profile else ""
try:
    apply_exam_overrides_for_student_session(student_username)
except Exception:
    pass

if features.get("lock_case_switch_exam", False) and admin_settings.get("app_mode") == "Exam" and not is_admin and is_student_logged_in:
    if "locked_exam_case_id" not in st.session_state:
        st.session_state["locked_exam_case_id"] = case_id
    else:
        if st.session_state["locked_exam_case_id"] != case_id:
            st.error("🔒 Exam mode: Case switching is locked. Please continue the started case.")
            st.stop()

# =============================
# Attempt limit enforcement
# =============================
student_username = student_profile.get("username", "") if student_profile else ""
# Rebuild autosave index on login/user change
if st.session_state.get("_autosave_index_user") != str(student_username or ""):
    st.session_state["_autosave_index_user"] = str(student_username or "")
    st.session_state["_progress_index_loaded"] = False
    st.session_state["_autosave_index"] = {}


# ✅ Basic exam protection (deterrent only)
if student_profile:
    # Best-effort deterrents against copy/print in BOTH practice and exam modes + watermark.
    try:
        ts = now_local().strftime("%Y-%m-%d %H:%M")
    except Exception:
        ts = ""
    disp = str(student_profile.get("display_name") or student_profile.get("username") or "").strip()
    sid = str(student_profile.get("student_id") or "").strip()
    wm = " | ".join([x for x in [disp, sid, ts] if x])
    try:
        pol = load_nclex_policy()
        wm_use = wm if bool(pol.get("watermark_enabled", True)) else ""
    except Exception:
        wm_use = wm
    inject_basic_exam_protection(bool(pol.get("exam_protection_enabled", True)), watermark_text=wm_use)

student_display_name = student_profile.get("display_name", "") if student_profile else ""
student_id = student_profile.get("student_id", "") if student_profile else ""
student_cohort = student_profile.get("cohort", "") if student_profile else ""

attempted = attempts_count_for(student_username, case_id) if student_username else 0
max_attempts = max_attempts_for_case(case_id)
max_attempts_int = None if str(max_attempts) == "unlimited" else safe_int(max_attempts, None)
remaining = None if max_attempts_int is None else max(0, max_attempts_int - attempted)

if is_student_logged_in and not is_admin:
    if max_attempts_int is None:
        st.sidebar.info(f"Attempts for this case: **{attempted}** / **Unlimited**")
    else:
        st.sidebar.info(f"Attempts for this case: **{attempted}** / **{max_attempts_int}** (Remaining: **{remaining}**)")

locked_out = (max_attempts_int is not None and attempted >= max_attempts_int and not is_admin)
if locked_out:
    st.error("🚫 Attempt limit reached for this case. Please contact faculty.")
    st.stop()

# =============================
# Gold standard targets + UI options
# =============================
gs_assess = get_gs_list(case, ["keyAssessments", "assessment"])
gs_prio = get_gs_list(case, ["priorities", "prioritize"])
gs_inter = get_gs_list(case, ["interventions"])
gs_reass = get_gs_list(case, ["reassessment", "reassess"])
gs_sbar = get_gs_sbar(case)

# Build distractor-enriched options (gold targets + distractors)
# NOTE: Some cases have 3 gold items; selection requirement adapts per case.
ui_assess, gold4_assess = build_domain_options('A', case, gs_assess, total=10, distractors=6)
ui_prio, gold4_prio = build_domain_options('B', case, gs_prio, total=10, distractors=6)
ui_inter, gold4_inter = build_domain_options('C', case, gs_inter, total=10, distractors=6)
ui_reass, gold4_reass = build_domain_options('D', case, gs_reass, total=10, distractors=6)

# Determine required selection counts per domain based on each case's gold-standard targets (fair scoring).
req_A = max(1, len(gold4_assess)) if isinstance(gold4_assess, list) and len(gold4_assess) > 0 else 4
req_B = max(1, len(gold4_prio)) if isinstance(gold4_prio, list) and len(gold4_prio) > 0 else 4
req_C = max(1, len(gold4_inter)) if isinstance(gold4_inter, list) and len(gold4_inter) > 0 else 4
req_D = max(1, len(gold4_reass)) if isinstance(gold4_reass, list) and len(gold4_reass) > 0 else 4


# Force Reassess to require exactly 2 selections when options are available (per course design)
if isinstance(ui_reass, list) and len(ui_reass) >= 2:
    req_D = 2
# If a case does not provide reassessment options, allow submission via notes (do not block workflow).
if isinstance(ui_reass, list):
    if len(ui_reass) == 0:
        req_D = 0
    elif req_D > len(ui_reass):
        req_D = len(ui_reass)

# Store max points so the summary can display cleanly
st.session_state["domain_max"] = {"A": req_A, "B": req_B, "C": req_C, "D": req_D, "E": 4}

# Randomize option order per student/session (deterrent; helps reduce sharing)
try:
    nclex_policy_tmp = load_nclex_policy()
    rand_student = bool(nclex_policy_tmp.get("randomize_per_student_session", True))
    seed_base = f"{student_username}|{case_id}|{st.session_state.get('attempt_started_epoch') or ''}|{admin_settings.get('app_mode','Practice')}"
    if rand_student:
        ui_assess = shuffle_if_needed(ui_assess, seed_str=seed_base + "|A", enabled=True)
        ui_prio   = shuffle_if_needed(ui_prio,   seed_str=seed_base + "|B", enabled=True)
        ui_inter  = shuffle_if_needed(ui_inter,  seed_str=seed_base + "|C", enabled=True)
        ui_reass  = shuffle_if_needed(ui_reass,  seed_str=seed_base + "|D", enabled=True)
except Exception:
    pass



# SBAR elements (for graded selection) + keep text boxes for practice writing
sbar_gold = [x for x in [gs_sbar.get('S',''), gs_sbar.get('B',''), gs_sbar.get('A',''), gs_sbar.get('R','')] if str(x).strip()]
ui_sbar_opts, gold4_sbar = build_domain_options('E', case, sbar_gold, total=10, distractors=6)

# Expected SBAR (used in rationales/feedback). Derived from this case's SBAR gold selections.
try:
    sbar_expected = '; '.join([str(x).strip() for x in (gold4_sbar or []) if str(x).strip()])
except Exception:
    sbar_expected = ''

# Determine required selection count for SBAR based on this case's gold-standard targets (fair scoring).
req_E = max(1, len(gold4_sbar)) if isinstance(gold4_sbar, list) and len(gold4_sbar) > 0 else 4
# Update domain max for SBAR now that gold targets are known
try:
    dm = st.session_state.get("domain_max", {}) or {}
    dm["E"] = req_E
    st.session_state["domain_max"] = dm
except Exception:
    pass

# Randomize SBAR option order per student/session
try:
    nclex_policy_tmp = load_nclex_policy()
    rand_student = bool(nclex_policy_tmp.get("randomize_per_student_session", True))
    seed_base = f"{student_username}|{case_id}|{st.session_state.get('attempt_started_epoch') or ''}|{admin_settings.get('app_mode','Practice')}"
    if rand_student:
        ui_sbar_opts = shuffle_if_needed(ui_sbar_opts, seed_str=seed_base + "|E", enabled=True)
except Exception:
    pass


with st.expander("📘 Rubric (How you will be graded)", expanded=False):
    st.markdown(score_explainer_markdown())

# =============================
# Timer: establish deadline for this attempt + show it
# =============================
mode = admin_settings.get("app_mode", "Practice")
timer_minutes = timer_minutes_for(case_policy, case_id, mode)

# Ensure attempt state exists + aligns with case
current_case_id = case_id
if "active_case_id" not in st.session_state:
    st.session_state["active_case_id"] = current_case_id
    reset_attempt_state()
    st.session_state["attempt_deadline_epoch"] = None
    st.session_state["attempt_started_epoch"] = st.session_state.get("attempt_started_epoch") or time.time()
elif st.session_state["active_case_id"] != current_case_id:
    st.session_state["active_case_id"] = current_case_id
    reset_attempt_state()
    st.session_state["attempt_deadline_epoch"] = None
    st.session_state["attempt_started_epoch"] = st.session_state.get("attempt_started_epoch") or time.time()

# Set deadline once per attempt if timed
if st.session_state.get("attempt_deadline_epoch") is None:
    if timer_minutes != "unlimited":
        # Ensure timer keys exist (avoid KeyError)
        started = st.session_state.get("attempt_started_epoch", None)
        if started is None:
            started = time.time()
            st.session_state["attempt_started_epoch"] = started
        if "attempt_deadline_epoch" not in st.session_state:
            st.session_state["attempt_deadline_epoch"] = None
        st.session_state["attempt_deadline_epoch"] = started + (int(timer_minutes) * 60)
# Compute remaining time
deadline = st.session_state.get("attempt_deadline_epoch")
time_left_sec = None
expired = False
if deadline is not None:
    time_left_sec = int(deadline - time.time())
    expired = time_left_sec <= 0

with st.sidebar.expander("⏱ Case Timer", expanded=True):
    if timer_minutes == "unlimited" or deadline is None:
        st.write("Unlimited (default)")
    else:
        st.write(f"Limit: **{timer_minutes} min**")
        st.write("Remaining:", f"**{format_seconds(max(0, time_left_sec))}**")
        if expired and not is_admin:
            st.error("Time is up (Exam submissions locked).")

# In EXAM, lock actions on expiry for students (admin can still test)
timer_lock = (expired and (timer_minutes != "unlimited") and (not is_admin))

# =============================
# Admin-only: ✅ NCLEX Validator + Auto-fix (Step 4)
# =============================
if is_admin and features.get("nclex_validator", True) and SHOW_ADMIN_PANELS_IN_SIDEBAR:
    with st.sidebar.expander("✅ NCLEX Validator (IDs + Auto-fix)", expanded=False):
        nclex = load_nclex_items()
        nclex_case_keys = sorted(list((nclex.get("cases") or {}).keys()))
        case_ids_all = sorted([str(c.get("id", "")).strip() for c in cases if str(c.get("id", "")).strip()])

        st.write("Cases in cases.json:", len(case_ids_all))
        st.write("Cases in nclex_items.json:", len(nclex_case_keys))

        missing = [cid for cid in case_ids_all if cid not in (nclex.get("cases") or {})]
        extra = [cid for cid in nclex_case_keys if cid not in set(case_ids_all)]

        if missing:
            st.error(f"Missing NCLEX packs for {len(missing)} case IDs")
            st.caption("Example missing IDs:")
            st.code("\n".join(missing[:10]))
        else:
            st.success("No missing NCLEX case packs ✅")

        if extra:
            st.warning(f"Extra NCLEX packs not found in cases.json: {len(extra)}")
            st.caption("Example extra IDs:")
            st.code("\n".join(extra[:10]))

        if st.button("🛠 Auto-create missing NCLEX packs (empty)", key="nclex_autofix_btn"):
            nclex.setdefault("cases", {})
            for cid in missing:
                nclex["cases"][cid] = {"items": []}
            save_nclex_items(nclex)
            st.success("Created empty packs for missing cases. Now paste/generate items for those cases.")
            st.rerun()


with st.expander("🛠️ Debug (click to open)", expanded=False):
    st.write("BASE_DIR:", str(BASE_DIR))
    st.write("cases.json exists:", CASES_PATH.exists())
    st.write("students.json exists:", STUDENTS_PATH.exists())
    st.write("attempts_policy.json exists:", ATTEMPT_POLICY_PATH.exists())
    st.write("admin_settings.json exists:", ADMIN_SETTINGS_PATH.exists())
    st.write("case_policy.json exists:", CASE_POLICY_PATH.exists())
    st.write("exam_access_policy.json exists:", EXAM_ACCESS_POLICY_PATH.exists())
    st.write("attempts_log.jsonl exists:", ATTEMPTS_PATH.exists())
    st.write("nclex_items.json exists:", NCLEX_ITEMS_PATH.exists())
    st.write("nclex_policy.json exists:", NCLEX_POLICY_PATH.exists())
    st.write("features.json exists:", FEATURES_PATH.exists())
    st.write("OPENAI_API_KEY set:", bool(os.environ.get("OPENAI_API_KEY", "").strip()))
    st.write("Mode:", admin_settings.get("app_mode"))
    st.write("AI enabled:", admin_settings.get("ai_enabled"), "AI debrief:", admin_settings.get("ai_debrief_enabled"))
    st.write("default_visibility:", case_policy.get("default_visibility"))
    st.write("default timers:", case_policy.get("default_timer_minutes_practice"), case_policy.get("default_timer_minutes_exam"))
    pe = load_exam_access_policy()
    st.write("Exam policy enabled:", pe.get("enabled"), "active_now:", is_exam_password_active(pe), "expires_at:", pe.get("expires_at"))
    st.write("is_admin:", is_admin, "student_logged_in:", is_student_logged_in)
    st.write("features:", features)

# =============================
# Tabs restore
# =============================
if effective_admin:
    inspector_tab, student_tab, analytics_tab, editor_tab = st.tabs(["🧑‍🏫 Admin View", "👩‍⚕️ Student View", "📊 Analytics", "🧩 Case Editor"])
else:
    inspector_tab = None
    student_tab = st.container()
    analytics_tab = None
    editor_tab = None

# =============================
# Admin View
# =============================
if effective_admin:
    with inspector_tab:
        st.subheader("🧑‍🏫 Admin View (Faculty)")

        st.markdown(
            "**Research-Grade Data (bonus)**\n\n"
            "Because answers are stored structurally, you can:\n"
            "- Analyze reasoning patterns\n"
            "- Compare cohorts\n"
            "- Publish education research\n"
            "- Defend AI-assisted learning outcomes\n"
        )

        st.divider()
        st.subheader("📤 Faculty Export (CSV)")
        st.caption("Exports all attempts from attempts_log.jsonl into a grading-ready CSV.")

        if st.button("🧾 Build CSV export file now"):
            cases_by_id = {str(c.get("id","")): c for c in cases}
            nclex_bank = load_nclex_items()
            nclex_cases = (nclex_bank.get("cases") or {})
            rows = []
            qrows = []  # detailed NCLEX per-question rows
            for rec in iter_attempts():
                cid = str(rec.get("caseId", ""))
                case = cases_by_id.get(cid, {}) if isinstance(cases_by_id, dict) else {}
                answers = rec.get("answers") or {}
        
                # --- A–E item analysis (gold vs selected) ---
                def _gold4(lst):
                    lst = lst or []
                    return list(lst)[:4] if isinstance(lst, list) else []
                gs_assess = _gold4(get_gs_list(case, ["keyAssessments", "assessment"]))
                gs_prio   = _gold4(get_gs_list(case, ["priorities", "prioritize"]))
                gs_inter  = _gold4(get_gs_list(case, ["interventions"]))
                gs_reass  = _gold4(get_gs_list(case, ["reassessment", "reassess"]))
                gs_sbar_d = get_gs_sbar(case) or {}
                gs_sbar   = _gold4([x for x in [gs_sbar_d.get("S",""), gs_sbar_d.get("B",""), gs_sbar_d.get("A",""), gs_sbar_d.get("R","")] if str(x).strip()])
        
                def _sel(dom):
                    a = answers.get(dom) or {}
                    if isinstance(a, dict):
                        return a.get("selected") or []
                    return []
        
                def _stats(selected, gold):
                    selected = list(selected or [])
                    gold = list(gold or [])
                    correct = [x for x in selected if x in gold]
                    wrong = [x for x in selected if x not in gold]
                    missed = [x for x in gold if x not in selected]
                    return len(correct), len(wrong), len(missed)
        
                a_ok,a_wrong,a_miss = _stats(_sel("A"), gs_assess)
                b_ok,b_wrong,b_miss = _stats(_sel("B"), gs_prio)
                c_ok,c_wrong,c_miss = _stats(_sel("C"), gs_inter)
                d_ok,d_wrong,d_miss = _stats(_sel("D"), gs_reass)
                e_ok,e_wrong,e_miss = _stats(_sel("E"), gs_sbar)
        
                # --- NCLEX item analysis ---
                nblob = rec.get("nclex") or {}
                ndetails = nblob.get("details") or []
                n_answers = rec.get("nclex_answers") or {}
                n_correct = 0
                n_total = 0
                for d in ndetails:
                    n_total += 1
                    if bool(d.get("correct")) or (str(d.get("points")) == str(d.get("max"))):
                        n_correct += 1
        
                row = {
                    "timestamp": rec.get("timestamp", ""),
                    "duration_seconds": rec.get("duration_seconds", ""),
                    "mode": rec.get("mode", ""),
                    "timer_minutes": rec.get("timer_minutes", ""),
                    "timer_expired": rec.get("timer_expired", ""),
                    "student_username": rec.get("student_username", ""),
                    "student_display_name": rec.get("student_display_name", ""),
                    "student_id": rec.get("student_id", ""),
                    "cohort": rec.get("student_cohort", ""),
                    "caseId": rec.get("caseId", ""),
                    "caseTitle": rec.get("caseTitle", ""),
                    "system": rec.get("system", ""),
                    "setting": rec.get("setting", ""),
                    "A": rec.get("scores", {}).get("A", ""),
                    "B": rec.get("scores", {}).get("B", ""),
                    "C": rec.get("scores", {}).get("C", ""),
                    "D": rec.get("scores", {}).get("D", ""),
                    "E": rec.get("scores", {}).get("E", ""),
                    "total": rec.get("total", ""),
                    "unsafe_total": rec.get("unsafe_total", ""),
                    "attempt_number_for_case": rec.get("attempt_number_for_case", ""),
                    "A_correct": a_ok, "A_wrong": a_wrong, "A_missed": a_miss,
                    "B_correct": b_ok, "B_wrong": b_wrong, "B_missed": b_miss,
                    "C_correct": c_ok, "C_wrong": c_wrong, "C_missed": c_miss,
                    "D_correct": d_ok, "D_wrong": d_wrong, "D_missed": d_miss,
                    "E_correct": e_ok, "E_wrong": e_wrong, "E_missed": e_miss,
                    "nclex_points": (nblob or {}).get("total_points", ""),
                    "nclex_max": (nblob or {}).get("total_max", ""),
                    "nclex_correct": n_correct,
                    "nclex_total": n_total,
                }
                rows.append(row)
        
                item_map = {}
                try:
                    pack = nclex_cases.get(cid, {}) or {}
                    for it in (pack.get("items") or []):
                        if isinstance(it, dict) and it.get("id"):
                            item_map[str(it.get("id"))] = it
                except Exception:
                    item_map = {}
        
                for d in ndetails:
                    qid = str(d.get("qid", ""))
                    meta = item_map.get(qid, {})
                    qrows.append({
                        "timestamp": rec.get("timestamp", ""),
                        "student_username": rec.get("student_username", ""),
                        "student_id": rec.get("student_id", ""),
                        "cohort": rec.get("student_cohort", ""),
                        "caseId": cid,
                        "caseTitle": rec.get("caseTitle", ""),
                        "qid": qid,
                        "type": d.get("type", ""),
                        "difficulty": meta.get("difficulty", ""),
                        "client_need": meta.get("client_need", ""),
                        "topic": meta.get("topic", ""),
                        "points": d.get("points", ""),
                        "max": d.get("max", ""),
                        "correct": d.get("correct", ""),
                        "student_answer": json.dumps(n_answers.get(qid, ""), ensure_ascii=False),
                        "correct_answer": json.dumps(meta.get("correct", ""), ensure_ascii=False),
                    })
        
            out = io.StringIO()
            if rows:
                w = csv.DictWriter(out, fieldnames=list(rows[0].keys()))
                w.writeheader()
                w.writerows(rows)
            ATTEMPTS_CSV_PATH.write_text(out.getvalue(), encoding="utf-8")
            st.success(f"CSV built: {ATTEMPTS_CSV_PATH.name} ({len(rows)} rows)")
        
            out2 = io.StringIO()
            if qrows:
                w2 = csv.DictWriter(out2, fieldnames=list(qrows[0].keys()))
                w2.writeheader()
                w2.writerows(qrows)
            NCLEX_ITEM_CSV_PATH.write_text(out2.getvalue(), encoding="utf-8")
            st.info(f"NCLEX item analysis built: {NCLEX_ITEM_CSV_PATH.name} ({len(qrows)} rows)")
        if ATTEMPTS_CSV_PATH.exists():
            st.download_button(
                "⬇️ Download attempts_export.csv",
                data=ATTEMPTS_CSV_PATH.read_bytes(),
                file_name="attempts_export.csv",
                mime="text/csv"
            )
        if NCLEX_ITEM_CSV_PATH.exists():
            st.download_button(
                "⬇️ Download nclex_item_analysis.csv",
                data=NCLEX_ITEM_CSV_PATH.read_bytes(),
                file_name="nclex_item_analysis.csv",
                mime="text/csv"
            )


        st.divider()
        st.subheader("🧾 NCLEX Policy Quick Controls")
        nclex_policy = load_nclex_policy()
        nclex_enabled = st.toggle("Enable NCLEX-style practice", value=bool(nclex_policy.get("enabled", False)))
        nclex_policy["enabled"] = bool(nclex_enabled)

        st.caption("Enable/disable NCLEX item types (advanced)")
        et = nclex_policy.get("enabled_types") or {}
        cols = st.columns(3)
        for i, tname in enumerate(["mcq", "sata", "ordered_response", "cloze", "matrix", "evolving_case"]):
            with cols[i % 3]:
                et[tname] = st.toggle(f"Enable {tname}", value=bool(et.get(tname, True)), key=f"nclex_type_{tname}")
        nclex_policy["enabled_types"] = et

        nclex_policy["show_correct_answers_after_submit"] = st.toggle(
            "Show correct answers after practical submit (Practice policy)",
            value=bool(nclex_policy.get("show_correct_answers_after_submit", True))
        )
        nclex_policy["show_rationales_after_submit"] = st.toggle(
            "Show rationales after practical submit (Practice policy)",
            value=bool(nclex_policy.get("show_rationales_after_submit", True))
        )
        nclex_policy["shuffle_options"] = st.toggle(
            "Shuffle options (where safe)",
            value=bool(nclex_policy.get("shuffle_options", True))
        )

        if st.button("💾 Save NCLEX policy"):
            NCLEX_POLICY_PATH.write_text(json.dumps(nclex_policy, ensure_ascii=False, indent=2), encoding="utf-8")
            st.success("Saved nclex_policy.json")
    
# --- NCLEX Rotation (Admin-controlled) ---
if effective_admin and SHOW_ADMIN_PANELS_IN_SIDEBAR:
    with st.sidebar.expander("🎛 NCLEX question bank rotation (Admin)", expanded=False):
        pol = load_nclex_policy()
        rot_enabled = st.checkbox("Enable rotation using an admin-generated active set (per case)", value=bool(pol.get("rotation_enabled", False)), key="nclex_rotation_enabled")
        rand_student = st.checkbox("Randomize NCLEX item order per student/session", value=bool(pol.get("randomize_per_student_session", True)), key="nclex_rand_student")
        one_at_time = st.checkbox("Student view: 1 NCLEX question at a time", value=bool(pol.get("one_question_at_a_time", False)), key="nclex_one_at_time")
        footer_enabled = st.checkbox("Show tiny footer session code (student view)", value=bool(pol.get("footer_session_code_enabled", True)), key="nclex_footer_enabled")
        if st.button("💾 Save NCLEX security/rotation settings", key="save_nclex_rot_settings"):
            pol["rotation_enabled"] = bool(rot_enabled)
            pol["randomize_per_student_session"] = bool(rand_student)
            pol["one_question_at_a_time"] = bool(one_at_time)
            pol["footer_session_code_enabled"] = bool(footer_enabled)
            # watermark removed
            pol["watermark_enabled"] = False
            NCLEX_POLICY_PATH.write_text(json.dumps(pol, ensure_ascii=False, indent=2), encoding="utf-8")
            st.success("Saved.")
            st.rerun()

        st.caption("Rotation uses the NCLEX bank already in nclex_items.json. You can generate a new active 30-question set per case at any time.")
        cid_pick = st.selectbox("Select case to generate a new active NCLEX set", options=[""] + case_ids_all, index=0, key="rot_case_pick")
        if cid_pick:
            pack = (nclex.get("cases") or {}).get(cid_pick) or {}
            bank_items = list(pack.get("items", []) or [])
            bank_items = [it for it in bank_items if (pol.get("enabled_types") or {}).get(it.get("type"), True)] if isinstance(pol.get("enabled_types"), dict) else bank_items
            st.write(f"Bank size for {cid_pick}: **{len(bank_items)}** items (after type filters).")
            active_sets = load_nclex_active_sets()
            active = (active_sets.get("by_case") or {}).get(cid_pick) or {}
            active_qids = active.get("qids", []) if isinstance(active, dict) else []
            if active_qids:
                st.info(f"Active set currently has {len(active_qids)} items (generated_at: {active.get('generated_at','')}).")
            else:
                st.warning("No active set generated yet for this case. (Defaults will be used until you generate one.)")

            # Preview tools (do not modify data unless you click Generate)
            with st.expander("🔎 Preview ACTIVE / candidate 30", expanded=False):
                k = nclex_items_per_case(pol, cid_pick)
                by_id = {str(it.get("id","")): it for it in bank_items}
                # Active set preview
                if active_qids:
                    st.markdown("**Active set (currently used by students):**")
                    active_items = [by_id[q] for q in active_qids if q in by_id][:k]
                    st.write(f"Items: {len(active_items)}/{k}")
                    for ii, it in enumerate(active_items, start=1):
                        st.markdown(f"**{ii}. {it.get('id','')}** ({it.get('type','')}, {it.get('difficulty','')})")
                        st.write(it.get('stem') or it.get('prompt') or '')
                        st.write("✅ Correct:", it.get('correct'))
                        if it.get('rationale'):
                            st.write("🧠 Rationale:", it.get('rationale'))
                        st.divider()
                else:
                    st.markdown("**No active set yet.** Students will see the default first 30 (after filters) until you generate one.")

                # Candidate preview (not saved)
                if st.button("🔁 Preview random candidate (no save)", key="preview_candidate_btn"):
                    qids_all = [str(it.get('id','')) for it in bank_items if str(it.get('id','')).strip()]
                    rnd_seed = sha256_hex(f"preview|{cid_pick}|{time.time()}|{secrets.token_hex(4)}")
                    rnd = random.Random(int(rnd_seed[:8], 16))
                    if len(qids_all) <= k:
                        cand = list(qids_all); rnd.shuffle(cand)
                    else:
                        cand = rnd.sample(qids_all, k)
                    st.session_state["_nclex_candidate_preview"] = cand

                cand = st.session_state.get("_nclex_candidate_preview")
                if cand:
                    st.markdown("**Candidate set (preview only):**")
                    cand_items = [by_id[q] for q in cand if q in by_id][:k]
                    st.write(f"Items: {len(cand_items)}/{k}")
                    for ii, it in enumerate(cand_items, start=1):
                        st.markdown(f"**{ii}. {it.get('id','')}** ({it.get('type','')}, {it.get('difficulty','')})")
                        st.write(it.get('stem') or it.get('prompt') or '')
                        st.write("✅ Correct:", it.get('correct'))
                        if it.get('rationale'):
                            st.write("🧠 Rationale:", it.get('rationale'))
                        st.divider()

            if st.button("🎲 Generate NEW active set now (random)", key="gen_active_set_btn"):
                k = nclex_items_per_case(pol, cid_pick)
                # Choose up to k unique items from the bank
                qids_all = [str(it.get("id","")) for it in bank_items if str(it.get("id","")).strip()]
                rnd_seed = sha256_hex(f"{cid_pick}|{time.time()}|{secrets.token_hex(4)}")
                rnd = random.Random(int(rnd_seed[:8], 16))
                if len(qids_all) <= k:
                    picked = list(qids_all)
                    rnd.shuffle(picked)
                else:
                    picked = rnd.sample(qids_all, k)
                active_sets.setdefault("by_case", {})
                active_sets.setdefault("by_case", {})
                prev = active_sets["by_case"].get(cid_pick)
                if not isinstance(prev, dict):
                    prev = {}
                hist = prev.get("history") if isinstance(prev.get("history"), list) else []
                entry = {
                    "qids": picked,
                    "generated_at": now_local().isoformat(),
                    "generated_by": "admin",
                    "seed": rnd_seed[:16],
                }
                hist.append(entry)
                # keep last 50 generations
                if len(hist) > 50:
                    hist = hist[-50:]
                active_sets["by_case"][cid_pick] = {
                    "qids": picked,
                    "generated_at": entry["generated_at"],
                    "generated_by": entry["generated_by"],
                    "seed": entry["seed"],
                    "history": hist,
                }
                save_nclex_active_sets(active_sets)
                st.success(f"Generated a new active {len(picked)}-item set for {cid_pick}.")


        active_sets = load_nclex_active_sets()

        # --- Rotation history for this case ---
        rec = (active_sets.get("by_case") or {}).get(cid_pick) or {}
        hist = rec.get("history") if isinstance(rec.get("history"), list) else []
        if hist:
            st.markdown("### 🕘 Rotation history (this case)")
            rows = []
            for j, h in enumerate(reversed(hist), start=1):
                rows.append({
                    "#": j,
                    "generated_at": h.get("generated_at", ""),
                    "count": len(h.get("qids") or []),
                    "seed": h.get("seed", ""),
                    "generated_by": h.get("generated_by", ""),
                })
            st.dataframe(rows, width="stretch", height=240)
            # CSV export of history
            try:
                out_csv = io.StringIO()
                w = csv.DictWriter(out_csv, fieldnames=list(rows[0].keys()))
                w.writeheader()
                for r in rows:
                    w.writerow(r)
                st.download_button(
                    "⬇️ Download rotation history (CSV)",
                    data=out_csv.getvalue().encode("utf-8"),
                    file_name=f"rotation_history_{cid_pick}.csv",
                    mime="text/csv",
                )
            except Exception:
                pass
        else:
            st.caption("No rotation history yet for this case.")

# =============================
# Debug panel (always visible to Admin + Students)
# =============================


# =============================
# Analytics dashboard (Step 9)
# =============================
if is_admin and analytics_tab is not None and features.get("analytics_dashboard", True):
    with analytics_tab:
        st.subheader("📊 Analytics (Read-only)")
        st.caption("Read-only summaries from attempts_log.jsonl. No changes to student flow.")

        # Aggregate attempts
        per_case = {}
        per_cohort = {}

        total_attempts = 0
        for rec in iter_attempts():
            total_attempts += 1
            cid = str(rec.get("caseId", ""))
            cohort = str(rec.get("student_cohort", rec.get("cohort", "")) or "")
            scores = rec.get("scores", {}) or {}
            total_score = rec.get("total", None)
            if total_score is None:
                try:
                    total_score = int(scores.get("A", 0)) + int(scores.get("B", 0)) + int(scores.get("C", 0)) + int(scores.get("D", 0)) + int(scores.get("E", 0))
                except Exception:
                    total_score = 0

            per_case.setdefault(cid, {"count": 0, "sum": 0})
            per_case[cid]["count"] += 1
            per_case[cid]["sum"] += int(total_score)

            per_cohort.setdefault(cohort, {"count": 0, "sum": 0})
            per_cohort[cohort]["count"] += 1
            per_cohort[cohort]["sum"] += int(total_score)

        st.write("Total attempts logged:", total_attempts)

        case_rows = []
        for cid, v in sorted(per_case.items(), key=lambda x: x[0]):
            avg = (v["sum"] / v["count"]) if v["count"] else 0
            case_rows.append({"caseId": cid, "attempts": v["count"], "avg_total_score": round(avg, 2)})
        st.subheader("Average score by case")
        st.dataframe(case_rows, width="stretch")

        cohort_rows = []
        for coh, v in sorted(per_cohort.items(), key=lambda x: x[0]):
            avg = (v["sum"] / v["count"]) if v["count"] else 0
            cohort_rows.append({"cohort": coh or "(blank)", "attempts": v["count"], "avg_total_score": round(avg, 2)})
        st.subheader("Average score by cohort")
        st.dataframe(cohort_rows, width="stretch")

# =============================
# Case Editor (Step 10) — minimal safe editor
# =============================
if is_admin and editor_tab is not None:
    with editor_tab:
        st.subheader("🧩 Case Editor (Admin-only)")
        st.caption("Edits are validated and saved back to cases.json only when you click Save.")

        case_ids_all = [str(c.get("id", "")).strip() for c in cases]
        edit_id = st.selectbox("Select case to edit", case_ids_all, index=case_ids_all.index(case_id) if case_id in case_ids_all else 0)
        edit_case = next((c for c in cases if str(c.get("id", "")).strip() == edit_id), None)

        if not edit_case:
            st.error("Case not found.")
        else:
            # -----------------------------
            # Instructor Key (Admin-only)
            # -----------------------------
            with st.expander("📌 Instructor Key (answers + NCLEX key)", expanded=False):
                st.caption("Instructor-only view of the gold standard (A–E + intake) and the active NCLEX set used for students.")

                                # Quick PDF export (Instructor Key)
                if REPORTLAB_AVAILABLE:
                    try:
                        _pol = load_nclex_policy()
                        _k = nclex_items_per_case(_pol, _cid)
                        _nclex_all = load_nclex_items()
                        _cid = str(edit_case.get("id", "")).strip()
                        _pack = (_nclex_all.get("cases") or {}).get(_cid) or {}
                        _bank_items = list(_pack.get("items", []) or [])
                
                        # ---- Export controls
                        st.markdown("### 📄 PDF Export")
                        _sec_labels = ["A–E Key", "Intake Key", "NCLEX Key"]
                        _sec_default = list(_sec_labels)
                        _sec_sel = st.multiselect(
                            "Include sections in PDF",
                            options=_sec_labels,
                            default=_sec_default,
                            key=f"pdf_sections_{_cid or 'case'}",
                        )
                        _sec_map = {"A–E Key": "ae", "Intake Key": "intake", "NCLEX Key": "nclex"}
                        _sec_codes = [_sec_map[s] for s in (_sec_sel or []) if s in _sec_map]
                        if not _sec_codes:
                            _sec_codes = ["ae", "intake", "nclex"]
                
                        _nclex_type_allow = None
                        _export_all_nclex = False
                        if "nclex" in _sec_codes:
                            _type_opts = ["All", "mcq", "sata", "ordered_response", "cloze", "matrix", "evolving_case"]
                            _type_sel = st.multiselect(
                                "NCLEX question types to include",
                                options=_type_opts,
                                default=["All"],
                                key=f"pdf_nclex_types_{_cid or 'case'}",
                                help="Choose specific types to export. Keep 'All' to include every type.",
                            )
                            if _type_sel and "All" not in _type_sel:
                                _nclex_type_allow = [t.strip().lower() for t in _type_sel if str(t).strip()]
                
                            _export_all_nclex = st.checkbox(
                                "Export ALL NCLEX items for this case (ignore items_per_case limit)",
                                value=False,
                                key=f"pdf_all_nclex_{_cid or 'case'}",
                            )
                
                        # ---- Choose NCLEX items: active set first (if exists), else bank
                        _active_sets_local = load_nclex_active_sets()
                        _rec = (_active_sets_local.get("by_case") or {}).get(_cid) or {}
                        _qids = list(_rec.get("qids") or [])
                        _by_id = {str(it.get("id", "")).strip(): it for it in _bank_items if isinstance(it, dict)}
                        if _qids:
                            _use_items = [_by_id[q] for q in _qids if q in _by_id]
                        else:
                            _use_items = list(_bank_items)
                
                        # Respect items_per_case unless admin chooses export-all
                        if not _export_all_nclex:
                            _use_items = _use_items[:_k]
                
                        _pdf_bytes = build_instructor_key_pdf_bytes(
                            edit_case,
                            (edit_case.get("intake_gold") or {}),
                            (edit_case.get("goldStandard") or {}),
                            "",
                            _use_items,
                            sections=_sec_codes,
                            nclex_type_allow=_nclex_type_allow,
                        )
                        st.download_button(
                            "📄 Export Instructor Key as PDF",
                            data=_pdf_bytes,
                            file_name=f"instructor_key_{_cid or 'case'}.pdf",
                            mime="application/pdf",
                            key=f"dl_instr_pdf_{_cid or 'case'}",
                        )
                    except Exception:
                        st.caption("PDF export failed (check ReportLab install).")
                else:
                    st.info("PDF export requires the 'reportlab' package. Install with: python -m pip install reportlab")



                gs = (edit_case.get("goldStandard") or {})
                intake_gold = (edit_case.get("intake_gold") or {})

                def _esc(s):
                    try:
                        s = "" if s is None else str(s)
                    except Exception:
                        s = ""
                    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

                def _render_bullets(section_title: str, items):
                    st.markdown(f"#### {section_title}")
                    if not items:
                        st.info("—")
                        return
                    # Clean + render as bullet list (no Python list indices)
                    clean = [str(x).strip() for x in (items or []) if str(x).strip()]
                    if not clean:
                        st.info("—")
                        return
                    lis = "".join([f"<li>{_esc(x)}</li>" for x in clean])
                    st.markdown(f"<ul style='margin-top:0.25rem; margin-bottom:1rem;'>{lis}</ul>", unsafe_allow_html=True)

                def _render_intake_kv(d: dict):
                    if not d:
                        st.info("No intake_gold found for this case.")
                        return
                    for k, v in d.items():
                        st.markdown(
                            f"<div style='margin:0.15rem 0;'><span class='kv-key'>{_esc(k)}:</span> <span class='kv-val'>{_esc(v)}</span></div>",
                            unsafe_allow_html=True
                        )

                tab1, tab2, tab3 = st.tabs(["A–E Key", "Intake Key", "NCLEX Key (active 30)"])
                with tab1:
                    st.markdown("### A–E Gold Standard")
                    ka = list(gs.get("keyAssessments", []) or gs.get("assessments", []) or [])
                    pr = list(gs.get("priorities", []) or [])
                    iv = list(gs.get("interventions", []) or [])
                    re_ = list(gs.get("reassessment", []) or gs.get("reassess", []) or [])
                    sb_raw = gs.get("sbar") or {}
                    if isinstance(sb_raw, dict):
                        _order = ["S","B","A","R"]
                        sb = []
                        used = set()
                        for k in _order:
                            if k in sb_raw:
                                sb.append(f"{k}: {sb_raw.get(k)}")
                                used.add(k)
                        for k, v in sb_raw.items():
                            if k in used: continue
                            sb.append(f"{k}: {v}")
                    elif isinstance(sb_raw, list):
                        sb = [str(x).strip() for x in sb_raw if str(x).strip()]
                    else:
                        sb = [str(sb_raw).strip()] if str(sb_raw).strip() else []
                    _render_bullets(f"Assessment (select exactly {len(ka)})", ka)
                    _render_bullets(f"Prioritize (select exactly {len(pr)})", pr)
                    _render_bullets(f"Interventions (select exactly {len(iv)})", iv)
                    _render_bullets(f"Reassess (select exactly {len(re_)})", re_)
                    _render_bullets("SBAR key elements", sb)

                with tab2:
                    st.markdown("### Intake Gold (keywords / expected content)")
                    if isinstance(intake_gold, dict):
                        _render_intake_kv(intake_gold)
                    else:
                        st.info("No intake_gold found for this case.")
                with tab3:
                    st.markdown("#### NCLEX Key (active set)")
                    try:
                        nclex_bank = load_nclex_items()
                        pol = load_nclex_policy()
                        pack = (nclex_bank.get("cases") or {}).get(str(edit_case.get("id",""))) or {}
                        bank_items = list(pack.get("items", []) or [])
                        # apply type filters
                        et = pol.get("enabled_types") or {}
                        if isinstance(et, dict) and et:
                            bank_items = [it for it in bank_items if et.get(it.get("type"), True)]
                        k = nclex_items_per_case(pol, cid_pick)
                        active_sets = load_nclex_active_sets()
                        active = (active_sets.get("by_case") or {}).get(str(edit_case.get("id",""))) or {}
                        qids = active.get("qids") if isinstance(active, dict) else None
                        use_items = []
                        if bool(pol.get("rotation_enabled", False)) and isinstance(qids, list) and qids:
                            by_id = {str(it.get("id","")): it for it in bank_items}
                            use_items = [by_id[q] for q in qids if q in by_id][:k]
                            st.info(f"Showing ACTIVE set ({len(use_items)}/{k}).")
                        else:
                            use_items = bank_items[:k]
                            st.info(f"Showing DEFAULT first {len(use_items)}/{k} (no active set).")
                        for ii, it in enumerate(use_items, start=1):
                            qid = str(it.get("id",""))
                            st.markdown(f"**{ii}. {qid}** ({it.get('type','')}, {it.get('difficulty','')})")
                            st.write(it.get("stem") or it.get("prompt") or "")
                            corr = it.get("correct")
                            st.write("✅ **Correct:**", corr)
                            rat = it.get("rationale")
                            if rat:
                                st.write("🧠 **Rationale:**", rat)
                            st.divider()

                    except Exception as e:
                        st.error("Could not render NCLEX key.")
                        st.write(e)

            title = st.text_input("Title", value=str(edit_case.get("title", "")))
            system = st.text_input("System", value=str(edit_case.get("system", edit_case.get("category", "")) or ""))
            setting = st.text_input("Setting", value=str(edit_case.get("setting", "")))
            chief = st.text_area("Chief Complaint", value=str(edit_case.get("chiefComplaint", "")), height=80)

            st.markdown("### Vitals (one per line: Key: Value)")
            _vitals_obj0 = edit_case.get("vitals", {}) or {}
            if isinstance(_vitals_obj0, dict):
                # Stable order for common vital keys
                _vital_order = ["T","Temp","HR","RR","BP","SpO2","O2","Pain","BG"]
                _lines = []
                used = set()
                for k in _vital_order:
                    if k in _vitals_obj0:
                        _lines.append(f"{k}: {_vitals_obj0.get(k)}")
                        used.add(k)
                for k, v in _vitals_obj0.items():
                    if k in used: continue
                    _lines.append(f"{k}: {v}")
                vitals_text = st.text_area("Vitals", value="\n".join(_lines), height=140)
            else:
                vitals_text = st.text_area("Vitals", value=str(_vitals_obj0), height=140)

            st.markdown("### Findings (one per line)")
            findings_list = edit_case.get("findings", []) or []
            findings_text = st.text_area("Findings", value="\n".join([str(x) for x in findings_list]), height=120)

            st.markdown("### Gold Standard Lists (one per line)")
            gs = edit_case.get("goldStandard", {}) or {}
            gs_assess_t = st.text_area("Gold: Assessments", value="\n".join(gs.get("keyAssessments", gs.get("assessment", [])) or []), height=120)
            gs_prio_t = st.text_area("Gold: Priorities", value="\n".join(gs.get("priorities", gs.get("prioritize", [])) or []), height=120)
            gs_inter_t = st.text_area("Gold: Interventions", value="\n".join(gs.get("interventions", []) or []), height=120)
            gs_reass_t = st.text_area("Gold: Reassess", value="\n".join(gs.get("reassessment", gs.get("reassess", [])) or []), height=120)

            st.markdown("### Gold: SBAR (one per line: Key: Value)")
            _sbar0 = (gs.get("sbar") or {})
            if isinstance(_sbar0, dict):
                _order = ["S","B","A","R"]
                _lines = []
                used = set()
                for k in _order:
                    if k in _sbar0:
                        _lines.append(f"{k}: {_sbar0.get(k)}")
                        used.add(k)
                for k, v in _sbar0.items():
                    if k in used: continue
                    _lines.append(f"{k}: {v}")
                sbar_text = st.text_area("Gold SBAR", value="\n".join(_lines), height=160)
            else:
                sbar_text = st.text_area("Gold SBAR", value=str(_sbar0), height=160)

            if st.button("💾 Save changes to cases.json"):
                # Validate Vitals + SBAR fields (line-based, JSON accepted as fallback)
                def _parse_kv_text(txt: str) -> dict:
                    txt = (txt or "").strip()
                    if not txt:
                        return {}
                    # Allow JSON paste as a fallback
                    if txt.lstrip().startswith("{"):
                        try:
                            obj = json.loads(txt)
                            return obj if isinstance(obj, dict) else {}
                        except Exception:
                            pass
                    out = {}
                    for raw in txt.splitlines():
                        line = (raw or "").strip()
                        if not line:
                            continue
                        # allow bullets
                        line = line.lstrip("•-	 ").strip()
                        if ":" not in line:
                            continue
                        k, v = line.split(":", 1)
                        k = k.strip()
                        v = v.strip()
                        if not k:
                            continue
                        out[k] = v
                    return out

                try:
                    vitals_obj = _parse_kv_text(vitals_text)
                    sbar_obj = _parse_kv_text(sbar_text)
                    if not isinstance(vitals_obj, dict):
                        raise ValueError("Vitals must be a dict.")
                    if not isinstance(sbar_obj, dict):
                        raise ValueError("SBAR must be a dict.")
                except Exception as e:
                    st.error(f"Validation error: {e}")
                    st.stop()
                # Backup cases.json before saving (safe)
                backup_file(CASES_PATH)

                # Apply changes
                for c in cases:
                    if str(c.get("id", "")).strip() == edit_id:
                        c["title"] = title
                        c["system"] = system
                        c["setting"] = setting
                        c["chiefComplaint"] = chief
                        c["vitals"] = vitals_obj
                        c["findings"] = [ln.strip() for ln in findings_text.splitlines() if ln.strip()]

                        c.setdefault("goldStandard", {})
                        c["goldStandard"]["keyAssessments"] = [ln.strip() for ln in gs_assess_t.splitlines() if ln.strip()]
                        c["goldStandard"]["priorities"] = [ln.strip() for ln in gs_prio_t.splitlines() if ln.strip()]
                        c["goldStandard"]["interventions"] = [ln.strip() for ln in gs_inter_t.splitlines() if ln.strip()]
                        c["goldStandard"]["reassessment"] = [ln.strip() for ln in gs_reass_t.splitlines() if ln.strip()]
                        c["goldStandard"]["sbar"] = {
                            "S": str(sbar_obj.get("S", "")),
                            "B": str(sbar_obj.get("B", "")),
                            "A": str(sbar_obj.get("A", "")),
                            "R": str(sbar_obj.get("R", "")),
                        }
                        break

                save_cases(cases)
                st.success("Saved cases.json (backup created). Refreshing…")
        

# =============================
# Student View
# =============================
with student_tab:
    # If a new-attempt reset was requested, perform it BEFORE widgets are instantiated.
    if st.session_state.get("_reset_pending", False):
        st.session_state["_reset_pending"] = False
        reset_full_attempt_for_case(case_id, student_username=student_username)
        st.rerun()

    # Start New Attempt (only after the current attempt is finalized, unless admin)
    can_force_reset = bool(is_admin)
    current_finalized = bool(st.session_state.get("nclex_finalized", False))
    max_attempts = max_attempts_for_case(case_id)
    max_attempts_int = None if str(max_attempts) == "unlimited" else safe_int(max_attempts, None)
    attempted = attempts_count_for(student_username, case_id) if student_username else 0

    # Disable reset if attempt still in progress (students must finish current attempt first)
    reset_disabled = False
    reset_help = ""
    if not can_force_reset:
        if not current_finalized:
            reset_disabled = True
            reset_help = "Finish and finalize the current attempt before starting a new one."
        elif max_attempts_int is not None and attempted >= max_attempts_int:
            reset_disabled = True
            reset_help = "Attempt limit reached for this case."

    if reset_help:
        st.sidebar.info(reset_help)

    if st.sidebar.button("🧼 Start New Attempt (Reset Answers)", disabled=reset_disabled):
        st.session_state["_reset_pending"] = True
        st.rerun()

    # Restore last autosave (best-effort) — useful if browser/server closed unexpectedly
    if st.sidebar.button("↩ Restore last autosave"):
        draft = load_last_autosave(student_username, case_id)
        if not draft:
            st.sidebar.warning("No autosave found for this student + case yet.")
        else:
            apply_restored_draft(draft)
            st.sidebar.success("Restored last autosave. Refreshing…")
            st.rerun()


    # Autosave every 60 seconds (safe)
    if "autosave_last_epoch" not in st.session_state:
        st.session_state["autosave_last_epoch"] = time.time()
    if time.time() - st.session_state["autosave_last_epoch"] > 60:
        autosave_draft(features, student_username, case_id, {
            "answers": st.session_state.get("answers", {}),
            "scores": st.session_state.get("scores", {}),
            "nclex_answers": st.session_state.get("nclex_answers", {}),
            "practical_submitted": bool(st.session_state.get("practical_submitted", False)),
            "nclex_finalized": bool(st.session_state.get("nclex_finalized", False)),
            "caseId": case_id,
            "mode": mode,
            "intake": st.session_state.get("intake", {}),
            "intake_score": int(st.session_state.get("intake_score", 0) or 0),
            "intake_breakdown": st.session_state.get("intake_breakdown", {}),
            "intake_submitted": bool(st.session_state.get("intake_submitted", False)),
        })
        st.session_state["autosave_last_epoch"] = time.time()

    # =============================
    # Introductory case video (optional; per case)
    # =============================
    intro_cfg = (admin_settings.get("intro_videos") or {}).get(str(case_id))
    if isinstance(intro_cfg, dict) and intro_cfg.get("path"):
        st.subheader("🎬 Introductory Video")
        vpath = Path(str(intro_cfg.get("path")))
        if vpath.exists():
            try:
                st.video(vpath.read_bytes())
            except Exception:
                # fallback: pass path directly
                st.video(str(vpath))
        else:
            st.warning("Intro video file is missing on the server. Please contact the administrator.")

        if bool(intro_cfg.get("require_watch", True)):
            k_ok = f"intro_video_confirmed__{case_id}"
            if k_ok not in st.session_state:
                st.session_state[k_ok] = False
            st.checkbox("I watched the introductory video.", key=k_ok)
            if not st.session_state.get(k_ok):
                st.info("Watch the video, then check the box to continue to the clinical scenario and activities.")
                st.stop()

    st.markdown("")
    # Clinical Scenario (always visible except NCLEX section)
    if not bool(st.session_state.get("nclex_in_progress", False)):
        
        trigger = build_scenario_trigger(case)
        st.markdown("<div style='background:#e8f4ff;border:1px solid rgba(0,0,0,.08);padding:12px 14px;border-radius:14px;'><span style='color:#b00020;font-weight:900;'>Clinical Scenario (Trigger)</span></div>", unsafe_allow_html=True)
        if trigger:
            # Render in a light-yellow scenario box (keep the text as-is)
            st.markdown(
                f"""<div style='background:#fff7cc;border:1px solid rgba(0,0,0,.08);padding:14px 16px;border-radius:14px;line-height:1.55;font-size:1.02rem;'>{trigger}</div>""",
                unsafe_allow_html=True
            )
        else:
            st.info("Scenario text is unavailable for this case.")

        st.markdown("<div style='height:14px'></div>", unsafe_allow_html=True)


    # =============================
    
    # Helper: light-blue guidance box with alarm icon
    def _guidance_box(msg: str):
        st.markdown(
            f"""
            <div style="background:#e8f4ff;border:1px solid rgba(0,0,0,.08);border-left:6px solid #2b78c5;padding:12px 14px;border-radius:12px;margin:10px 0;">
              <div style="font-weight:800;margin-bottom:6px;">🚨 Guidance</div>
              <div style="line-height:1.55;">{msg}</div>
            </div>
            """,
            unsafe_allow_html=True
        )

# Student Intake (5 marks)
    # =============================
    st.markdown("<div style='background:#e8f4ff;border:1px solid rgba(0,0,0,.08);padding:12px 14px;border-radius:14px;'><span style='color:#b00020;font-weight:900;'>📝 Student Intake</span></div>", unsafe_allow_html=True)
    st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)
    with st.expander("📝 Student Intake (5 marks)", expanded=(bool(is_admin) or (not bool(st.session_state.get("intake_submitted", False))))):
        # Initialize intake dict if missing (safety)
        if "intake" not in st.session_state or not isinstance(st.session_state.get("intake"), dict):
            st.session_state["intake"] = {"age": "", "setting": "", "chief_complaint": "", "signs_symptoms": "", "findings": "", "history": ""}
            st.session_state["intake_score"] = 0
            st.session_state["intake_breakdown"] = {}
            st.session_state["intake_submitted"] = False

        # Use case-specific widget keys to prevent cross-case bleed
        _k = str(case_id)
        intake_final = (not is_admin) and bool(st.session_state.get("intake_submitted", False))

        colA, colB = st.columns(2)
        with colA:
            age_val = st.text_input("Age", value=str(st.session_state["intake"].get("age","")), key=f"intake_age_{_k}", disabled=bool(timer_lock or intake_final))
            setting_val = st.text_input("Setting", value=str(st.session_state["intake"].get("setting","")), key=f"intake_setting_{_k}", disabled=bool(timer_lock or intake_final))
            cc_val = st.text_input("Chief Complaint", value=str(st.session_state["intake"].get("chief_complaint","")), key=f"intake_cc_{_k}", disabled=bool(timer_lock or intake_final))
        with colB:
            sx_val = st.text_area("Major signs & symptoms", value=str(st.session_state["intake"].get("signs_symptoms","")), height=120, key=f"intake_sx_{_k}", disabled=bool(timer_lock or intake_final))
            findings_val = st.text_area("Findings (assessment/exam/labs)", value=str(st.session_state["intake"].get("findings","")), height=120, key=f"intake_findings_{_k}", disabled=bool(timer_lock or intake_final))

        hist_val = st.text_area("History (PMH/meds/allergies/social/HPI)", value=str(st.session_state["intake"].get("history","")), height=120, key=f"intake_hist_{_k}", disabled=bool(timer_lock or intake_final))

        # Persist to session state (so autosave can pick it up)
        st.session_state["intake"] = {
            "age": age_val,
            "setting": setting_val,
            "chief_complaint": cc_val,
            "signs_symptoms": sx_val,
            "findings": findings_val,
            "history": hist_val,
        }

        c1, c2, c3 = st.columns([1,1,2])
        with c1:
            if st.button("💾 Save Intake", key=f"intake_save_{_k}", disabled=bool(timer_lock or intake_final)):
                # lightweight autosave right away
                autosave_draft(features, student_username, case_id, {
                    "answers": st.session_state.get("answers", {}),
                    "scores": st.session_state.get("scores", {}),
                    "nclex_answers": st.session_state.get("nclex_answers", {}),
                    "practical_submitted": bool(st.session_state.get("practical_submitted", False)),
                    "nclex_finalized": bool(st.session_state.get("nclex_finalized", False)),
                    "caseId": case_id,
                    "mode": mode,
                    "intake": st.session_state.get("intake", {}),
                    "intake_score": int(st.session_state.get("intake_score", 0) or 0),
                    "intake_breakdown": st.session_state.get("intake_breakdown", {}),
                    "intake_submitted": bool(st.session_state.get("intake_submitted", False)),
                })
                st.success("Intake saved.")
        with c2:
            if st.button("✅ Submit Intake (Score /5)", key=f"intake_submit_{_k}", disabled=bool(timer_lock or intake_final)):
                s, br = score_intake(case, st.session_state.get("intake", {}))
                st.session_state["intake_score"] = int(s)
                st.session_state["intake_breakdown"] = br or {}
                st.session_state["intake_submitted"] = True

        
                st.session_state["ae_focus"] = "A"
                st.session_state["ae_guidance_done"] = False
                st.rerun()

        if st.session_state.get("intake_submitted"):
            # Score in a yellow box
            _score = int(st.session_state.get('intake_score', 0) or 0)
            st.markdown(
                f"""<div style='background:#fff7cc;border:1px solid rgba(0,0,0,.08);padding:10px 14px;border-radius:12px;display:inline-block;'>
                      <span style='color:#b00020;font-weight:900;'>Intake Score:</span>
                      <span style='font-weight:900;'> {_score} / 5</span>
                    </div>""",
                unsafe_allow_html=True
            )

            br = st.session_state.get("intake_breakdown", {}) or {}
            rows = [
                {"Component": "Age", "Score": f"{br.get('age', 0)}/1"},
                {"Component": "Setting", "Score": f"{br.get('setting', 0)}/1"},
                {"Component": "Chief Complaint", "Score": f"{br.get('chief_complaint', 0)}/1"},
                {"Component": "Signs/Symptoms + Findings", "Score": f"{br.get('signs_symptoms_findings', 0)}/1"},
                {"Component": "History", "Score": f"{br.get('history', 0)}/1"},
            ]

            # Render a light-green table with red/bold headers (Streamlit st.table cannot style headers)
            try:
                header_html = """<tr>
                    <th style='padding:8px 10px;border-bottom:1px solid rgba(0,0,0,.08);color:#b00020;font-weight:900;text-align:left;'>Component</th>
                    <th style='padding:8px 10px;border-bottom:1px solid rgba(0,0,0,.08);color:#b00020;font-weight:900;text-align:left;'>Score</th>
                </tr>"""
                body_html = "".join([
                    f"<tr><td style='padding:8px 10px;border-bottom:1px solid rgba(0,0,0,.06);'>{r['Component']}</td>"
                    f"<td style='padding:8px 10px;border-bottom:1px solid rgba(0,0,0,.06);font-weight:700;'>{r['Score']}</td></tr>"
                    for r in rows
                ])
                st.markdown(
                    f"""<div style='background:#e9f7ef;border:1px solid rgba(0,0,0,.08);padding:10px 12px;border-radius:14px;margin-top:10px;'>
                           <table style='width:100%;border-collapse:collapse;background:#e9f7ef;'>
                             <thead>{header_html}</thead>
                             <tbody>{body_html}</tbody>
                           </table>
                         </div>""",
                    unsafe_allow_html=True
                )
            except Exception:
                st.table(rows)

    init_case_state_if_needed(case_id, case.get('title','') if isinstance(case, dict) else '')
    ensure_section_locks(case_id)
    ensure_widget_defaults()

    # =============================
    # Gate 
    # Guidance screen between Intake submission and A–E workflow (students)
    if (not is_admin) and bool(st.session_state.get("intake_submitted", False)) and (not bool(st.session_state.get("ae_guidance_done", False))):
        _guidance_box(
            "Now you will be forwarded to a nursing activity for this clinical scenario that includes <b>Assessment</b>, <b>Prioritize</b>, <b>Interventions</b>, <b>Reassessment</b>, and <b>SBAR</b>. "
            "You must complete these parts <b>one by one</b>. You <b>cannot jump</b> to another activity until the current activity shows completion. "
            "After you submit each activity, your <b>grade</b>, <b>correct/incorrect selections</b>, <b>rationale</b>, and (if enabled) the <b>AI coach</b> will appear automatically. "
            "When you are ready, click <b>Next</b> to start."
        )
        if st.button("Next ▶", key=f"ae_guidance_next_{case_id}"):
            st.session_state["ae_guidance_done"] = True
            st.session_state["ae_focus"] = "A"
            st.rerun()
        st.stop()

    # A–E until Intake is submitted (students only)
    if (not is_admin) and (not bool(st.session_state.get("intake_submitted", False))):
        st.info("✅ Please complete and **submit** the Intake to unlock the Clinical Reasoning activities (Assessment → SBAR).")
        st.stop()

    # A–E workflow (sequential, vertical layout)
    # =============================
    def _ae_done(dom: str) -> bool:
        fb = (st.session_state.get("last_feedback", {}) or {}).get(dom)
        return fb is not None

    # Students must complete in order; admins can open everything for review.
    allow_free_nav = bool(is_admin)

    unlock_A = True
    unlock_B = allow_free_nav or _ae_done("A")
    unlock_C = allow_free_nav or _ae_done("B")
    unlock_D = allow_free_nav or _ae_done("C")
    unlock_E = allow_free_nav or _ae_done("D")

    def _step_header(title: str, dom: str, unlocked: bool):
        done = _ae_done(dom)
        status = "✅ Completed" if done else ("🔓 Unlocked" if unlocked else "🔒 Locked")
        st.markdown(
            f"""
            <div class='ae-step-card'>
                <div class='ae-step-top'>
                    <div class='ae-step-title'>{title}</div>
                    <div class='ae-step-status'>{status}</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

    # Small CSS for the stacked "icon boxes"
    st.markdown(
        """
        <style>
          .ae-step-card{border:1px solid rgba(0,0,0,.08); border-radius:14px; padding:10px 12px; margin:10px 0 6px 0; background:#ffffff;}
          .ae-step-top{display:flex; align-items:center; justify-content:space-between; gap:10px;}
          .ae-step-title{font-weight:800; font-size:1.05rem;}
          .ae-step-status{font-weight:700; font-size:.95rem; opacity:.9;}
          .ae-locked-note{padding:8px 10px; border-radius:12px; border:1px dashed rgba(0,0,0,.18); background:rgba(0,0,0,.02); margin:6px 0 10px 0;}
                  .ae-next-box{padding:10px 12px; border-radius:14px; border:1px solid rgba(0,0,0,.10); background:rgba(46,134,193,.10); margin:10px 0 6px 0;}
          .ae-next-box b{font-weight:800;}
          .major-title-box{background:#e8f4ff;border:1px solid rgba(0,0,0,.08);border-left:6px solid #2b78c5;padding:10px 12px;border-radius:14px;margin:8px 0 6px 0;font-weight:900;font-size:1.15rem;}

          /* === A–E visual boxes (requested) === */
          .ae-section-header{background:#e8f4ff;border:1px solid rgba(0,0,0,.08);border-radius:14px;padding:10px 12px;margin:8px 0 6px 0;display:flex;align-items:center;gap:10px;}
          .ae-letter-badge{background:#ffd966;border-radius:10px;padding:4px 10px;font-weight:900;color:#ffffff;display:inline-block;min-width:34px;text-align:center;}
          .ae-section-title{font-weight:900;color:#c00000;font-size:1.10rem;line-height:1.2;}
          .ae-review-box{background:#fff2cc;border:1px solid rgba(0,0,0,.10);border-radius:14px;padding:12px 14px;margin:10px 0 8px 0;}
          .ae-ai-box{background:#e8f4ff;border:1px solid rgba(0,0,0,.10);border-radius:14px;padding:12px 14px;margin:10px 0 8px 0;}
          .ae-score-badge{background:#ffd966;border-radius:10px;padding:4px 10px;font-weight:900;color:#c00000;display:inline-block;}
          .ae-review-title{font-weight:900;color:#c00000;margin-top:8px;margin-bottom:4px;}
          .ae-review-text{color:#000000;}
          .ae-review-box ul{margin-top:4px;margin-bottom:8px;}
          .ae-review-box li{color:#000000;}
          .ae-ai-box li{color:#000000;}

</style>
        """,
        unsafe_allow_html=True
    )

    # A–E workflow (sequential, vertical layout)
    # =============================
    def _ae_done(dom: str) -> bool:
        fb = (st.session_state.get("last_feedback", {}) or {}).get(dom)
        return fb is not None

    # Students must complete in order; admins can open everything for review.
    allow_free_nav = bool(is_admin)

    unlock_A = True
    unlock_B = allow_free_nav or _ae_done("A")
    unlock_C = allow_free_nav or _ae_done("B")
    unlock_D = allow_free_nav or _ae_done("C")
    unlock_E = allow_free_nav or _ae_done("D")

    def _step_header(title: str, dom: str, unlocked: bool):
        done = _ae_done(dom)
        status = "✅ Completed" if done else ("🔓 Unlocked" if unlocked else "🔒 Locked")
        st.markdown(
            f"""
            <div class='ae-step-card'>
                <div class='ae-step-top'>
                    <div class='ae-step-title'>{title}</div>
                    <div class='ae-step-status'>{status}</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

    # Small CSS for the stacked "icon boxes"
    st.markdown(
        """
        <style>
          .ae-step-card{border:1px solid rgba(0,0,0,.08); border-radius:14px; padding:10px 12px; margin:10px 0 6px 0; background:#ffffff;}
          .ae-step-top{display:flex; align-items:center; justify-content:space-between; gap:10px;}
          .ae-step-title{font-weight:800; font-size:1.05rem;}
          .ae-step-status{font-weight:700; font-size:.95rem; opacity:.9;}
          .ae-locked-note{padding:8px 10px; border-radius:12px; border:1px dashed rgba(0,0,0,.18); background:rgba(0,0,0,.02); margin:6px 0 10px 0;}
        </style>
        """,
        unsafe_allow_html=True
    )

    
    # Single-activity wizard view for students (prevents screen crowding).
    _focus = st.session_state.get("ae_focus", "A") or "A"
    if _focus not in ["A","B","C","D","E","E_DONE"]:
        _focus = "A"
        st.session_state["ae_focus"] = "A"

    if not allow_free_nav:
        # Keep the student on the current step after submission so they can see
        # grade + review + rationale + AI coach. Only snap focus if the current
        # focus is not unlocked (e.g., deep-link / stale state).
        unlock_map = {"A": unlock_A, "B": unlock_B, "C": unlock_C, "D": unlock_D, "E": unlock_E}
        if not unlock_map.get(_focus, False):
            for k in ["A","B","C","D","E"]:
                if unlock_map.get(k, False) and (not _ae_done(k)):
                    _focus = k
                    st.session_state["ae_focus"] = k
                    break
        titles = {"A":"🅰️ Assessment","B":"🅱️ Prioritize","C":"🅲 Interventions","D":"🅳 Reassess","E":"🅴 SBAR"}
        _step_header(titles.get(_focus,"🅰️ Assessment"), _focus if _focus in titles else "A", unlock_map.get(_focus, True))
    else:
        # Admin overview
        _step_header("🅰️ Assessment", "A", unlock_A)
        _step_header("🅱️ Prioritize", "B", unlock_B)
        _step_header("🅲 Interventions", "C", unlock_C)
        _step_header("🅳 Reassess", "D", unlock_D)
        _step_header("🅴 SBAR", "E", unlock_E)

# AI rules
    def ai_allowed_now() -> bool:
        return (admin_settings.get("app_mode") == "Practice") and bool(admin_settings.get("ai_enabled", False))

    def debrief_allowed_now() -> bool:
        return (admin_settings.get("app_mode") == "Practice") and bool(admin_settings.get("ai_debrief_enabled", False))

    def maybe_run_ai_coach(domain_key: str, student_text: str, matched, missed, unsafe_hits):
        if not ai_allowed_now():
            return
        model = (admin_settings.get("ai_model") or "gpt-5.2").strip() or "gpt-5.2"
        prompt = build_domain_coach_prompt(domain_key, case, student_text, matched, missed, unsafe_hits)
        try:
            with st.spinner("AI coach is generating guidance..."):
                text = openai_responses_call(model, AI_SYSTEM_PROMPT, prompt)
            st.session_state["ai_coach"][domain_key] = text
        except Exception as e:
            st.warning(f"AI coach unavailable: {e}")

    def show_ai_coach(domain_key: str):
        if admin_settings.get("app_mode") == "Exam":
            return
        text = st.session_state.get("ai_coach", {}).get(domain_key)
        fb = st.session_state.get("last_feedback", {}).get(domain_key)
        unsafe_hits = []
        if isinstance(fb, dict):
            unsafe_hits = fb.get("unsafe", []) or []
        elif isinstance(fb, (list, tuple)) and len(fb) >= 4:
            unsafe_hits = fb[3] or []

        if text:
            # Show safety alert only when we have real (non-empty) unsafe hits
            unsafe_hits = [str(u).strip() for u in (unsafe_hits or []) if str(u).strip()]
            if unsafe_hits:
                st.warning("⚠️ Safety alert: potentially unsafe actions detected. Review patient safety first.")
            _ai_html = format_text_with_red_titles(text)
            _ai_html = _strip_md_emphasis_markers(_ai_html)
            st.markdown(
                "<div class='ae-ai-box'>"
                "<div class='ae-review-title'>AI Coach</div>"
                "<div class='ae-review-text'>" + str(_ai_html or "—") + "</div>"
                "</div>",
                unsafe_allow_html=True,
            )

    submit_disabled = locked_out or timer_lock

    if timer_lock:
        st.error("⏱ Exam time is finished. Submissions are locked for students.")

    # A — Assessment
    if allow_free_nav or (st.session_state.get("ae_focus","A")=="A"):
        with st.container():
                st.markdown(render_ae_section_header_html('A','Assessment'), unsafe_allow_html=True)
                st.markdown("<div style='color:#b00020;font-weight:900;'>Select exactly {n} options. Your score is based on correct selections (minus wrong selections), capped at {n}.</div>".format(n=req_A), unsafe_allow_html=True)
                locked_A = False  # per-section locking disabled; timer controls locking
                done_A = _ae_done("A")
                disable_A = submit_disabled or (done_A and (not allow_free_nav))
                if done_A and (not allow_free_nav):
                    st.success("✅ Assessment already submitted. You can review your selections, but resubmission is locked.")
                st.multiselect("Assessment checklist (select exactly {n})".format(n=req_A), ui_assess, key="A_selected", disabled=disable_A)
                st.text_area("Focused assessment notes (optional)", key="A_notes", height=120, disabled=disable_A)

                if st.button("✅ Submit Assessment", disabled=disable_A, key=f"submit_A_{case_id}"):
                    selected = st.session_state.get("A_selected", [])
                    if len(selected) != req_A:
                        st.warning(f"Please select exactly {req_A} options before submitting.")
                        st.stop()

                    notes = st.session_state.get("A_notes", "")
                    student_text = "\n".join(selected) + "\n" + (notes or '')
                    correct_selected, wrong_selected, missed_correct = diff_selected_vs_gold(selected, gold4_assess)
                    unsafe_hits = detect_unsafe("\n".join(selected) + "\n" + (notes or ''))
                    raw = score_select4(selected, gold4_assess)
                    score = apply_unsafe_penalty(raw, unsafe_hits)

                    st.session_state.scores["A"] = max(st.session_state.scores["A"], int(score))
                    st.session_state.answers["A"] = {"selected": selected, "notes": notes}
                    st.session_state["last_feedback"]["A"] = {"mode":"select","score": score, "correct": correct_selected, "wrong": wrong_selected, "missed": missed_correct, "unsafe": unsafe_hits}

                    maybe_run_ai_coach("A", student_text, correct_selected, missed_correct, unsafe_hits)
                    st.rerun()
                fb = st.session_state["last_feedback"].get("A")
                if fb:
                    if isinstance(fb, dict) and fb.get("mode") == "select":
                        st.markdown(render_select_feedback_html("Assessment", "A", fb.get("score",0), fb.get("correct",[]), fb.get("wrong",[]), fb.get("missed",[]), int((st.session_state.get("domain_max",{}) or {}).get("A", 4) or 4), case=case, gold_items=gold4_assess), unsafe_allow_html=True)
                        show_ai_coach("A")
                        if not is_admin:
                            _next_disabled = (not _ae_done("A"))
                            st.markdown("<div class=\'ae-next-box\'>✅ <b>Proceed to Prioritize (B)</b> When you\'re ready, click <b>Next</b> to continue.</div>", unsafe_allow_html=True)
                            if st.button("Next ▶", key=f"ae_next_A_{case_id}", disabled=_next_disabled):
                                st.session_state["ae_focus"] = "B"
                                st.rerun()
                    else:
                        score, matched, missed, unsafe_hits = fb
                        st.markdown(feedback_markdown("Assessment", score, matched, missed, unsafe_hits))
                        show_ai_coach("A")

    # B — Prioritize
    if allow_free_nav or (st.session_state.get("ae_focus","A")=="B"):
        with st.container():
                st.markdown(render_ae_section_header_html('B','Prioritize'), unsafe_allow_html=True)
                st.caption("Select exactly {n} priorities. Any wrong selection gives 0 for this section.".format(n=req_B))
                locked_B = False  # per-section locking disabled; timer controls locking
                done_B = _ae_done("B")
                disable_B = submit_disabled or ((not unlock_B) and (not allow_free_nav)) or (done_B and (not allow_free_nav))
                if done_B and (not allow_free_nav):
                    st.success("✅ Prioritize already submitted. Resubmission is locked.")
                if not unlock_B and not allow_free_nav:
                    st.markdown("<div class='ae-locked-note'>Complete <b>Assessment</b> and submit to unlock <b>Prioritize</b>.</div>", unsafe_allow_html=True)
                st.multiselect("Select priorities (up to 4) (select exactly {n})".format(n=req_B), ui_prio, key="B_selected", disabled=disable_B)
                st.text_area("Brief rationale (optional)", key="B_rationale", height=120, disabled=disable_B)

                if st.button("✅ Submit Priorities", disabled=disable_B, key=f"submit_B_{case_id}"):
                    selected = st.session_state.get("B_selected", [])
                    if len(selected) != req_B:
                        st.warning(f"Please select exactly {req_B} options before submitting.")
                        st.stop()

                    rationale = st.session_state.get("B_rationale", "")
                    student_text = "\n".join(selected) + "\n" + (rationale or "")
                    correct_selected, wrong_selected, missed_correct = diff_selected_vs_gold(selected, gold4_prio)
                    unsafe_hits = detect_unsafe(student_text)
                    raw = score_select4(selected, gold4_prio)
                    score = apply_unsafe_penalty(raw, unsafe_hits)

                    st.session_state.scores["B"] = max(st.session_state.scores["B"], int(score))
                    st.session_state.answers["B"] = {"selected": selected, "rationale": rationale}
                    st.session_state["last_feedback"]["B"] = {"mode":"select","score": score, "correct": correct_selected, "wrong": wrong_selected, "missed": missed_correct, "unsafe": unsafe_hits}

                    maybe_run_ai_coach("B", student_text, correct_selected, missed_correct, unsafe_hits)
                    st.rerun()
                fb = st.session_state["last_feedback"].get("B")
                if fb:
                    if isinstance(fb, dict) and fb.get("mode") == "select":
                        st.markdown(render_select_feedback_html("Prioritize", "B", fb.get("score",0), fb.get("correct",[]), fb.get("wrong",[]), fb.get("missed",[]), int((st.session_state.get("domain_max",{}) or {}).get("B", 4) or 4), case=case, gold_items=gold4_prio), unsafe_allow_html=True)
                        show_ai_coach("B")
                        if not is_admin:
                            _next_disabled = (not _ae_done("B"))
                            st.markdown("<div class=\'ae-next-box\'>✅ <b>Proceed to Interventions (C)</b> When you\'re ready, click <b>Next</b> to continue.</div>", unsafe_allow_html=True)
                            if st.button("Next ▶", key=f"ae_next_B_{case_id}", disabled=_next_disabled):
                                st.session_state["ae_focus"] = "C"
                                st.rerun()
                    else:
                        score, matched, missed, unsafe_hits = fb
                        st.markdown(feedback_markdown("Prioritize", score, matched, missed, unsafe_hits))
                        show_ai_coach("B")

    # C — Interventions
    if allow_free_nav or (st.session_state.get("ae_focus","A")=="C"):
        with st.container():
                st.markdown(render_ae_section_header_html('C','Interventions'), unsafe_allow_html=True)
                st.markdown("<div style='color:#b00020;font-weight:900;'>Select exactly {n} options. Your score is based on correct selections (minus wrong selections), capped at {n}.</div>".format(n=req_B), unsafe_allow_html=True)
                locked_C = False  # per-section locking disabled; timer controls locking
                done_C = _ae_done("C")
                disable_C = submit_disabled or ((not unlock_C) and (not allow_free_nav)) or (done_C and (not allow_free_nav))
                if done_C and (not allow_free_nav):
                    st.success("✅ Interventions already submitted. Resubmission is locked.")
                if not unlock_C and not allow_free_nav:
                    st.markdown("<div class='ae-locked-note'>Complete <b>Prioritize</b> and submit to unlock <b>Interventions</b>.</div>", unsafe_allow_html=True)
                st.multiselect("Interventions checklist", ui_inter, key="C_selected", disabled=disable_C)
                st.text_area("Rationale (optional)", key="C_rationale", height=120)

                if st.button("✅ Submit Interventions", disabled=disable_C, key=f"submit_C_{case_id}"):
                    selected = st.session_state.get("C_selected", [])
                    if len(selected) != req_C:
                        st.warning(f"Please select exactly {req_C} options before submitting.")
                        st.stop()

                    rationale = st.session_state.get("C_rationale", "")
                    student_text = "\n".join(selected) + "\n" + (rationale or '')
                    correct_selected, wrong_selected, missed_correct = diff_selected_vs_gold(selected, gold4_inter)
                    unsafe_hits = detect_unsafe("\n".join(selected) + "\n" + (rationale or ''))
                    raw = score_select4(selected, gold4_inter)
                    score = apply_unsafe_penalty(raw, unsafe_hits)

                    st.session_state.scores["C"] = max(st.session_state.scores["C"], int(score))
                    st.session_state.answers["C"] = {"selected": selected, "rationale": rationale}
                    st.session_state["last_feedback"]["C"] = {"mode":"select","score": score, "correct": correct_selected, "wrong": wrong_selected, "missed": missed_correct, "unsafe": unsafe_hits}

                    maybe_run_ai_coach("C", student_text, correct_selected, missed_correct, unsafe_hits)
                    st.rerun()
                fb = st.session_state["last_feedback"].get("C")
                if fb:
                    if isinstance(fb, dict) and fb.get("mode") == "select":
                        st.markdown(render_select_feedback_html("Interventions", "C", fb.get("score",0), fb.get("correct",[]), fb.get("wrong",[]), fb.get("missed",[]), int((st.session_state.get("domain_max",{}) or {}).get("C", 4) or 4), case=case, gold_items=gold4_inter), unsafe_allow_html=True)
                        show_ai_coach("C")
                        if not is_admin:
                            _next_disabled = (not _ae_done("C"))
                            st.markdown("<div class=\'ae-next-box\'>✅ <b>Proceed to Reassess (D)</b> When you\'re ready, click <b>Next</b> to continue.</div>", unsafe_allow_html=True)
                            if st.button("Next ▶", key=f"ae_next_C_{case_id}", disabled=_next_disabled):
                                st.session_state["ae_focus"] = "D"
                                st.rerun()
                    else:
                        score, matched, missed, unsafe_hits = fb
                        st.markdown(feedback_markdown("Interventions", score, matched, missed, unsafe_hits))
                        show_ai_coach("C")

    # D — Reassess
    if allow_free_nav or (st.session_state.get("ae_focus","A")=="D"):
        with st.container():
                st.markdown(render_ae_section_header_html('D','Reassess'), unsafe_allow_html=True)
                st.markdown("<div style='color:#b00020;font-weight:900;'>Select exactly {n} options. Your score is based on correct selections (minus wrong selections), capped at {n}.</div>".format(n=req_C), unsafe_allow_html=True)
                locked_D = False  # per-section locking disabled; timer controls locking
                done_D = _ae_done("D")
                disable_D = submit_disabled or ((not unlock_D) and (not allow_free_nav)) or (done_D and (not allow_free_nav))
                if done_D and (not allow_free_nav):
                    st.success("✅ Reassess already submitted. Resubmission is locked.")
                timing_options = ["Immediately", "5 minutes", "15 minutes", "30 minutes", "1 hour", "4 hours", "PRN / as needed"]
                if not unlock_D and not allow_free_nav:
                    st.markdown("<div class='ae-locked-note'>Complete <b>Interventions</b> and submit to unlock <b>Reassess</b>.</div>", unsafe_allow_html=True)
                st.selectbox("When will you reassess?", timing_options, key="D_time", disabled=(submit_disabled or (not unlock_D)))
                if not ui_reass:
                    st.info("No preset reassessment options found for this case. Enter your reassessment notes below and submit.")
                st.multiselect("What will you monitor?", ui_reass, key="D_selected", disabled=disable_D)
                st.text_area("Reassessment notes (optional)", key="D_notes", height=120, disabled=(submit_disabled or (not unlock_D)))

                if st.button("✅ Submit Reassess Plan", disabled=(submit_disabled or (not unlock_D)), key=f"submit_D_{case_id}"):
                    selected = st.session_state.get("D_selected", [])
                    if req_D > 0 and len(selected) != req_D:
                        st.warning(f"Please select exactly {req_D} options before submitting.")
                        st.stop()
                    if req_D == 0 and not (selected or (st.session_state.get('D_notes','') or '').strip() or (st.session_state.get('D_time','') or '').strip()):
                        st.warning("Please enter reassessment notes (or select options) before submitting.")
                        st.stop()

                    notes = st.session_state.get("D_notes", "")
                    timing = st.session_state.get("D_time", "")
                    student_text = "\n".join(selected) + "\n" + (notes or '') + (timing or '')
                    correct_selected, wrong_selected, missed_correct = diff_selected_vs_gold(selected, gold4_reass)
                    unsafe_hits = detect_unsafe("\n".join(selected) + "\n" + (notes or '') + (timing or ''))
                    raw = score_select4(selected, gold4_reass)
                    score = apply_unsafe_penalty(raw, unsafe_hits)

                    st.session_state.scores["D"] = max(st.session_state.scores["D"], int(score))
                    st.session_state.answers["D"] = {"selected": selected, "notes": notes, "timing": timing}
                    st.session_state["last_feedback"]["D"] = {"mode":"select","score": score, "correct": correct_selected, "wrong": wrong_selected, "missed": missed_correct, "unsafe": unsafe_hits}

                    maybe_run_ai_coach("D", student_text, correct_selected, missed_correct, unsafe_hits)
                    st.rerun()
                fb = st.session_state["last_feedback"].get("D")
                if fb:
                    if isinstance(fb, dict) and fb.get("mode") == "select":
                        st.markdown(render_select_feedback_html("Reassess", "D", fb.get("score",0), fb.get("correct",[]), fb.get("wrong",[]), fb.get("missed",[]), int((st.session_state.get("domain_max",{}) or {}).get("D", 4) or 4), case=case, gold_items=gold4_reass), unsafe_allow_html=True)
                        show_ai_coach("D")
                        if not is_admin:
                            _next_disabled = (not _ae_done("D"))
                            st.markdown("<div class=\'ae-next-box\'>✅ <b>Proceed to SBAR (E)</b> When you\'re ready, click <b>Next</b> to continue.</div>", unsafe_allow_html=True)
                            if st.button("Next ▶", key=f"ae_next_D_{case_id}", disabled=_next_disabled):
                                st.session_state["ae_focus"] = "E"
                                st.rerun()
                    else:
                        score, matched, missed, unsafe_hits = fb
                        st.markdown(feedback_markdown("Reassess", score, matched, missed, unsafe_hits))
                        show_ai_coach("D")

    # E — SBAR
    if allow_free_nav or (st.session_state.get("ae_focus","A")=="E"):
        with st.container():
                st.subheader("SBAR (Structured)")
                st.markdown("<div style='color:#b00020;font-weight:900;'>Select exactly {n} options. Your score is based on correct selections (minus wrong selections), capped at {n}.</div>".format(n=req_E), unsafe_allow_html=True)
                locked_E = False  # per-section locking disabled; timer controls locking
                done_E = _ae_done("E")
                disable_E = submit_disabled or ((not unlock_E) and (not allow_free_nav)) or (done_E and (not allow_free_nav))
                if done_E and (not allow_free_nav):
                    st.success("✅ SBAR already submitted. Resubmission is locked.")
                if not unlock_E and not allow_free_nav:
                    st.markdown("<div class='ae-locked-note'>Complete <b>Reassess</b> and submit to unlock <b>SBAR</b>.</div>", unsafe_allow_html=True)

                # --- SBAR (E) visual wrapper (UI-only; no logic changes) ---
                st.markdown(
                    "<div class='sbar-box'>"
                    "<div style='display:flex;align-items:center;gap:8px;margin-bottom:10px;'>"
                    "<span class='sbar-e-badge'>E</span>"
                    "<span style='color:#b00020;font-weight:900;'>SBAR (Structured)</span>"
                    "</div>",
                    unsafe_allow_html=True
                )

                st.markdown("<div class='sbar-title'>S – Situation</div>", unsafe_allow_html=True)
                st.text_area("", key="E_S", disabled=(submit_disabled or (not unlock_E)), height=100, label_visibility="collapsed")

                st.markdown("<div class='sbar-title'>B – Background</div>", unsafe_allow_html=True)
                st.text_area("", key="E_B", disabled=(submit_disabled or (not unlock_E)), height=100, label_visibility="collapsed")

                st.markdown("<div class='sbar-title'>A – Assessment</div>", unsafe_allow_html=True)
                st.text_area("", key="E_A", disabled=(submit_disabled or (not unlock_E)), height=100, label_visibility="collapsed")

                st.markdown("<div class='sbar-title'>R – Recommendation</div>", unsafe_allow_html=True)
                st.text_area("", key="E_R", disabled=(submit_disabled or (not unlock_E)), height=100, label_visibility="collapsed")

                st.caption("SBAR grading: select up to 4 key elements (gold standard = 4). Any wrong selection gives 0.")

                st.markdown("<div class='sbar-key'>SBAR key elements (up to 4)</div>", unsafe_allow_html=True)
                st.multiselect("", ui_sbar_opts, key="E_selected", disabled=disable_E, label_visibility="collapsed")

                st.markdown("</div>", unsafe_allow_html=True)

                if st.button("✅ Submit SBAR", disabled=disable_E, key=f"submit_E_{case_id}"):
                    S = st.session_state.get("E_S", "")
                    B = st.session_state.get("E_B", "")
                    A = st.session_state.get("E_A", "")
                    R = st.session_state.get("E_R", "")
                    selected = st.session_state.get("E_selected", [])
                    student_text = "S: " + S + "\nB: " + B + "\nA: " + A + "\nR: " + R + "\n\nSelected elements:\n" + "\n".join(selected)
                    correct_selected, wrong_selected, missed_correct = diff_selected_vs_gold(selected, gold4_sbar)
                    unsafe_hits = detect_unsafe(student_text)
                    raw = score_select4(selected, gold4_sbar)
                    score = apply_unsafe_penalty(raw, unsafe_hits)

                    st.session_state.scores["E"] = max(st.session_state.scores["E"], int(score))
                    st.session_state.answers["E"] = {"S": S, "B": B, "A": A, "R": R, "selected_elements": selected}
                    st.session_state["last_feedback"]["E"] = {"mode":"select","score": score, "correct": correct_selected, "wrong": wrong_selected, "missed": missed_correct, "unsafe": unsafe_hits}

                    maybe_run_ai_coach("E", student_text, correct_selected, missed_correct, unsafe_hits)
                    st.session_state["ae_focus"] = "E"
                    st.rerun()
                fb = st.session_state["last_feedback"].get("E")
                if fb:
                    if isinstance(fb, dict) and fb.get("mode") == "select":
                        st.markdown(render_select_feedback_html("SBAR", "E", fb.get("score",0), fb.get("correct",[]), fb.get("wrong",[]), fb.get("missed",[]), int((st.session_state.get("domain_max",{}) or {}).get("E", 4) or 4), case=case, gold_items=gold4_sbar, sbar_expected=sbar_expected), unsafe_allow_html=True)
                        show_ai_coach("E")
                        if not is_admin:
                            _next_disabled = (not _ae_done("E"))
                            st.markdown("<div class=\'ae-next-box\'>✅ <b>Proceed to Summary</b> When you\'re ready, click <b>Next</b> to continue.</div>", unsafe_allow_html=True)
                            if st.button("Next ▶", key=f"ae_next_E_{case_id}", disabled=_next_disabled):
                                st.session_state["ae_focus"] = "E_DONE"
                                st.rerun()
                    else:
                        score, matched, missed, unsafe_hits = fb
                        st.markdown(feedback_markdown("SBAR", score, matched, missed, unsafe_hits))
                        show_ai_coach("E")


    # Keep student view uncluttered: show totals/debrief only after SBAR is completed and the student clicks Next.
    if (not is_admin):
        _completed_all = all((st.session_state.get('last_feedback', {}) or {}).get(k) is not None for k in ['A','B','C','D','E'])
        if (not _completed_all) or (st.session_state.get('ae_focus') != 'E_DONE'):
            st.stop()

    # End-of-case AI Debrief
    st.divider()
    total = sum(st.session_state.scores.values())
    intake_pts = int(st.session_state.get("intake_score", 0) or 0)
    domain_max = st.session_state.get("domain_max", {}) or {}
    total_max = int(sum([int(domain_max.get(k, 0) or 0) for k in ["A","B","C","D","E"]]) or 20)
    overall_max = int(total_max + 5)
    # =============================
    # Score summary (styled)
    # =============================
    _kpi_box = """<div style="background:#fff7cc;border:1px solid rgba(0,0,0,.08);padding:14px 16px;border-radius:14px;margin:10px 0;display:inline-block;min-width:260px;">
      <span style="color:#b00020;font-weight:900;font-size:2.1rem;">{label}: {num} / {den}</span>
    </div>"""
    st.markdown(_kpi_box.format(label="Score", num=_html.escape(str(total)), den=_html.escape(str(total_max))), unsafe_allow_html=True)
    st.markdown(_kpi_box.format(label="Intake", num=_html.escape(str(intake_pts)), den="5"), unsafe_allow_html=True)
    st.markdown(_kpi_box.format(label="Overall", num=_html.escape(str(total + intake_pts)), den=_html.escape(str(overall_max))), unsafe_allow_html=True)

    # Friendly domain score display (light green table; no index column)
    _rows = []
    for k, label in [("A","Assessment"),("B","Prioritize"),("C","Interventions"),("D","Reassess"),("E","SBAR")]:
        mx = int(domain_max.get(k, 4) or 4)
        sc = int((st.session_state.get("scores", {}) or {}).get(k, 0) or 0)
        _rows.append((f"{k} — {label}", f"{sc}/{mx}"))

    _table_rows_html = ""
    for d, s in _rows:
        _table_rows_html += (
            "<tr>"
            f"<td style='padding:10px 12px;border-bottom:1px solid rgba(0,0,0,.08);'>{_html.escape(str(d))}</td>"
            f"<td style='padding:10px 12px;border-bottom:1px solid rgba(0,0,0,.08);'>{_html.escape(str(s))}</td>"
            "</tr>"
        )

    _table_html = (
        "<div style='background:#dff5e1;border:1px solid rgba(0,0,0,.08);padding:10px 12px;border-radius:12px;margin:8px 0 14px 0;'>"
        "<table style='width:100%;border-collapse:collapse;'>"
        "<thead>"
        "<tr>"
        "<th style='text-align:left;padding:10px 12px;color:#b00020;font-weight:900;border-bottom:1px solid rgba(0,0,0,.12);'>Domain</th>"
        "<th style='text-align:left;padding:10px 12px;color:#b00020;font-weight:900;border-bottom:1px solid rgba(0,0,0,.12);'>Score</th>"
        "</tr>"
        "</thead>"
        "<tbody>"
        + _table_rows_html +
        "</tbody>"
        "</table>"
        "</div>"
    )
    st.markdown(_table_html, unsafe_allow_html=True)

    completed = all(st.session_state["last_feedback"].get(k) is not None for k in ["A", "B", "C", "D", "E"])

    if debrief_allowed_now():
        if completed:
            if st.session_state.get("ai_debrief") is None:
                # Show a simple guidance step before generating the AI debrief (Practice only)
                if not bool(st.session_state.get("ai_debrief_requested", False)):
                    _guidance_box(
                        "To see your <b>AI debrief</b> (strengths, weaknesses, and what to focus on next), click <b>Next</b>."
                    )
                    if st.button("Next ▶", key=f"debrief_next_{case_id}"):
                        st.session_state["ai_debrief_requested"] = True
                        st.rerun()
                    st.stop()

                missed_by_domain = {}
                unsafe_by_domain = {}
                for key, name in [("A", "Assessment"), ("B", "Prioritize"), ("C", "Interventions"), ("D", "Reassess"), ("E", "SBAR")]:
                    fb = st.session_state["last_feedback"].get(key)
                    if fb:
                        missed = []
                        unsafe_hits = []
                        if isinstance(fb, dict):
                            missed = fb.get('missed', []) or []
                            unsafe_hits = fb.get('unsafe', []) or []
                        elif isinstance(fb, (list, tuple)) and len(fb) >= 4:
                            missed = fb[2] or []
                            unsafe_hits = fb[3] or []
                        missed_by_domain[name] = list(missed)[:4]
                        unsafe_by_domain[name] = list(unsafe_hits)[:4]
                    else:
                        missed_by_domain[name] = []
                        unsafe_by_domain[name] = []

                prompt = build_debrief_prompt(case, st.session_state.scores, missed_by_domain, unsafe_by_domain)
                model = (admin_settings.get("ai_model") or "gpt-5.2").strip() or "gpt-5.2"
                try:
                    with st.spinner("AI is generating your end-of-case debrief..."):
                        text = openai_responses_call(model, AI_SYSTEM_PROMPT, prompt)
                    st.session_state["ai_debrief"] = text
                except Exception as e:
                    st.warning(f"AI debrief unavailable: {e}")

            if st.session_state.get("ai_debrief"):
                st.markdown('🧠 <span class="nr-title">AI End-of-Case Debrief</span> (coaching, not answers):', unsafe_allow_html=True)
                st.markdown("<div class='ai-debrief-box'>" + (render_ai_debrief_html(st.session_state.get('ai_debrief','')) or '—') + "</div>", unsafe_allow_html=True)
        else:
            st.info("To get AI debrief, submit all 5 domains (A–E) at least once.")
    else:
        if admin_settings.get("app_mode") == "Exam":
            st.info("Exam mode: AI debrief is disabled.")

    # =============================
    # NCLEX-style practice should appear AFTER A–E completed
    # =============================
    if completed:
        # Gate NCLEX section behind a guidance screen + Next button
        if "nr_nclex_ready" not in st.session_state:
            st.session_state["nr_nclex_ready"] = False
        if not st.session_state.get("nr_nclex_ready", False):
            st.session_state["nclex_in_progress"] = False
            st.markdown(
                """
                <div style="background:#eaf4ff;border:1px solid rgba(11,34,51,0.18);padding:12px;border-radius:12px;">
                  <div style="font-weight:900;color:#0b2233;font-size:1.02rem;">🚨 NCLEX-style practice/exam — before you start</div>
                  <div style="margin-top:6px;line-height:1.55;color:#0b2233;font-size:0.98rem;">
                    You are now moving to the NCLEX-style section. Read each question carefully:
                    <ul style="margin:6px 0 0 18px;">
                      <li><b>MCQ:</b> choose one best answer.</li>
                      <li><b>SATA:</b> select all that apply.</li>
                      <li><b>Cloze/Drop-down:</b> choose the best option for each blank.</li>
                    </ul>
                    Your answers are saved while you work. When finished, click <b>Submit & Finalize</b> to lock the attempt, then click <b>Show Review</b> to view your report (per policy settings).
                  </div>
                </div>
                """,
                unsafe_allow_html=True
            )
            if st.button("▶ Next", key=f"nr_next_to_nclex_{case_id}"):
                st.session_state["nr_nclex_ready"] = True
                st.rerun()
        else:
            st.session_state["nclex_in_progress"] = True
            nclex_policy = load_nclex_policy()
            nclex_items = load_nclex_items()
            render_nclex_practical(case_id, nclex_policy, nclex_items, features, mode, timer_lock)
    else:
        st.session_state["nclex_in_progress"] = False
        st.info("🧾 NCLEX-style practice will appear after you submit all 5 domains (A–E) at least once.")

    # =============================
    # Save Attempt (includes NCLEX)
    # =============================
    def build_attempt_record():
        started = st.session_state.get("attempt_started_epoch", time.time())
        duration_sec = int(time.time() - started)

        # NCLEX-style practice score blob (if available)
        nclex_blob = st.session_state.get("nclex_scored") or None
        nclex_answers_blob = st.session_state.get("nclex_answers") or {}

        # KPI policy (controls what we store for research/teaching)
        kpi = load_kpi_policy()

        # Unsafe flags count (per domain)
        unsafe_counts = {}
        for dom in ["A", "B", "C", "D", "E"]:
            fb = st.session_state.get("last_feedback", {}).get(dom)

            # last_feedback can be either:
            # - dict (newer "select" mode): {"unsafe": [...], ...}
            # - tuple/list (older mode): (score, matched, missed, unsafe_hits)
            unsafe_hits = []
            if isinstance(fb, dict):
                unsafe_hits = fb.get("unsafe", []) or []
            elif isinstance(fb, (list, tuple)) and len(fb) >= 4:
                unsafe_hits = fb[3] or []

            unsafe_counts[dom] = len(unsafe_hits) if isinstance(unsafe_hits, list) else 0
        unsafe_total = sum(unsafe_counts.values())

        # Attempts per case (based on existing log)
        attempts_for_case = attempts_count_for(student_username, case_id) if student_username else None

        rec = {
            "timestamp": utc_now_iso(),
            "mode": admin_settings.get("app_mode"),
            "timer_minutes": timer_minutes,
            "timer_expired": bool(expired),
            "student_username": student_username,
            "student_display_name": student_display_name,
            "student_id": student_id,
            "student_cohort": student_profile.get("cohort", "") if student_profile else "",
            "caseId": case.get("id"),
            "caseTitle": case.get("title"),
            "system": safe_get_system(case),
            "setting": safe_get_setting(case),
            "ai_enabled": bool(admin_settings.get("ai_enabled")) and (admin_settings.get("app_mode") == "Practice"),
            "ai_debrief_enabled": bool(admin_settings.get("ai_debrief_enabled")) and (admin_settings.get("app_mode") == "Practice"),
        }

        # Store KPIs depending on policy
        if kpi.get("time_to_completion", True):
            rec["duration_seconds"] = duration_sec

        if kpi.get("domain_profile", True):
            rec["scores"] = st.session_state.scores

        if kpi.get("total_score", True):
            rec["total"] = total
            rec["intake_score"] = int(st.session_state.get("intake_score", 0) or 0)
            rec["total_with_intake"] = int(total) + int(st.session_state.get("intake_score", 0) or 0)
            rec["intake"] = st.session_state.get("intake", {})

        # Always keep raw answers (needed for grading/review); you can turn this off later if desired
        rec["answers"] = st.session_state.answers

        if kpi.get("nclex_score", True):
            rec["nclex"] = nclex_blob
            rec["nclex_answers"] = nclex_answers_blob

        if kpi.get("unsafe_flags", True):
            rec["unsafe_counts"] = unsafe_counts
            rec["unsafe_total"] = unsafe_total

        if kpi.get("attempts_per_case", True):
            rec["attempt_number_for_case"] = attempts_for_case

        # --- Research payload (optional; does not affect teaching flow) ---

        try:

            rp = load_research_policy()

            if should_collect_research():

                research = {}

                research["consented"] = bool(st.session_state.get("research_consent", False))

                research["reflection"] = str(st.session_state.get("research_reflection","")).strip()

                research["reflection_word_count"] = len(research["reflection"].split()) if research["reflection"] else 0


                if bool(rp.get("collect_answer_change", True)):

                    research["nclex_track"] = st.session_state.get("nclex_track", {})


                if bool(rp.get("collect_section_performance", True)):

                    research["performance_by_section"] = {

                        "Intake": {"score": int(st.session_state.get("intake_score", 0) or 0)},

                        "A": {"score": int(st.session_state.get("scores", {}).get("A", 0) if isinstance(st.session_state.get("scores", {}), dict) else 0)},

                        "B": {"score": int(st.session_state.get("scores", {}).get("B", 0) if isinstance(st.session_state.get("scores", {}), dict) else 0)},

                        "C": {"score": int(st.session_state.get("scores", {}).get("C", 0) if isinstance(st.session_state.get("scores", {}), dict) else 0)},

                        "D": {"score": int(st.session_state.get("scores", {}).get("D", 0) if isinstance(st.session_state.get("scores", {}), dict) else 0)},

                        "E": {"score": int(st.session_state.get("scores", {}).get("E", 0) if isinstance(st.session_state.get("scores", {}), dict) else 0)},

                        "NCLEX": {"score": int(st.session_state.get("nclex_score", 0) or 0), "total": int(st.session_state.get("nclex_total", 0) or 0)}

                    }


                rec["research"] = research


                # Separate research log stream (append-only)

                try:

                    student_username = rec.get("student_username","")

                    if bool(rp.get("anonymize_student_id", True)):

                        student_username = _anon_student(student_username)

                                        # Build de-identified research dataset row (separate from teaching attempts)
                    try:
                        if should_collect_research():
                            pid_src = str(rec.get('student_id','') or rec.get('student_username','') or '')
                            participant_id = _hash_participant(pid_src) if bool(rp.get('anonymize_student_id', True)) else pid_src
                            research_row = {
                                'participant_id': participant_id,
                                'submitted_at': rec.get('timestamp',''),
                                'caseId': rec.get('caseId',''),
                                'caseTitle': rec.get('caseTitle',''),
                                'mode': rec.get('mode',''),
                                'cohort': rec.get('cohort',''),
                                'total_score': rec.get('total_score', None),
                                'total_with_intake': rec.get('total_with_intake', None),
                                'nclex_score': rec.get('nclex_score', None),
                                'nclex_total': rec.get('nclex_total', None),
                                'domain_scores': rec.get('scores', {}),
                                'intake_score': rec.get('intake_score', None),
                                'duration_seconds': rec.get('duration_seconds', None),
                            }
                            # Optional fields
                            if bool(rp.get('collect_answer_change', True)):
                                research_row['nclex_changes'] = rec.get('nclex_changes', None)
                            if bool(rp.get('collect_section_performance', True)):
                                research_row['performance_by_section'] = rec.get('performance_by_section', None)
                            if bool(rp.get('collect_reflection', False)):
                                research_row['reflection'] = str(st.session_state.get('research_reflection','') or '')
                            append_research_dataset_row(research_row)
                    except Exception:
                        pass

                    append_research_event({

                        "event": "attempt_saved",

                        "ts": now_local().isoformat(),

                        "student": student_username,

                        "caseId": rec.get("caseId",""),

                        "mode": rec.get("mode",""),

                        "total_score": rec.get("total_score", None),

                        "nclex_score": rec.get("nclex_score", None),

                    })

                except Exception:

                    pass

        except Exception:

            pass

        return rec

    # Step 7: auto-submit on expiry (optional)
    if features.get("auto_submit_on_expiry", False) and mode == "Exam" and expired and is_student_logged_in and not is_admin:
        key = f"{student_username}::{case_id}"
        autosub = st.session_state.get("autosubmitted_case", {})
        if not autosub.get(key, False):
            st.warning("⏱ Time expired — auto-saving attempt now.")
            save_attempt(build_attempt_record_from_state(case_id))
            autosub[key] = True
            st.session_state["autosubmitted_case"] = autosub
            st.success("Auto-saved to attempts_log.jsonl")
            st.rerun()

# Research reflection (optional, research-only)
try:
    rp = load_research_policy()
    if should_collect_research() and bool(rp.get("collect_reflection", False)):
        st.markdown("### 📝 Reflection (research)")
        st.session_state["research_reflection"] = st.text_area(
            "Optional: briefly describe your reasoning and what you learned (2–4 sentences).",
            value=str(st.session_state.get("research_reflection","")),
            height=120,
            key="research_reflection_text"
        )
except Exception:
    pass